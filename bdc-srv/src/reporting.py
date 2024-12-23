import logging
import dataclasses
import json
from werkzeug.utils import secure_filename
import os
from io import BytesIO
import zipfile
import tempfile

import jsonpickle # serialize objects to json
import filetype # figure out the type of image files looking at content

import numpy as np
import cv2
import doc2pdf

import os
from io import BytesIO

from sqlalchemy.orm import Session
from sqlalchemy import select
from models import Certificate, Measurement, MeasurementAnnotationFile, MeasurementImageFile, VTShot, VTShotImageFile, ValidatedMeasurementAnnotationFile, db, Inspection, Image, ImageFile, Blade

import xlsxwriter
import uuid

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

import plotly.express as px
import pandas as pd
import matplotlib.pyplot as plt
from dateutil.parser import parse

from datetime import datetime, timedelta, timezone
from datetime import date
import csv

from PIL import Image as PImage
from PIL import ImageDraw


# defect specific configuration
from config import defect_bgr_colors
from config import defect_severity_color
from config import defect_severity

from concurrent.futures import ThreadPoolExecutor, wait
THREAD_POOL_SIZE = 50

logger = logging.getLogger()

COLOR_SKY = "59C3C9"
COLOR_DARK_SKY = "439297"
COLOR_EVER_GREEN = "005E60"




def add_heading(doc, text, level):
    # heading is a Paragraph
    heading = doc.add_heading('', level=level)
    run = heading.add_run(text)
    run.font.color.rgb = RGBColor.from_string(COLOR_EVER_GREEN)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading





def add_certificate_report_header(doc):
    header = doc.sections[0].header
    logo_path = os.path.abspath('./src/report/GE-Vernova-Emblem.png')
    banner_path = os.path.abspath('./src/report/certificate_title.png')
    heading_table = header.add_table(1, 2, Inches(7))
    heading_cells = heading_table.rows[0].cells
    set_table_cell_bg_color(heading_cells[0], COLOR_EVER_GREEN)
    set_table_cell_bg_color(heading_cells[1], COLOR_EVER_GREEN)

    heading_cells[0].width = Inches(6)
    heading_cells[0].text = "\n\n"
    run0 = heading_cells[0].paragraphs[0].runs[0]
    run0.font.size = Pt(5)
    run0.add_picture(banner_path, width=Inches(3))
    heading_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # heading_cells[0].width = Inches(5)
    # heading_cells[0].text = f'Blade Quality Certificate'
    # run0 = heading_cells[0].paragraphs[0].runs[0]
    # run0.font.bold = True
    # run0.font.color.rgb = RGBColor(255, 255, 255)
    # run0.font.name = "Sons Condensed"
    # run0.font.size = Pt(24)
    # # run0.font.name = "Arial Nova"
    # # run0.font.size = Pt(14)
    # heading_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    heading_cells[1].width = Inches(1)
    heading_cells[1].text = "\n\n\n"
    # cell1_para = heading_cells[1].add_paragraph()
    # run1 = cell1_para.add_run()
    run1 = heading_cells[1].paragraphs[0].runs[0]
    run1.font.size = Pt(5)
    run1.add_picture(logo_path, width=Inches(1.25))
    heading_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    header.is_linked_to_previous = False  # only on front page


def add_certificate_report_footer(doc, footer_text=''):
    footer = doc.sections[0].footer
    paragraph = footer.paragraphs[0]
    run = paragraph.add_run(footer_text)
    run.font.color.rgb = RGBColor.from_string(COLOR_EVER_GREEN)
    run.font.size = Pt(10)
    # paragraph.style = document.styles["Footer"]


def get_defect_bgr_color(label, is_ai=False):
    color = (0, 234, 255)

    if is_ai:
        #color = (245,66,245) # BGR bright pink
        color = (0,0,255) # BGR red 
        return color

    if label is not None:
        label_str = label.replace("/", "-")
        if label_str in defect_bgr_colors:
            color = defect_bgr_colors[label_str]
    return color


def get_defect_severity(defect_name):
    if defect_name == '':
        return 'NONE'
    label = defect_name.replace('_', ' ')
    return defect_severity.get(label, 'NONE')


def get_defect_severity_color(severity):
    if severity == '' or severity is None:
        return 'green'
    return defect_severity_color.get(severity, 'green')


def get_disposition_color(disp):
    if disp is not None:
        if 'No Repair' in disp:
            return '#D3D3D3'  # LightGrey
        elif 'Out of' in disp:
            return '#90EE90'  # LightGreen
        elif 'AI' in disp:
            return '#FF00FF'  # Magenta
        elif 'Pending' in disp:
            return '#FFA07A'  # LightSalmon
    return '#D3D3D3'  # LightGrey matches '' as well

# shade should be in RGB e.g. "005E60"


def set_table_cell_bg_color(cell, shade):
    # cell = tbl.cell(0, 0)
    tcPr = cell._tc.get_or_add_tcPr()
    tcVAlign = OxmlElement('w:shd')
    tcVAlign.set(qn('w:fill'), shade)
    tcPr.append(tcVAlign)


def set_table_cell_bg_color_ALTERNATIVE(cell, shade):
    xml_fragment = r'<w:shd {} w:fill="'+shade+r'"/>'
    shading_elm = parse_xml(xml_fragment.format(nsdecls('w')))
    cell._tc.get_or_add_tcPr().append(shading_elm)


def get_section_name(sect):
    # logging.info(f'get_section_name called for: {sect}')
    resp = sect
    if sect is not None:
        if 'cw' in sect or 'cent' in sect.lower():
            resp = 'Center Web'
        elif 'le' in sect or 'leading' in sect.lower():
            resp = 'Leading Edge'
        elif 'te' in sect or 'trailing' in sect.lower():
            resp = 'Trailing Edge'
        elif 'cs' in sect or 'stiff' in sect.lower():
            resp = 'C Stiffener'
        elif 'tw' in sect or 'third' in sect.lower():
            resp = 'Third Web'

    # logging.info(f'get_section_name for: {sect} -> {resp}')
    return resp


def get_severity(category):
    severity = defect_severity.get(category.replace('_', ' '))
    if severity is None:
        severity = 'NONE'
    return severity

# Horizontal side view of the blade from root to tip with 3 sections


def make_cell_bold(cell):
    cell.paragraphs[0].runs[0].font.bold = True


def set_cell_font_size(cell, font_size):
    cell.paragraphs[0].runs[0].font.size = Pt(font_size)

# overwrite global document styles
def configure_styles(document):
    style = document.styles['Normal']
    font = style.font
    # font.name = 'Inter'
    font.name = 'Liberation Sans'
    font.size = Pt(10)


# Green Vernova banner on top of each page
def add_header_old(doc, esn):
    header = doc.sections[0].header
    banner_path = os.path.abspath('./src/report/GE-Vernova-Emblem.png')
    heading_table = header.add_table(1, 2, Inches(7))
    heading_cells = heading_table.rows[0].cells
    set_table_cell_bg_color(heading_cells[0], COLOR_EVER_GREEN)
    set_table_cell_bg_color(heading_cells[1], COLOR_EVER_GREEN)
    ht0 = heading_cells[0].add_paragraph()
    kh = ht0.add_run()
    kh.add_picture(banner_path, width=Inches(1.5))
    para = heading_cells[1].add_paragraph()
    run = para.add_run(f'In-Factory Inspection Report\nBlade ID: {esn}')
    run.font.color.rgb = RGBColor(255, 255, 255)
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT


# Green Vernova banner on top of each page
def add_report_header(doc, esn, title=None):
    header = doc.sections[0].header

    logo_path = os.path.abspath('./src/report/GE-Vernova-Emblem.png')
    banner_path = os.path.abspath('./src/report/inspection_report_title.png')
    heading_table = header.add_table(1, 2, Inches(7))
    heading_cells = heading_table.rows[0].cells
    set_table_cell_bg_color(heading_cells[0], COLOR_EVER_GREEN)
    set_table_cell_bg_color(heading_cells[1], COLOR_EVER_GREEN)

    heading_cells[0].text = ""
    run0 = heading_cells[0].paragraphs[0].runs[0]
    if title is None:
        title = 'Blade Quality Report'
        # heading_cells[0].width = Inches(5)
        # heading_cells[0].text = "\n\n"
        # run0.font.size = Pt(5)
        # run0.add_picture(banner_path, width=Inches(3))

    heading_cells[0].width = Inches(5)
    run0.add_text(title)
    run0.font.size = Pt(24)
    run0.font.bold = True
    run0.font.color.rgb = RGBColor(255, 255, 255)
    run0.font.name = 'Liberation Sans Narrow'

    heading_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # heading_cells[0].width = Inches(5)
    # heading_cells[0].text = f'BLADE QUALITY INSPECTION'
    # run0 = heading_cells[0].paragraphs[0].runs[0]
    # run0.font.bold = True
    # run0.font.color.rgb = RGBColor(255, 255, 255)
    # run0.font.name = "Sons Condensed"
    # run0.font.size = Pt(24)
    # # run0.font.name = "Arial Nova"
    # # run0.font.size = Pt(14)
    # heading_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    heading_cells[1].width = Inches(2)
    heading_cells[1].text = "\n\n\n"
    # cell1_para = heading_cells[1].add_paragraph()
    # run1 = cell1_para.add_run()
    run1 = heading_cells[1].paragraphs[0].runs[0]
    run1.font.size = Pt(3)
    run1.add_picture(logo_path, width=Inches(1.25))
    heading_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    header.is_linked_to_previous = False  # only on front page


def add_report_footer(doc, blade_id, title=None):
    if title is None:
        title = 'In-factory Post Molding Internal Inspection'
    footer = doc.sections[0].footer
    paragraph = footer.paragraphs[0]
    run = paragraph.add_run(
        f"{title} \t\t Blade ID: {blade_id.upper()}")
    run.font.color.rgb = RGBColor.from_string(COLOR_EVER_GREEN)
    run.font.size = Pt(10)
    # paragraph.style = document.styles["Footer"]


def add_donut_chart(doc, file_path, category_names, category_data):
    fig = px.pie(values=category_data, names=category_names, hole=0.3)
    fig.write_image(file_path, format='png')  # save the figure to file
    # add_line_break(doc)
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_picture(os.path.abspath(file_path), width=Inches(4))


# TODO: Remove, not being used.
# plotly does not allow rotation, so we could not position the labels correctly
def add_semi_circle_chart_plotly(doc, file_path, category_names, category_data):
    logging.info('add_semi_circle_chart_plotly() called')

    sum = 0
    for value in category_data:
        sum += value
    category_names.append("")
    # this makes the pie 50/50 with the other values on top
    category_data.append(sum)

    logging.info(f'category_names: {category_names}')
    logging.info(f'category_data: {category_data}')

    fig = px.pie(values=category_data, names=category_names, hole=0.5)
    fig.update_layout(showlegend=False)
    fig.write_image(file_path, format='png')  # save the figure to file
    # add_line_break(doc)
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_picture(os.path.abspath(file_path), width=Inches(4))


# produces a png image file with a semi-circle of categories, data, using optional colors array
def generate_semi_circle_chart(file_path, category_names, category_data, colors_data=None):
    logging.info('generate_semi_circle_chart() called')

    # append the sum of all data so we can divide the chart in two sections
    sum = 0
    for value in category_data:
        sum += value
    category_names.append("")
    # this makes the pie 50/50 with the other values on top
    category_data.append(sum)

    # break names and add value to labels
    category_names2 = []
    for i, cat in enumerate(category_names):
        category_names2.append(cat.replace(' ', '\n') +
                            '\n'+str(category_data[i]))

    if (len(colors_data) == 0):
        colors_data = None

    logging.info(f'category_names2: {category_names2}')
    logging.info(f'category_data: {category_data}')
    logging.info(f'colors_data: {colors_data}')

    # fig = px.pie(values=category_data, names=category_names, hole=0.5)
    # fig.update_layout(showlegend=False)
    # fig.write_image(file_path, format='png')  # save the figure to file

    plt.figure(figsize=(8, 6), dpi=300)
    wedges, labels = plt.pie(category_data, wedgeprops=dict(
        width=0.4, edgecolor='w'), labels=category_names2, colors=colors_data, labeldistance=0.7)
    wedges[-1].set_visible(False)

    plt.savefig(fname=file_path, bbox_inches='tight')

    # crop the bottom of the image, where the total half is plotted.
    # the result is a semi-circle with the category_names
    img = PImage.open(file_path)
    width, height = img.size
    crop_img = img.crop((0, 0, width, abs(height/2)))
    crop_img.save(file_path)


# This method uses matplotlib to plot a semi-circle chart
# colors_data array is optional
def add_semi_circle_chart(doc, file_path, category_names, category_data, colors_data=None):

    generate_semi_circle_chart(
        file_path, category_names, category_data, colors_data)

    para = doc.add_paragraph()
    run = para.add_run()
    run.add_picture(os.path.abspath(file_path), width=Inches(4))


# category severity is an array of HIGH, LOW, SAFETY, NONE values for each category
# category names come from the labeling stats
# category data is the total per category. also comes from stats
def add_sunburst_chart(doc, file_path, category_names, category_data, category_severity):
    logging.info(
        f'add_sunburst_chart called with {category_names} {category_data} {category_severity}')
    center_cats = []
    for i in range(0, len(category_names)):
        center_cats.append('Failure Probability')

    # the data frame is like a csv file with colum names = colum values. The length of each list must match
    df = pd.DataFrame(
        dict(center=center_cats, names=category_names,
            data=category_data, severity=category_severity)
    )

    # the path, from center to circumference, indicates a parent -> son relationship, ending with the numeric values parameter
    fig = px.sunburst(df, path=['center', 'severity', 'names'], values='data', color='severity',
                    color_discrete_map=defect_severity_color)
    fig.update_traces(textinfo="label+value+percent entry")
    fig.write_image(file_path, format='png')  # save the figure to file
    # add_line_break(doc)
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_picture(os.path.abspath(file_path), width=Inches(5))


def add_line_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break()


def add_table_and_semi_circle_chart_content(doc, temp_folder_path, title, tbl_category_names, tbl_category_data, tbl_category_color,
                                            chart_category_names, chart_category_data, chart_category_color):
    logging.info(f'add_table_and_semi_circle_chart_content() called')
    logging.info(f'category_names: {tbl_category_names}')
    logging.info(f'category_data: {tbl_category_data}')

    unique_name = str(uuid.uuid4())
    chart_filename = f'report_chart_'+unique_name+'.png'
    img_file_path = os.path.join(temp_folder_path, chart_filename)

    # -------------- setup left and right layout areas ----------
    number_rows = 1
    number_cols = 2  # data table, chart
    logging.info(f'table dimensions: {number_rows},{number_cols}')
    table = doc.add_table(rows=number_rows, cols=number_cols)
    table.style = 'Table Grid'
    table.autofit = True
    table.allow_autofit = True
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    left_cell = table.cell(0, 0)
    right_cell = table.cell(0, 1)

    # ------------------------ left side ------------------------

    # add space for header and total rows
    number_rows = len(tbl_category_names)+2
    number_cols = 3  # color, label, data
    # one extra row for the header and one for total
    left_table = left_cell.add_table(rows=number_rows, cols=number_cols)
    left_table.style = 'Table Grid'
    left_table.autofit = True
    left_table.allow_autofit = True

    for row in left_table.rows:
        logging.info(f'row cells: {row.cells}')
        cell = row.cells[0]
        cell.width = Pt(50)
        cell = row.cells[1]
        cell.width = Pt(300)
        cell = row.cells[2]
        cell.width = Pt(100)

    cell = left_table.cell(0, 0)
    cell.text = ' '
    cell.width = Pt(50)
    set_table_cell_bg_color(cell, COLOR_SKY)
    for i, color in enumerate(tbl_category_color):
        logging.info(f'row:{i+1},{color}')
        left_table.cell(i+1, 0).text = ' '
        left_table.cell(i+1, 0).width = Pt(50)
        set_table_cell_bg_color(left_table.cell(i+1, 0), color)

    cell = left_table.cell(0, 1)
    cell.text = title
    make_cell_bold(cell)
    set_table_cell_bg_color(cell, COLOR_SKY)
    for i, category in enumerate(tbl_category_names):
        logging.info(f'row:{i+1},{category}')
        left_table.cell(i+1, 1).text = category
    cell = left_table.cell(number_rows-1, 1)
    cell.text = "Total"
    make_cell_bold(cell)
    # set_table_cell_bg_color(cell, "439297")

    total = 0
    # mid is the unique identifier of the defect
    cell = left_table.cell(0, 2)
    cell.text = "Number"
    make_cell_bold(cell)
    set_table_cell_bg_color(cell, COLOR_SKY)
    for i, data in enumerate(tbl_category_data):
        logging.info(f'row:{i+1},{data}')
        left_table.cell(i+1, 2).text = str(data)
        total += data
    cell = left_table.cell(number_rows-1, 2)
    cell.text = str(total)
    make_cell_bold(cell)
    # set_table_cell_bg_color(cell, "439297")

    # ------------------- right sie -----------------------------
    right_run = right_cell.paragraphs[0].add_run()

    generate_semi_circle_chart(
        img_file_path, chart_category_names, chart_category_data, chart_category_color)
    right_run.add_picture(img_file_path, width=Pt(200))


# It is a table with left-side properties and right-side blade side view and cross section images, and the actual photo at the bottom
def add_measurement_image_table(doc, distance, finding_label, section, measurement, inspection, measurement_image_path, include_cad_models=True):
    logging.info(f'measurement: {measurement}')

    severity = get_defect_severity(finding_label)
    severity_color = get_defect_severity_color(severity)

    logging.info(
        f'finding_label: {finding_label}, severity: {severity}, color: {severity_color}')

    number_rows = 1
    number_cols = 2
    if not include_cad_models and ('stiff' not in section.lower() and 'third' not in section.lower()):
        number_cols = 1
    logging.info(f'image table dimensions: {number_rows},{number_cols}')
    meas_table = doc.add_table(rows=number_rows, cols=number_cols)
    meas_table.style = 'Table Grid'
    meas_table.autofit = True
    meas_table.allow_autofit = True
    meas_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    left_cell = meas_table.cell(0, 0)
    right_cell = None
    if include_cad_models:
        right_cell = meas_table.cell(0, 1)

    left_table = left_cell.add_table(rows=13, cols=2)
    left_table.style = 'Table Grid'
    left_table.autofit = True
    left_table.allow_autofit = True

    table_cell_width = 300
    for row in meas_table.rows:
        cell = row.cells[0]
        cell.width = Pt(table_cell_width)

    left_table.cell(0, 0).text = "Defect ID"
    left_table.cell(1, 0).text = "Blade Section"
    left_table.cell(2, 0).text = "Distance from Root"
    left_table.cell(3, 0).text = "Anomaly Type"
    left_table.cell(4, 0).text = "Anomaly Length (mm)"
    left_table.cell(5, 0).text = "Anomaly Width (mm)"

    left_table.cell(6, 0).text = "Design Tolerance Spec (mm)"
    left_table.cell(7, 0).text = "Disposition"
    left_table.cell(8, 0).text = "Disposition Provided by"
    left_table.cell(9, 0).text = "Repair Date"
    left_table.cell(10, 0).text = "Repair Report ID"
    left_table.cell(11, 0).text = "Repair Approved by"
    left_table.cell(12, 0).text = "Description"

    for idx in range(0, 13):
        cell = left_table.cell(idx, 0)
        make_cell_bold(cell)

    # mid is the unique identifier of the defect
    left_table.cell(0, 1).text = inspection.esn+'-'+str(measurement.id)
    left_table.cell(1, 1).text = str(section) + \
        " Internal Cavity"  # Blade Section
    left_table.cell(2, 1).text = str(distance)  # Distance from Root
    # Anomaly Type -> TODO: update to measurement.finding_type
    left_table.cell(3, 1).text = finding_label
    left_table.cell(4, 1).text = str(
        measurement.length * 1000) if measurement.length is not None else ''
    left_table.cell(5, 1).text = str(
        measurement.width * 1000) if measurement.width is not None else ''

    left_table.cell(6, 1).text = measurement.design_tolerance * 1000 if measurement.design_tolerance is not None else ''   # design_tolerance
    left_table.cell(7, 1).text = str(measurement.ge_disposition) if measurement.ge_disposition is not None else ''
    left_table.cell(8, 1).text = measurement.disposition_provided_by if measurement.disposition_provided_by is not None else ''   # disposition_provided_by
    left_table.cell(9, 1).text = str(measurement.repair_date.strftime("%Y-%m-%d")) if measurement.repair_date is not None else ''   # repair_date
    left_table.cell(10, 1).text = measurement.repair_report_id if measurement.repair_report_id is not None else ''   # repair_report_id
    left_table.cell(11, 1).text = measurement.repair_approved_by if measurement.repair_approved_by is not None else ''  # repair_approved_by
    left_table.cell(12, 1).text = measurement.description if measurement.description is not None else ''  # comments


    if include_cad_models and ('stiff' not in section.lower() and 'third' not in section.lower()):
        right_run = right_cell.paragraphs[0].add_run()

        blade_view_path = measurement_image_path.replace('.jpg','_blade.jpg').replace('.png','_blade.jpg')
        generate_blade_side_view_position(
            distance, section, severity_color, blade_view_path)
        # blade_view_path = os.path.abspath('./src/report/blade_side_view.png')
        right_run.add_picture(blade_view_path, width=Pt(200))
        # right_run.add_picture(convert_file_to_jpg(blade_view_path), width=Pt(200))

        right_run.add_break()
        right_run.add_break()

        cross_section_path = measurement_image_path.replace('.jpg','_cross.jpg').replace('.png','_cross.jpg')
        generate_blade_cross_section(section, severity_color, cross_section_path)
        # cross_section_path = os.path.abspath('./src/report/blade_cross_section.png')
        right_run.add_picture(cross_section_path, width=Pt(200))

    add_line_break(doc)  # between table and picture

    # doc.add_picture(measurement_image_path, width=Pt(400))
    # saves space. reduce the size of report in 10x
    doc.add_picture(measurement_image_path, width=Pt(400))
    # doc.add_picture(convert_file_to_jpg_cv2(measurement_image_path), width=Pt(400))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


# Horizontal side view of the blade from root to tip with 3 sections
def generate_blade_side_view_position(distance, location, color, output_filename):

    logging.info(
        f'generate_blade_side_view_position called for {distance}, {location}')
    
    skipDrawing = False

    if distance == None:
        distance = 0.0
        logging.info(f'adjusted distance: {distance}')
    if location is None:
        location = ''
        logging.info(f'adjusted location: {location}')
    if color == None:
        color = 'white'
        logging.info(f'adjusted color: {color}')

    blade_length = 70  # in meters
    blade_view_path = os.path.abspath('./src/report/blade_side_view.jpg')
    image = PImage.open(blade_view_path)

    # the 0 mark in the image starts at 30 points from left
    offset = 25 - (float(distance)/10) * 3

    if float(distance) > blade_length:
        distance = blade_length

    # x,y is the circle center
    if float(distance) < 0:
        distance = 1
    x = int(float(distance) * int(image.width / blade_length)) + offset
    if x < 0:
        x = offset
    y = image.height / 4  # up part of figure is for 'leading edge'

    loc = location.lower()
    if 'cent' in loc or 'ce' in loc:
        y = image.height * (3/8)
    elif 'leading' in loc or 'le' in loc:
        y = image.height / 6
    elif 'trailing' in loc or 'te' in loc:
        y = image.height / 2  # lower in the figure
    else:
        skipDrawing = True

    # skip drawing if cavities are not shown in the current CAD model
    if not skipDrawing:
        circ_width = 30
        circ_height = 30
        bbox = (x - circ_width/2, y - circ_height/2,
                x + circ_width/2, y + circ_width/2)
        draw = ImageDraw.Draw(image)
        draw.ellipse(bbox, fill=color)

    output_path = os.path.abspath(output_filename)
    image.save(output_path)


# two lists of distances and their respective location strings
# Overall view of the blade with color circles indicating defect categories and severity
def generate_total_blade_view(distance_list, location_list, category_severity_list, disposition_list, output_filename, rotation=90):
    logging.info(
        f'generate_total_blade_view called with: {distance_list}, {location_list} {category_severity_list} {output_filename}')
    blade_length = 70  # in meters
    blade_view_path = os.path.abspath('./src/report/blade_side_view.jpg')
    image = PImage.open(blade_view_path)
    draw = ImageDraw.Draw(image)

    logging.info(f'distance_list len: {len(distance_list)}')
    logging.info(f'location_list len: {len(location_list)}')
    logging.info(f'category_severity_list len: {len(category_severity_list)}')

    for i in range(0, len(distance_list)):
        distance = float(distance_list[i])
        location = location_list[i]
        # x,y is the circle center
        if distance < 0:
            distance = 1
        x = int(distance * int(image.width / blade_length))
        if x < 0:
            x = 10
        y = image.height / 4  # up part of figure is for 'leading edge'
        loc = get_section_name(location)
        if 'trailing' in loc.lower():
            y = image.height / 2  # lower in the figure
        elif 'cent' in loc.lower():
            y = image.height / 3
        elif 'leading' in loc.lower():
            y = image.height / 4
        circ_width = 30
        circ_height = 30
        bbox = (x - circ_width/2, y - circ_height/2,
                x + circ_width/2, y + circ_width/2)
        # color = "green"
        fill_color = get_defect_severity_color(category_severity_list[i])
        outline_color = get_disposition_color(disposition_list[i])
        # logging.info(f'severity: {category_severity[i]}, color: {color}')
        draw.ellipse(bbox, fill=fill_color, outline=outline_color, width=5)

    output_path = os.path.abspath(output_filename)
    image = image.rotate(rotation, expand=True)
    image.save(output_path)


# location is either 'Trailing Edge', 'Leading Edge' or 'Center Web'
# each defect has a severity color
# Blade cross section image with defect circles
def generate_blade_cross_section(location, color, output_filename):
    logging.info(f'generate_cross_section called with: {location} {color}')
    
    skipDraw = False
    cross_section_path = os.path.abspath(
        './src/report/blade_cross_section.jpg')
    image = PImage.open(cross_section_path)
    y_offset = 10
    x_offset = 0

    # x,y is the circle center
    x = (1/4) * image.width
    y = image.height / 2 + y_offset  # Vertical Center
    loc = get_section_name(location)
    if 'trailing' in loc.lower():
        x = (3/4) * image.width + x_offset  # right in the figure
    elif 'leading' in loc.lower():
        x = (1/4) * image.width + x_offset  # left in the figure
    elif 'cent' in loc.lower():
        x = (3/8) * image.width + x_offset  # center section of figure
    else:
        skipDraw = True

    if not skipDraw:
        circ_width = 20
        circ_height = 20
        bbox = (x - circ_width/2, y - circ_height/2,
                x + circ_width/2, y + circ_width/2)
        draw = ImageDraw.Draw(image)
        draw.ellipse(bbox, fill=color)

    output_path = os.path.abspath(output_filename)
    image.save(output_path)



# in this function, an indicator can be a measurement or a defect
def add_indicators_table_and_side_blade(doc, inspection, indicator_list, total_blade_image_path):
    logging.info(f'indicator_list len: {len(indicator_list)}')

    indicator_labels = [
        'Defect ID',
        'Distance from Root (m)',
        'Blade Cavity',
        'Defect Type',
        'Defect Status'
    ]

    full_indicator_labels = indicator_labels.copy()
    full_indicator_labels.append('Probability of Failure')
    full_indicator_labels.append('Disposition')

    indicator_data_list = []
    full_indicator_data_list = []  # list of indicators with 2 extra fields
    for indicator in indicator_list:
        defect_id = inspection.esn + '-' + str(indicator.id)
        distance_from_root = str(indicator.root_face_distance)
        blade_cavity = indicator.location
        defect_type = indicator.finding_type
        defect_status = indicator.status
        normalized_defect_type = 'Other'
        if defect_type is not None:
            normalized_defect_type = defect_type.replace('_', ' ')
        probability_of_failure = defect_severity.get(normalized_defect_type)
        disposition = indicator.ge_disposition

        # Note: all must be strings
        indicator_row = []
        indicator_row.append(defect_id)
        indicator_row.append(distance_from_root)
        indicator_row.append(blade_cavity)
        indicator_row.append(defect_type)
        indicator_row.append(defect_status)
        indicator_data_list.append(indicator_row)

        full_indicator_row = indicator_row.copy()
        full_indicator_row.append(probability_of_failure)
        full_indicator_row.append(disposition)
        full_indicator_data_list.append(full_indicator_row)

    number_rows = 1
    number_cols = 1
    if total_blade_image_path is not None:
        number_cols = 2
    logging.info(f'image table dimensions: {number_rows},{number_cols}')
    table = doc.add_table(rows=number_rows, cols=number_cols)
    table.style = 'Table Grid'
    table.autofit = True
    table.allow_autofit = True
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    left_cell = table.cell(0, 0)
    

    add_data_table(left_cell, indicator_labels, indicator_data_list,
                labels_on_top=True,  data_width=200, font_size=9)
    
    if total_blade_image_path is not None:
        right_cell = table.cell(0, 1)
        right_run = right_cell.paragraphs[0].add_run()
        right_run.add_picture(total_blade_image_path, width=Inches(2.6))




# Pure text table with head text labels and text data
# labels is a list of column or row labels e.g. ['l1','l2','l3','l4']
# table_data_list is an array of data arrays e.g. [['a','b','c'],['d','e','f']]
# labels match data_list by array index. both list must have the same size.
def add_data_table(doc, labels, table_data_list, labels_on_top=True, label_width=200, data_width=200, font_size=10):
    logging.info(f'table labels: {labels}')
    logging.info(f'table data: {table_data_list}')
    logging.info(f'labels_on_top: {labels_on_top}')

    # if len(labels) != len(table_data_list[0]):
    #     logging.info(f'Error: Length of labels: {len(labels)} and data lists: {len(table_data_list[0])} do not match.')

    if labels_on_top:
        number_rows = len(table_data_list) + 1
        number_cols = len(labels)
    else:  # labels on left
        number_rows = len(labels)
        number_cols = len(table_data_list) + 1

    logging.info(f'table dimensions: {number_rows},{number_cols}')
    table = doc.add_table(rows=number_rows, cols=number_cols)
    table.style = 'Table Grid'
    table.autofit = True
    table.allow_autofit = True
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # initialize size
    for row in table.rows:
        for cell in row.cells:
            cell.width = Pt(200)

    if labels_on_top:
        logging.info(f'labels: {labels}')
        # make_rows_bold(table.rows[0])
        for c, label in enumerate(labels):
            cell = table.cell(0, c)
            cell.text = str(label)
            make_cell_bold(cell)
            set_cell_font_size(cell, font_size)
            cell.width = Pt(label_width)

        logging.info(f'h_processing: {table_data_list}')
        # table_data_list is an array of data
        # data is an array of strings
        for r, data in enumerate(table_data_list):
            # fill horizontally
            # logging.info(f'horizontal_table_data: {data}')
            for c, el in enumerate(data):
                cell = table.cell(r+1, c)
                cell.text = str(el)
                set_cell_font_size(cell, font_size)
                cell.width = Pt(data_width)

    else:  # labels on left column 0, fill labels vertically
        for c, label in enumerate(labels):
            cell = table.cell(c, 0)
            cell.text = str(label)
            set_cell_font_size(cell, font_size)
            make_cell_bold(cell)
            cell.width = Pt(label_width)
            set_table_cell_bg_color(cell, COLOR_SKY)

        logging.info(f'v_processing table_data_list: {table_data_list}')
        for c, data in enumerate(table_data_list):
            # fill vertically
            # logging.info(f'vertical_table_data {data}')
            for r, el in enumerate(data):
                cell = table.cell(r, c+1)
                cell.text = str(el)
                set_cell_font_size(cell, font_size)
                cell.width = Pt(data_width)


# 1st row of labels, second row with data
def add_simple_table(doc, labels, table_data_list):
    logging.info(f'table labels: {labels}')
    logging.info(f'table data: {table_data_list}')
    number_rows = 2
    number_cols = len(labels)
    logging.info(f'table dimensions: {number_rows},{number_cols}')
    table = doc.add_table(rows=number_rows, cols=number_cols)
    table.style = 'Table Grid'
    table.autofit = True
    table.allow_autofit = True
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # initialize size
    for row in table.rows:
        for cell in row.cells:
            cell.width = Pt(80)
    c = 0
    for label in labels:
        logging.info(f'label: {label}')
        cell = table.cell(0, c)
        cell.text = str(label)
        make_cell_bold(cell)
        c += 1
    c = 0
    for data in table_data_list:
        logging.info(f'data: {data}')
        cell = table.cell(1, c)
        cell.text = str(data)
        c += 1


def add_severity_total_table(doc, category_names, category_data):
    severity_dic = {}
    i = 0
    for category in category_names:
        severity = get_severity(category)
        total = severity_dic.get(severity, 0)
        total += category_data[i]
        severity_dic[severity] = total
        i += 1
    severity_labels = severity_dic.keys()
    severity_data = []
    for label in severity_labels:
        severity_data.append(severity_dic[label])
    add_simple_table(doc, severity_labels, severity_data)


# ------------------------------- Image File format manipulation functions --------------------------
def convert_png_file_to_jpg(png_file_path):
    output_path = png_file_path.replace('.png', '.jpg')
    # if the file is already a .jpg does nothing...
    if output_path != png_file_path:
        im = PImage.open(png_file_path).convert(
            "RGB")  # open file as an image object
        # im.save(output_path, quality=80, optimize=True) # save image as jpg with options
        # default is faster and smaller. quality=75 and optimize=False
        im.save(output_path)
    else:
        logging.info(f'Skip png2pdf for: {png_file_path}')
    return output_path


def convert_png_content_to_jpg(png_file_content):
    im = PImage.open(BytesIO(png_file_content)).convert("RGB")  # open file as an image object
    # im.save(output_path, quality=80, optimize=True) # save image as jpg with options
    # default is faster and smaller. quality=75 and optimize=False
    stream = BytesIO()
    im.save(stream, quality=95, format='JPEG')
    return stream.getvalue()


def convert_png_file_to_jpg_cv2(png_file_path):
    jpg_file_path = png_file_path.replace('.png', '.jpg')
    image = cv2.imread(png_file_path)
    cv2.imwrite(jpg_file_path, image, [int(cv2.IMWRITE_JPEG_QUALITY), 75])

# ---------------------------------------------------------------------------------------------------

# return a dictionary from the original png file to the new jpg file
def convert_png_to_jpg_in_parallel(path_list):
    logging.info('convert_png_to_jpg_in_parallel()')
    futures = []
    with ThreadPoolExecutor(THREAD_POOL_SIZE) as executor:
        # for each 360 image record...
        for image_file_path in path_list:
            logger.info(f'call png2jpg: {image_file_path}')
            # Note: function name is a prop in executor, followed by its attributes
            futures.append(executor.submit(
                convert_png_file_to_jpg, image_file_path))
            # add_image_measurements_to_document(doc, session, image, inspection, inspection_images_temp_path)
        wait(futures)
    result_dict = {}
    for idx, future in enumerate(futures):
        result_dict[str(path_list[idx])] = str(future.result())
    return result_dict
