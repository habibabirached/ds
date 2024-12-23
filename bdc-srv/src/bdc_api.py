#!flask/bin/python
from reporting import add_data_table, add_heading, add_indicators_table_and_side_blade, add_line_break, add_measurement_image_table, add_report_footer, add_report_header, add_table_and_semi_circle_chart_content, configure_styles, convert_png_content_to_jpg, generate_blade_cross_section, generate_blade_side_view_position, generate_total_blade_view, get_defect_bgr_color, get_defect_severity, get_defect_severity_color, get_disposition_color, get_section_name, get_severity, set_table_cell_bg_color
import copy  # copy json objects
import os
import shutil
import pathlib
#from measurement import ping_measurement
from measurement import ping_measurement_httpx
#from measurement import get_measurements_json
from measurement import get_measurements_json_httpx
#from measurement import get_measurements_str

from s3_access import get_file_content_from_s3
from s3_access import download_file_from_s3
from s3_access import upload_content_to_s3
from s3_access import get_s3_access_url
from s3_access import get_inspection_s3key


from dateutil.parser import parse
import matplotlib.pyplot as plt
import pandas as pd
import plotly.express as px
from config import defect_severity
from config import defect_severity_color
from config import defect_bgr_colors
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.oxml.shared import OxmlElement, qn
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, RGBColor
from docx import Document
import uuid
import xlsxwriter
from models import Certificate, Defect, DefectAnnotationFragment, InspectionLogs, Measurement, MeasurementAnnotationFile, MeasurementImageFile, OriginalMeasurementAnnotationFile, RepairEvidenceFile, VTShot, VTShotImageFile, ValidatedMeasurementAnnotationFile, db, Inspection, Image, ImageFile, Blade, BladeMonitoring, BladeDefectModel, DefectModel, BladeType
from sqlalchemy import select, distinct, func, or_
from sqlalchemy.orm import Session
import doc2pdf
import cv2
import numpy as np
from concurrent.futures import ThreadPoolExecutor, wait
from operator import and_
import time
from flask import Flask, jsonify, request, abort, make_response, send_file, json, redirect
from flask import after_this_request, send_from_directory
from flask_restful import Api, Resource, reqparse, fields, marshal
from flask_httpauth import HTTPBasicAuth
from flasgger import Swagger
from flask_wtf.csrf import CSRFProtect
from get_userinfo import get_user_claims, is_sso_in_dl_group, generate_dl_members_list, logout_user, get_user_dl_group, get_esn_prefix
from service_now import create_service_request, get_all_service_incident, create_service_incident, get_all_service_tasks
from download_input_files import download_s3_input_files

#import requests
import httpx
import shutil
import threading

# from flask.ext import excel
from datetime import datetime, timedelta, timezone
from datetime import date
import csv

import logging
import dataclasses
import json
from werkzeug.utils import secure_filename
import os
from io import BytesIO
import zipfile
import tempfile

import jsonpickle  # serialize objects to json
import filetype  # figure out the type of image files looking at content

from PIL import Image as PImage
from PIL import ImageDraw

# support for truncated image files
from PIL import ImageFile as PImageFile
PImageFile.LOAD_TRUNCATED_IMAGES = True

# For S3 upload
import re
import boto3
import botocore
from botocore.client import Config
import ros_bag_processor
import subprocess
#import sqlite3
from functools import reduce

# defect specific configuration


from sqlalchemy.ext.declarative import DeclarativeMeta
import json

def create_blade_if_not_exist():
    with Session(db) as session:
        blade = session.query(Blade).get(1)
        if blade is None:
            new_blade = Blade(
                length=12.5,
                manufacturer="TPI MX3",
                model="GE 68.7 Root",
                number="1",
                serial_number="TPI-51179",
                set_number="378",
                tpi_ncr_number="1234"
            )
            session.add(new_blade)
            session.commit()


logger = logging.getLogger()

COLOR_SKY = "59C3C9"
COLOR_DARK_SKY = "439297"
COLOR_EVER_GREEN = "005E60"

UPLOAD_FOLDER = 'upload'
ALLOWED_EXTENSIONS = {'txt', 'csv', 'docx', 'pdf', 'png', 'jpg', 'jpeg', 'gif', 'json'}

FILE_CACHE_FOLDER = 'cache'
if not os.path.exists(FILE_CACHE_FOLDER):
    os.makedirs(FILE_CACHE_FOLDER)

FILE_CACHE_TIMEOUT = 1 * 60 * 60 # 1 hour

THREAD_POOL_SIZE = 50

# ---------------------------------------- S3 Access Flag -----------------------------------
USE_S3 = False

# that's the way to read boolean from env in python!
SAVE_FILES_TO_DB = os.getenv(
    "SAVE_FILES_TO_DB", 'False').lower() in ('true', '1', 't')
if SAVE_FILES_TO_DB is not None:
    print(f'read SAVE_FILES_TO_DB from env: {SAVE_FILES_TO_DB}')
    USE_S3 = not SAVE_FILES_TO_DB

print(f'USE_S3: {USE_S3}')

# -------------------------------------- Data Cache ----------------------------------------

from diskcache import Cache
cache = Cache(FILE_CACHE_FOLDER)
CACHE_TTL = 1800 # 30 min

USE_INSPECTION_CACHE = False

# ------------------------------------------------------------------------------------------

# test measurement service
# ping_resp = ping_measurement()
# print(f'ping_resp: {ping_resp}')

# test measurement
# example_annotation = {
#     "version": "5.2.1",
#     "flags": {},
#     "shapes": [
#         {
#             "points": [
#             [ 802,  488],
#             [1159,  453],
#             [1178,  483],
#             [ 807,  521]
#         ],
#             "group_id": None,
#             "description": "",
#             "shape_type": "polygon",
#             "flags": {},
#             "label": "CoreGap"
#         }
#     ],
#     "imagePath": "image_3.png",
#     "imageHeight": 1200,
#     "imageWidth": 1600,
#     "imageData": None,
#     "imageYfov": 90,
#     "imageHfov": 90,
#     "imagePitch": 0,
#     "imageYaw": -50
# }


# example_frame = {
#     "sect": "Central_Web",
#     "eqHeight": 2688,
#     "eqWidth": 5376,
#     "r": [[0.92388349, -0.38258549, -0.00822412], [0.38266053, 0.92381736, 0.01150656], [0.00319534, -0.01377777, 0.99989998]],
#     "tf": [
#     [ 0.92388349,0.38266053  ,0.00319534  ,0.32184721],
#     [-0.38258549  ,0.92381736 ,-0.01377777 ,-0.41491946],
#     [-0.00822412  ,0.01150656  ,0.99989998 , 1.37315703],
#     [ 0.0 ,0.0 ,0.0 ,1.0 ]
#         ],
#     "pos_lidar": [0.42132799, -0.638, 18.31]
# }


# measurement_resp = get_measurements_json(example_annotation, example_frame)
# print(f'measurement_resp: {measurement_resp}')


# Imported .PNG files will be converted to .JPG to save DB space
# We use high quality preservation compression (95%)
COMPRESS_INCOMING_IMAGES = True
logging.info(f'COMPRESS_INCOMING_MESSAGES: {COMPRESS_INCOMING_IMAGES}')

# protect against Cross-Site Request Forgery (CSRF)
csrf = CSRFProtect()
app = Flask(__name__, static_url_path="")
csrf.init_app(app)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['CACHE_FOLDER'] = FILE_CACHE_FOLDER

if not os.path.exists(app.config['UPLOAD_FOLDER']):
    os.makedirs(app.config['UPLOAD_FOLDER'])

if not os.path.exists(app.config['CACHE_FOLDER']):
    os.makedirs(app.config['CACHE_FOLDER'])

api = Api(app)
# swagger = Swagger(app) # this is called in the app.js instead

# We need at least one blade record #1 for the import to work.
create_blade_if_not_exist()

# ==================================== Authorization API ==============================
auth = HTTPBasicAuth()


@auth.get_password
def get_password(username):
    if username == 'miguel':
        return 'python'
    return None


@auth.error_handler
def unauthorized():
    # return 403 instead of 401 to prevent browsers from displaying the default
    # auth dialog
    return make_response(jsonify({'message': 'Unauthorized access'}), 403)


class UserInfoAPI(Resource):

    def __init__(self):
        super(UserInfoAPI, self).__init__()

    def get(self):
        """
        Get the user information from the header
        ---
        responses:
          200:
            description: Message with claims
            schema:
              id: ClaimsObject
              properties:
                message:
                  type: string
                  example: "Hello"
        """

        claims = get_user_claims(request)
        if claims:
            # Process extracted claims (e.g., personalize content)
            return {'message': f'Welcome, {claims}!'}
        else:
            # Handle case where headers are not found
            return {'message': 'User not authenticated'}, 401

# ========================== IDM Integration ============================


class IDMGroupAPI(Resource):
    esn_prefix = None

    def __init__(self):
        super(IDMGroupAPI, self).__init__()

    def get(self, id):
        """
        Integrating with IDM to check if the user is part of DL
        ---
        parameters:
          - in: path
            name: id
            schema:
              type: string
            description: user sso id
            example: 503419305

        responses:
          200:
            description: user is part of DL or not

        """
        if is_sso_in_dl_group(id):
            IDMGroupAPI.esn_prefix = get_esn_prefix(id)
            return True
        else:
            return False, 401


# ========================== IDM DL Group  ============================

class WhichIDMGroupAPI(Resource):

    def __init__(self):
        super(WhichIDMGroupAPI, self).__init__()

    def get(self, id):
        """
        Integrating with IDM to check user is part of which DL
        ---
        parameters:
          - in: path
            name: id
            schema:
              type: string
            description: user sso id
            example: 503419305

        responses:
          200:
            description: user is part of DL or not

        """
        return get_user_dl_group(id)

# ========================== IDM Integration ============================


class IDMDLMembersAPI(Resource):

    def __init__(self):
        super(IDMDLMembersAPI, self).__init__()

    def get(self):
        """
        Integrating with IDM to check if the user is part of DL
        ---
        parameters:
          - in: path
            name: id
            schema:
              type: string
            description: user sso id
            example: 503419305

        responses:
          200:
            description: user is part of DL or not
        """
        return generate_dl_members_list("g03042086")

# ========================== LogOff User API ============================


class LogOffUserAPI(Resource):

    def __init__(self):
        super(LogOffUserAPI, self).__init__()

    def get(self):
        """
        Integrating with Logoff user sso url
        ---
        responses:
          200:
            description: user is logged off.
        """
        # return logout_user()
        try:
            logoutUrl = os.getenv('LOGOUT_URL')
            response = make_response(redirect(logoutUrl))
            # Delete the AWSELBAuthSessionCookie-0
            response.delete_cookie('AWSELBAuthSessionCookie-0', domain=None, path='/')
            # Delete any other session cookies if necessary
            response.delete_cookie('session', domain=None, path='/')
            return response
        except Exception as e:
            logger.error(f"Error during logout api call: {str(e)}", exc_info=True)
            return {'status': 'error', 'message': f'Error during logout: {str(e)}'}, 500

# ========================== Ping ================================


class PingAPI(Resource):
    def __init__(self):
        super(PingAPI, self).__init__()

    def get(self):
        """
        Ping this server to check whether it is alive
        ---
        responses:
          200:
            description: Ping message
            schema:
              id: PingObject
              properties:
                message:
                  type: string
                  example: "Hello"
        """

        return {'message': "it's alive!"}


# ===========================Async Tasks API ========================

# -------------------------------- Task Utils -----------------------
CLEAN_TEMP_INTERVAL = 60 * 5  # seconds
DOWNLOAD_FILES_TIMEOUT = 60 * 30  # 30 min in seconds
TEMP_DOWNLOAD_FILES_DIR = '/tmp'

# Global task registry
TASK_STATUS_REGISTRY = {}

# async calls task executor
TASK_EXECUTOR = ThreadPoolExecutor(THREAD_POOL_SIZE)

# cleans up potential files left behind


def clean_task_status_registry():
    while (True):
        now_timestamp = datetime.now().timestamp()
        logging.info(f"Clean up TASK_STATUS_REGISTRY ...")
        for task_id in TASK_STATUS_REGISTRY:
            task = TASK_STATUS_REGISTRY[task_id]
            if task is not None:
                delta = now_timestamp - task['timestamp']
                logging.info(f'examining task {task} got delta: {delta}')
                if delta > DOWNLOAD_FILES_TIMEOUT:
                    delete_task(task_id)

        logging.info(
            f"Clean up {TEMP_DOWNLOAD_FILES_DIR} folder now...")  # NOSONAR
        for root, dirs, files in os.walk(TEMP_DOWNLOAD_FILES_DIR, topdown=False):
            for name in files:
                filepath = os.path.join(TEMP_DOWNLOAD_FILES_DIR, name)
                lastUpdate = 0
                try:
                    if os.path.isfile(filepath):
                        file_stat = os.stat(filepath)
                        lastUpdate = file_stat.st_mtime  # last modified in seconds
                except Exception as e:
                    logger.info(e)
                    pass

                if lastUpdate > 0 and (time.time() - lastUpdate) > DOWNLOAD_FILES_TIMEOUT:
                    if os.path.isfile(filepath):
                        os.unlink(filepath)
                        logger.info(f"REMOVED: {filepath}")

            for name in dirs:
                dirpath = os.path.join(TEMP_DOWNLOAD_FILES_DIR, name)
                lastUpdate = 0
                try:
                    if os.path.exists(dirpath):
                        file_stat = os.stat(dirpath)
                        lastUpdate = file_stat.st_mtime  # last modified in seconds
                except Exception as e:
                    logger.info(e)
                    pass
                if lastUpdate > 0 and (time.time() - lastUpdate) > DOWNLOAD_FILES_TIMEOUT:
                    if os.path.exists(dirpath):
                        shutil.rmtree(dirpath)
                        logger.info(f"REMOVED DIR: {dirpath}")

        time.sleep(CLEAN_TEMP_INTERVAL)


# safe net to cleanup left behind temp files
CLEANUP_TASK = threading.Thread(target=clean_task_status_registry)
CLEANUP_TASK.start()


def delete_task(id):
    logging.info(f'delete_task() called for {id}')
    if id in TASK_STATUS_REGISTRY:
        task = TASK_STATUS_REGISTRY[id]
        logging.info(f'found task: {task}')
        task_id = task['id']
        if 'path' in task:
            file_path = task['path']
            if os.path.isfile(file_path):
                os.unlink(file_path)
                TASK_STATUS_REGISTRY[id] = None
                logging.info(f'task: {task_id} deleted')

# --------------------------------- Async Task API -------------------------------------


class TaskStatusAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):
        super(TaskStatusAPI, self).__init__()

    def get(self, id):
        """
        Read the status of an existing task by its id
        ---
        parameters:
          - in: path
            required: true
            name: id
            schema:
              type: string
            description: Task id

        responses:
          200:
            description: Existing task metadata
            schema:
              id: TaskStatusObject
              properties:
                id:
                  type: number
                status:
                  type: string
                  example: COMPLETE ACCEPTED NOT_FOUND
                filename:
                  type: string

        """

        logging.info(f'Get task status')

        if id in TASK_STATUS_REGISTRY:
            task = TASK_STATUS_REGISTRY[id]
            return task

        return {'id': id,
                'status': 'NOT_FOUND',
                'filename': None}

    def delete(self, id):
        """
        Deletes an existing task
        ---
        parameters:
          - in: path
            required: true
            name: id
            schema:
              type: string
            description: Cancels and cleans up an existing task
        responses:
          200:
            description: Status message
            schema:
              message:
                type: string
        """
        logger.info(f'DELETE task id: {id}')

        if id in TASK_STATUS_REGISTRY:
            task = TASK_STATUS_REGISTRY[id]
            logging.info(f'found task: {task}')
            file_path = task['path']
            if os.path.isfile(file_path):
                os.unlink(file_path)
                TASK_STATUS_REGISTRY[id] = None
                return {'message': f'Tasks id {id} successfully deleted.'}
            else:
                return {'message':f"File for id: {id} not found."}, 404
        else:
            return {'message': f"Task id: {id} not found."}, 404


class TaskStatusListAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):
        super(TaskStatusListAPI, self).__init__()

    def get(self):
        """
        Read the status of an existing task by its id
        ---
    
        responses:
          200:
            description: Existing image measurements metadata
            schema:
              id: TaskStatusList
              type: array
              items:
                schema:
                  id: TaskStatusObject

        """

        logging.info(f'Get task status')

        resp = []
        for key in TASK_STATUS_REGISTRY.keys():
            task = TASK_STATUS_REGISTRY[key]
            resp.append(task)

        return resp

   

class TaskFileAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):
        super(TaskFileAPI, self).__init__()

    def get(self, id):
        """
        Reads a file output produced by an asynchronous task
        ---
        parameters:
          - in: path
            required: true
            name: id
            schema:
              type: number
            description: Task id

        produces:
          - application/octet-stream
        responses:
          200:
            description: file produced by an asynchronous task
            schema:
              type: file

        """

        logging.info(f'Get task file')

        if id in TASK_STATUS_REGISTRY:
            task = TASK_STATUS_REGISTRY[id]
            logging.info(f'found task: {task}')
            file_path = task['path']
            file_name = task['filename']
            if os.path.isfile(file_path):
                logging.info(f'sending file: {file_path}')
                return send_file(file_path,  download_name=file_name, as_attachment=True)
            else:
                return {'message':f"File for id: {id} not found."}, 404

        return {'message': f"Task id: {id} not found."}, 404

    def delete(self, id):
        """
        Deletes an existing tasks file
        ---
        parameters:
          - in: path
            required: true
            name: id
            schema:
              type: number
            description: Cancels and cleans up an existing task deleting its file
        responses:
          200:
            description: Status message
            schema:
              message:
                type: string
        """
        logger.info(f'DELETE task id: {id}')

        if id in TASK_STATUS_REGISTRY:
            task = TASK_STATUS_REGISTRY[id]
            logging.info(f'found task: {task}')
            file_path = task['path']
            dir_path = os.path.split(os.path.abspath(file_path))[0]
            if os.path.isfile(file_path):
                os.unlink(file_path)
                if dir_path != TEMP_DOWNLOAD_FILES_DIR:
                    shutil.rmtree(dir_path)
                    logging.info(f'Removed: {dir_path}')
                TASK_STATUS_REGISTRY[id] = None
                return {'message': f'Tasks id {id} successfully deleted.'}
            else:
                return {'message':f"File for id: {id} not found."}, 404
        else:
            return {'message': f"Task id: {id} not found."}, 404


# =========================== Data Model API ========================

# ---------------------------- ImageMeasurement  ---------------------------


# Remember measurements represent label-me files with many defects in different categories
class ImageMeasurementsAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):
        super(ImageMeasurementsAPI, self).__init__()

    def get(self, id):
        """
        Read an existing measurement record for an image id
        ---
        parameters:
          - in: path
            name: id
            schema:
              type: number
            description: existing image id
            example: 1

        responses:
          200:
            description: Existing image measurements metadata
            schema:
              id: MeasurementObjectList
              type: array
              items:
                schema:
                  id: MeasurementObject

        """
        logging.info(f'Look for image metadata for id {id}')
        measurement_list = []
        resp = []
        with Session(db) as session:
            image = session.query(Image).get(id)
            if image is None:
                return {'message': f'No image id# {id} found in DB.'}

            measurement_list = session.scalars(
                select(Measurement).where(Measurement.image_id == id)
            ).all()

            if len(measurement_list) > 0:
                resp = Measurement.toJsonList(measurement_list)

        return resp

# Remember defects are a sub-group of shapes from a measurement under the same category.


class ImageDefectsAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):
        super(ImageDefectsAPI, self).__init__()

    def get(self, id):
        """
        Read an existing defect records for an image id
        ---
        parameters:
          - in: path
            name: id
            schema:
              type: number
            description: existing image id
            example: 1

        responses:
          200:
            description: Existing defects metadata for image
            schema:
              id: DefectObjectList
              type: array
              items:
                schema:
                  id: DefectObject

        """
        logging.info(f'Look for image metadata for id {id}')
        defect_list = []
        resp = []
        with Session(db) as session:
            image = session.query(Image).get(id)
            if image is None:
                return {'message': f'No image id# {id} found in DB.'}

            defect_list = session.scalars(
                select(Defect).where(Defect.image_id == id)
            ).all()

            if len(defect_list) > 0:
                resp = Defect.toJsonList(defect_list)

        return resp

# ---------------------------- VTShots  ---------------------------


class ImageVTShotsAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):
        super(ImageVTShotsAPI, self).__init__()

    def get(self, id):
        """
        Read an existing vtshot record for an image id
        ---
        parameters:
          - in: path
            name: id
            schema:
              type: number
            description: existing image id
            example: 1

        responses:
          200:
            description: Existing image vtshot metadata
            schema:
              id: VTShotObjectList
              type: array
              items:
                schema:
                  id: VTShotObject

        """
        logging.info(f'Look for image metadata for id {id}')
        vtshot_list = []
        resp = []
        with Session(db) as session:
            image = session.query(Image).get(id)
            if image is None:
                return {'message': f'No measurements for image {id} in DB.'}

            vtshot_list = session.scalars(
                select(VTShot).where(VTShot.image_id == id)
            ).all()

            if len(vtshot_list) > 0:
                resp = VTShot.toJsonList(vtshot_list)

        return resp


class ImageAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):

        self.reqparse = reqparse.RequestParser()
        self.reqparse.add_argument('timestamp', type=str, location='json')
        self.reqparse.add_argument('distance', type=float, location='json')
        self.reqparse.add_argument('inspection_id', type=int, location='json')
        self.reqparse.add_argument('blade_id', type=int, location='json')
        self.reqparse.add_argument('image_file_id', type=int, location='json')
        self.reqparse.add_argument(
            'defect_severity', type=str, location='json')
        self.reqparse.add_argument(
            'defect_location', type=str, location='json')
        self.reqparse.add_argument('defect_size', type=float, location='json')
        self.reqparse.add_argument('defect_desc', type=str, location='json')
        self.reqparse.add_argument('frame', type=str, location='json')

        super(ImageAPI, self).__init__()

    def get(self, id):
        """
        Read an existing image meta-data
        ---
        parameters:
          - in: path
            required: true
            name: id
            schema:
              type: number
            description: Existing image id

        responses:
          200:
            description: Existing image metadata
            schema:
              id: ImageObject
              properties:
                id:
                  type: number
                timestamp:
                  type: string
                  example: "Mon, 16 Aug 2021 00:00:00 GMT"
                distance:
                  type: number
                inspection_id:
                  type: number
                blade_id:
                  type: number
                image_file_id:
                  type: number
                defect_severity:
                  type: string
                  enum: [1,2,3,4,5]
                  example: 5
                defect_location:
                  type: string
                  enum: [le_uw, le_dw, te_uw, te_dw, sw_uw, sw_dw]
                  example: le_uw
                defect_size:
                  type: number
                  example: 10.5
                defect_desc:
                  type: string
                  example: "defect description"
                frame:
                  type: string
                  example: "json representing frame data used for measurement calculation"

        """

        logging.info(f'Look for image metadata for id {id}')
        with Session(db) as session:
            image = session.query(Image).get(id)
            if image is not None:
                jsonResp = image.toJson()
                return jsonResp

        return {'message': f'Image {id} not found.'}

    def delete(self, id):
        """
        Deletes an existing image
        ---
        parameters:
          - in: path
            required: true
            name: id
            schema:
              type: number
            description: Existing image id
        responses:
          200:
            description: Existing image metadata
            schema:
              message:
                type: string
        """
        logger.info(f'DELETE image id: {id}')
        with Session(db) as session:
            image = session.query(Image).get(id)
            if image is None:
                return {'message': f'image id {id} not found'}
            msg = delete_image_dependencies(id)
            session.delete(image)
            session.commit()
            return {'message': f'image id: {id} successfully deleted. {msg}'}

    def post(self, id):
        """
        Update an existing image meta-data
        ---
        parameters:
          - in: path
            required: true
            name: id
            schema:
              type: number
            description: Existing image id
          - in: body
            required: true
            name: UpdateImageBody
            schema:
              id: ImageObject

        responses:
          200:
            description: Updated image metadata
            schema:
              id: ImageObject

        """

        args = self.reqparse.parse_args()
        logger.info(args.items())

        logging.info(f'Look for image metadata for image id# {id}')
        with Session(db) as session:
            image = session.query(Image).get(id)
            if image is not None:

                if args['timestamp']:
                    image.timestamp = args['timestamp']
                if args['distance']:
                    image.distance = args['distance']
                if args['inspection_id']:
                    image.inspection_id = args['inspection_id']
                if args['blade_id']:
                    image.blade_id = args['blade_id']
                if args['image_file_id']:
                    image.image_file_id = args['image_file_id']
                if args['defect_severity']:
                    image.defect_severity = args['defect_severity']
                if args['defect_location']:
                    image.defect_location = args['defect_location']
                if args['defect_size']:
                    image.defect_size = args['defect_size']
                if args['defect_desc']:
                    image.defect_desc = args['defect_desc']
                # since frame is a json object, we store it as a string
                if args['frame']:
                    image.frame = None if (
                        args['frame'] == 'null' or args['frame'] == None) else json.dumps(args['frame'])

                session.commit()
                session.refresh(image)
                return image.toJson()

        return {'message': f'Image {id} not found.'}

# The content of an Image stored in the ImageFile record


class ImageFileAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):
        super(ImageFileAPI, self).__init__()

    def allowed_file(self, filename):
        return '.' in filename and \
            filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

    def get(self, id):
        """
        Read an existing image file record, representing either a 360 or 2d still image
        ---
        parameters:
          - in: path
            required: true
            name: id
            schema:
              type: number
            description: Existing image id

        responses:
          200:
            description: image file content for the provided image id
            schema:
              id: ImageFileContent
              type: string
              format: binary
        """
        logger.info(f'Reading image_file content for image.id {id}')
        with Session(db) as session:
            # image_results = session.scalars(select(ImageFile).where(ImageFile.image_id == id)).all()
            image_results = session.query(ImageFile).filter(
                ImageFile.image_id == id
            ).all()

            if image_results is not None and len(image_results) > 0:
                image_file = image_results[0]
                file_content = read_file_record_content(image_file)

                if file_content is not None:
                    return send_file_content_as_jpg(image_file.filename, file_content)

        return {'message': f'File for Image id: {id} not found.'}

    def post(self, id):
        """
        Updates an existing 360 image
        ---
        consumes:
          - multipart/form-data
        parameters:
          - in: path
            required: true
            name: id
            schema:
              type: number
            description: Existing image id
          - in: formData
            required: true
            name: image_file
            type: file
            description: image file
        produces:
          - application/json
        responses:
          200:
            schema:
              id: UploadResponse
              properties:
                message:
                  type: string
        """
        logger.info(f'Updating image_file content for image.id {id}')
        with Session(db) as session:

            logging.info(f'Image id: {id}')
            image = session.query(Image).get(id)
            if image is None:
                return {'message': f'Image record: {id} not found. Existing image id required.'}

            form = request.form.to_dict()
            logger.info(f'Form data: {form}')

            if 'image_file' not in request.files:
                return {'message': 'File not found or no image_file provided'}

            msg = ''
            files_list = []

            image_id = id
            image_file_id = -1  # create or replace

            image_file_list = session.query(ImageFile).filter(
                ImageFile.image_id == id).all()
            new_image_file_rec = None
            if image_file_list is None or len(image_file_list) == 0:
                new_image_file_rec = ImageFile(
                    image_id=image_id,
                    filename=None,
                    content=None,
                    thumbnail=None)
            else:
                new_image_file_rec = image_file_list[0]

            # -------------------------------- formData image_file ---------------------------------
            # Then we read the image_file associated with the image record
            image_file_attachment = request.files.get("image_file")
            image_file_json = {}
            if image_file_attachment and self.allowed_file(image_file_attachment.filename):
                filename = secure_filename(image_file_attachment.filename)
                unique_filename = str(uuid.uuid4())+'_'+filename
                image_file_path = os.path.join(
                    app.config['UPLOAD_FOLDER'], unique_filename)
                image_file_attachment.save(image_file_path)

                if image_id > 0:
                    logging.info('Processing image_file...')
                    bin_content = None
                    with open(image_file_path, mode="rb") as f:
                        bin_content = f.read()
                    thumb_bin_content = get_thumbnail_content(image_file_path)
                    if os.path.isfile(image_file_path):
                        os.remove(image_file_path)

                    # add new or replace existing content
                    if COMPRESS_INCOMING_IMAGES:
                        filename = filename.replace('.png', '.jpg')
                        bin_content = convert_png_content_to_jpg(bin_content)

                    new_image_file_rec.filename = filename
                    new_image_file_rec.content = bin_content
                    new_image_file_rec.thumbnail = thumb_bin_content  # it is already jpg

                    session.add(new_image_file_rec)
                    session.commit()
                    session.refresh(new_image_file_rec)
                    image_file_id = new_image_file_rec.id

                    new_image_file_rec.content = None
                    new_image_file_rec.thumbnail = None
                    logging.info(f'new_image_file_rec: {new_image_file_rec}')

                    image_file_json = new_image_file_rec.toJson()
                    logging.info(f'image_file_json: {image_file_json}')

                msg += f'Image file {image_file_attachment.filename} uploaded successfully. '
                files_list.append(image_file_attachment.filename)
            else:
                msg += 'Image file not found, or no image_file provided.'

            resp = {'message': msg,
                    'image_id': image_id,
                    'image_file_id': image_file_id,
                    'files_list': files_list}

            logging.info(f'resp: {resp}')

            return make_response(jsonify(resp), 200)


class ImageThumbnailAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):
        super(ImageThumbnailAPI, self).__init__()

    def get(self, id):
        """
        Read an existing image file thumbnail, representing either a 360 or 2d still image
        ---
        parameters:
          - in: path
            required: true
            name: id
            schema:
              type: number
            description: Existing image id
          - in: query
            required: false
            name: includeAnnotations
            schema:
              type: boolean
            description: Whether or not to include annotations on the image

        responses:
          200:
            description: image file content for the provided image id
            schema:
              id: ImageThumbnailContent
              type: string
              format: binary
        """

        include_annotations_str = request.args.get('includeAnnotations')
        if include_annotations_str is not None:
            include_annotations_str = include_annotations_str.capitalize()
        logger.info(f'include_annotations_str: {include_annotations_str}')
        include_annotations = False
        if include_annotations_str == 'True':
            include_annotations = True

        logger.info(f'Reading image_file thumbnail for image.id {id}')
        logger.info(f'include_annotations: {include_annotations}')
        with Session(db) as session:
            # image_results = session.scalars(select(ImageFile).where(ImageFile.image_id == id)).all()
            image_thumbnail_results = session.query(
                ImageFile.id, ImageFile.filename, ImageFile.thumbnail).filter(
                    ImageFile.image_id == id
            ).all()
            if image_thumbnail_results is not None and len(image_thumbnail_results) > 0:
                image_thumbnail = image_thumbnail_results[0]
                thumbnail_content = image_thumbnail.thumbnail
                if thumbnail_content is not None:

                    if include_annotations == True:
                        logging.info('including annotations...')
                        annotation_file_results = session.scalars(
                            select(MeasurementAnnotationFile).where(MeasurementAnnotationFile.measurement_id == id)).all()
                        if annotation_file_results is not None and len(annotation_file_results) > 0:
                            annotation_file = annotation_file_results[0]
                            annotation_str = annotation_file.content
                            annotation_json = json.loads(annotation_str)

                            # overwrite the original content with the new one including annotations
                            thumbnail_content = draw_annotations_on_image(
                                thumbnail_content, annotation_json, 1, False)

                    return send_file_content_as_jpg(image_thumbnail.filename, thumbnail_content)

        return {'message': f'Thumbnail for Image id: {id} not found.'}


# Read ImageFile for the provided id and returns is local path
def download_image_file(id, session, use_thumbnail=False):
    temp_image_path = None
    image_results = session.scalars(
        select(ImageFile).where(ImageFile.image_id == id)).all()
    if image_results is not None and len(image_results) > 0:
        image_file = image_results[0]
        temp_image_path = save_image_file(image_file, use_thumbnail)
    return temp_image_path


#image_file, defect_evidence_file, measurement_annotation_file, etc, have similar filename, s3key, content props
def read_file_record_content(file_record):
    logging.info(f'read_file_content() for file_record.filename: {file_record.filename}')
    file_content = file_record.content
    if USE_S3 and file_record.s3key is not None:
        file_content = get_file_content_from_s3(file_record.s3key)
    else:
        logging.info('reading file content from DB record')
    return file_content


# Save image_file.content or image_file.s3key to a temp local path
def save_image_file(image_file, use_thumbnail=False):
    temp_image_path = None

    if use_thumbnail is True:
        # logging.info('saving thumbnail')
        image_file_content = image_file.thumbnail
    else:
        # logging.info('saving content')
        image_file_content = read_file_record_content(image_file)

    if image_file_content is not None:
        unique_name = str(uuid.uuid4())
        image_filename = unique_name+'_'+image_file.filename
        temp_image_path = os.path.join(
            app.config['UPLOAD_FOLDER'], image_filename)
        with open(temp_image_path, "wb") as img_file:
            img_file.write(image_file_content)

    return temp_image_path

# ============================================ Download Inspection (and snapshots) ZIP ==========================================


def zip_folder_and_return_memory_file(path, zip_file_name):
    memory_file = BytesIO()
    # zip_file_name = 'inspection_data.zip'
    # file_path = '/home/data/'

    resolutions = {}
    with zipfile.ZipFile(memory_file, 'w', zipfile.ZIP_DEFLATED, compresslevel=1) as zipf:
        for root, dirs, files in os.walk(path):
            # First pass: store resolutions and zip images
            for file in files:
                file_path = os.path.join(root, file)
                archive_path = str(file_path).replace(str(path), '')
                logging.info(f'zipping: {file_path} as {archive_path}')
                zipf.write(file_path, archive_path)
    memory_file.seek(0)
    # return send_file(memory_file,
    #                  attachment_filename=zip_file_name,
    #                  as_attachment=True)
    # logging.info(f'returning: {zip_file_name}')
    # return send_file(memory_file, download_name=zip_file_name, as_attachment=True)
    return memory_file


class InspectionZipAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):
        # self.TASK_EXECUTOR = ThreadPoolExecutor(THREAD_POOL_SIZE)

        super(InspectionZipAPI, self).__init__()

    def __download_measurement_and_annotations(self, measurement, image_basename, image_measurements_path):
        measurement_id = measurement[0]
        measurement_pitch = measurement[1]
        measurement_yaw = measurement[2]
        measurement_hfov = measurement[3]
        measurement_file_content = measurement[4]
        measurement_s3key = measurement[5]
        measurement_annotation = measurement[6]

        if USE_S3 and measurement_s3key is not None:
            measurement_file_content = get_file_content_from_s3(
                measurement_s3key)
        else:
            logging.info('fallback using measurement_file_content from DB')

        # round_distance = str(
        #     round(image.distance, 1)).replace('.', '_')
        # measurement_id = measurement.id # parent record id of measurement image files and optional annotations

        extension = 'png'
        mime_type = get_content_mime_type(measurement_file_content)
        if ('jpeg' in mime_type):
            extension = 'jpg'

        # measurement_filename = f'{image_basename}-{loc}-z{round_distance}-s{idx}_mid{measurement_id}.{extension}'
        measurement_filename = f'{image_basename}_mid{measurement_id}.{extension}'

        measurement_image_file_path = os.path.join(
            image_measurements_path, measurement_filename)
        logging.info(measurement_image_file_path)
        with open(measurement_image_file_path, "wb") as meas_img_file:
            # measurement_file_content = measurement.content
            meas_img_file.write(measurement_file_content)

        # measurement_annotation_file_results = session.query(MeasurementAnnotationFile).filter(MeasurementAnnotationFile.measurement_id == measurement_id).all()
        # logging.info(f'measurement_annotation_file_results: {measurement_annotation_file_results}')
        if measurement_annotation is not None:
            metadata_file_json = json.loads(
                measurement_annotation)
            metadata_file_json['imagePath'] = measurement_image_file_path.split(
                '/')[-1]
            metadata_file_json['imagePitch'] = measurement_pitch
            metadata_file_json['imageYaw'] = measurement_yaw
            metadata_file_json['imageHfov'] = measurement_hfov
            new_content = json.dumps(
                metadata_file_json)
            measurement_annotation_file_path = measurement_image_file_path.replace(
                '.png', '.json').replace('.jpg', '.json')
            with open(measurement_annotation_file_path, "w") as vts_anno_file:
                vts_anno_file.write(new_content)

    def __download_vtshot(self, vtshot, image_basename, image_vtshots_path):
        vtshot_id = vtshot[0]
        vtshot_pitch = vtshot[1]
        vtshot_yaw = vtshot[2]
        vtshot_hfov = vtshot[3]
        vtshot_date = vtshot[4]
        vtshot_root_face_distance = vtshot[5]
        vtshot_file_content = vtshot[6]
        vtshot_s3key = vtshot[7]

        if USE_S3 and vtshot_s3key is not None:
            vtshot_file_content = get_file_content_from_s3(vtshot_s3key)
        else:
            logging.info('Using measurement_file_content from DB')

        round_distance = str(
            round(image.distance, 1)).replace('.', '_')
        # measurement_id = measurement.id # parent record id of measurement image files and optional annotations

        extension = 'png'
        mime_type = get_content_mime_type(vtshot_file_content)
        if ('jpeg' in mime_type):
            extension = 'jpg'

        # vtshot_filename = f'{image_basename}-{loc}-z{round_distance}-s{idx}_vtsid{vtshot_id}.{extension}'
        vtshot_filename = f'{image_basename}_vtsid{vtshot_id}.{extension}'
        vtshot_image_file_path = os.path.join(
            image_vtshots_path, vtshot_filename)
        logging.info(vtshot_image_file_path)
        with open(vtshot_image_file_path, "wb") as vts_img_file:
            # measurement_file_content = measurement.content
            vts_img_file.write(
                convert_png_content_to_jpg(vtshot_file_content))

        # measurement_annotation_file_results = session.query(MeasurementAnnotationFile).filter(MeasurementAnnotationFile.measurement_id == measurement_id).all()
        # logging.info(f'measurement_annotation_file_results: {measurement_annotation_file_results}')

        metadata_file_json = {}
        metadata_file_json['imagePath'] = vtshot_image_file_path.split(
            '/')[-1]
        metadata_file_json['imagePitch'] = vtshot_pitch
        metadata_file_json['imageYaw'] = vtshot_yaw
        metadata_file_json['imageHfov'] = vtshot_hfov
        metadata_file_json['date'] = str(vtshot_date)
        metadata_file_json['distance'] = vtshot_root_face_distance
        new_content = json.dumps(
            metadata_file_json)
        metadata_file_path = vtshot_image_file_path.replace(
            '.png', '.json').replace('.jpg', '.json')
        with open(metadata_file_path, "w") as vts_anno_file:
            vts_anno_file.write(new_content)

    def __download_image_file_and_metadata(self, image, image_file, image_file_path, tempdirname):
        image_file_content = read_file_record_content(image_file)

        if image_file_content is not None:
            # image_basename = image_filename.split('.')[0]
            # image_ext = image_filename.split('.')[1]
            # image_file_path = f'{inspection_images_path}/{image_basename}-{loc}-z{round_distance}.{image_ext}'

            image_file_meta_path = image_file_path.replace(
                '.png', '.json').replace('.jpg', '.json')
            logging.info(image_file_path)
            with open(image_file_path, "wb") as f:
                f.write(image_file_content)

            frame_json = None
            if image.frame is not None and image.frame != '':
                frame_json = json.loads(image.frame)
            image_meta_json = {"image_id": image.id,
                               "image_ts": str(datetime.now()) if image.timestamp is None else str(image.timestamp),
                               "image_distance": image.distance,
                               "defect_severity": image.defect_severity,
                               "defect_location": image.defect_location,
                               "defect_size": image.defect_size,
                               "defect_desc": image.defect_desc,
                               "image_path": image_file_path.replace(tempdirname, ''),
                               "frame": frame_json,
                               }
            with open(image_file_meta_path, 'w') as f:
                json.dump(image_meta_json, f)

    def __handle_request(self, id, request):
        snapshots_only_str = request.args.get('snapshotsOnly')
        if snapshots_only_str is not None:
            snapshots_only_str = snapshots_only_str.capitalize()
        logger.info(f'snapshotOnlyStr: {snapshots_only_str}')
        snapshots_only = False
        if snapshots_only_str == 'True':
            snapshots_only = True
        logger.info(f'snapshotOnly: {snapshots_only}')

        snapshots_and_360_str = request.args.get('snapshotsAnd360Only')
        if snapshots_and_360_str is not None:
            snapshots_and_360_str = snapshots_and_360_str.capitalize()
        logger.info(f'snapshotsAnd360OnlyStr: {snapshots_and_360_str}')
        snapshots_and_360_only = False
        if snapshots_and_360_str == 'True':
            snapshots_and_360_only = True
        logger.info(f'snapshotsAnd360Only: {snapshots_and_360_only}')

        async_str = request.args.get('async')
        if async_str is not None:
            async_str = async_str.capitalize()
        logger.info(f'asyncStr: {async_str}')
        is_async = False
        if async_str == 'True':
            is_async = True
        logger.info(f'is_async: {is_async}')

        zip_file_name = "inspection_files.zip"

        task_id = uuid.uuid4().hex
        if is_async:
            TASK_STATUS_REGISTRY[task_id] = {
                'id': task_id,
                'timestamp': datetime.now().timestamp(),
                  'status': 'RUNNING',
                  'filename': None,
            }
            TASK_EXECUTOR.submit(self.__run_inspection_zip, app,
                                 request.environ, id, task_id,  snapshots_only, snapshots_and_360_only, is_async)
            return {'id': task_id,
                    'status': 'RUNNING'}, 202

        else:
            return self.__run_inspection_zip(app,
                                             request.environ, id, task_id, snapshots_only, snapshots_and_360_only, is_async)

    def __run_inspection_zip(self, current_app, environ, inspection_id, task_id, snapshots_only, snapshots_and_360_only, is_async):

        with current_app.request_context(environ):
            with tempfile.TemporaryDirectory(dir=UPLOAD_FOLDER, prefix='temp-', suffix='-folder') as tempdirname:

                if task_id in TASK_STATUS_REGISTRY:
                    logging.info(f'task status: {TASK_STATUS_REGISTRY[task_id]}')

                logging.info(f'created temporary directory {tempdirname}')
                # unique_name = str(uuid.uuid4());
                # response_filename = 'inspection_'+unique_name+'.zip'
                # logger.info(f'Creating temp zip file: {response_filename}')

                # zip_file_path = os.path.join(
                #   app.config['UPLOAD_FOLDER'], response_filename)

                # logging.info(f'temp zip filename: {zip_file_path}')

                inspection_images_path = tempdirname+'/images'
                snapshot_images_path = tempdirname+'/2d_snapshots'

                if not snapshots_only:
                    logging.info(f'creating images folder: {inspection_images_path}')
                    os.makedirs(inspection_images_path, exist_ok=True)

                if snapshots_only or snapshots_and_360_only:
                    logging.info(f'creating 2d_snapshots folder: {snapshot_images_path}')
                    os.makedirs(snapshot_images_path, exist_ok=True)

                logger.info(f'root url: {request.url_root}')
                # logger.info(f'Reading inspection content for id: {id}')
                with Session(db) as session:
                    inspection_list = session.scalars(
                        select(Inspection).where(
                            Inspection.id == inspection_id)
                    ).all()
                    if inspection_list is None or len(inspection_list) == 0:
                        return {'message': f'Inspection record for {inspection_id} not found'}
                    inspection = inspection_list[0]

                    zip_file_name = f"{inspection.esn}_{inspection.sect}_files.zip"

                    inspection_esn = inspection.esn  # serial number
                    inspection_sect = inspection.sect  # section
                    loc = inspection_esn+'-'+inspection_sect

                    inspection_filename = f'{tempdirname}/metadata.json'
                    inspection_file_json = json.loads(
                        jsonpickle.encode(inspection, unpicklable=False))
                    del inspection_file_json["_sa_instance_state"]
                    inspection_file_content = json.dumps(inspection_file_json)
                    logging.info(inspection_filename)
                    with open(inspection_filename, "w") as f:
                        f.write(inspection_file_content)

                    inspection_image_list = session.scalars(
                        select(Image).where(Image.inspection_id ==
                                            inspection_id).order_by(Image.distance)
                    ).all()

                    logging.info('processing images in parallel...')
                    futures = []
                    with ThreadPoolExecutor(THREAD_POOL_SIZE) as executor:
                        # for each 360 image...
                        for image in inspection_image_list:
                            # logger.info(f'image: {str(image)}')

                            round_distance = str(
                                round(image.distance, 1)).replace('.', '_')

                            image_file_results = session.scalars(
                                select(ImageFile).where(
                                    ImageFile.image_id == image.id)
                            ).all()
                            if image_file_results is not None and len(image_file_results) > 0:
                                # only one image_file for each image
                                image_file = image_file_results[0]
                                image_filename = image_file.filename

                                # ------------------------------ If include 360 images -----------------------
                                if not snapshots_only or snapshots_and_360_only:
                                    image_basename = image_filename.split('.')[0]
                                    image_ext = image_filename.split('.')[1]
                                    image_file_path = f'{inspection_images_path}/{image_basename}-{loc}-z{round_distance}.{image_ext}'
                                    futures.append(executor.submit(
                                        self.__download_image_file_and_metadata, image, image_file, image_file_path, tempdirname))
                                    # self.__download_image_file_and_metadata(image, image_file, image_file_path, tempdirname)
                                else:
                                    logging.info(f'skipping: {image_filename}')

                                # ---------------------- Include 360 image measurements + annotations ------------------

                                # list measurement image file + annotation for a given image.id
                                measurement_image_file_results = session.query(
                                    Measurement.id,
                                    Measurement.image_pitch,
                                    Measurement.image_yaw,
                                    Measurement.image_hfov,
                                    MeasurementImageFile.content,
                                    MeasurementImageFile.s3key,
                                    MeasurementAnnotationFile.content.label(
                                        'annotation')
                                ).join(
                                    MeasurementImageFile, MeasurementImageFile.measurement_id == Measurement.id
                                ).join(
                                    MeasurementAnnotationFile, MeasurementAnnotationFile.measurement_id == Measurement.id,
                                    isouter=True
                                ).filter(Measurement.image_id == image.id).all()

                                logging.info(
                                    f'found {len(measurement_image_file_results)} results for image: {image.id}')

                                if measurement_image_file_results is not None and len(measurement_image_file_results) > 0:
                                    image_basename = image_filename.split('.')[
                                        0]

                                    # in the snapshots_only mode we do not create individual folders for _measurements
                                    image_measurements_path = f'{snapshot_images_path}'
                                    if not snapshots_only and not snapshots_and_360_only:
                                        image_measurements_path = f'{inspection_images_path}/{image_basename}-{loc}-z{round_distance}_measurements'
                                        os.makedirs(image_measurements_path)

                                    for idx, measurement in enumerate(measurement_image_file_results):
                                        futures.append(executor.submit(self.__download_measurement_and_annotations, measurement, f'{image_basename}-{loc}-z{round_distance}-s{idx}', image_measurements_path))
                                        # self.__download_measurement_and_annotations(measurement, f'{image_basename}-{loc}-z{round_distance}-s{idx}', image_measurements_path)
                                    
                                    #Write image.frame content to a file within image_measurements_path
                                    if image.frame is not None and image.frame != '':
                                        frame_filename = f'{image_basename}-{loc}-z{round_distance}_frame.json'
                                        frame_file_path = image_measurements_path+'/'+frame_filename
                                        with open(frame_file_path, "w") as frame_file:
                                            json_frame = json.loads(image.frame)
                                            json.dump(json_frame, frame_file)


                                # ---------------------- Include 360 image vtshots + metadata ------------------

                                # include virtualtour data only if in regular mode
                                if not snapshots_only:
                                    # list measurement image file + annotation for a given image.id
                                    vtshot_image_file_results = session.query(
                                        VTShot.id,
                                        VTShot.image_pitch,
                                        VTShot.image_yaw,
                                        VTShot.image_hfov,
                                        VTShot.date,
                                        VTShot.root_face_distance,
                                        VTShotImageFile.content,
                                        VTShotImageFile.s3key,
                                    ).join(
                                        VTShotImageFile, VTShotImageFile.vtshot_id == VTShot.id

                                    ).filter(VTShot.image_id == image.id).all()

                                    logging.info(
                                        f'found {len(vtshot_image_file_results)} results for image: {image.id}')

                                    if vtshot_image_file_results is not None and len(vtshot_image_file_results) > 0:
                                        image_basename = image_filename.split('.')[
                                            0]

                                        # in the snapshots_only mode we do not create individual folders
                                        image_vtshots_path = f'{inspection_images_path}/{image_basename}_virtualtour'
                                        os.makedirs(image_vtshots_path)

                                        for idx, vtshot in enumerate(vtshot_image_file_results):
                                            futures.append(executor.submit(self.__download_vtshot, vtshot, f'{image_basename}-{loc}-z{round_distance}', image_vtshots_path))
                                            # self.__download_vtshot(vtshot, f'{image_basename}-{loc}-z{round_distance}', image_vtshots_path)

                        wait(futures)

                memory_file = zip_folder_and_return_memory_file(
                    tempdirname, zip_file_name)

                if is_async:
                    output_filename = task_id+'_'+zip_file_name
                    output_file_path = os.path.join(TEMP_DOWNLOAD_FILES_DIR, output_filename)
                    with open(output_file_path, "wb") as f:
                        f.write(memory_file.getbuffer())
                    TASK_STATUS_REGISTRY[task_id] = {
                        'id': task_id,
                        'timestamp': datetime.now().timestamp(),
                        'status': 'COMPLETE',
                        'filename': zip_file_name,
                        'path': os.path.abspath(output_file_path)
                    }
                    logging.info(f'Set TASK_STATUS: for task_id: {task_id} to: {TASK_STATUS_REGISTRY[task_id]}')
                else:
                    logging.info(f'returning: {zip_file_name}')
                    return send_file(memory_file, download_name=zip_file_name, as_attachment=True)

    # post method has a larger timeout

    def post(self, id):
        """
        Generates a zip file with the inspection data - metadata and images - for the provided inspection id
        ---
        parameters:
          - in: path
            required: true
            name: id
            schema:
              type: number
            description: Existing inspection id
          - in: query
            required: false
            name: snapshotsOnly
            schema:
              type: boolean
            description: Whether or not to include snapshots only
          - in: query
            required: false
            name: snapshotsAnd360Only
            schema:
              type: boolean
            description: Whether or not to include snapshots and annotated 360 images only
          - in: query
            required: false
            name: async
            schema:
              type: boolean
            description: Allow the use of this method in asynchronous mode
            default: false
        consumes:
          - application/json
        produces:
          - application/x-zip
        responses:
          200:
            description: zip file with inspection 360 images and 2d snapshots
            schema:
              type: file
        """

        return self.__handle_request(id, request)

    def get(self, id):
        """
        Generates a zip file with the inspection data - metadata and images - for the provided inspection id
        ---
        parameters:
          - in: path
            required: true
            name: id
            schema:
              type: number
            description: Existing inspection id
          - in: query
            required: false
            name: snapshotsOnly
            schema:
              type: boolean
            description: Whether or not to include snapshots only
          - in: query
            required: false
            name: snapshotsAnd360Only
            schema:
              type: boolean
            description: Whether or not to include snapshots and annotated 360 images only
          - in: query
            required: false
            name: async
            schema:
              type: boolean
            description: Use this method in asynchronous mode
            default: false
        consumes:
          - application/json
        produces:
          - application/x-zip
        responses:
          200:
            description: zip file with inspection 360 images and 2d snapshots
            schema:
              type: file
        """

        return self.__handle_request(id, request)


class SelectedDefectsZipAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):

        # self.reqparse = reqparse.RequestParser()
        # self.reqparse.add_argument('idList', action='append')

        super(SelectedDefectsZipAPI, self).__init__()

    def __download_defect_files(self, defect, idx, inspection_images_path, tempdirname):
        image_distance = defect.image_distance
        defect_id = defect.defect_id
        measurement_id = defect.measurement_id
        measurement_file_content = defect.measurement_file_content
        image360_file_content = defect.image360Content
        #defect_annotation_original_content = defect.annotation_original_content

        logging.info(f'download image and annotation files for defect id#{defect_id}')

        if USE_S3 and defect.measurement_file_s3key is not None:
            # use setter
            # measurement.measurement_file_content
            measurement_file_content = get_file_content_from_s3(
                defect.measurement_file_s3key)
            image360_file_content = get_file_content_from_s3(
                defect.image360_file_s3key)
        else:
            logging.info('Using measurement_file_content from DB')

        measurement_file_content = convert_png_content_to_jpg(measurement_file_content)
        image360_file_content = convert_png_content_to_jpg(image360_file_content)
        defect_annotation_content = defect.annotation_content
        loc = defect.inspection_sect
        inspection_esn = defect.inspection_esn

        round_distance = str(round(image_distance, 1)).replace('.', '_')

        # add that, TPI-60030-884

        measurement_filename = f'img-{inspection_esn}-{loc}-z{round_distance}-s{idx}_mid{measurement_id}_id{defect_id}.jpg'
        measurement_image_file_path = os.path.join(inspection_images_path, measurement_filename)
        logging.info(f'writing: {measurement_image_file_path}')
        with open(measurement_image_file_path, "wb") as mes_img_file:
            # measurement_file_content = measurement.content
            mes_img_file.write(measurement_file_content)

        # write the 360 image at the root level
        # Ensure the 'images360' directory exists within the temporary directory
        images360_dir = os.path.join(tempdirname, 'images360')
        os.makedirs(images360_dir, exist_ok=True)

        # Construct the file path within the 'images360' directory
        image_360_file_name = f'img-{inspection_esn}-{loc}-z{round_distance}-s{idx}_mid{measurement_id}.jpg'
        root_360_image_file_path = os.path.join(
            images360_dir, image_360_file_name)
        with open(root_360_image_file_path, "wb") as img360file:
            img360file.write(image360_file_content)

        # measurement_annotation_file_results = session.query(MeasurementAnnotationFile).filter(MeasurementAnnotationFile.measurement_id == measurement_id).all()
        # logging.info(f'measurement_annotation_file_results: {measurement_annotation_file_results}')
        if defect_annotation_content is not None:
            annotation_file_json = json.loads(defect_annotation_content)
            annotation_file_json['imagePath'] = measurement_image_file_path.split('/')[-1]
            annotation_file_json['image360Path'] = image_360_file_name.split('/')[-1]
            modified_annotation_json_content = json.dumps(annotation_file_json)
            measurement_annotation_file_path = measurement_image_file_path.replace(
                '.png', '.json').replace('.jpg', '.json')
            with open(measurement_annotation_file_path, "w") as mes_anno_file:
                mes_anno_file.write(modified_annotation_json_content)
        else:
            logging.info('no defect_annotation_content found.')


    def __run_selected_defects_zip(self, current_app, environ, task_id, id_list, is_async):

        with current_app.request_context(environ):
            zip_file_name = "filtered_defect_files.zip"

            with tempfile.TemporaryDirectory(dir=UPLOAD_FOLDER, prefix='temp-', suffix='-folder') as tempdirname:
                logging.info(f'created temporary directory {tempdirname}')

                inspection_images_path = tempdirname+'/imagesDefects'
                #inspection_images_path_validated = tempdirname+'/imagesValidated'
                logging.info(f'creating images folder: {inspection_images_path}')
                os.makedirs(inspection_images_path, exist_ok=True)
                # per Arpit's ask, he needed the validated images to be in a different directory.
                #os.makedirs(inspection_images_path_validated, exist_ok=True)

                # ---------------------- Include 360 image measurements + annotations ------------------

                with Session(db) as session:
                    defect_list = session.query(
                        Inspection.id.label('inspection_id'),
                        Inspection.esn.label('inspection_esn'),
                        Inspection.sect.label('inspection_sect'),
                        Inspection.blade_type.label('inspection_blade_type'),
                        Inspection.upload_date.label('inspection_upload_date'),
                        Inspection.sso.label('inspection_sso'),

                        Image.id.label('image_id'),
                        Image.distance.label('image_distance'),

                        Defect.id.label('defect_id'),
                        Defect.ge_disposition.label(
                            'defect_ge_disposition'),
                        Defect.status.label('defect_status'),
                        Defect.finding_type.label(
                            'defect_finding_type'),
                        Defect.is_manual.label('defect_is_manual'),

                        Measurement.id.label('measurement_id'),

                        MeasurementImageFile.id.label('measurement_file_id'),
                        MeasurementImageFile.filename.label(
                            'measurement_file_filename'),
                        MeasurementImageFile.content.label(
                            'measurement_file_content'),  # bin content
                        MeasurementImageFile.s3key.label(
                            'measurement_file_s3key'),

                        # JSON content from MeasurementAnnotationFile
                        DefectAnnotationFragment.content.label(
                            'annotation_content'),
                        ImageFile.content.label(
                            'image360Content'),  # bin content
                        ImageFile.s3key.label('image360_file_s3key'),

                    ).join(Image, Image.inspection_id == Inspection.id,  isouter=False
                           ).join(Defect, Defect.image_id == Image.id,  isouter=False
                                  ).join(Measurement, Measurement.id == Defect.measurement_id
                                  ).join(MeasurementImageFile, MeasurementImageFile.measurement_id == Measurement.id
                                  ).join(DefectAnnotationFragment, DefectAnnotationFragment.defect_id == Defect.id, isouter=True
                                  ).join(ImageFile, ImageFile.image_id == Image.id, isouter=False
                                  ).filter(Defect.id.in_(id_list)
                                    ).all()

                    logging.info(f'found {len(defect_list)} results. ')
                    #logging.info(f'defect_list: {defect_list}')

                    if defect_list is not None and len(defect_list) > 0:
                        
                        logging.info('download files in parallel...')
                        futures = []
                        with ThreadPoolExecutor(THREAD_POOL_SIZE) as executor:
                            for idx, defect in enumerate(defect_list):
                                futures.append(executor.submit(self.__download_defect_files,
                                                               defect, idx,
                                                               inspection_images_path, tempdirname))
                            wait(futures)
                        
                        # logging.info('download files in series...')
                        # for idx, defect in enumerate(defect_list):
                        #     self.__download_defect_files(defect, idx,inspection_images_path, tempdirname)

                memory_file = zip_folder_and_return_memory_file(
                    tempdirname, zip_file_name)

                if is_async:
                    # save in /tmp instead of tempdirname so the output .zip does not get cleaned up upon return.
                    output_filename = task_id+'_'+zip_file_name
                    output_file_path = os.path.join(TEMP_DOWNLOAD_FILES_DIR, output_filename)
                    with open(output_file_path, "wb") as f:
                        f.write(memory_file.getbuffer())
                    TASK_STATUS_REGISTRY[task_id] = {
                        'id': task_id,
                        'timestamp': datetime.now().timestamp(),
                        'status': 'COMPLETE',
                        'filename': zip_file_name,
                        'path': os.path.abspath(output_file_path)
                    }
                    logging.info(f'Set TASK_STATUS: for task_id: {task_id} to: {TASK_STATUS_REGISTRY[task_id]}')
                else:
                    logging.info(f'returning: {zip_file_name}')
                    return send_file(memory_file, download_name=zip_file_name, as_attachment=True)

    def post(self):
        """
        Generates a zip file with a list of selected defect annotations and images for a set of defect ids
        ---
        parameters:
          - in: body
            required: true
            name: idListBody
            description: List of defect_id
            schema:
              type: object
              properties:
                is_async:
                  type: boolean
                  default: false
                id_list:
                  type: array
                  items:
                    type: number
                  example: [601,664,497,647,681,437,292]
        consumes:
          - application/json
        produces:
          - application/x-zip
        responses:
          200:
            description: zip file with inspection 360 images and defect 2d snapshots with annotations
            schema:
              type: file
        """
        body = request.json
        logging.info(f'body: {body}')

        id_list = body['id_list']
        logging.info(f'id_list: {id_list}')

        is_async = False
        if 'is_async' in body:
            is_async = body['is_async']
            logging.info(f'is_async: {is_async}')

        task_id = uuid.uuid4().hex
        if is_async:
            TASK_STATUS_REGISTRY[task_id] = {
                'id': task_id,
                'timestamp': datetime.now().timestamp(),
                  'status': 'RUNNING',
                  'filename': None,
            }
            TASK_EXECUTOR.submit(self.__run_selected_defects_zip, app,
                                 request.environ, task_id, id_list, is_async)
            return {'id': task_id,
                    'status': 'RUNNING'}, 202

        else:
            return self.__run_selected_defects_zip(app, request.environ, task_id, id_list, is_async)


class SelectedImagesZipAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):

        # self.reqparse = reqparse.RequestParser()
        # self.reqparse.add_argument('idList', action='append')

        super(SelectedImagesZipAPI, self).__init__()


    def __download_image_files(self, image, idx, inspection_images_path, tempdirname):
        image_distance = image.image_distance
        image360_file_content = image.image360Content
        #defect_annotation_original_content = defect.annotation_original_content

        logging.info(f'download image file for image id# {image.image_id}')

        if USE_S3 and image.image360_file_s3key is not None:
            # use setter
            # measurement.measurement_file_content
            image360_file_content = get_file_content_from_s3(
                image.image360_file_s3key)
        else:
            logging.info('Using measurement_file_content from DB')

        image360_file_content = convert_png_content_to_jpg(image360_file_content)
        loc = image.inspection_sect
        inspection_esn = image.inspection_esn

        round_distance = str(round(image_distance, 1)).replace('.', '_')

        # write the 360 image at the root level
        # Ensure the 'images360' directory exists within the temporary directory
        images360_dir = os.path.join(tempdirname, 'images360')
        os.makedirs(images360_dir, exist_ok=True)

        # Construct the file path within the 'images360' directory
        image_360_file_name = f'img-{inspection_esn}-{loc}-z{round_distance}-s{idx}.jpg'
        root_360_image_file_path = os.path.join(
            images360_dir, image_360_file_name)
        with open(root_360_image_file_path, "wb") as img360file:
            img360file.write(image360_file_content)


    def __run_selected_images_zip(self, current_app, environ, task_id, id_list, is_async):

        with current_app.request_context(environ):
            zip_file_name = "filtered_image_files.zip"

            with tempfile.TemporaryDirectory(dir=UPLOAD_FOLDER, prefix='temp-', suffix='-folder') as tempdirname:
                logging.info(f'created temporary directory {tempdirname}')

                inspection_images_path = tempdirname+'/images'
                #inspection_images_path_validated = tempdirname+'/imagesValidated'
                logging.info(f'creating images folder: {inspection_images_path}')
                os.makedirs(inspection_images_path, exist_ok=True)
                # per Arpit's ask, he needed the validated images to be in a different directory.
                #os.makedirs(inspection_images_path_validated, exist_ok=True)

                # ---------------------- Include 360 image measurements + annotations ------------------

                with Session(db) as session:
                    image_list = session.query(
                        Inspection.id.label('inspection_id'),
                        Inspection.esn.label('inspection_esn'),
                        Inspection.sect.label('inspection_sect'),
                        Inspection.blade_type.label('inspection_blade_type'),
                        Inspection.upload_date.label('inspection_upload_date'),
                        Inspection.sso.label('inspection_sso'),

                        Image.id.label('image_id'),
                        Image.distance.label('image_distance'),

                        ImageFile.content.label(
                            'image360Content'),  # bin content
                        ImageFile.s3key.label('image360_file_s3key'),

                    ).join(Image, Image.inspection_id == Inspection.id,  isouter=False  
                    ).join(ImageFile, ImageFile.image_id == Image.id, isouter=False
                    ).filter(Image.id.in_(id_list)
                    ).all()

                    logging.info(f'found {len(image_list)} results. ')
                    #logging.info(f'defect_list: {defect_list}')

                    if image_list is not None and len(image_list) > 0:
                        
                        logging.info('download files in parallel...')
                        futures = []
                        with ThreadPoolExecutor(THREAD_POOL_SIZE) as executor:
                            for idx, image in enumerate(image_list):
                                futures.append(executor.submit(self.__download_image_files,
                                                               image, idx,
                                                               inspection_images_path, tempdirname))
                            wait(futures)
                        
                        # logging.info('download files in series...')
                        # for idx, defect in enumerate(defect_list):
                        #     self.__download_defect_files(defect, idx,inspection_images_path, tempdirname)

                memory_file = zip_folder_and_return_memory_file(
                    tempdirname, zip_file_name)

                if is_async:
                    # save in /tmp instead of tempdirname so the output .zip does not get cleaned up upon return.
                    output_filename = task_id+'_'+zip_file_name
                    output_file_path = os.path.join(TEMP_DOWNLOAD_FILES_DIR, output_filename)
                    with open(output_file_path, "wb") as f:
                        f.write(memory_file.getbuffer())
                    TASK_STATUS_REGISTRY[task_id] = {
                        'id': task_id,
                        'timestamp': datetime.now().timestamp(),
                        'status': 'COMPLETE',
                        'filename': zip_file_name,
                        'path': os.path.abspath(output_file_path)
                    }
                    logging.info(f'Set TASK_STATUS: for task_id: {task_id} to: {TASK_STATUS_REGISTRY[task_id]}')
                else:
                    logging.info(f'returning: {zip_file_name}')
                    return send_file(memory_file, download_name=zip_file_name, as_attachment=True)

    def post(self):
        """
        Generates a zip file with a list of selected defect annotations and images for a set of defect ids
        ---
        parameters:
          - in: body
            required: true
            name: idListBody
            description: List of defect_id
            schema:
              type: object
              properties:
                is_async:
                  type: boolean
                  default: false
                id_list:
                  type: array
                  items:
                    type: number
                  example: [601,664,497,647,681,437,292]
        consumes:
          - application/json
        produces:
          - application/x-zip
        responses:
          200:
            description: zip file with inspection 360 images and defect 2d snapshots with annotations
            schema:
              type: file
        """
        body = request.json
        logging.info(f'body: {body}')

        id_list = body['id_list']
        logging.info(f'id_list: {id_list}')

        is_async = False
        if 'is_async' in body:
            is_async = body['is_async']
            logging.info(f'is_async: {is_async}')

        task_id = uuid.uuid4().hex
        if is_async:
            TASK_STATUS_REGISTRY[task_id] = {
                'id': task_id,
                'timestamp': datetime.now().timestamp(),
                  'status': 'RUNNING',
                  'filename': None,
            }
            TASK_EXECUTOR.submit(self.__run_selected_images_zip, app,
                                 request.environ, task_id, id_list, is_async)
            return {'id': task_id,
                    'status': 'RUNNING'}, 202

        else:
            return self.__run_selected_images_zip(app, request.environ, task_id, id_list, is_async)



# --------------------------------------------- Validated annotations API --------------------------------
class ValidatedAnnotationsZipAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):

        # self.reqparse = reqparse.RequestParser()
        # self.reqparse.add_argument('idList', action='append')

        super(ValidatedAnnotationsZipAPI, self).__init__()

    def __download_measurement_files(self, measurement, idx, inspection_images_path_validated, inspection_images_path, tempdirname):
        image_distance = measurement.image_distance
        measurement_id = measurement.measurement_id
        measurement_file_content = measurement.measurement_file_content
        image360_file_content = measurement.image360Content
        measurement_annotation_original_content = measurement.annotation_original_content

        if USE_S3 and measurement.measurement_file_s3key is not None:
            # use setter
            # measurement.measurement_file_content
            measurement_file_content = get_file_content_from_s3(
                measurement.measurement_file_s3key)
            image360_file_content = get_file_content_from_s3(
                measurement.image360_file_s3key)
        else:
            logging.info('Using measurement_file_content from DB')

        measurement_file_content = convert_png_content_to_jpg(
            measurement_file_content)
        image360_file_content = convert_png_content_to_jpg(
            image360_file_content)
        measurement_annotation = measurement.annotation_content
        loc = measurement.inspection_sect
        inspection_esn = measurement.inspection_esn

        round_distance = str(
            round(image_distance, 1)).replace('.', '_')

        # measurement_id = measurement.id # parent record id of measurement image files and optional annotations
        # write the validated images and json ***
        # write the validated images and json ***
        # write the validated images and json ***
        measurement_filename = f'img-{inspection_esn}-{loc}-z{round_distance}-s{idx}_mid{measurement_id}.jpg'
        measurement_image_file_path_validated = os.path.join(
            inspection_images_path_validated, measurement_filename)
        logging.info(measurement_image_file_path_validated)
        with open(measurement_image_file_path_validated, "wb") as mes_img_file:
            # measurement_file_content = measurement.content
            mes_img_file.write(measurement_file_content)

        # add that, TPI-60030-884

        # measurement_annotation_file_results = session.query(MeasurementAnnotationFile).filter(MeasurementAnnotationFile.measurement_id == measurement_id).all()
        # logging.info(f'measurement_annotation_file_results: {measurement_annotation_file_results}')
        if measurement_annotation is not None:
            annotation_file_json = json.loads(measurement_annotation)
            annotation_file_json['imagePath'] = measurement_image_file_path_validated.split(
                '/')[-1]
            modified_annotation_json_content = json.dumps(annotation_file_json)
            measurement_annotation_file_path = measurement_image_file_path_validated.replace(
                '.png', '.json').replace('.jpg', '.json')
            with open(measurement_annotation_file_path, "w") as mes_anno_file:
                mes_anno_file.write(modified_annotation_json_content)

        # do the same for originals ***
        # do the same for originals ***
        # do the same for originals ***
        measurement_filename = f'img-{inspection_esn}-{loc}-z{round_distance}-s{idx}_mid{measurement_id}.jpg'
        measurement_image_file_path = os.path.join(
            inspection_images_path, measurement_filename)
        logging.info(measurement_image_file_path)
        with open(measurement_image_file_path, "wb") as mes_img_file:
            # measurement_file_content = measurement.content
            mes_img_file.write(measurement_file_content)

        # write the 360 image at the root level
        # Ensure the 'images360' directory exists within the temporary directory
        images360_dir = os.path.join(tempdirname, 'images360')
        os.makedirs(images360_dir, exist_ok=True)

        # Construct the file path within the 'images360' directory
        image_360_file_name = f'img-{inspection_esn}-{loc}-z{round_distance}-s{idx}_mid{measurement_id}.360.jpg'
        root_360_image_file_path = os.path.join(
            images360_dir, image_360_file_name)
        with open(root_360_image_file_path, "wb") as img360file:
            img360file.write(image360_file_content)

        # that's the current annotation content
        if measurement_annotation_original_content is not None:
            annotation_file_json = json.loads(
                measurement_annotation_original_content)
            annotation_file_json['imagePath'] = measurement_image_file_path.split(
                '/')[-1]
            modified_annotation_json_content = json.dumps(annotation_file_json)
            measurement_annotation_file_path = measurement_image_file_path.replace(
                '.png', '.json').replace('.jpg', '.json')
            with open(measurement_annotation_file_path, "w") as mes_anno_file:
                mes_anno_file.write(modified_annotation_json_content)

    def __run_validated_annotations_zip(self, current_app, environ, task_id, id_list, is_async):

        with current_app.request_context(environ):
            zip_file_name = "validated_measurement_files.zip"

            with tempfile.TemporaryDirectory(dir=UPLOAD_FOLDER, prefix='temp-', suffix='-folder') as tempdirname:
                logging.info(f'created temporary directory {tempdirname}')

                inspection_images_path = tempdirname+'/imagesOriginals'
                inspection_images_path_validated = tempdirname+'/imagesValidated'
                logging.info(f'creating images folder: {inspection_images_path}')
                os.makedirs(inspection_images_path, exist_ok=True)
                # per Arpit's ask, he needed the validated images to be in a different directory.
                os.makedirs(inspection_images_path_validated, exist_ok=True)

                # ---------------------- Include 360 image measurements + annotations ------------------

                with Session(db) as session:
                    measurement_list = session.query(
                        Inspection.id.label('inspection_id'),
                        Inspection.esn.label('inspection_esn'),
                        Inspection.sect.label('inspection_sect'),
                        Inspection.blade_type.label('inspection_blade_type'),
                        Inspection.upload_date.label('inspection_upload_date'),
                        Inspection.sso.label('inspection_sso'),

                        Image.id.label('image_id'),
                        Image.distance.label('image_distance'),

                        Measurement.id.label('measurement_id'),
                        Measurement.ge_disposition.label(
                            'measurement_ge_disposition'),
                        Measurement.status.label('measurement_status'),
                        Measurement.finding_type.label(
                            'measurement_finding_type'),
                        Measurement.is_manual.label('measurement_is_manual'),

                        MeasurementImageFile.id.label('measurement_file_id'),
                        MeasurementImageFile.filename.label(
                            'measurement_file_filename'),
                        MeasurementImageFile.content.label(
                            'measurement_file_content'),  # bin content
                        MeasurementImageFile.s3key.label(
                            'measurement_file_s3key'),

                        ValidatedMeasurementAnnotationFile.id.label(
                            'annotation_id'),
                        ValidatedMeasurementAnnotationFile.content.label(
                            'annotation_content'),  # json content
                        ValidatedMeasurementAnnotationFile.validated_by.label(
                            'annotation_validated_by'),
                        ValidatedMeasurementAnnotationFile.validation_status.label(
                            'annotation_validation_status'),
                        ValidatedMeasurementAnnotationFile.validation_timestamp.label(
                            'annotation_validation_timestamp'),

                        # JSON content from MeasurementAnnotationFile
                        MeasurementAnnotationFile.content.label(
                            'annotation_original_content'),
                        ImageFile.content.label(
                            'image360Content'),  # bin content
                        ImageFile.s3key.label('image360_file_s3key'),

                    ).join(Image, Image.inspection_id == Inspection.id,  isouter=False
                           ).join(Measurement, Measurement.image_id == Image.id,  isouter=False
                                  ).join(MeasurementImageFile, MeasurementImageFile.measurement_id == Measurement.id
                                  ).join(ValidatedMeasurementAnnotationFile, ValidatedMeasurementAnnotationFile.measurement_id == Measurement.id, isouter=True
                                  ).join(MeasurementAnnotationFile, MeasurementAnnotationFile.measurement_id == Measurement.id, isouter=True
                                  ).join(ImageFile, ImageFile.image_id == Image.id, isouter=False
                                  ).filter(Measurement.id.in_(id_list)
                                    ).all()

                    logging.info(f'found {len(measurement_list)} results. ')

                    if measurement_list is not None and len(measurement_list) > 0:
                        futures = []
                        # download files in parallel
                        with ThreadPoolExecutor(THREAD_POOL_SIZE) as executor:
                            for idx, measurement in enumerate(measurement_list):
                                futures.append(executor.submit(self.__download_measurement_files,
                                                               measurement, idx, inspection_images_path_validated,
                                                               inspection_images_path, tempdirname))
                            wait(futures)

                memory_file = zip_folder_and_return_memory_file(
                    tempdirname, zip_file_name)

                if is_async:
                    # save in /tmp instead of tempdirname so the output .zip does not get cleaned up upon return.
                    output_filename = task_id+'_'+zip_file_name
                    output_file_path = os.path.join(TEMP_DOWNLOAD_FILES_DIR, output_filename)
                    with open(output_file_path, "wb") as f:
                        f.write(memory_file.getbuffer())
                    TASK_STATUS_REGISTRY[task_id] = {
                        'id': task_id,
                        'timestamp': datetime.now().timestamp(),
                        'status': 'COMPLETE',
                        'filename': zip_file_name,
                        'path': os.path.abspath(output_file_path)
                    }
                    logging.info(f'Set TASK_STATUS: for task_id: {task_id} to: {TASK_STATUS_REGISTRY[task_id]}')
                else:
                    logging.info(f'returning: {zip_file_name}')
                    return send_file(memory_file, download_name=zip_file_name, as_attachment=True)

    def post(self):
        """
        Generates a zip file with a list of validated annotations and images for a set of measurement ids
        ---
        parameters:
          - in: body
            required: true
            name: idListBody
            description: List of measurement_id
            schema:
              type: object
              properties:
                is_async:
                  type: boolean
                  default: false
                id_list:
                  type: array
                  items:
                    type: number
                  example: [601,664,497,647,681,437,292]
        consumes:
          - application/json
        produces:
          - application/x-zip
        responses:
          200:
            description: zip file with inspection 360 images and 2d snapshots
            schema:
              type: file
        """
        body = request.json
        logging.info(f'body: {body}')

        id_list = body['id_list']
        logging.info(f'id_list: {id_list}')

        is_async = False
        if 'is_async' in body:
            is_async = body['is_async']
            logging.info(f'is_async: {is_async}')

        task_id = uuid.uuid4().hex
        if is_async:
            TASK_STATUS_REGISTRY[task_id] = {
                'id': task_id,
                'timestamp': datetime.now().timestamp(),
                  'status': 'RUNNING',
                  'filename': None,
            }
            TASK_EXECUTOR.submit(self.__run_validated_annotations_zip, app,
                                 request.environ, task_id, id_list, is_async)
            return {'id': task_id,
                    'status': 'RUNNING'}, 202

        else:
            return self.__run_validated_annotations_zip(app, request.environ, task_id, id_list, is_async)


# =============================== XLSX Report ==================================


def createExcelWorksheet(filename):
    excel_file_path = os.path.join(
        app.config['UPLOAD_FOLDER'], filename)
    workbook = xlsxwriter.Workbook(excel_file_path)
    worksheet = workbook.add_worksheet()
    return worksheet

# ---------------------------------- Inspection Report XLSX -----------------------------


class InspectionXlsAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):
        super(InspectionXlsAPI, self).__init__()

    def get(self, id):
        """
        Generates a xls file with the inspection data
        ---
        parameters:
          - in: path
            required: true
            name: id
            schema:
              type: number
            description: Existing inspection id
        consumes:
          - application/json
        produces:
          - application/vnd.openxmlformats-officedocument.spreadsheetml.sheet
        responses:
          200:
            description: excel file with the inspection, images and defects data
            schema:
              type: file
        """

        include_360_images = False
        image_file_list = []
        horizontal_image_distance = 4

        if include_360_images:
            horizontal_image_distance = 15

        unique_name = str(uuid.uuid4())
        report_filename = 'report_'+unique_name+'.xlsx'
        logger.info(f'Creating temp excel file: {report_filename}')

        excel_file_path = os.path.join(
            app.config['UPLOAD_FOLDER'], report_filename)
        workbook = xlsxwriter.Workbook(excel_file_path)
        worksheet = workbook.add_worksheet()
        logging.info(f'temp report filename: {excel_file_path}')

        row = 0
        col = 0
        bold = workbook.add_format({'bold': True})
        worksheet.write(row, col, 'Inspection Report', bold)
        row += 2
        logger.info(f'root url: {request.url_root}')
        # logger.info(f'Reading inspection content for id: {id}')
        with Session(db) as session:
            inspection_list = session.scalars(
                select(Inspection).where(Inspection.id == id)
            ).all()
            if inspection_list is None:
                return {'message': f'Inspection record for {id} not found'}
            inspection = inspection_list[0]

            worksheet.write(row, col, 'Customer:', bold)
            worksheet.write(row, col+1, inspection.customer_name)
            row += 1

            worksheet.write(row, col, 'Location:', bold)
            worksheet.write(row, col+1, inspection.location)
            row += 1

            worksheet.write(row, col, 'Date:', bold)
            worksheet.write(row, col+1, str(inspection.date))
            row += 1

            worksheet.write(row, col, 'Engine Type:', bold)
            worksheet.write(row, col+1, inspection.engine_type)
            row += 1

            worksheet.write(row, col, 'ESN:', bold)
            worksheet.write(row, col+1, inspection.esn)
            row += 1

            # Inspection properties...

            # "app_type": "crawler-thetav",
            # "customer_name": "Garden City Lay Down Yard",
            # "date": "Mon, 16 Aug 2021 00:00:00 GMT",
            # "disp": "30",
            # "engine_type": "Blade Crawler-THETA-V",
            # "esn": "J80812",
            # "id": 2,
            # "location": "NA",
            # "misc": "J80812-07561-W860",
            # "sect": "te_uw",

            row += 2

            worksheet.write(row, col, 'Distance', bold)
            worksheet.write(row, col+1, 'Location', bold)
            worksheet.write(row, col+2, 'Severity', bold)
            worksheet.write(row, col+3, 'Size', bold)
            worksheet.write(row, col+4, 'Description', bold)

            if include_360_images:
                worksheet.write(row, col+5, 'Image Url', bold)
            else:
                worksheet.write(row, col+5, 'Findings', bold)

            # Image data format:
            # "blade_id": 0,
            # "defect_desc": "defect description",
            # "defect_location": "le_uw",
            # "defect_severity": 5,
            # "defect_size": 10.5,
            # "distance": 0,
            # "id": 0,
            # "image_file_id": 0,
            # "inspection_id": 0,
            # "timestamp": "Mon, 16 Aug 2021 00:00:00 GMT"

            base_url = 'http://vrn1masda.crd.ge.com:3000'

            row += 1
            inspection_image_list = session.scalars(
                select(Image).where(Image.inspection_id ==
                                    id).order_by(Image.distance)
            ).all()

            scale = 0.1
            use_thumbnail = False
            col_img_step = 2

            # logging.info(f'images for inspection {inspection_image_list}')
            for image in inspection_image_list:
                logger.info(f'image id: {str(image.id)}')

                # all measurement files extracted from the current image
                measurement_image_file_list = session.scalars(
                    select(MeasurementImageFile).where(
                        MeasurementImageFile.image_id == image.id)
                ).all()

                if len(measurement_image_file_list) > 0:
                    worksheet.write(row, col, image.distance)
                    worksheet.write(row, col+1, image.defect_location)
                    worksheet.write(row, col+2, image.defect_severity)
                    worksheet.write(row, col+3, image.defect_size)
                    worksheet.write(row, col+4, image.defect_desc)

                if include_360_images:
                    worksheet.write(
                        row, col+5, f'{base_url}/api/image/{image.id}/file')
                    row += 1
                    if use_thumbnail is True:
                        scale = 1.0
                        horizontal_image_distance = 5
                    image_file_path = download_image_file(
                        image.id, session, use_thumbnail)
                    if image_file_path is not None:
                        worksheet.insert_image(
                            row, col+5, image_file_path, {"x_scale": scale, "y_scale": scale})
                        # save so we can delete later
                        image_file_list.append(image_file_path)
                    col_img_step = 10

                col += col_img_step

                for measurement_image_file in measurement_image_file_list:
                    logger.info(
                        f'measurement_image_file: {str(measurement_image_file.filename)}')

                    measurement_image_file_path = save_image_file(
                        measurement_image_file, use_thumbnail)
                    if measurement_image_file_path is not None:
                        col += 3  # shapshot size
                        worksheet.write(
                            row, col, f'{measurement_image_file.filename}')
                        worksheet.insert_image(
                            row+1, col, measurement_image_file_path, {"x_scale": scale, "y_scale": scale})
                        # save so we can delete later
                        image_file_list.append(measurement_image_file_path)

                col = 0
                row += horizontal_image_distance

        workbook.close()

        @after_this_request
        def remove_file(response):
            logging.info(f'deleting file: {excel_file_path}')
            try:
                os.remove(excel_file_path)
                for image_file in image_file_list:
                    os.remove(image_file)
            except Exception as error:
                app.logger.error(
                    "Error removing or closing downloaded file:", error)

            return response

        return send_file(os.path.abspath(excel_file_path), as_attachment=True)


# -------------------------- Cross Section and Blade Position ----------------------------------


def send_image_file(filepath, download_filename):
    with open(filepath, mode='rb') as file:  # b is important -> binary
        file_content = file.read()
        if file_content is not None:
            return send_file(
                BytesIO(file_content),
                mimetype='image/jpeg',
                as_attachment=True,
                download_name=download_filename)

    return {'message': f'file: {filepath} not found'}


class BladeCrossSectionPositionAPI(Resource):
    def __init__(self):
        self.reqparse = reqparse.RequestParser()

        self.reqparse.add_argument('distance', type=float, location='args')
        self.reqparse.add_argument('location', type=str, location='args')
        self.reqparse.add_argument('color', type=str, location='args')

        super(BladeCrossSectionPositionAPI, self).__init__()

    def get(self):
        """
        Returns an image of a blade cross section with a dot marking a defect position
        ---
        consumes:
          - application/json
        produces:
          - application/json
        parameters:
          - in: query
            required: true
            name: distance
            schema:
              type: number
            description: Distance from blade root
            example: 3.5
          - in: query
            required: true
            name: location
            schema:
              type: string
              enum: ['Leading Edge', 'Center Web', 'Trailing Edge', 'Third Web', 'C Stiffener']
            description: Location in the blade
            example: 'Trailing Edge'
          - in: query
            required: true
            name: color
            schema:
              type: string
            description: Color of the dot to be painted at location and distance
            example: 'red'
        responses:
          200:
            description: image file of a blade cross section with a color dot indicating defect position
            schema:
              id: CrossSectionFileContent
              type: string
              format: binary
        """

        args = self.reqparse.parse_args()
        logger.info(args.items())

        location = args['location']
        color = args['color']
        # TODO: use different cross section images based on the distance from root
        distance = args['distance']

        with tempfile.TemporaryDirectory(dir=UPLOAD_FOLDER, prefix='temp-', suffix='-folder') as tempdirname:
            # image with the location dot that we will generate
            output_filename = os.path.join(
                tempdirname, 'blade_cross_section.jpg')
            generate_blade_cross_section(location, color, output_filename)

            download_filename = 'blade_cross_section.jpg'
            return send_image_file(output_filename, download_filename)


class BladeSideViewPositionAPI(Resource):
    def __init__(self):
        self.reqparse = reqparse.RequestParser()

        self.reqparse.add_argument('distance', type=float, location='args')
        self.reqparse.add_argument('location', type=str, location='args')
        self.reqparse.add_argument('color', type=str, location='args')

        super(BladeSideViewPositionAPI, self).__init__()

    def get(self):
        """
        Returns an image of a blade with a dot marking a defect position
        ---
        consumes:
          - application/json
        produces:
          - application/json
        parameters:
          - in: query
            required: true
            name: distance
            schema:
              type: number
            description: Distance from blade root
            example: 3.5
          - in: query
            required: true
            name: location
            schema:
              type: string
              enum: ['Leading Edge', 'Center Web', 'Trailing Edge', 'Third Web', 'C Stiffener']
            description: Location in the blade
            example: 'Trailing Edge'
          - in: query
            required: true
            name: color
            schema:
              type: string
            description: Color of the dot to be painted at location and distance
            example: 'red'
        responses:
          200:
            description: image file of a blade with a color dot indicating defect position
            schema:
              id: BladePositionFileContent
              type: string
              format: binary
        """
        args = self.reqparse.parse_args()
        logger.info(args.items())

        distance = args['distance']
        location = args['location']
        color = args['color']

        with tempfile.TemporaryDirectory(dir=UPLOAD_FOLDER, prefix='temp-', suffix='-folder') as tempdirname:
            # image with the location dot that we will generate
            output_filename = os.path.join(tempdirname, 'blade_position.jpg')
            generate_blade_side_view_position(
                distance, location, color, output_filename)

            download_filename = 'blade_side_view.jog'
            return send_image_file(output_filename, download_filename)


# -------------------------------- Defect Severity Table -----------------------------
class DefectSeverityAPI(Resource):
    def __init__(self):
        super(DefectSeverityAPI, self).__init__()

    def get(self):
        """
        Returns table of defects and their severity
        ---
        consumes:
          - application/json
        produces:
          - application/json
        responses:
          200:
            description: Table of defect types and their severity
        """
        return defect_severity

# -------------------------------- Defect Colors Table -----------------------------


class DefectColorsAPI(Resource):
    def __init__(self):
        super(DefectColorsAPI, self).__init__()

    def get(self):
        """
        Returns table of defects and their colors
        ---
        consumes:
          - application/json
        produces:
          - application/json
        responses:
          200:
            description: Table of defect types and their colors
        """
        return defect_bgr_colors


# we assume inspection[prop] is of type string
def merge_inspection_list(inspection_list, inspection_props):
    logging.info(f'merging common data for {len(inspection_list)} inspections')
    merged_inspection = Inspection()
    for inspection in inspection_list:
        for prop in inspection_props:
            if getattr(merged_inspection, prop) == '' or getattr(merged_inspection, prop) is None:
                if getattr(inspection, prop) is not None and len(str(getattr(inspection, prop))) > 0:
                    setattr(merged_inspection, prop, getattr(inspection, prop))
    return merged_inspection

# --------------------------------------- docx virtual tour report --------------------------------


class VirtualTourReportDocxAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):
        self.reqparse = reqparse.RequestParser()
        self.reqparse.add_argument('esn', type=str, location='args')
        self.reqparse.add_argument(
            'manufacture_stage', type=str, location='args')
        super(VirtualTourReportDocxAPI, self).__init__()

    def get(self):
        """
        Generates a docx report with the virtual blade tour data for an esn
        ---
        parameters:
          - in: query
            required: true
            name: esn
            schema:
              type: string
            description: Blade serial number ESN
            example: J80812
          - in: query
            required: true
            name: manufacture_stage
            schema:
              type: string
            description: Produce report only for that stage
            example: Post Molding
        consumes:
          - application/json
        produces:
          - application/vnd.openxmlformats-officedocument.wordprocessingml.document
        responses:
          200:
            description: docx file with the virtual tour images of a blade
            schema:
              type: file
        """
        # return {'message':'Not yet implemented'}
        args = self.reqparse.parse_args()
        logger.info(args.items())

        esn = args['esn']
        logging.info(f'esn: {esn}')

        manufacture_stage = "%" if (
            args['manufacture_stage'] is None or args['manufacture_stage'] == '') else args['manufacture_stage']
        logging.info(f'manufacture_stage: {manufacture_stage}')

        with tempfile.TemporaryDirectory(dir=UPLOAD_FOLDER, prefix='temp-', suffix='-folder') as tempdirname:
            logging.info(f'created temporary directory {tempdirname}')
            docx_file_path, msg = generate_docx_virtual_tour_report(
                esn, manufacture_stage, tempdirname)
            if docx_file_path is not None:
                return send_file(os.path.abspath(docx_file_path), as_attachment=True, download_name=f'VirtualTour-{esn}.docx')
            else:
                return msg


class VirtualTourReportPdfAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):
        # self.TASK_EXECUTOR = ThreadPoolExecutor(THREAD_POOL_SIZE)
        self.reqparse = reqparse.RequestParser()
        self.reqparse.add_argument('esn', type=str, location='args')
        self.reqparse.add_argument(
            'manufacture_stage', type=str, location='args')
        super(VirtualTourReportPdfAPI, self).__init__()

    def __run_generate_vt_pdf_report(self, current_app, environ, task_id, is_async, esn, manufacture_stage):
        with current_app.request_context(environ):
            with tempfile.TemporaryDirectory(dir=UPLOAD_FOLDER, prefix='temp-', suffix='-folder') as tempdirname:

                if is_async:
                    # we download lots of images included in the report, so we need a separate dir
                    tempdirname = f'{TEMP_DOWNLOAD_FILES_DIR}/{task_id}'
                    if not os.path.exists(tempdirname):
                        os.makedirs(tempdirname)

                logging.info(f'created temporary directory {tempdirname}')
                docx_file_path, msg = generate_docx_virtual_tour_report(
                    esn, manufacture_stage, tempdirname)

                if docx_file_path is None:
                    return msg

                pdf_file_path = docx_file_path.replace('docx', 'pdf')
                doc2pdf.convert(docx_file_path, pdf_file_path)
                if manufacture_stage == '%':
                    manufacture_stage = 'All_Stages'
                output_filename = f'VirtualTourReport-{esn}-{manufacture_stage}.pdf'

                if is_async:
                    TASK_STATUS_REGISTRY[task_id] = {
                        'id': task_id,
                        'timestamp': datetime.now().timestamp(),
                        'status': 'COMPLETE',
                        'filename': output_filename,
                        'path': os.path.abspath(pdf_file_path)
                    }
                    logging.info(f'Set TASK_STATUS: for task_id: {task_id} to: {TASK_STATUS_REGISTRY[task_id]}')
                else:
                    return send_file(os.path.abspath(pdf_file_path), as_attachment=True, download_name=output_filename)

    def get(self):
        """
        Generates a pdf report with the virtual tour images for a blade  esn
        ---
        parameters:
          - in: query
            required: true
            name: esn
            schema:
              type: string
            description: Blade serial number ESN
            example: J80812
          - in: query
            required: true
            name: manufacture_stage
            schema:
              type: string
            description: Produce report only for that stage
            example: Post Molding
          - in: query
            required: false
            name: async
            schema:
              type: boolean
            description: use asynchronous strategy with background tasks
            default: false
        consumes:
          - application/json
        produces:
          - application/pdf
        responses:
          200:
            description: pdf file with the virtual tour images of a blade
            schema:
              type: file
        """
        # return {'message':'Not yet implemented'}
        args = self.reqparse.parse_args()
        logger.info(args.items())

        esn = args['esn']
        logging.info(f'esn: {esn}')

        # booleans are handled differently in python, so we use string inputs
        async_str = request.args.get('async')
        if async_str is not None:
            async_str = async_str.capitalize()
        logger.info(f'asyncStr: {async_str}')
        is_async = False
        if async_str == 'True':
            is_async = True
        logger.info(f'is_async: {is_async}')

        manufacture_stage = "%" if (
            args['manufacture_stage'] is None or args['manufacture_stage'] == '') else args['manufacture_stage']
        logging.info(f'manufacture_stage: {manufacture_stage}')

        # parameter and inspection validation
        with Session(db) as session:
            inspection_list = session.scalars(
                select(Inspection).filter(
                    Inspection.esn == esn).filter(
                        Inspection.manufacture_stage.ilike(manufacture_stage))
            ).all()

            for inspection in inspection_list:
                logging.info(f'inspection.id: {inspection.id}')

            if inspection_list is None or len(inspection_list) == 0:
                return {'message': f'No inspections for esn: {esn} and manufacture_stage: {manufacture_stage} were found.'}

        task_id = uuid.uuid4().hex
        if is_async:
            TASK_STATUS_REGISTRY[task_id] = {
                'id': task_id,
                'timestamp': datetime.now().timestamp(),
                  'status': 'RUNNING',
                  'filename': None,
            }
            TASK_EXECUTOR.submit(self.__run_generate_vt_pdf_report, app,
                                 request.environ, task_id, is_async, esn, manufacture_stage)
            resp = {'id': task_id,
                    'status': 'RUNNING'}
            logging.info(f'returning: {resp}')
            return resp, 202

        else:
            return self.__run_generate_vt_pdf_report(app,
                                                  request.environ, task_id, is_async, esn, manufacture_stage)

        # with tempfile.TemporaryDirectory(dir=UPLOAD_FOLDER, prefix='temp-', suffix='-folder') as tempdirname:
        #     logging.info(f'created temporary directory {tempdirname}')
        #     docx_file_path, msg = generate_docx_virtual_tour_report(
        #         esn, manufacture_stage, tempdirname)

        #     if docx_file_path is not None:
        #         pdf_file_path = docx_file_path.replace('docx', 'pdf')
        #         doc2pdf.convert(docx_file_path, pdf_file_path)
        #         return send_file(os.path.abspath(pdf_file_path), as_attachment=True, download_name=f'VirtualTour-{esn}.pdf')
        #     else:
        #         return msg


def generate_docx_virtual_tour_report(esn, manufacture_stage, tempdirname):
    inspection_images_temp_path = tempdirname+'/images'
    logging.info(f'creating images folder: {inspection_images_temp_path}')
    os.makedirs(inspection_images_temp_path, exist_ok=True)

    unique_name = str(uuid.uuid4())
    report_filename = f'virtual-tour_{esn}_'+unique_name+'.docx'
    logger.info(f'Creating temp docx file: {report_filename}')

    docx_file_path = os.path.join(tempdirname, report_filename)
    image_file_list = []  # save path of temp image files we download

    with Session(db) as session:
        inspection_list = session.scalars(
            select(Inspection).filter(
                Inspection.esn == esn).filter(
                    Inspection.manufacture_stage.ilike(manufacture_stage))
        ).all()

        for inspection in inspection_list:
            logging.info(f'inspection.id: {inspection.id}')

        if inspection_list is None or len(inspection_list) == 0:
            return (None, {'message': f'No inspections for esn: {esn} were found.'})

        blade_areas_set = set()
        for inspection in inspection_list:
            sect = inspection.sect
            sec_name = get_section_name(sect)
            if 'Leading Edge' == sec_name:
                blade_areas_set.add(
                    'Leading_Edge_Internal_Cavity (Crawler - Visual)')
            elif 'Trailing Edge' == sec_name:
                blade_areas_set.add(
                    'Trailing_Edge_Internal_Cavity (Crawler - Visual)')
            elif 'Center Web' == sec_name:
                blade_areas_set.add(
                    'Central_Web_Internal_Cavity (Crawler - Visual)')
            elif 'Third Web' == sec_name:
                blade_areas_set.add(
                    'Third_Web_Internal_Cavity (Crawler - Visual)')
            elif 'C Stiffener' == sec_name:
                blade_areas_set.add(
                    'C_Stiffener_Internal_Cavity (Crawler - Visual)')
        areas_inspected = ''
        for area in blade_areas_set:
            if len(areas_inspected) > 0:
                areas_inspected += '\n'
            areas_inspected += area

        # Create a new document
        doc = Document()

        configure_styles(doc)

        # Add a title
        # title = doc.add_heading('Digital Blade Inspection Report', level=1)
        # title = add_heading(doc,  'Digital Blade Inspection Report', 1)
        # title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        add_report_header(doc, inspection.esn, 'Blade Internal Images')
        add_report_footer(doc, inspection.esn, 'In-Factory Internal Images')

        # Add a paragraph with bold and italic text
        # paragraph = doc.add_paragraph('This is a sample document created using the python-docx library.')
        # run = paragraph.runs[0]
        # run.bold = True
        # run.italic = True

        # -------------------------------------------- Front Page ------------------------------------
        # Add a heading
        add_heading(doc, 'Inspection Details', level=2)
        # first_inspection = inspection_list[0]

        # Note we excluded sect -> Section which we will use to distinguish each inspection
        inspection_table_labels = [
            "Blade Serial Number",
            "Blade Model",
            "Supplier Name",
            "Factory Location",
            "Factory Name",
            "Manufacture Date",
            "Manufacturing Inspection Stage",
            "Imaging Equipment",
            "Blade Areas Imaged",
            "Report Creation Date"]
        inspection_props = [
            "esn",
            "blade_type",
            "customer_name",
            "location",
            "factory_name",
            "manufacture_date",
            "manufacture_stage",
            "app_type",
            "sect",
            "date"]

        # make sure if the user fixes a missing prop for one inspection, it gets seen in this combined report
        merged_inspeciton = merge_inspection_list(
            inspection_list, inspection_props)

        inspection_data_list = []
        inspection_data = []
        for prop in inspection_props:
            # replace the sect property with the list of areas inspected
            if prop == 'esn':
                inspection_data.append(
                    str(getattr(merged_inspeciton, prop)).upper())
            elif prop == 'sect':
                inspection_data.append(areas_inspected.replace('_', " "))
            elif prop == 'date':
                # inspection_data.append( datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
                now = datetime.now(timezone.utc)
                dt_string = now.strftime("%Y-%m-%d %H:%M:%S UTC")
                inspection_data.append(dt_string)
            else:
                attr = str(getattr(merged_inspeciton, prop))
                if 'date' in prop:
                    attr = attr.split(' ')[0]  # remove time part from date
                # logging.info(f'{prop}:{getattr(inspection, prop)}')
                inspection_data.append(attr)
        inspection_data_list.append(inspection_data)

        # logging.info(f'inspection_table_labels: {inspection_table_labels}')
        # logging.info(f'inspection_data_list: {inspection_data_list}')
        add_data_table(doc, inspection_table_labels, inspection_data_list,
                       labels_on_top=False, label_width=200, data_width=400)

        category_names, category_data, distance_list, category_list, location_list, category_severity_list, disposition_list, measurement_list = gether_stats(
            session, inspection_list)

        measurement_id_list = []
        for vshot in measurement_list:
            measurement_id_list.append(vshot.id)

        # --------------------------------- Individual 2d image pages -------------------------------
        for inspection in inspection_list:
            logging.info(f'adding images for inspection: {inspection}')

            section_name = get_section_name(inspection.sect)
            logging.info(f'section name: {section_name} for sect: {inspection.sect}')

            doc.add_page_break()
            # doc.add_heading(f'{section_name}', level=2)
            add_heading(doc, section_name, 2)

            inspection_image_list = session.scalars(
                select(Image).where(Image.inspection_id ==
                                    inspection.id).order_by(Image.distance)
            ).all()

            if inspection_image_list is None or len(inspection_image_list) == 0:
                add_line_break(doc)
                doc.add_paragraph(
                    'No findings recorded for this blade section.')

            logging.info(
                f'adding {len(inspection_image_list)} images from inspection {inspection.id}')

            futures = []
            with ThreadPoolExecutor(THREAD_POOL_SIZE) as executor:
                # for each 360 image record...
                for image in inspection_image_list:
                    # logger.info(f'image: {str(image)}')
                    futures.append(executor.submit(download_360image_vtshots_to_temp_folder,
                                                   session, image, inspection, inspection_images_temp_path))
                    # add_image_measurements_to_document(doc, session, image, inspection, inspection_images_temp_path)
                wait(futures)

            # futures is an array with elements for each image
            # each element is an array of images on the same round_distance (ideally 8 elements)
            # each element has the properties:
            # ('round_distance', 'finding_label', 'measurement', 'measurement_image_file_path')
            # where measurement is the DB record with all its properties

            # png_jpg_image_dict = {}
            # image_path_list = []
            # for future in futures:
            #     if future.result() is not None and len(future.result()) > 0:
            #         results_list = future.result()  # list of annotated 2d snapshots
            #         for idx, result in enumerate(results_list):
            #             image_path_list.append(
            #                 result['vshot_image_file_path'])
            # png_jpg_image_dict = convert_png_to_jpg_in_parallel(
            #     image_path_list)

            for future in futures:
                if future.result() is not None and len(future.result()) > 0:
                    results_list = future.result()  # list of annotated 2d snapshots
                    image_distance = results_list[0]['round_distance']
                    doc.add_paragraph(
                        f'{section_name} - Z distance: {image_distance}')

                    img_table = doc.add_table(rows=1, cols=2)
                    img_table.style = 'Table Grid'
                    img_table.autofit = True
                    img_table.allow_autofit = True
                    img_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    last_paragraph = doc.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    left_cell = img_table.cell(0, 0)
                    right_cell = img_table.cell(0, 1)
                    left_run = left_cell.paragraphs[0].add_run()
                    right_run = right_cell.paragraphs[0].add_run()

                    finding_label = ""
                    severity = get_defect_severity(finding_label)
                    severity_color = get_defect_severity_color(severity)
                    blade_view_path = inspection_images_temp_path + \
                        f'/blade_view_{section_name}_{image_distance}.jpg'
                    generate_blade_side_view_position(
                        image_distance, inspection.sect, severity_color, blade_view_path)
                    left_run.add_picture(blade_view_path, width=Pt(200))

                    cross_section_path = inspection_images_temp_path + \
                        f'/cross_section_{section_name}_{image_distance}.jpg'
                    generate_blade_cross_section(
                        inspection.sect, severity_color, cross_section_path)
                    right_run.add_picture(cross_section_path, width=Pt(200))

                    number_cols = 2
                    number_rows = len(results_list) // number_cols + 1

                    logging.info(f'image table dimensions: {number_rows},{number_cols}')
                    table = doc.add_table(rows=number_rows, cols=number_cols)
                    table.style = 'Table Grid'
                    table.autofit = True
                    table.allow_autofit = True
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    last_paragraph = doc.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    logging.info(f'results_list len: {len(results_list)}')
                    for idx, result in enumerate(results_list):
                        col = idx % number_cols
                        row = idx // number_cols
                        logging.info(f'row,col: {row},{col}')

                        image_file_path = result['vshot_image_file_path']
                        vshot = result['vshot']
                        cell = table.cell(row, col)
                        run = cell.paragraphs[0].add_run()

                        run.add_picture(image_file_path, width=Inches(3))
                        # run.add_picture(convert_file_to_jpg(image_file_path), width=Inches(3))

                        run.add_text(f'pitch: {round(vshot.image_pitch, 2)}, yaw:{round(vshot.image_yaw, 2)}, hfov:{round(vshot.image_hfov, 2)}')

                    doc.add_page_break()

    # save in the context of the temp folder. It will be deleted when the response is delivered
    doc.save(docx_file_path)
    return (docx_file_path, {'message': 'document successfully generated.'})


# # return a dictionary from the original png file to the new jpg file
# def convert_png_to_jpg_in_parallel(path_list):
#     logging.info('convert_png_to_jpg_in_parallel()')
#     futures = []
#     with ThreadPoolExecutor(THREAD_POOL_SIZE) as executor:
#         # for each 360 image record...
#         for image_file_path in path_list:
#             logger.info(f'call png2jpg: {image_file_path}')
#             # Note: function name is a prop in executor, followed by its attributes
#             futures.append(executor.submit(
#                 convert_png_file_to_jpg, image_file_path))
#             # add_image_measurements_to_document(doc, session, image, inspection, inspection_images_temp_path)
#         wait(futures)
#     result_dict = {}
#     for idx, future in enumerate(futures):
#         result_dict[str(path_list[idx])] = str(future.result())
#     return result_dict

# --------------------------------------- docx inspection report ----------------------------------


class InspectionReportDocxAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):
        self.reqparse = reqparse.RequestParser()
        self.reqparse.add_argument('esn', type=str, location='args')
        self.reqparse.add_argument(
            'manufacture_stage', type=str, location='args')
        super(InspectionReportDocxAPI, self).__init__()

    def get(self):
        """
        Generates a docx report with the inspection data for an esn
        ---
        parameters:
          - in: query
            required: true
            name: esn
            schema:
              type: string
            description: Blade serial number ESN
            example: J80812
          - in: query
            required: true
            name: manufacture_stage
            schema:
              type: string
            description: Include only inspections for the provided Manufacture Stage
            example: Post Molding
        consumes:
          - application/json
        produces:
          - application/vnd.openxmlformats-officedocument.wordprocessingml.document
        responses:
          200:
            description: docx file with the inspection, images and defects data
            schema:
              type: file
        """
       
        # return {'message':'Not yet implemented'}
        args = self.reqparse.parse_args()
        logger.info(args.items())

        esn = args['esn']
        logging.info(f'esn: {esn}')

        manufacture_stage = "%" if (
            args['manufacture_stage'] is None or args['manufacture_stage'] == '') else args['manufacture_stage']
        logging.info(f'manufacture_stage: {manufacture_stage}')

        with tempfile.TemporaryDirectory(dir=UPLOAD_FOLDER, prefix='temp-', suffix='-folder') as tempdirname:
            logging.info(f'created temporary directory {tempdirname}')
            docx_file_path, msg = generate_docx_inspection_report(
                esn, manufacture_stage, tempdirname)
            if docx_file_path is not None:
                if manufacture_stage == '%':
                    manufacture_stage = 'All_Stages'
                return send_file(os.path.abspath(docx_file_path), as_attachment=True, download_name=f'InspectionReport-{esn}-{manufacture_stage}.docx')
            else:
                return msg


class InspectionReportPdfAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):
        # self.TASK_EXECUTOR = ThreadPoolExecutor(THREAD_POOL_SIZE)
        # self.reqparse = reqparse.RequestParser()
        # self.reqparse.add_argument('esn', type=str, location='args')
        # self.reqparse.add_argument('manufacture_stage', type=str, location='args')
        # self.reqparse.add_argument('async', type=str, location='args')
        super(InspectionReportPdfAPI, self).__init__()

    def __run_generate_insp_pdf_report(self, current_app, environ, task_id, is_async, esn, manufacture_stage):
        with current_app.request_context(environ):
            with tempfile.TemporaryDirectory(dir=UPLOAD_FOLDER, prefix='temp-', suffix='-folder') as tempdirname:

                if task_id in TASK_STATUS_REGISTRY:
                    logging.info(f'task status: {TASK_STATUS_REGISTRY[task_id]}')

                logging.info(f'created temporary directory {tempdirname}')

                output_folder = tempdirname
                if is_async:
                    output_folder = f'{TEMP_DOWNLOAD_FILES_DIR}/{task_id}'
                    if not os.path.exists(output_folder):
                        os.makedirs(output_folder)

                docx_file_path, msg = generate_docx_inspection_report(
                    esn, manufacture_stage, output_folder)

                if docx_file_path is None:
                    return msg

                pdf_file_path = docx_file_path.replace('docx', 'pdf')

                logging.info(f'pdf_file_path: {pdf_file_path}')
                logging.info(f'docx_file_path: {docx_file_path}')

                doc2pdf.convert(docx_file_path, pdf_file_path)
                if manufacture_stage == '%':
                    manufacture_stage = 'All_Stages'
                output_filename = f'InspectionReport-{esn}-{manufacture_stage}.pdf'

                if is_async:
                    TASK_STATUS_REGISTRY[task_id] = {
                        'id': task_id,
                        'timestamp': datetime.now().timestamp(),
                        'status': 'COMPLETE',
                        'filename': output_filename,
                        'path': os.path.abspath(pdf_file_path)
                    }
                    logging.info(f'Set TASK_STATUS: for task_id: {task_id} to: {TASK_STATUS_REGISTRY[task_id]}')
                else:
                    logging.info(f'returning file from: {pdf_file_path}')
                    return send_file(os.path.abspath(pdf_file_path), as_attachment=True, download_name=output_filename)

    def __handle_request(self, request):
        
        esn = request.args.get('esn')
        logging.info(f'esn: {esn}')

        manufacture_stage_str = request.args.get('manufacture_stage')
        logging.info(f'manufacture_stage_str: {manufacture_stage_str}')

        # booleans are handled differently in python, so we use string inputs
        async_str = request.args.get('async')
        if async_str is not None:
            async_str = async_str.capitalize()
        logger.info(f'asyncStr: {async_str}')
        is_async = False
        if async_str == 'True':
            is_async = True
        logger.info(f'is_async: {is_async}')

        manufacture_stage = "%" if (
           manufacture_stage_str is None or manufacture_stage_str == '') else manufacture_stage_str
        logging.info(f'manufacture_stage: {manufacture_stage}')

        # parameter and inspection validation
        with Session(db) as session:
            inspection_list = session.scalars(
                select(Inspection).filter(
                    Inspection.esn == esn).filter(
                        Inspection.manufacture_stage.ilike(manufacture_stage))
            ).all()

            for inspection in inspection_list:
                logging.info(f'inspection.id: {inspection.id}')

            if inspection_list is None or len(inspection_list) == 0:
                return {'message': f'No inspections for esn: {esn} and manufacture_stage: {manufacture_stage} were found.'}

        task_id = uuid.uuid4().hex
        if is_async:
            TASK_STATUS_REGISTRY[task_id] = {
                'id': task_id,
                'timestamp': datetime.now().timestamp(),
                  'status': 'RUNNING',
                  'filename': None,
            }
            TASK_EXECUTOR.submit(self.__run_generate_insp_pdf_report, app,
                                 request.environ, task_id, is_async, esn, manufacture_stage)
            resp = {'id': task_id,
                    'status': 'RUNNING'}
            logging.info(f'returning: {resp}')
            return resp, 202

        else:
            logging.info('return file synchronously.')
            return self.__run_generate_insp_pdf_report(app,
                                                  request.environ, task_id, is_async, esn, manufacture_stage)

    def get(self):
        """
        Generates a pdf report with the inspection data for an esn
        ---
        parameters:
          - in: query
            required: true
            name: esn
            schema:
              type: string
            description: Blade serial number ESN
            example: J80812
          - in: query
            required: true
            name: manufacture_stage
            schema:
              type: string
            description: Include only inspections for the provided Manufacture Stage
            example: Post Molding
          - in: query
            required: false
            name: async
            schema:
              type: boolean
            description: use asynchronous strategy with background tasks
            default: false
        consumes:
          - application/json
        produces:
          - application/pdf
        responses:
          200:
            description: pdf file with the inspection, images and defects data
            schema:
              type: file
        """
        return self.__handle_request(request)
    
    def post(self):
        """
        Generates a pdf report with the inspection data for an esn
        ---
        parameters:
          - in: query
            required: true
            name: esn
            schema:
              type: string
            description: Blade serial number ESN
            example: J80812
          - in: query
            required: true
            name: manufacture_stage
            schema:
              type: string
            description: Include only inspections for the provided Manufacture Stage
            example: Post Molding
          - in: query
            required: false
            name: async
            schema:
              type: boolean
            description: use asynchronous strategy with background tasks
            default: false
        consumes:
          - application/json
        produces:
          - application/pdf
        responses:
          200:
            description: pdf file with the inspection, images and defects data
            schema:
              type: file
        """
        return self.__handle_request(request)


def get_measurement_list_for_inspection(inspection, session):
    measurement_list = []
    inspection_image_list = session.scalars(
        select(Image).where(Image.inspection_id ==
                            inspection.id).order_by(Image.distance)
    ).all()
    for image in inspection_image_list:
        image_measurement_list = session.scalars(
            select(Measurement).where(Measurement.image_id == image.id)
        ).all()
        for measurement in image_measurement_list:
            measurement_list.append(measurement)
    return measurement_list


def generate_docx_inspection_report(esn, manufacture_stage, output_folder):

    inspection_images_temp_path = output_folder+'/images'
    logging.info(f'creating images folder: {inspection_images_temp_path}')
    os.makedirs(inspection_images_temp_path, exist_ok=True)

    unique_name = str(uuid.uuid4())
    report_filename = f'Report_{esn}_'+unique_name+'.docx'
    logger.info(f'Creating temp docx file: {report_filename}')

    docx_file_path = os.path.join(output_folder, report_filename)
    image_file_list = []  # save path of temp image files we download

    with Session(db) as session:
        inspection_list = session.scalars(
            select(Inspection).filter(Inspection.esn == esn).filter(
                Inspection.manufacture_stage.ilike(manufacture_stage))
        ).all()

        for inspection in inspection_list:
            logging.info(f'inspection.id: {inspection.id}')

        if inspection_list is None or len(inspection_list) == 0:
            return (None, {'message': f'No inspections for esn: {esn} were found.'})

        blade_areas_set = set()
        for inspection in inspection_list:
            sect = inspection.sect
            sec_name = get_section_name(sect)
            if 'Leading Edge' == sec_name:
                blade_areas_set.add(
                    'Leading_Edge_Internal_Cavity (Crawler - Visual)')
            elif 'Trailing Edge' == sec_name:
                blade_areas_set.add(
                    'Trailing_Edge_Internal_Cavity (Crawler - Visual)')
            elif 'Center Web' == sec_name:
                blade_areas_set.add(
                    'Central_Web_Internal_Cavity (Crawler - Visual)')
            elif 'Third Web' == sec_name:
                blade_areas_set.add(
                    'Third_Web_Internal_Cavity (Crawler - Visual)')
            elif 'C Stiffener' == sec_name:
                blade_areas_set.add(
                    'C_Stiffener_Internal_Cavity (Crawler - Visual)')
        areas_inspected = ''
        for area in blade_areas_set:
            if len(areas_inspected) > 0:
                areas_inspected += '\n'
            areas_inspected += area

        include_cad_models = True
        if 'stiff' in areas_inspected.lower() or 'third' in areas_inspected.lower():
            include_cad_models = False

        # Create a new document
        doc = Document()

        configure_styles(doc)

        # Add a title
        # title = doc.add_heading('Digital Blade Inspection Report', level=1)
        # title = add_heading(doc,  'Digital Blade Inspection Report', 1)
        # title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        add_report_header(doc, inspection.esn)
        add_report_footer(doc, inspection.esn)

        # Add a paragraph with bold and italic text
        # paragraph = doc.add_paragraph('This is a sample document created using the python-docx library.')
        # run = paragraph.runs[0]
        # run.bold = True
        # run.italic = True

        # -------------------------------------------- Front Page ------------------------------------
        # Add a heading
        add_heading(doc, 'Inspection Details', level=2)
        # first_inspection = inspection_list[0]

        # Note we excluded sect -> Section which we will use to distinguish each inspection
        inspection_table_labels = [
            "Blade Serial Number",
            "Blade Model",
            "Supplier Name",
            "Factory Location",
            "Factory Name",
            "Manufacture Date",
            "Manufacturing Inspection Stage",
            "Inspection Modality & HW",
            "Blade Areas Inspected",
            "Inspection Date",
            "Inspectors",
            "Quality Certification Status"]
        inspection_props = [
            "esn",
            "blade_type",
            "customer_name",
            "location",
            "factory_name",
            "manufacture_date",
            "manufacture_stage",
            "app_type",
            "sect",
            "date",
            "inspector_name",
            "certification_status"]

        # make sure if the user fixes a missing prop for one inspection, it gets seen in this combined report
        merged_inspection = merge_inspection_list(
            inspection_list, inspection_props)

        inspection_data_list = []
        inspection_data = []
        for prop in inspection_props:
            # replace the sect property with the list of areas inspected
            if prop == 'esn':
                inspection_data.append(
                    str(getattr(merged_inspection, prop)).upper())
            elif prop == 'sect':
                inspection_data.append(areas_inspected.replace('_', " "))
            else:
                # logging.info(f'{prop}:{getattr(inspection, prop)}')
                attr = str(getattr(merged_inspection, prop))
                if 'date' in prop:
                    attr = attr.split(' ')[0]  # remove time from date
                    logging.info(f'{prop}: {attr}')
                inspection_data.append(attr)

        now = datetime.now(timezone.utc)
        dt_string = now.strftime("%Y-%m-%d %H:%M:%S UTC")
        inspection_table_labels.append('Report Creation Date')
        inspection_data.append(dt_string)

        inspection_data_list.append(inspection_data)

        # logging.info(f'inspection_table_labels: {inspection_table_labels}')
        # logging.info(f'inspection_data_list: {inspection_data_list}')
        add_data_table(doc, inspection_table_labels, inspection_data_list,
                       labels_on_top=False, label_width=200, data_width=400)

        # ------------------- summary semi-circle charts ---------------------
        inspection_heading = add_heading(
            doc, 'Inspection Summary', level=2)
        inspection_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        category_names, category_data, distance_list, category_list, location_list, category_severity_list, disposition_list, measurement_list = gether_stats(
            session, inspection_list)

        # measurement_id_list = []
        # for measurement in measurement_list:
        #     measurement_id_list.append(measurement.id)

        # ---------------- probability of failure stats chart -------------------
        category_severity_dict = {}

        tbl_category_names = []
        tbl_category_data = []
        tbl_category_color = []
        chart_category_names = category_names
        chart_category_data = category_data
        chart_category_color = []

        # compute severity category totals
        logging.info(f'category names: {category_names}')
        for i, measurement in enumerate(category_names):
            status = get_defect_severity(measurement)
            if category_severity_dict.get(status) is None:
                category_severity_dict[status] = 0
            category_severity_dict[status] = category_severity_dict.get(
                status) + category_data[i]

        # table has HIGH, MEDIUM, LOW, SAFETY
        for key, value in category_severity_dict.items():
            tbl_category_names.append(key)
            tbl_category_data.append(value)
            tbl_category_color.append(get_defect_severity_color(key))

        # chart pie sections have colors based on severity
        for status in category_severity_list:
            chart_category_color.append(get_defect_severity_color(status))

        add_table_and_semi_circle_chart_content(
            doc, inspection_images_temp_path, "Indication Severity",
            tbl_category_names, tbl_category_data, tbl_category_color,
            chart_category_names, chart_category_data, chart_category_color)

      # --------------------- indication status chart -------------------------------
        add_line_break(doc)
        measurement_status_dict = {}

        for measurement in measurement_list:
            disp = measurement.ge_disposition
            disposition = disp if (
                disp is not None and disp.strip() != '') else 'Pending'
            if measurement_status_dict.get(disposition) is None:
                measurement_status_dict[disposition] = 0
            measurement_status_dict[disposition] = measurement_status_dict.get(
                disposition) + 1

        tbl_category_names = []
        tbl_category_data = []
        tbl_category_color = []
        for key, value in measurement_status_dict.items():
            tbl_category_names.append(key)
            tbl_category_data.append(value)
            tbl_category_color.append(get_disposition_color(key))

        chart_category_names = tbl_category_names
        chart_category_data = tbl_category_data
        chart_category_color = tbl_category_color

        add_table_and_semi_circle_chart_content(
            doc, inspection_images_temp_path, "Indication Status",
            tbl_category_names, tbl_category_data, tbl_category_color,
            chart_category_names, chart_category_data, chart_category_color)

        # -------------------------------------- Old charts ----------------------------------------
        # chart_filename = f'report_chart_{esn}_'+unique_name+'.png'
        # chart_file_path = os.path.join(inspection_images_path, chart_filename)

        # add_semi_circle_chart(doc, chart_file_path, category_names, category_data)
        # add_donut_chart(doc, chart_file_path, category_names, category_data)
        # add_sunburst_chart(doc, chart_file_path, category_names, category_data, category_severity)

        # doc.add_page_break()
        # -------------------------------------------- Summary Page ---------------------------------

        # number of defects per severity category
        # add_severity_total_table(doc, category_names, category_data)

        # simplified version of the sunburst, with defect counts by type only, no severity.
        # add_donut_chart(doc, chart_file_path, category_names, category_data)

        # total defect by category table
        # add_simple_table(doc, category_names, category_data)

        # doc.add_picture(total_blade_image_path, width=Inches(3))

        # doc.add_page_break()
        # ------------------------------------- New Defect Table of Contents -----------------------------

        # doc.add_heading('Defect Table of Contents', level=2)
        # add_heading(doc, 'Defect Table of Contents', 2)

        # page_counter = 4
        # toc_data_list = []
        # toc_labels = [
        #     'Defect ID',
        #     'Distance from Root (m)',
        #     'Defect Type',
        #     'Probability of Failure',
        #     'Defect Status',
        #     'Defect Details Page'
        # ]
        # logging.info(f'distance_list: {distance_list}')
        # for i in range(0, len(distance_list)):
        #     indication_row = []
        #     # mid is the unique identifier of the defect
        #     defect_id = inspection.esn+'-'+str(measurement_id_list[i])
        #     indication_row.append(defect_id),
        #     indication_row.append(str(round(distance_list[i], 1)))
        #     category = category_list[i]
        #     indication_row.append(category)
        #     status = defect_severity.get(category.replace('_', ' '))
        #     indication_row.append(status)
        #     indication_row.append('')  # status
        #     indication_row.append(str(page_counter))

        #     toc_data_list.append(indication_row)
        #     page_counter += 1

        # add_data_table(doc, toc_labels, toc_data_list,
        #                labels_on_top=True,  data_width=200)
        # doc.add_page_break()

        # ------------------------- Single total blade view -------------------
        # total_blade_image_path = f'{
        #     inspection_images_temp_path}/total_blade.png'
        # generate_total_blade_view(distance_list, location_list, category_severity_list, disposition_list, total_blade_image_path, 270)
        # ------------------------- End Single Total blade view ----------------

        # -------------------------------- New indicators TOC with blade view -----------------------
        # each section has one total blade view

        idx = 0
        for inspection in inspection_list:
            section_name = get_section_name(inspection.sect)
            add_heading(doc, f'{section_name} Defect Table of Contents', 2)

            category_names, category_data, distance_list, category_list, location_list, category_severity_list, disposition_list, measurement_list = gether_stats(
                session, [inspection])

            total_blade_image_path = None
            if include_cad_models:
                total_blade_image_path = f'{inspection_images_temp_path}/total_blade_{idx}.jpg'
                generate_total_blade_view(
                    distance_list, location_list, category_severity_list, disposition_list, total_blade_image_path, 270)

            # if total_blade_image_path is None this method will not include cad model
            add_indicators_table_and_side_blade(
                doc, inspection, measurement_list, total_blade_image_path)
            idx += 1

        # --------------------------------- Individual 2d image pages -------------------------------
        logging.info(f'2d image pages for {len(inspection_list)} inspections:')
        for inspection in inspection_list:
            logging.info(f'inspection.id: {inspection.id}')

        for inspection in inspection_list:
            logging.info(f'adding images for inspection: {inspection}')

            section_name = get_section_name(inspection.sect)
            logging.info(f'section name: {section_name} for sect: {inspection.sect}')

            doc.add_page_break()
            # doc.add_heading(f'{section_name}', level=2)
            add_heading(doc, section_name, 2)

            inspection_image_list = session.scalars(
                select(Image).where(Image.inspection_id ==
                                    inspection.id).order_by(Image.distance)
            ).all()
            logging.info(f'Found {len(inspection_image_list)} 360 images in inspection {inspection.id}')

            if inspection_image_list is None or len(inspection_image_list) == 0:
                add_line_break(doc)
                doc.add_paragraph(
                    'No findings recorded for this section.')

            logging.info(
                f'adding {len(inspection_image_list)} images for inspection id: {inspection.id}')

            # logging.info('running in series...')
            # for image in inspection_image_list:
            #     results_list = download_360image_defects_with_annotations_to_temp_folder(
            #                                        session, image, inspection, inspection_images_temp_path)
            #     for result in results_list:
            #             # add_line_break(doc)
            #             image_file_path = result['measurement_image_file_path']
            #             # result props: 'round_distance''finding_label','measurement','measurement_image_file_path'
            #             add_measurement_image_table(
            #                 doc,
            #                 result['round_distance'],
            #                 result['finding_label'],
            #                 section_name,  # from inspection
            #                 result['measurement'],
            #                 inspection,
            #                 result['measurement_image_file_path'],
            #                 include_cad_models)
            #             # add_line_break(doc)
            #             doc.add_page_break()

            logging.info('running in parallel...')
            futures = []
            with ThreadPoolExecutor(THREAD_POOL_SIZE) as executor:
                # for each 360 image record...
                for image in inspection_image_list:
                    # logger.info(f'image: {str(image)}')
                    # produces a list of dictionaries with the following
                    # props: {'round_distance','finding_label','measurement','measurement_image_file_path'}

                    # futures.append(executor.submit(download_360image_measurements_with_annotations_to_temp_folder,
                    #                               session, image, inspection, inspection_images_temp_path))
                    futures.append(executor.submit(download_360image_defects_with_annotations_to_temp_folder,
                                                   session, image, inspection, inspection_images_temp_path))
                    # add_image_measurements_to_document(doc, session, image, inspection, inspection_images_temp_path)
                wait(futures)

            for future in futures:
                if future.result() is not None:
                    # for each call, unpack the return values
                    results_list = future.result()

                    # result_list is a list of dictionaries with the following
                    # props: {'round_distance','finding_label','measurement','measurement_image_file_path'}
                    for result in results_list:
                        # add_line_break(doc)
                        image_file_path = result['measurement_image_file_path']
                        # result props: 'round_distance''finding_label','measurement','measurement_image_file_path'
                        add_measurement_image_table(
                            doc,
                            result['round_distance'],
                            result['finding_label'],
                            section_name,  # from inspection
                            result['measurement'],
                            inspection,
                            result['measurement_image_file_path'])
                        # add_line_break(doc)
                        doc.add_page_break()

    # save in the context of the temp folder. It will be deleted when the response is delivered
    doc.save(docx_file_path)
    return (docx_file_path, {'message': 'document successfully generated.'})


def download_360image_measurements_with_annotations_to_temp_folder(session, image, inspection, inspection_images_temp_path):
    section_name = get_section_name(inspection.sect)
    round_distance = str(round(image.distance, 1))

    # image_file with only what we need: id and filename
    image_file_results = session.query(
        ImageFile.id, ImageFile.filename).filter(
            ImageFile.image_id == image.id
    ).all()

    # image file with full content
    # image_file_results = session.scalars(
    #     select(ImageFile).where(ImageFile.image_id == image.id)
    # ).all()

    if image_file_results is not None and len(image_file_results) > 0:

        resp_list = []

        # only one image_file for each image
        image_file = image_file_results[0]
        image_filename = image_file.filename
        image_basename = image_filename.split('.')[0]

        # where measurements will reside
        # Note: in the snapshotsOnly mode we do not create these individual folders
        image_measurements_path = f'{inspection_images_temp_path}'

        # measurements meta-data. Each measurement has a 2d image and json annotations (optional)
        measurement_results = session.scalars(
            select(Measurement).where(
                Measurement.image_id == image.id)
        ).all()

        logging.info(f'found {len(measurement_results)} measurement records for image.id: {image.id}')

        # there may be many measurements per 360 image....
        if measurement_results is None or len(measurement_results) == 0:
            # doc.add_paragraph("None")
            pass
        else:
            # doc.add_heading(f'Distance: {round_distance}',level=3)
            finding_label = 'Other'

            # for each measurement meta-data look for 2d snapshots and respective annotations
            for idx, measurement in enumerate(measurement_results):

                measurement_image_file_results = session.scalars(
                    select(MeasurementImageFile).where(
                        MeasurementImageFile.measurement_id == measurement.id)
                ).all()

                # if a measurement exists, a 2d snapshot measurement_image_file should exist as well
                if measurement_image_file_results is not None and len(measurement_image_file_results) > 0:
                    measurement_image_file = measurement_image_file_results[0]

                    annotation_json_content = None
                    # Annotation files are optional, but if they exist, there must be only one per measurement entry
                    measurement_annotation_file_results = session.scalars(
                        select(MeasurementAnnotationFile).where(
                            MeasurementAnnotationFile.measurement_id == measurement.id)
                    ).all()
                    if measurement_annotation_file_results is not None and len(measurement_annotation_file_results) > 0:
                        annotation = measurement_annotation_file_results[0]

                        # Parse the annotation content looking for the finding label
                        annotation_file_content = annotation.content
                        annotation_json_content = json.loads(
                            annotation_file_content)
                        # logging.info(f'annotation json_content: {json_content}') # list of objects
                        if isinstance(annotation_json_content, dict) and 'shapes' in annotation_json_content:
                            shapes = annotation_json_content['shapes']
                            if shapes is not None:
                                for el in shapes:
                                    #logging.info(f'element: {el}')
                                    if 'label' in el:
                                        finding_label = el['label']
                                        logging.info(f'using annotation file finding_label: {finding_label}')
                                        break
                                    else:
                                        finding_label = 'Other'

                    # build a measurement .png file name based on the 360 image file basename
                    measurement_filename = f'{image_basename}-{section_name}-z{round_distance}-s{idx}.jpg'
                    measurement_image_file_path = os.path.join(
                        image_measurements_path, measurement_filename)
                    logging.info(measurement_image_file_path)
                    with open(measurement_image_file_path, "wb") as mes_img_file:

                        if USE_S3 and measurement_image_file.s3key is not None:
                            measurement_image_file.content = get_file_content_from_s3(
                                measurement_image_file.s3key)
                        else:
                            logging.info(
                                'Using measurement_file_content from DB')

                        measurement_file_content = convert_png_content_to_jpg(
                            measurement_image_file.content)
                        if annotation_json_content is not None:
                            logging.info(f'drawing annotations on measurement image file: {measurement_filename}')
                            measurement_file_content = draw_annotations_on_image(
                                measurement_file_content, annotation_json_content)

                        mes_img_file.write(measurement_file_content)

                    # # Add an image
                    # doc.add_heading('Section 3: Image', level=2)
                    # doc.add_picture(measurement_image_file_path, width=Pt(400))

                    resp_list.append({
                        'round_distance': round_distance,
                        'finding_label': finding_label,
                        'measurement': measurement,
                        'measurement_image_file_path': measurement_image_file_path
                    })

                    # add_line_break(doc)
                    # add_measurement_image_content(
                    #     doc, round_distance, finding_label, section_name, measurement, inspection, measurement_image_file_path)
                    # # add_line_break(doc)
                    # doc.add_page_break()
        return resp_list  # list of images generated


def download_360image_defects_with_annotations_to_temp_folder(session, image, inspection, inspection_images_temp_path):
    section_name = get_section_name(inspection.sect)
    round_distance = str(round(image.distance, 1))

    # image_file with only what we need: id and filename
    image_file_results = session.query(
        ImageFile.id, ImageFile.filename).filter(
            ImageFile.image_id == image.id
    ).all()

    # image file with full content
    # image_file_results = session.scalars(
    #     select(ImageFile).where(ImageFile.image_id == image.id)
    # ).all()

    if image_file_results is not None and len(image_file_results) > 0:

        resp_list = []

        # only one image_file for each image
        image_file = image_file_results[0]
        image_filename = image_file.filename
        image_basename = image_filename.split('.')[0]

        # where measurements will reside
        # Note: in the snapshotsOnly mode we do not create these individual folders
        image_defects_path = f'{inspection_images_temp_path}'

        # measurements meta-data. Each measurement has a 2d image and json annotations (optional)
        defect_results = session.scalars(
            select(Defect).where(
                Defect.image_id == image.id)
        ).all()

        logging.info(f'found {len(defect_results)} defect records for image.id: {image.id}')

        # there may be many measurements per 360 image....
        if defect_results is None or len(defect_results) == 0:
            # doc.add_paragraph("None")
            pass
        else:
            # doc.add_heading(f'Distance: {round_distance}',level=3)
            finding_label = 'Other'

            # for each measurement meta-data look for 2d snapshots and respective annotations
            for idx, defect in enumerate(defect_results):

                measurement_image_file_results = session.scalars(
                    select(MeasurementImageFile).where(
                        MeasurementImageFile.measurement_id == defect.measurement_id)
                ).all()

                # if a measurement exists, a 2d snapshot measurement_image_file should exist as well
                if measurement_image_file_results is not None and len(measurement_image_file_results) > 0:
                    measurement_image_file = measurement_image_file_results[0]

                    annotation_json_content = None
                    # Annotation files are optional, but if they exist, there must be only one per measurement entry
                    defect_annotation_fragment_results = session.scalars(
                        select(DefectAnnotationFragment).where(
                            DefectAnnotationFragment.defect_id == defect.id)
                    ).all()
                    if defect_annotation_fragment_results is not None and len(defect_annotation_fragment_results) > 0:
                        annotation = defect_annotation_fragment_results[0]

                        # Parse the annotation content looking for the finding label
                        annotation_file_content = annotation.content
                        annotation_json_content = json.loads(
                            annotation_file_content)
                        # logging.info(f'annotation json_content: {json_content}') # list of objects
                        if isinstance(annotation_json_content, dict) and 'shapes' in annotation_json_content:
                            shapes = annotation_json_content['shapes']
                            if shapes is not None:
                                for el in shapes:
                                    #logging.info(f'element: {el}')
                                    if 'label' in el:
                                        finding_label = el['label']
                                        logging.info(f'using annotation file finding_label: {finding_label}')
                                        break
                                    else:
                                        finding_label = 'Other'

                    # build a measurement .png file name based on the 360 image file basename
                    defect_filename = f'{image_basename}-{section_name}-z{round_distance}-s{idx}.jpg'
                    defect_image_file_path = os.path.join(
                        image_defects_path, defect_filename)
                    logging.info(defect_image_file_path)
                    with open(defect_image_file_path, "wb") as def_img_file:

                        if USE_S3 and measurement_image_file.s3key is not None:
                            measurement_image_file.content = get_file_content_from_s3(
                                measurement_image_file.s3key)
                        else:
                            logging.info(
                                'Using measurement_file_content from DB')

                        defect_file_content = convert_png_content_to_jpg(
                            measurement_image_file.content)
                        if annotation_json_content is not None:
                            logging.info(f'drawing annotations on defect image file: {defect_filename}')
                            defect_file_content = draw_annotations_on_image(
                                defect_file_content, annotation_json_content)

                        def_img_file.write(defect_file_content)

                    # # Add an image
                    # doc.add_heading('Section 3: Image', level=2)
                    # doc.add_picture(measurement_image_file_path, width=Pt(400))

                    resp_list.append({
                        'round_distance': round_distance,
                        'finding_label': finding_label,
                        'measurement': defect,
                        'measurement_image_file_path': defect_image_file_path
                    })

                    # add_line_break(doc)
                    # add_measurement_image_content(
                    #     doc, round_distance, finding_label, section_name, measurement, inspection, measurement_image_file_path)
                    # # add_line_break(doc)
                    # doc.add_page_break()
        return resp_list  # list of images generated


# downloads all 360 image 2d measurement files into a temp folder, adding annotation layer to the 2d images when applicable
def download_360image_vtshots_to_temp_folder(session, image, inspection, inspection_images_temp_path):
    section_name = get_section_name(inspection.sect)
    round_distance = str(round(image.distance, 1))

    # image_file with only what we need: id and filename
    image_file_results = session.query(
        ImageFile.id, ImageFile.filename).filter(
            ImageFile.image_id == image.id
    ).all()

    # image file with full content
    # image_file_results = session.scalars(
    #     select(ImageFile).where(ImageFile.image_id == image.id)
    # ).all()

    if image_file_results is not None and len(image_file_results) > 0:

        resp_list = []

        # only one image_file for each image
        image_file = image_file_results[0]
        image_filename = image_file.filename
        image_basename = image_filename.split('.')[0]

        # where measurements will reside
        # Note: in the snapshotsOnly mode we do not create these individual folders
        image_vtshots_path = f'{inspection_images_temp_path}'

        # measurements meta-data. Each measurement has a 2d image and json annotations (optional)
        vtshot_results = session.scalars(
            select(VTShot).where(
                VTShot.image_id == image.id)
        ).all()

        logging.info(f'found {len(vtshot_results)} vtshot records for image.id: {image.id}')

        # there may be many measurements per 360 image....
        if vtshot_results is None or len(vtshot_results) == 0:
            # doc.add_paragraph("None")
            pass
        else:
            # doc.add_heading(f'Distance: {round_distance}',level=3)
            finding_label = 'Other'

            # for each measurement meta-data look for 2d snapshots and respective annotations
            for idx, vtshot in enumerate(vtshot_results):

                vtshot_image_file_results = session.scalars(
                    select(VTShotImageFile).where(
                        VTShotImageFile.vtshot_id == vtshot.id)
                ).all()

                # if a measurement exists, a 2d snapshot measurement_image_file should exist as well
                if vtshot_image_file_results is not None and len(vtshot_image_file_results) > 0:
                    vtshot_image_file = vtshot_image_file_results[0]

                    # build a measurement .png file name based on the 360 image file basename
                    vshot_filename = f'{image_basename}-{section_name}-z{round_distance}-s{idx}.jpg'
                    vshot_image_file_path = os.path.join(
                        image_vtshots_path, vshot_filename)
                    logging.info(vshot_image_file_path)
                    with open(vshot_image_file_path, "wb") as vshot_img_file:

                        if USE_S3 and vtshot_image_file.s3key is not None:
                            vtshot_image_file.content = get_file_content_from_s3(
                                vtshot_image_file.s3key)
                        else:
                            logging.info(
                                'Using measurement_file_content from DB')

                        vshot_file_content = vtshot_image_file.content
                        vshot_img_file.write(
                            convert_png_content_to_jpg(vshot_file_content))

                    resp_list.append({
                        'round_distance': round_distance,
                        'vshot': vtshot,
                        'vshot_image_file_path': vshot_image_file_path
                    })

        return resp_list  # list of images generated


# This method returns a set of lists
def gether_stats(session, inspection_list):
    logging.info(f'gather_stats called for {len(inspection_list)} inspections')
    category_names = ['A', 'B', 'C']
    category_data = [100, 1000, 2000]
    distance_list = []
    category_list = []
    location_list = []
    category_severity_list = []
    disposition_list = []
    measurement_list = []

    labels_dict = {}

    def update_stats_helper(distance, location, measurement, finding_label):
        count = labels_dict.get(finding_label, 0)
        count += 1
        labels_dict[finding_label] = count
        distance_list.append(distance)
        category_list.append(finding_label)
        location_list.append(location)
        category_severity_list.append(get_severity(finding_label))
        disposition_list.append(measurement.ge_disposition)
        measurement_list.append(measurement)


    for inspection in inspection_list:
        logging.info(f'read stats for inspection: id# {inspection.id} sect: {inspection.sect}')
        location = inspection.sect  # section of the blade
        inspection_image_list = session.scalars(
            select(Image).where(Image.inspection_id ==
                                inspection.id).order_by(Image.distance)
        ).all()
        logging.info(
            f'found {len(inspection_image_list)} images for inspection id: {inspection.id}')
        if inspection_image_list is not None and len(inspection_image_list) > 0:
            
            # logging.info('executing in series...')
            # for image in inspection_image_list:
            #     logger.info(f'image: {str(image)}')
            #     result_list = gather_image_defect_stats(image)
            #     logging.info(f'result: {result_list}')
            #     if result_list is not None:
            #         for result in result_list:
            #             update_stats_helper(result['distance'], location, result['measurement'], result['finding_label'])    

            logging.info('executing in parallel...')
            futures = []
            with ThreadPoolExecutor(THREAD_POOL_SIZE) as executor:
                # for each 360 image record...
                for image in inspection_image_list:
                    # logger.info(f'image: {str(image)}')
                    # futures.append(executor.submit(gather_image_measurement_stats, session, image))
                    futures.append(executor.submit(gather_image_defect_stats, image))
                wait(futures)
            for future in futures:
                if future.result() is not None:
                    result_list = future.result()
                    for result in result_list:
                        update_stats_helper( result['distance'], location, result['measurement'], result['finding_label'])

        
    category_names = []
    category_data = []
    logging.info(f'stats: {labels_dict}')
    for key, val in labels_dict.items():
        category_names.append(key.replace('_', ' '))
        category_data.append(val)

    logging.info(f'returning {category_list} {location_list} {category_severity_list}')
    return (category_names, category_data, distance_list, category_list, location_list, category_severity_list, disposition_list, measurement_list)


# compute {distance, measurement, finding_label} for each measurement of an image
# use the saved annotation content as the ground truth
def gather_image_measurement_stats(session, image):
    resp_list = []
    distance = image.distance
    measurement_list = session.scalars(
        select(Measurement).where(
            Measurement.image_id == image.id)
    ).all()
    logging.info(f'found {len(measurement_list)} measurements for image id: {image.id}')
    if measurement_list is not None and len(measurement_list) > 0:
        for measurement in measurement_list:
            logging.info(f'processing measurement id: {measurement.id}')
            measurement_annotation_file_list = session.scalars(
                select(MeasurementAnnotationFile).where(
                    MeasurementAnnotationFile.measurement_id == measurement.id)
            ).all()
            if measurement_annotation_file_list is not None and len(measurement_annotation_file_list) > 0:
                annotation_file = measurement_annotation_file_list[0]
                annotation_file_content = annotation_file.content
                json_content = json.loads(annotation_file_content)
                # list of objects
                # logging.info(f'annotation json_content: {json_content}')
                if isinstance(json_content, dict) and 'shapes' in json_content:
                    shapes = json_content['shapes']
                    if shapes is not None and len(shapes) > 0:
                        for el in shapes:
                            logging.info(f'element: {el}')
                            finding_label = 'Other'
                            if 'label' in el:
                                finding_label = el['label']
                                finding_label = finding_label.replace(' ', '_')
                                logging.info(f'found annotation file label. will classify defect as: {finding_label}')
                            resp_list.append({'distance': float(
                                distance), 'measurement': measurement, 'finding_label': finding_label})
                    else:
                        finding_label = measurement.finding_type if measurement.finding_type is not None else 'Other'
                        logging.info(
                            f'defect has no shapes array, will classify it as {finding_label}')
                        resp_list.append(
                            {'distance': float(distance), 'measurement': measurement, 'finding_label': finding_label})
            else:
                finding_label = measurement.finding_type if measurement.finding_type is not None else 'Other'
                logging.info(
                    f'Measurement has no annotations, will classify it as {finding_label}')
                resp_list.append(
                    {'distance': float(distance), 'measurement': measurement, 'finding_label': finding_label})

    logging.info('Image has no measurements. Returning None')
    return resp_list


# a defect is a group of shapes from a measurement that share a common label.
# measurements can have more than one defect.
# return a dictionary in the form: {defect, measurement, finding_label}
def gather_image_defect_stats(image):
    logging.info(f'gether_image_defect_stats called for image id#{image.id}')
    resp_list = []
    distance = image.distance

    with Session(db) as session:
        defect_list = session.scalars(
            select(Defect).where(
                Defect.image_id == image.id)
        ).all()
        logging.info(f'found {len(defect_list)} defects for image id: {image.id}')
        if defect_list is not None and len(defect_list) > 0:
            for defect in defect_list:
                logging.info(f'processing defect id: {defect.id}')
                defect_annotation_fragment_list = session.scalars(
                    select(DefectAnnotationFragment).where(
                        DefectAnnotationFragment.defect_id == defect.id)
                ).all()
                if defect_annotation_fragment_list is not None and len(defect_annotation_fragment_list) > 0:
                    annotation_fragment = defect_annotation_fragment_list[0]
                    annotation_fragment_content = annotation_fragment.content
                    json_content = json.loads(annotation_fragment_content)
                    # list of objects
                    # logging.info(f'annotation json_content: {json_content}')
                    if isinstance(json_content, dict) and 'shapes' in json_content:
                        shapes = json_content['shapes']
                        if shapes is not None and len(shapes) > 0:
                            for el in shapes:
                                logging.info(f'element: {el}')
                                finding_label = 'Other'
                                if 'label' in el:
                                    finding_label = el['label']
                                    finding_label = finding_label.replace(' ', '_')
                                    logging.info(f'found annotation file label. will classify defect as: {finding_label}')
                                resp_list.append({'distance': float(
                                    distance), 'measurement': defect, 'finding_label': finding_label})
                        else:
                            finding_label = defect.finding_type if defect.finding_type is not None else 'Other'
                            logging.info(
                                f'defect has no shapes array, will classify it as {finding_label}')
                            resp_list.append(
                                {'distance': float(distance), 'measurement': defect, 'finding_label': finding_label})
                else:
                    finding_label = defect.finding_type if defect.finding_type is not None else 'Other'
                    logging.info(
                        f'Measurement has no annotations, will classify it as {finding_label}')
                    resp_list.append(
                        {'distance': float(distance), 'measurement': defect, 'finding_label': finding_label})

    logging.info(f'Image has no defects. Returning: {resp_list}')
    return resp_list


# ----------------------------------------- Measurement ------------------------------------

class MeasurementImageThumbnailAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):
        super(MeasurementImageThumbnailAPI, self).__init__()

    def get(self, id):
        """
        Read an existing measurement image file thumbnail with or without annotations
        ---
        parameters:
          - in: path
            required: true
            name: id
            schema:
              type: number
            description: Existing measurement id
          - in: query
            required: false
            name: includeAnnotations
            schema:
              type: boolean
            description: Whether or not to include annotations on the image

        responses:
          200:
            description: image file content for the provided measurement id
            schema:
              id: MeasurementImageThumbnailContent
              type: string
              format: binary
        """
        include_annotations_str = request.args.get('includeAnnotations')
        if include_annotations_str is not None:
            include_annotations_str = include_annotations_str.capitalize()
        logger.info(f'include_annotations_str: {include_annotations_str}')
        include_annotations = False
        if include_annotations_str == 'True':
            include_annotations = True

        logger.info(f'Reading image_file content for image.id {id}')
        logger.info(f'include_annotations: {include_annotations}')
        with Session(db) as session:
            image_results = session.scalars(
                select(MeasurementImageFile).where(MeasurementImageFile.measurement_id == id)).all()
            if image_results is not None and len(image_results) > 0:
                image_file = image_results[0]
                thumbnail_content = image_file.thumbnail
                if thumbnail_content is not None:

                    if include_annotations == True:
                        logging.info('including annotations...')
                        annotation_file_results = session.scalars(
                            select(MeasurementAnnotationFile).where(MeasurementAnnotationFile.measurement_id == id)).all()
                        if annotation_file_results is not None and len(annotation_file_results) > 0:
                            annotation_file = annotation_file_results[0]
                            annotation_str = annotation_file.content  # json annotaiton content
                            annotation_json = json.loads(annotation_str)

                            # overwrite the original content with the new one including annotations
                            thumbnail_content = draw_annotations_on_image(
                                thumbnail_content, annotation_json, 1, False)

                    return send_file_content_as_jpg(image_file.filename, thumbnail_content)

        return {'message': f'Thumbnail for Measurement id: {id} not found.'}


# ------------------------------------------ Measurement -----------------------------

# at the end of the annotation file we can have properties as pitch, yaw, hfov that we parse
# them here and store the results in the measurement record holding this annotation file

def parse_annotation_file_and_save_props(annotation_file_path, measurement_rec, session):
    annotation_json_content = None

    with open(annotation_file_path) as f:
        annotation_json_content = json.load(f)
        logging.info(f'annotation_file_content: {annotation_json_content}')

        # read coordinates from file if they are included
        if isinstance(annotation_json_content, dict):
            if 'imagePitch' in annotation_json_content:
                measurement_rec.image_pitch = annotation_json_content['imagePitch']
            if 'imageYaw' in annotation_json_content:
                measurement_rec.image_yaw = annotation_json_content['imageYaw']
            if 'imageHfov' in annotation_json_content:
                measurement_rec.image_hfov = annotation_json_content['imageHfov']
            if 'AI' in annotation_json_content:
                measurement_rec.is_manual = bool(
                    str(annotation_json_content['AI']) == 'false')
                logging.info(f'is_manual: {measurement_rec.is_manual}')

            # extract finding_type from annotation
            # TODO: this assumes all findings with an image are all the same
            if 'shapes' in annotation_json_content and len(annotation_json_content['shapes']) > 0:
                shapes = annotation_json_content['shapes']
                for el in shapes:
                    logging.info(f'element: {el}')
                    if 'label' in el:
                        finding_label = el['label'].replace('_', ' ')
                        logging.info(
                            f'Updating {measurement_rec.id} measurement rec with finding_type: {finding_label}')
                        measurement_rec.finding_type = finding_label
            else:
                logging.info('Could not find shapes array')

            session.add(measurement_rec)
            session.commit()
            session.refresh(measurement_rec)
            logging.info(
                f'Found coordinates in .json. Updated new_measurement_rec: {measurement_rec}')
        else:
            logging.info(
                'annotation file content is not json. Could not parse.')

    return annotation_json_content



def add_snapshot_coordinates(annotation_json_content, measurement_rec):
    logging.info(f'add_snapshot_coordinates() called using measurement id# {measurement_rec.id}')
    # add coordinates from db if they are not included
    if annotation_json_content is not None and isinstance(annotation_json_content, dict):
        if not 'imagePitch' in annotation_json_content:
            annotation_json_content['imagePitch'] = measurement_rec.image_pitch
        if not 'imageYaw' in annotation_json_content:
            annotation_json_content['imageYaw'] = measurement_rec.image_yaw
        if not 'imageHfov' in annotation_json_content:
            annotation_json_content['imageHfov'] = measurement_rec.image_hfov
        if not 'AI' in annotation_json_content:
            annotation_json_content['AI'] = False
    return annotation_json_content


def are_measurements_included_in_shape(shape):
    for key in ['area','width','length']:
        if not key in shape:
            return False
    return True


# measurement Martin call.
def measure_annotation_file_and_update_json_and_measurement_rec(annotation_json_content, measurement_rec, session, force_compute=False):
    logging.info(f'measure_annotation_file_and_update_record() for measurement id# {measurement_rec.id}')
    
    if annotation_json_content is None:
        logging.info('annotation_json_content is None')
        return annotation_json_content

    image_rec = session.query(Image).get(measurement_rec.image_id)
    if image_rec is None:
        logging.info(f'no frame information for 360 image {measurement_rec.image_id}')
        return annotation_json_content
    else:
        logging.info(f'using image_rec frame: {image_rec.frame}')
    
    frame_str = image_rec.frame

    # if there is a frame (icp ran on Martin's code) -- and if the transform of that frame is not ''
    # then compute measurements
    if frame_str is not None and frame_str != '':
        frame = json.loads(frame_str)
        if 'transform' in frame and frame['transform'] != '':
            shapes = []
            if 'shapes' in annotation_json_content and len(annotation_json_content['shapes']) > 0:
                shapes = annotation_json_content['shapes']
            
            do_compute_measurements = False
            if force_compute == True:
                do_compute_measurements = True
            else:
                for shape in shapes:
                    if not are_measurements_included_in_shape(shape):
                        logging.info(f'measurements are NOT included in shape: {shape}')
                        do_compute_measurements = True

            # used to update the measurement record
            total_area = 0.0
            max_width = 0.0
            max_length = 0.0

            logging.info(f'force_compute: {force_compute}')
            logging.info(f'do_compute_measurements: {do_compute_measurements}')               
            
            if do_compute_measurements:
                # calculate using measurement service then update totals and the annotation_json_content
                logging.info(f'calling get_measurements using frame: {frame}')
                measurements_resp = get_measurements_json_httpx(annotation_json_content, frame)
                logging.info(f'get_measurement() resp: {measurements_resp}')

                # measurements_resp will be None if it cannot measure the file
                if measurements_resp is None:
                    logging.info(f'error calculating measurements for defect id# {measurement_rec.id}')
                    return annotation_json_content

                # for each returned prop from the measurement, copy it to the shape within the json file
                for key, value in measurements_resp.items():
                    logging.info(f'processing measurement resp: {key}: {value}')
                    if type(value).__name__ == 'list':
                        for idx, value_at_idx in enumerate(value):
                            if idx < len(shapes):
                                shape = shapes[idx]
                                shape[key] = value_at_idx

                                if key == 'width_blob':
                                    max_width = max(max_width, value_at_idx)

                                if key == 'area' and value_at_idx > 0:
                                    total_area += value_at_idx
                                if key == 'width':
                                    max_width = max(max_width, value_at_idx)
                                if key == 'length':
                                    max_length = max(max_length, value_at_idx)
                    else: # value is scalar (old format)
                        if len(shapes) > 0:
                            shape = shapes[0]
                            shape[key] = value

                            # alternative width algoritym
                            if key == 'width_blob' and value > 0:
                                max_width = max(max_width, value)

                            if key == 'area' and value > 0:
                                total_area += value
                            if key == 'width':
                                max_width = max(max_width, value)
                            if key == 'length':
                                max_length = max(max_length, value)

            else:
                logging.info('Use measurements from annotaiton file. Measurement Service will NOT be called.')
                # calculate totals from the properties already in each shape
                for shape in shapes:
                    for key in ['width_blob','area','width','length']:
                        if key in shape:
                            value = shape[key]
                            if key == 'width_blob':
                                max_width = max(max_width, value)
                            if key == 'area' and value > 0:
                                total_area += value
                            if key == 'width':
                                max_width = max(max_width, value)
                            if key == 'length':
                                max_length = max(max_length, value)

            measurement_rec.area = total_area
            measurement_rec.length = max_length
            measurement_rec.width = max_width

            session.add(measurement_rec)
            session.commit()
            session.refresh(measurement_rec)
            logging.info(f'updated measurement rec area to: {measurement_rec.area}')
            logging.info(f'updated measurement rec length to: {measurement_rec.length}')
            logging.info(f'updated measurement rec width to: {measurement_rec.width}')
    else:
        logging.info('Found no frame info in image_rec')  

    logging.info(f'Updated json content with measurements: {annotation_json_content}')
    return annotation_json_content


# alternatively we can pass parameters in the body of the annotation
# this method reads these props passed in the validated annotation body
def read_validated_annotation_props(validated_annotation_json_content):
    # logging.info(f'reading props from: {validated_annotation_json_content}')
    props = {
        'validationStatus': None,
        'validatedBy': None,
        'validationTimestamp': None,
        'label': None,
    }

    # read coordinates from file if they are included
    if validated_annotation_json_content is not None and isinstance(validated_annotation_json_content, dict):
        if 'validationStatus' in validated_annotation_json_content:
            props['validationStatus'] = validated_annotation_json_content['validationStatus']
        if 'validatedBy' in validated_annotation_json_content:
            props['validatedBy'] = validated_annotation_json_content['validatedBy']
        if 'validationTimestamp' in validated_annotation_json_content:
            props['validationTimestamp'] = validated_annotation_json_content['validationTimestamp']
        if 'shapes' in validated_annotation_json_content:
            shapes = validated_annotation_json_content['shapes']
            if len(shapes) > 0:
                # sometimes label is not in shapes... and it was throwing an error. 
                if 'label' in shapes[0]:
                    props['label'] = shapes[0]['label']
                else:                    
                    props['label'] = None 

    return props


# Uploads both the json content and the .jpg image.
class UploadMeasurementImageAndAnnotationAPI(Resource):

    def __init__(self):
        super(UploadMeasurementImageAndAnnotationAPI, self).__init__()

    def allowed_file(self, filename):
        return '.' in filename and \
            filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

    # This is the method called by the initial import, so all annotations are considered UI generated.
    def post(self, image_id):
        """
        Upload new measurement performed on a 360 image. Reads 2d png image and json annotations file. Any annotation in _measurement folders are recorded AI generated.
        ---
        consumes:
          - multipart/form-data
        parameters:
          - in: path
            required: true
            name: image_id
            schema:
              type: number
          - in: formData
            name: annotation_file
            required: true
            type: file
          - in: formData
            required: true
            name: image_file
            type: file
            description: image file
        responses:
          200:
            description: New measurement object ids
            schema:
              id: MeasurementUploadResponse
              properties:
                id:
                  type: number
                image_id:
                  type: number
                measurement_file_id:
                  type: number
                annotation_file_id:
                  type: number

        """

        form = request.form.to_dict()
        logger.info(f'Form data {form}')

        post_image_file = request.files.get("image_file")
        if post_image_file is None or not self.allowed_file(post_image_file.filename):
            msg = 'Image file not found, or no image_file provided.'
            return {'message': msg}

        with Session(db) as session:
            logger.info(f'image_id: {image_id}')
            image = session.query(Image).get(image_id)
            if image is None:
                return {'message': f'Image {image_id} not found. Existing image id required.'}
            logger.info(f'using inspection: {image}')

            inspection = session.query(Inspection).get(image.inspection_id)

            msg = ''
            files_list = []

            distance = round(image.distance, 1)
            location = inspection.sect

            new_measurement_rec = Measurement(
                image_id=image_id,
                root_face_distance=distance,
                location=location,
                is_manual=False)
            session.add(new_measurement_rec)
            session.commit()
            session.refresh(new_measurement_rec)
            measurement_id = new_measurement_rec.id

            logging.info(f'new_measurement_rec: {new_measurement_rec}')
            measurement_json = new_measurement_rec.toJson()
            logging.info(f'measurement_rec_json: {measurement_json}')

            # -------------------------- optional annotation_file -----------------------------
            # metadata must be read first, so Image may be created
            post_annotation_file = request.files.get("annotation_file")
            measurement_annotation_file_id = None
            if post_annotation_file is None or not self.allowed_file(post_annotation_file.filename):
                msg += 'Measurement Annotation file not found.'
            else:
                annotation_filename = secure_filename(
                    post_annotation_file.filename)
                unique_annotation_filename = str(
                    uuid.uuid4())+'_'+annotation_filename
                annotation_file_path = os.path.join(
                    app.config['UPLOAD_FOLDER'], unique_annotation_filename)
                post_annotation_file.save(annotation_file_path)
                msg += f'Measurement Annotation file: {post_annotation_file.filename} uploaded successfully.'
                files_list.append(post_annotation_file.filename)

                annotation_json_content = None
                if os.path.isfile(annotation_file_path):
                    annotation_json_content = parse_annotation_file_and_save_props(
                        annotation_file_path, new_measurement_rec, session)
                    if annotation_json_content is not None:
                        annotation_json_content = measure_annotation_file_and_update_json_and_measurement_rec(
                            annotation_json_content, new_measurement_rec, session, False)
                    os.remove(annotation_file_path)

                str_annotation_json_content = None
                str_annotation_json_content = json.dumps(
                    annotation_json_content)
                logging.info(f'annotation_file_content str: {str_annotation_json_content}')

                new_measurement_annotation_file_rec = MeasurementAnnotationFile(
                    image_id=image_id,
                    measurement_id=measurement_id,
                    filename=annotation_filename,
                    s3key=None,
                    content=str_annotation_json_content)
                session.add(new_measurement_annotation_file_rec)
                session.commit()
                # to read the id
                session.refresh(new_measurement_annotation_file_rec)
                measurement_annotation_file_id = new_measurement_annotation_file_rec.id

                parse_annotation_file_defects(new_measurement_rec, new_measurement_annotation_file_rec)

                # new_measurement_annotation_file_rec.content = None
                # logging.info(f'new_measurement_annotation_file_rec: {new_measurement_annotation_file_rec}')
                measurement_annotation_file_rec_json = new_measurement_annotation_file_rec.toJson()
                logging.info(f'measurement_annotation_file_rec_json: {measurement_annotation_file_rec_json}')

            # -------------------------------- required formData image_file ---------------------------------
            # Then we read the image_file associated with the image record

            measurement_image_file_json = {}
            measurement_image_filename = secure_filename(
                post_image_file.filename)
            unique_image_filename = str(
                uuid.uuid4())+'_'+measurement_image_filename
            image_file_path = os.path.join(
                app.config['UPLOAD_FOLDER'], unique_image_filename)
            post_image_file.save(image_file_path)

            logging.info('Processing image_file...')
            bin_content = None
            with open(image_file_path, mode="rb") as f:
                bin_content = f.read()
            thumb_bin_content = get_thumbnail_content(image_file_path)
            if os.path.isfile(image_file_path):
                os.remove(image_file_path)

            if COMPRESS_INCOMING_IMAGES:
                measurement_image_filename = measurement_image_filename.replace(
                    '.png', '.jpg')
                bin_content = convert_png_content_to_jpg(bin_content)

            new_measurement_image_file_rec = MeasurementImageFile(
                image_id=image_id,
                measurement_id=measurement_id,
                filename=measurement_image_filename,
                content=bin_content,
                thumbnail=thumb_bin_content)

            if USE_S3:
                image_file_rec = session.scalars(
                    select(ImageFile).filter(ImageFile.image_id ==
                                             new_measurement_image_file_rec.image_id)
                ).all()[0]
                # name of image where this measurement came from
                pan_image_filename = image_file_rec.filename
                measurements_folder_name = pan_image_filename.split('.')[0]+'_measurements'
                s3key = get_inspection_s3key(inspection)+f'/images/{measurements_folder_name}/{measurement_image_filename}'
                response = upload_content_to_s3(bin_content, s3key)
                logging.info(f'upload to s3 resp: {response}')
                if response.status_code != 200:
                    return {'message': f'Error uploading file: {measurement_image_filename} to S3'}
                new_measurement_image_file_rec.s3key = s3key
                new_measurement_image_file_rec.content = None

            session.add(new_measurement_image_file_rec)
            session.commit()
            session.refresh(new_measurement_image_file_rec)
            measurement_image_file_id = new_measurement_image_file_rec.id

            new_measurement_image_file_rec.content = None
            new_measurement_image_file_rec.thumbnail = None
            logging.info(f'new_measurement_image_file_rec: {new_measurement_image_file_rec}')
            measurement_image_file_json = new_measurement_image_file_rec.toJson()
            logging.info(f'measurement_image_file_json: {measurement_image_file_json}')

            msg += f' Measurement Image file {post_image_file.filename} uploaded successfully. '
            files_list.append(post_image_file.filename)

            resp = {'message': msg,
                    'files_list': files_list,
                    'image_id': image_id,
                    'measurement_id': measurement_id,
                    'measurement_annotation_file_id': measurement_annotation_file_id,
                    'measurement_image_file_id': measurement_image_file_id}
            logging.info(f'resp: {resp}')

            inspection_id = image.inspection_id
            invalidate_inspection_cache(inspection_id)

            return make_response(jsonify(resp), 200)


# ------------------------------- VTShot - Virtual Tour Snapshot -----------------------------

class VTShotImageThumbnailAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):
        super(VTShotImageThumbnailAPI, self).__init__()

    def get(self, id):
        """
        Read an existing vtshot image file thumbnail, representing a 2d still image from a 360 image
        ---
        parameters:
          - in: path
            required: true
            name: id
            schema:
              type: number
            description: Existing vtshot id

        responses:
          200:
            description: image file content for the provided vtshot id
            schema:
              id: VTShotImageThumbnailContent
              type: string
              format: binary
        """
        logger.info(f'Reading vtshot_image_file content for image.id {id}')

        with Session(db) as session:
            vtshot_image_results = session.scalars(
                select(VTShotImageFile).where(VTShotImageFile.vtshot_id == id)).all()
            if vtshot_image_results is not None and len(vtshot_image_results) > 0:
                image_file = vtshot_image_results[0]
                thumbnail_content = image_file.thumbnail
                if thumbnail_content is not None:
                    return send_file_content_as_jpg(image_file.filename, thumbnail_content)

        return {'message': f'Thumbnail for VTShot id: {id} not found.'}


class VTShotAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):
        super(VTShotAPI, self).__init__()

    def delete(self, id):
        """
        Deletes an existing virtual tour snapshot
        ---
        parameters:
          - in: path
            required: true
            name: id
            schema:
              type: number
            description: Existing blade id

        responses:
          200:
            description: Success or failure message
            schema:
              properties:
                message:
                  type: string

        """
        with Session(db) as session:
            vtshot = session.query(VTShot).get(id)
            if vtshot is not None:
                logger.info(
                    f'get vtshot_image_file for measurement.id #{id}')
                vtshot_image_file_list = session.scalars(
                    select(VTShotImageFile).where(
                        VTShotImageFile.measurement_id == id)
                ).all()

                # delete all vtshot_image_file dependences under vtshot_image
                if vtshot_image_file_list is not None and len(vtshot_image_file_list) > 0:
                    vtshot_image_file_rec = vtshot_image_file_list[0]
                    session.delete(vtshot_image_file_rec)

                # finally deletes current record
                session.delete(vtshot)
                session.commit()

                return {'message': f'Successfully deleted vtshot id #{id}'}

        return {'message': f'VTShot id #{id} not found.'}

    def get(self, id):
        """
        Read an existing virtual tour snapshot record meta-data
        ---
        parameters:
          - in: path
            required: true
            name: id
            schema:
              type: number
            description: Existing virtual tour snapshot id

        responses:
          200:
            description: Existing virtual tour snapshot metadata
            schema:
              id: MeasurementObject

        """

        logging.info(f'Look for measurement record id {id}')
        with Session(db) as session:
            vtshot = session.query(VTShot).get(id)

            if vtshot is None:
                return {'message': f'VTShot {id} not found.'}

            jsonResp = vtshot.toJson()
            return jsonResp


class CreateVTShotAPI(Resource):

    def __init__(self):
        self.reqparse = reqparse.RequestParser()
        self.reqparse.add_argument('image_id', type=int, location='json')
        self.reqparse.add_argument('date', type=str, location='json')
        self.reqparse.add_argument(
            'root_face_distance', type=float, location='json')
        self.reqparse.add_argument('image_pitch', type=float, location='json')
        self.reqparse.add_argument('image_yaw', type=float, location='json')
        self.reqparse.add_argument('image_hfov', type=float, location='json')

        super(CreateVTShotAPI, self).__init__()

    def post(self):
        """
        Create a new virtual tour snapshot meta-data record. With this record one can attach a vt snapshot_image
        ---
        parameters:
          - name: CreateVTShotRequest
            in: body
            required: true
            schema:
              id: CreateVTShotBody
              properties:
                image_id:
                  type: number
                  example: 1
                date:
                  type: string
                  example: "2023-11-01"
                root_face_distance:
                  type: number
                image_pitch:
                  type: number
                image_yaw:
                  type: number
                image_hfov:
                  type: number

        responses:
          200:
            description: Created virtual tour snapshot meta-data record
            schema:
              id: VTShotObject
              properties:
                id:
                  type: number
                image_id:
                  type: number
                date:
                  type: string
                root_face_distance:
                  type: number
                image_pitch:
                  type: number
                image_yaw:
                  type: number
                image_hfov:
                  type: number

        """
        args = self.reqparse.parse_args()
        logger.info(args.items())

        if args['date'] == '':

            # current_date = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            current_date = datetime.now().isoformat()
            args['date'] = current_date

        image_id = args['image_id']
        with Session(db) as session:
            image = session.query(Image).get(image_id)
            if image is None:
                return {'message': f'Could not create virtual tour snapshot. Image record for image_id: {image_id} not found.'}

        vtShot = VTShot(
            image_id=args['image_id'],
            date=args['date'],
            root_face_distance=args['root_face_distance'],
            image_pitch=args['image_pitch'],
            image_yaw=args['image_yaw'],
            image_hfov=args['image_hfov']
        )

        image_distance = round(image.distance, 1)
        if vtShot.root_face_distance != image_distance:
            logging.info(f'Using distance from image. z distance: {image.distance}')
            vtShot.root_face_distance = image_distance

        with Session(db) as session:
            session.add(vtShot)
            session.commit()
            session.refresh(vtShot)  # to read the id

        logger.info(f'Created vt snapshot: {str(vtShot)}')

        # resp = jsonify(inspection.serialize())
        resp = vtShot.toJson()
        logging.info(f'resp: {resp}')
        return resp


class VTShotImageFileAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):
        super(VTShotImageFileAPI, self).__init__()

    def allowed_file(self, filename):
        return '.' in filename and \
            filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

    def post(self, id):
        """
        Update existing or create new virtual tour image file
        ---
        consumes:
          - multipart/form-data
        parameters:
          - in: path
            required: true
            name: id
            schema:
              type: number
            description: Existing vtshot record id
          - in: formData
            required: true
            name: image_file
            type: file
            description: new vtshot image file
        responses:
          200:
            description: Success or failure message
            schema:
              properties:
                message:
                  type: string
                id:
                  type: number

        """
        # formData image_file
        post_image_file = request.files.get("image_file")
        if post_image_file is None or not self.allowed_file(post_image_file.filename):
            msg = 'Image file not found, or no image_file provided.'
            return {'message': msg}

        vtshot_image_filename = secure_filename(post_image_file.filename)
        unique_vtshot_image_filename = str(
            uuid.uuid4())+'_'+vtshot_image_filename
        vtshot_image_file_path = os.path.join(
            app.config['UPLOAD_FOLDER'], unique_vtshot_image_filename)
        post_image_file.save(vtshot_image_file_path)

        logging.info('Reading image_file formData content...')
        bin_content = None
        with open(vtshot_image_file_path, mode="rb") as f:
            bin_content = f.read()
        thumb_bin_content = get_thumbnail_content(vtshot_image_file_path)
        if os.path.isfile(vtshot_image_file_path):
            os.remove(vtshot_image_file_path)

        with Session(db) as session:
            vtshot_rec = session.query(VTShot).get(id)
            if vtshot_rec is None:
                msg = f'VTShot record id: {id} not found.'
                return {'message': msg}

            logger.info(
                f'Reading vtshot_image_file content for measurement.id #{id}')
            vtshot_image_file_list = session.scalars(
                select(VTShotImageFile).where(
                    VTShotImageFile.vtshot_id == id)
            ).all()
            vtshot_image_file_rec = None
            if len(vtshot_image_file_list) == 0:

                if COMPRESS_INCOMING_IMAGES:
                    vtshot_image_filename = vtshot_image_filename.replace(
                        '.png', '.jpg')
                    bin_content = convert_png_content_to_jpg(bin_content)

                vtshot_image_file_rec = VTShotImageFile(
                    image_id=vtshot_rec.image_id,
                    vtshot_id=id,
                    filename=vtshot_image_filename,
                    s3key=None,
                    content=bin_content,
                    thumbnail=thumb_bin_content
                )
                session.add(vtshot_image_file_rec)
            else:
                vtshot_image_file_rec = vtshot_image_file_list[0]
                vtshot_image_file_rec.filename = vtshot_image_filename
                vtshot_image_file_rec.content = bin_content
                vtshot_image_file_rec.thumbnail = thumb_bin_content

            if USE_S3:
                # Obtain image, image_file and inspection for the measurement
                image_rec = session.query(Image).get(
                    vtshot_rec.image_id)  # required to obtain inspection
                image_file_rec = session.scalars(
                    select(ImageFile).filter(
                        ImageFile.image_id == vtshot_rec.image_id)
                ).all()[0]
                inspection_rec = session.query(
                    Inspection).get(image_rec.inspection_id)
                # name of image where this measurement came from
                pan_image_filename = image_file_rec.filename
                virtualtour_folder_name = pan_image_filename.split('.')[
                    0]+'_virtualtour'
                s3key = get_inspection_s3key(
                    inspection_rec)+f'/images/{virtualtour_folder_name}/{vtshot_image_filename}'
                response = upload_content_to_s3(bin_content, s3key)
                logging.info(f'upload to s3 resp: {response}')
                if response.status_code != 200:
                    return {'message': f'Error uploading file: {vtshot_image_filename} to S3'}
                vtshot_image_file_rec.s3key = s3key
                vtshot_image_file_rec.content = None

            session.commit()
            session.refresh(vtshot_image_file_rec)

            vtshot_image_file_rec.content = None
            vtshot_image_file_rec.thumbnail = None
            logging.info(
                f'new_vtshot_image_file_rec: {vtshot_image_file_rec}')
            vtshot_image_file_json = vtshot_image_file_rec.toJson()
            logging.info(
                f'vtshot_image_file_json: {vtshot_image_file_json}')

            resp = {
                'message': f' VTShot Image file {post_image_file.filename} uploaded successfully.',
                'id': vtshot_image_file_rec.id
            }
            return resp

    def get(self, id):
        """
        Read an existing vtshot image png file
        ---
        parameters:
          - in: path
            required: true
            name: id
            schema:
              type: number
            description: Existing vtshot id
        responses:
          200:
            description: image file content for the provided vtshot record id
            schema:
              id: VTShotImageFileContent
              type: string
              format: binary
        """

        logger.info(f'Reading image_file content for vtshot.id {id}')

        with Session(db) as session:
            vtshot_image_results = session.scalars(
                select(VTShotImageFile).where(VTShotImageFile.vtshot_id == id)).all()

            if vtshot_image_results is not None and len(vtshot_image_results) > 0:
                vtshot_image_file = vtshot_image_results[0]
                file_content = vtshot_image_file.content

                if USE_S3 and vtshot_image_file.s3key is not None:
                    file_content = get_file_content_from_s3(
                        vtshot_image_file.s3key)
                else:
                    logging.info('Using measurement_file_content from DB')

                if file_content is not None:
                    return send_file_content_as_jpg(vtshot_image_file.filename, file_content)
            else:
                logging.info(f'vtshot_image_results for vtshot_id {id} not found.')

        return {'message': f'VTShot Image File for VTShot id: {id} not found.'}


# ----------------------------- Measurement --------------------------


def extract_labels(validation_content):
    if (validation_content == None):
        return None
    print("measurement [17] = ", validation_content)
    # Parse the JSON string
    data = json.loads(validation_content)

    # Extract labels from the shapes
    labels = [shape.get('label', '') for shape in data.get('shapes', [])
              if shape.get('label')]

    return labels

# Example usage
# validation_content_1 = "{\"version\": \"5.2.1\", \"flags\": {}, \"shapes\": [{\"label\": \"LPS Cable Damage\", \"points\": [[492, 155], [415, 290], [650, 289], [652, 217]], \"description\": \"LPS Cable Damage\", \"shape_type\": \"polygon\", \"group_id\": null, \"flags\": {}}], \"imagePath\": \"snap-zunknown-s1.png\", \"imageData\": null, \"imageHeight\": 768, \"imageWidth\": 1024, \"polygons\": [{\"points\": [[492, 155], [415, 290], [650, 289], [652, 217]], \"flattenedPoints\": [492, 155, 415, 290, 650, 289, 652, 217, 490, 145], \"isComplete\": true, \"category\": \"LPS Cable Damage\", \"defectName\": \"LPS Cable Damage\", \"color\": \"#DDDDDD\"}]}"
# validation_content_2 = "{\"version\": \"5.2.1\", \"flags\": {}, \"shapes\": [{\"label\": \"\", \"points\": [], \"description\": \"\", \"shape_type\": \"polygon\", \"group_id\": null, \"flags\": {}}, {\"label\": \"\", \"points\": [], \"description\": \"\", \"shape_type\": \"polygon\", \"group_id\": null, \"flags\": {}}, {\"label\": \"Layers Overlap\", \"points\": [[406.36557564660103, 204.5390625], [398.3654535744259, 335.5390625], [631.3690089265278, 335.5390625]], \"description\": \"


class SearchMeasurementAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):
        super(SearchMeasurementAPI, self).__init__()

    def get(self):
        """
        Read an existing measurement list meta-data according to filter parameters
        ---
        parameters:
          - in: query
            required: false
            name: esn
            schema:
              type: string
          - in: query
            required: false
            name: blade_section
            schema:
              type: string
          - in: query
            required: false
            name: root_face_distance
            schema:
              type: string
          - in: query
            required: false
            name: validation_status

        responses:
          200:
            description: Existing measurement metadata and content
            schema:
              id: MeasurementSearchResponse
              type: array
              items:
                type: object
                properties:
                  defect_id:
                    type: string

                  inspection_id:
                    type: number
                  inspection_esn:
                    type: string
                  inspection_sect:
                    type: string

                  image_id:
                    type: number
                  image_distance:
                    type: number

                  id:
                    type: number
                  measurement_status:
                    type: string
                  measurement_disposition:
                    type: string
                  measurement_finding_type:
                    type: string
                  measurement_is_manual:
                    type: boolean

                  validated_measurement_annotation_file_id:
                    type: number
                  validated_by:
                    type: string
                  validation_status:
                    type: string
                  validation_timestamp:
                    type: string
                  validation_labels:
                    type: array
                    items:
                      type: string
                    example: [CoreGap, Lamination]

        """

        # Inspection filter props -- toto1
        esn = request.args.get('esn', '%')
        blade_section = request.args.get('blade_section', '%')

        # Image filter props
        root_face_distance = request.args.get('root_face_distance', None)

        # measurement filter props
        validation_status = request.args.get('validation_status', None)

        logging.info(f'esn: {esn}')
        logging.info(f'blade_section: {blade_section}')
        logging.info(f'root_face_distance: {root_face_distance}')
        logging.info(f'validation_status: {validation_status}')

        logging.info(f'Search for measurement records...')
        with Session(db) as session:
            measurement_list = session.query(
                Inspection.id.label('inspection_id'),
                Inspection.esn.label('inspection_esn'),
                Inspection.sect.label('inspection_sect'),
                Inspection.blade_type.label('inspection_blade_type'),
                Inspection.upload_date.label('inspection_upload_date'),
                Inspection.sso.label('inspection_sso'),

                Image.id.label('image_id'),
                Image.distance.label('image_distance'),

                Measurement.id.label('measurement_id'),
                Measurement.ge_disposition.label('measurement_disposition'),
                Measurement.status.label('measurement_status'),
                Measurement.finding_type.label('measurement_finding_type'),
                Measurement.is_manual.label('measurement_is_manual'),

                ValidatedMeasurementAnnotationFile.id.label(
                    'validated_measurement_annotation_file_id'),
                ValidatedMeasurementAnnotationFile.validated_by.label(
                    'validated_by'),
                ValidatedMeasurementAnnotationFile.validation_status.label(
                    'validation_status'),
                ValidatedMeasurementAnnotationFile.validation_timestamp.label(
                    'validation_timestamp'),
                ValidatedMeasurementAnnotationFile.content.label(
                    'validation_content'),

                OriginalMeasurementAnnotationFile.id.label(
                    'original_measurement_annotation_file_id'),
                OriginalMeasurementAnnotationFile.replaced_by.label(
                    'original_replaced_by'),
                OriginalMeasurementAnnotationFile.replaced_timestamp.label(
                    'original_replaced_timestamp'),
                OriginalMeasurementAnnotationFile.content.label(
                    'original_content')

            ).join(
                Image, Image.inspection_id == Inspection.id,  isouter=False
            ).join(Measurement, Measurement.image_id == Image.id,  isouter=False
                   ).join(ValidatedMeasurementAnnotationFile,
                          ValidatedMeasurementAnnotationFile.measurement_id == Measurement.id,
                          isouter=True
                          ).join(OriginalMeasurementAnnotationFile,
                          OriginalMeasurementAnnotationFile.measurement_id == Measurement.id,
                          isouter=True
                          ).order_by(Inspection.esn
                              ).filter(and_(Inspection.esn.like(esn),
                                     Inspection.sect.like(blade_section))
                                ).all()
            # logging.info(f'measurement_list: {measurement_list}')

            resp_list = []
            for measurement in measurement_list:

                measurement_id = measurement[6]
                defect_disp_list = session.query(Defect.ge_disposition).distinct(Defect.ge_disposition).filter(Defect.measurement_id == measurement_id).all()
                #logging.info(f'defect_list: {defect_disp_list}')
                dispositionSet = set()
                for defect in defect_disp_list:
                    dispositionSet.add(defect.ge_disposition)

                row = {
                    # inspection.esn + measurement.id
                    "id": measurement[1],
                    "measurement_id": str(measurement[1])+'-'+str(measurement[8]),

                    "inspection_id": measurement[0],
                    "inspection_esn": measurement[1],
                    "inspection_sect": measurement[2],
                    "inspection_blade_type": measurement[3],
                    "inspection_upload_date": measurement[4],
                    "inspection_sso": measurement[5],

                    "image_id": measurement[6],
                    "image_distance": measurement[7],

                    # measurement.id is the primary key of this view
                    "id": measurement[8],
                    "measurement_disposition": measurement[9],
                    "measurement_status": measurement[10],
                    "measurement_finding_type": measurement[11],
                    "measurement_is_manual": measurement[12],

                    "validated_measurement_annotation_file_id": measurement[13],
                    "validated_by": measurement[14],
                    "validation_status": measurement[15],
                    "validation_timestamp": measurement[16],
                    "validation_labels": extract_labels(measurement[17]),

                    "original_measurement_annotation_file_id": measurement[18],
                    "original_replaced_by": measurement[19],
                    "original_replaced_timestamp": measurement[20],
                    "original_labels": extract_labels(measurement[21]),

                    "defect_disposition_list": list(dispositionSet)

                }


                if len(dispositionSet) > 0 and row["measurement_disposition"]  == '' :
                    row["measurement_disposition"] = next(iter(dispositionSet))

                resp_list.append(row)

            # filter by validation status
            filtered_data = []
            if validation_status != None:
                logging.info(f'filter validation_status == {validation_status}')
                for el in resp_list:
                    if el['validation_status'] == validation_status:
                        filtered_data.append(el)
                resp_list = filtered_data
                logging.info(f'filtered by validation_status: {resp_list}')

            # filter distance
            filtered_data = []
            if root_face_distance != None:
                logging.info(f'filter root_face_distance == {root_face_distance}')
                for el in resp_list:
                    if str(el['image_distance']) == str(root_face_distance):
                        filtered_data.append(el)
                resp_list = filtered_data
                logging.info(f'filtered by distance: {resp_list}')

            logging.info(f'returning {len(resp_list)} measurements.')
            return jsonify(resp_list)


class MeasurementAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):
        super(MeasurementAPI, self).__init__()

    def delete(self, id):
        """
        Deletes an existing measurement with its companion measurement_image and measurement_annotation files if any
        ---
        parameters:
          - in: path
            required: true
            name: id
            schema:
              type: number
            description: Existing measurement id

        responses:
          200:
            description: Success or failure message
            schema:
              properties:
                message:
                  type: string

        """
        with Session(db) as session:

            measurement = session.query(Measurement).get(id)
            if measurement is not None:

                delete_measurement_dependencies(measurement.id)

                # Then finally delete Measurement (it can only be deelted if there is no back references via measurement_id)
                session.delete(measurement)
                session.commit()
                logging.info(f'Successfully deleted measurement id# {measurement.id}')

                return {'message': f'Successfully deleted measurement id# {id}'}

        return {'message': f'Measurement id #{id} not found.'}

    def get(self, id):
        """
        Read an existing measurement meta-data
        ---
        parameters:
          - in: path
            required: true
            name: id
            schema:
              type: number
            description: Existing measurement id

        responses:
          200:
            description: Existing measurement metadata
            schema:
              id: MeasurementObject

        """

        logging.info(f'Look for measurement record id {id}')
        with Session(db) as session:
            measurement = session.query(Measurement).get(id)
            if measurement is None:
                return {'message': f'Measurement {id} not found.'}
            jsonResp = measurement.toJson()
            return jsonResp


class CreateMeasurementAPI(Resource):

    def __init__(self):
        self.reqparse = reqparse.RequestParser()

        self.reqparse.add_argument('image_id', type=int, location='json')

        self.reqparse.add_argument('finding_code', type=str, location='json')
        self.reqparse.add_argument(
            'submission_code', type=str, location='json')
        self.reqparse.add_argument('date', type=str, location='json')
        self.reqparse.add_argument('component', type=str, location='json')
        self.reqparse.add_argument('reference', type=str, location='json')
        self.reqparse.add_argument(
            'position_in_blade', type=str, location='json')
        self.reqparse.add_argument('location', type=str, location='json')

        self.reqparse.add_argument(
            'root_face_distance', type=float, location='json')
        self.reqparse.add_argument(
            'edge_distance', type=float, location='json')
        self.reqparse.add_argument('le_distance', type=float, location='json')
        self.reqparse.add_argument('te_distance', type=float, location='json')
        self.reqparse.add_argument(
            'span_wise_length', type=float, location='json')
        self.reqparse.add_argument(
            'chord_wise_width', type=float, location='json')
        self.reqparse.add_argument('depth', type=float, location='json')
        self.reqparse.add_argument('height', type=float, location='json')
        self.reqparse.add_argument('width', type=float, location='json')
        self.reqparse.add_argument('length', type=float, location='json')
        self.reqparse.add_argument('aspect_ratio', type=float, location='json')
        self.reqparse.add_argument('area', type=float, location='json')
        self.reqparse.add_argument('percent_area', type=float, location='json')

        self.reqparse.add_argument('finding_type', type=str, location='json')
        self.reqparse.add_argument(
            'finding_category', type=str, location='json')
        self.reqparse.add_argument(
            'finding_reference', type=str, location='json')
        self.reqparse.add_argument('ge_disposition', type=str, location='json')
        self.reqparse.add_argument(
            'ge_disposition_response', type=str, location='json')
        self.reqparse.add_argument('dnv_response', type=str, location='json')
        self.reqparse.add_argument('is_priority', type=bool, location='json')
        self.reqparse.add_argument('description', type=str, location='json')

        self.reqparse.add_argument('image_pitch', type=float, location='json')
        self.reqparse.add_argument('image_yaw', type=float, location='json')
        self.reqparse.add_argument('image_hfov', type=float, location='json')

        self.reqparse.add_argument(
            'design_tolerance', type=str, location='json')
        self.reqparse.add_argument(
            'disposition_provided_by', type=str, location='json')
        self.reqparse.add_argument('status', type=str, location='json')
        self.reqparse.add_argument('repair_date', type=str, location='json')
        self.reqparse.add_argument(
            'repair_report_id', type=str, location='json')
        self.reqparse.add_argument(
            'repair_approved_by', type=str, location='json')
        self.reqparse.add_argument('is_manual', type=bool, location='json')
        self.reqparse.add_argument(
            'sso', type=str, location='json')

        super(CreateMeasurementAPI, self).__init__()

    def post(self):
        """
        Create a new measurement record , so one can later attach a 2d measurement image and an annotation file
        ---
        parameters:
          - name: CreateMeasurementRequest
            in: body
            required: true
            schema:
              id: CreateMeasurementBody
              properties:
                image_id:
                  type: number
                  example: 1
                finding_code:
                  type: string
                submission_code:
                  type: string
                date:
                  type: string
                  example: "2023-11-01"
                component:
                  type: string
                reference:
                  type: string
                position_in_blade:
                  type: string
                location:
                  type: string
                root_face_distance:
                  type: number
                edge_distance:
                  type: number
                le_distance:
                  type: number
                te_distance:
                  type: number
                span_wise_length:
                  type: number
                chord_wise_width:
                  type: number
                depth:
                  type: number
                height:
                  type: number
                width:
                  type: number
                length:
                  type: number
                aspect_ratio:
                  type: number
                area:
                  type: number
                percent_area:
                  type: number
                finding_type:
                  type: string
                finding_category:
                  type: string
                finding_reference:
                  type: string
                ge_disposition:
                  type: string
                ge_disposition_response:
                  type: string
                dnv_response:
                  type: string
                is_priority:
                  type: boolean
                description:
                  type: string
                image_pitch:
                  type: number
                image_yaw:
                  type: number
                image_hfov:
                  type: number
                design_tolerance:
                  type: string
                disposition_provided_by:
                  type: string
                status:
                  type: string
                repair_date:
                  type: string
                  example: "2023-11-01"
                repair_report_id:
                  type: string
                repair_approved_by:
                  type: string
                is_manual:
                  type: boolean
                sso:
                  type: string
        responses:
          200:
            description: Created inspection
            schema:
              id: MeasurementObject
              properties:
                id:
                  type: number

                image_id:
                  type: number

                finding_code:
                  type: string
                submission_code:
                  type: string
                date:
                  type: string
                component:
                  type: string
                reference:
                  type: string
                position_in_blade:
                  type: string
                location:
                  type: string

                root_face_distance:
                  type: number
                edge_distance:
                  type: number
                le_distance:
                  type: number
                te_distance:
                  type: number
                span_wise_length:
                  type: number
                chord_wise_width:
                  type: number
                depth:
                  type: number
                height:
                  type: number
                width:
                  type: number
                length:
                  type: number
                aspect_ratio:
                  type: number
                area:
                  type: number
                percent_area:
                  type: number

                finding_type:
                  type: string
                finding_category:
                  type: string
                finding_reference:
                  type: string
                ge_disposition:
                  type: string
                ge_disposition_response:
                  type: string
                dnv_response:
                  type: string
                is_priority:
                  type: boolean
                description:
                  type: string

                image_pitch:
                  type: number
                image_yaw:
                  type: number
                image_hfov:
                  type: number

                design_tolerance:
                  type: string
                disposition_provided_by:
                  type: string
                status:
                  type: string
                repair_date:
                  type: string
                repair_report_id:
                  type: string
                repair_approved_by:
                  type: string
                is_manual:
                  type: boolean
                sso:
                  type: string

        """
        args = self.reqparse.parse_args()
        logger.info(args.items())

        if args['date'] == '':
            # current_date = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            current_date = datetime.now().isoformat()
            args['date'] = current_date

        image_id = args['image_id']
        with Session(db) as session:
            image = session.query(Image).get(image_id)
            if image is None:
                return {'message': f'Could not create measurement. Image record for image_id: {image_id} not found.'}

        measurement = Measurement(
            image_id=args['image_id'],

            finding_code=args['finding_code'],
            submission_code=args['submission_code'],
            date=args['date'],
            component=args['component'],
            reference=args['reference'],
            position_in_blade=args['position_in_blade'],
            location=args['location'],

            root_face_distance=args['root_face_distance'],
            edge_distance=args['edge_distance'],
            le_distance=args['le_distance'],
            te_distance=args['te_distance'],
            span_wise_length=args['span_wise_length'],
            chord_wise_width=args['chord_wise_width'],
            depth=args['depth'],
            height=args['height'],
            width=args['width'],
            length=args['length'],
            aspect_ratio=args['aspect_ratio'],
            area=args['area'],
            percent_area=args['percent_area'],

            finding_type=args['finding_type'],
            finding_category=args['finding_category'],
            finding_reference=args['finding_reference'],
            ge_disposition=args['ge_disposition'],
            ge_disposition_response=args['ge_disposition_response'],
            dnv_response=args['dnv_response'],

            is_priority=args['is_priority'],
            description=args['description'],

            image_pitch=args['image_pitch'],
            image_yaw=args['image_yaw'],
            image_hfov=args['image_hfov'],

            design_tolerance=args['design_tolerance'],
            disposition_provided_by=args['disposition_provided_by'],
            status=args['status'],
            repair_date=args['repair_date'],
            repair_report_id=args['repair_report_id'],
            repair_approved_by=args['repair_approved_by'],
            is_manual=args['is_manual'],
            sso=args['sso'],
        )

        with Session(db) as session:
            session.add(measurement)
            session.commit()
            session.refresh(measurement)  # to read the id

            image = session.query(Image).get(measurement.image_id)
            inspection_id = image.inspection_id
            invalidate_inspection_cache(inspection_id)

            logger.info(f'Created measurement: {str(measurement)}')

        # resp = jsonify(inspection.serialize())
        resp = measurement.toJson()
        logging.info(f'resp: {resp}')
        return resp


class UpdateMeasurementAPI(Resource):

    def __init__(self):
        self.reqparse = reqparse.RequestParser()

        # self.reqparse.add_argument('image_id', type=int, location='json')

        self.reqparse.add_argument('finding_code', type=str, location='json')
        self.reqparse.add_argument(
            'submission_code', type=str, location='json')
        self.reqparse.add_argument('date', type=str, location='json')
        self.reqparse.add_argument('component', type=str, location='json')
        self.reqparse.add_argument('reference', type=str, location='json')
        self.reqparse.add_argument(
            'position_in_blade', type=str, location='json')
        self.reqparse.add_argument('location', type=str, location='json')

        self.reqparse.add_argument(
            'root_face_distance', type=float, location='json')
        self.reqparse.add_argument(
            'edge_distance', type=float, location='json')
        self.reqparse.add_argument('le_distance', type=float, location='json')
        self.reqparse.add_argument('te_distance', type=float, location='json')
        self.reqparse.add_argument(
            'span_wise_length', type=float, location='json')
        self.reqparse.add_argument(
            'chord_wise_width', type=float, location='json')
        self.reqparse.add_argument('depth', type=float, location='json')
        self.reqparse.add_argument('height', type=float, location='json')
        self.reqparse.add_argument('width', type=float, location='json')
        self.reqparse.add_argument('length', type=float, location='json')
        self.reqparse.add_argument('aspect_ratio', type=float, location='json')
        self.reqparse.add_argument('area', type=float, location='json')
        self.reqparse.add_argument('percent_area', type=float, location='json')

        self.reqparse.add_argument('finding_type', type=str, location='json')
        self.reqparse.add_argument(
            'finding_category', type=str, location='json')
        self.reqparse.add_argument(
            'finding_reference', type=str, location='json')
        self.reqparse.add_argument('ge_disposition', type=str, location='json')
        self.reqparse.add_argument(
            'ge_disposition_response', type=str, location='json')
        self.reqparse.add_argument('dnv_response', type=str, location='json')
        self.reqparse.add_argument('is_priority', type=bool, location='json')
        self.reqparse.add_argument('description', type=str, location='json')

        self.reqparse.add_argument('image_pitch', type=float, location='json')
        self.reqparse.add_argument('image_yaw', type=float, location='json')
        self.reqparse.add_argument('image_hfov', type=float, location='json')

        self.reqparse.add_argument(
            'design_tolerance', type=str, location='json')
        self.reqparse.add_argument(
            'disposition_provided_by', type=str, location='json')
        self.reqparse.add_argument('status', type=str, location='json')
        self.reqparse.add_argument('repair_date', type=str, location='json')
        self.reqparse.add_argument(
            'repair_report_id', type=str, location='json')
        self.reqparse.add_argument(
            'repair_approved_by', type=str, location='json')
        self.reqparse.add_argument('is_manual', type=bool, location='json')

        super(UpdateMeasurementAPI, self).__init__()

    def post(self, id):
        """
        Update an existing measurement record meta-data
        ---
        parameters:
          - in: path
            required: true
            name: id
            schema:
              type: number
            description: Existing measurement id
          - name: UpdateMeasurementRequest
            in: body
            required: true
            schema:
              id: CreateMeasurementBody

        responses:
          200:
            description: Updated inspection
            schema:
              id: MeasurementObject

        """
        args = self.reqparse.parse_args()
        logger.info(f'body: {args.items()}')

        # TODO: validate date instead
        if args['date'] is None or args['date'] == '':
            # current_date = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            current_date = datetime.now().isoformat()
            args['date'] = current_date

        with Session(db) as session:
            measurement = session.query(Measurement).get(id)
            if measurement is None:
                return {'message': f'Could not find measurement id#: {id}.'}

            # NOTE: We do not want to update the image id for this measurement directly, one should delete and re-create the measurement for that
            # if args['image_id'] is not None:
            #     image_id = args['image_id']
            #     image = session.query(Image).get(image_id)
            #     if image is None:
            #         return {'message': f'Invalid image id#: {image_id}.'}
            #     measurement.image_id = image_id

            if args['finding_code'] is not None:
                measurement.finding_code = args['finding_code']

            if args['submission_code'] is not None:
                measurement.submission_code = args['submission_code']

            if args['date'] is not None and args['date'] != '':
                measurement.date = args['date']

            if args['component'] is not None:
                measurement.component = args['component']

            if args['reference'] is not None:
                measurement.reference = args['reference']

            if args['position_in_blade'] is not None:
                measurement.position_in_blade = args['position_in_blade']

            if args['location'] is not None:
                measurement.location = args['location']

            if args['root_face_distance'] is not None:
                measurement.root_face_distance = args['root_face_distance']

            if args['edge_distance'] is not None:
                measurement.edge_distance = args['edge_distance']

            if args['le_distance'] is not None:
                measurement.le_distance = args['le_distance']

            if args['te_distance'] is not None:
                measurement.te_distance = args['te_distance']

            if args['span_wise_length'] is not None:
                measurement.span_wise_length = args['span_wise_length']

            if args['chord_wise_width'] is not None:
                measurement.chord_wise_width = args['chord_wise_width']

            if args['depth'] is not None:
                measurement.depth = args['depth']

            if args['height'] is not None:
                measurement.height = args['height']

            if args['width'] is not None:
                measurement.width = args['width']

            if args['length'] is not None:
                measurement.length = args['length']

            if args['aspect_ratio'] is not None:
                measurement.aspect_ratio = args['aspect_ratio']

            if args['area'] is not None:
                measurement.area = args['area']

            if args['percent_area'] is not None:
                measurement.percent_area = args['percent_area']

            if args['finding_type'] is not None:
                measurement.finding_type = args['finding_type']

            if args['finding_category'] is not None:
                measurement.finding_category = args['finding_category']

            if args['finding_reference'] is not None:
                measurement.finding_reference = args['finding_reference']

            if args['ge_disposition'] is not None:
                measurement.ge_disposition = args['ge_disposition']

            if args['ge_disposition_response'] is not None:
                measurement.ge_disposition_response = args['ge_disposition_response']

            if args['dnv_response'] is not None:
                measurement.dnv_response = args['dnv_response']

            if args['is_priority'] is not None:
                measurement.is_priority = args['is_priority']

            if args['description'] is not None:
                measurement.description = args['description']

            if args['image_pitch'] is not None:
                measurement.image_pitch = args['image_pitch']
            if args['image_yaw'] is not None:
                measurement.image_yaw = args['image_yaw']
            if args['image_hfov'] is not None:
                measurement.image_hfov = args['image_hfov']

            if args['design_tolerance'] is not None:
                measurement.design_tolerance = args['design_tolerance']
            if args['disposition_provided_by'] is not None:
                measurement.disposition_provided_by = args['disposition_provided_by']
            if args['status'] is not None:
                measurement.status = args['status']

            if args['repair_date'] is not None:
                # if one explicitly sets the date to empty, we remove the date
                if args['repair_date'].strip() == '':
                    logging.info('Setting repair_date to None')
                    measurement.repair_date = None
                else:
                    measurement.repair_date = args['repair_date']

            if args['repair_report_id'] is not None:
                measurement.repair_report_id = args['repair_report_id']
            if args['repair_approved_by'] is not None:
                measurement.repair_approved_by = args['repair_approved_by']
            if args['is_manual'] is not None:
                measurement.is_manual = args['is_manual']

            # logger.info(f'Measurement before update: {str(measurement)}')
            session.commit()
            session.refresh(measurement)

            image = session.query(Image).get(measurement.image_id)
            inspection_id = image.inspection_id
            invalidate_inspection_cache(inspection_id)

        logger.info(f'Updated measurement: {str(measurement)}')

        # resp = jsonify(inspection.serialize())
        resp = measurement.toJson()
        logging.info(f'resp: {resp}')
        return resp

# ----------------------------- Measurement Image File -------------------


class MeasurementImageFileAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):
        super(MeasurementImageFileAPI, self).__init__()

    def allowed_file(self, filename):
        return '.' in filename and \
            filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

    def post(self, id):
        """
        Update an existing measurement image file or create a new one in case the measurement has no image file
        ---
        consumes:
          - multipart/form-data
        parameters:
          - in: path
            required: true
            name: id
            schema:
              type: number
            description: Existing measurement record id
          - in: formData
            required: true
            name: image_file
            type: file
            description: new image file
        responses:
          200:
            description: Success or failure message
            schema:
              properties:
                message:
                  type: string
                id:
                  type: number

        """
        # formData image_file
        measurement_image_file = request.files.get("image_file")
        if measurement_image_file is None or not self.allowed_file(measurement_image_file.filename):
            msg = 'Measurement image content not found, or no image_file provided.'
            return {'message': msg}

        measurement_image_filename = secure_filename(
            measurement_image_file.filename)
        unique_image_filename = str(uuid.uuid4()) + \
        '_'+measurement_image_filename
        measurement_image_file_path = os.path.join(
            app.config['UPLOAD_FOLDER'], unique_image_filename)
        measurement_image_file.save(measurement_image_file_path)

        logging.info('Reading image_file formData content...')
        bin_content = None
        with open(measurement_image_file_path, mode="rb") as f:
            bin_content = f.read()
        logging.info('Generating image_file formData thumbnail...')
        thumb_bin_content = get_thumbnail_content(measurement_image_file_path)
        if os.path.isfile(measurement_image_file_path):
            os.remove(measurement_image_file_path)

        with Session(db) as session:
            measurement_rec = session.query(Measurement).get(id)
            if measurement_rec is None:
                msg = f'Measurement record id: {id} not found.'
                return {'message': msg}

            logger.info(
                f'Reading measurement_image_file content for measurement.id #{id}')
            measurement_image_file_list = session.scalars(
                select(MeasurementImageFile).where(
                    MeasurementImageFile.measurement_id == id)
            ).all()
            measurement_image_file_rec = None

            if COMPRESS_INCOMING_IMAGES:
                measurement_image_filename = measurement_image_filename.replace(
                    '.png', '.jpg')  # new filename
                bin_content = convert_png_content_to_jpg(bin_content)

            # inexisting record, create new
            if len(measurement_image_file_list) == 0:
                measurement_image_file_rec = MeasurementImageFile(
                    image_id=measurement_rec.image_id,
                    measurement_id=id,
                    filename=measurement_image_filename,
                    s3key=None,
                    content=bin_content,
                    thumbnail=thumb_bin_content
                )
                session.add(measurement_image_file_rec)
            else:  # update existing record
                measurement_image_file_rec = measurement_image_file_list[0]
                measurement_image_file_rec.filename = measurement_image_filename
                measurement_image_file_rec.content = bin_content
                measurement_image_file_rec.thumbnail = thumb_bin_content

            if USE_S3:
                # Obtain image, image_file and inspection for the measurement
                image_rec = session.query(Image).get(measurement_rec.image_id)
                image_file_rec = session.scalars(
                    select(ImageFile).filter(
                        ImageFile.image_id == measurement_rec.image_id)
                ).all()[0]
                inspection_rec = session.query(
                    Inspection).get(image_rec.inspection_id)
                # name of image where this measurement came from
                pan_image_filename = image_file_rec.filename
                measurements_folder_name = pan_image_filename.split('.')[
                    0]+'_measurements'
                s3key = get_inspection_s3key(
                    inspection_rec)+f'/images/{measurements_folder_name}/{measurement_image_filename}'
                response = upload_content_to_s3(bin_content, s3key)
                logging.info(f'upload to s3 resp: {response}')
                if response.status_code != 200:
                    return {'message': f'Error uploading file: {measurement_image_filename} to S3'}
                measurement_image_file_rec.s3key = s3key
                measurement_image_file_rec.content = None

            session.commit()
            session.refresh(measurement_image_file_rec)

            measurement_image_file_rec.content = None
            measurement_image_file_rec.thumbnail = None
            logging.info(
                f'new_measurement_image_file_rec: {measurement_image_file_rec}')
            measurement_image_file_json = measurement_image_file_rec.toJson()
            logging.info(
                f'measurement_image_file_json: {measurement_image_file_json}')

            resp = {
                'message': f' Measurement Image file {measurement_image_file.filename} uploaded successfully.',
                'id': measurement_image_file_rec.id
            }
            return resp

    def get(self, id):
        """
        Read an existing measurement image file
        ---
        parameters:
          - in: path
            required: true
            name: id
            schema:
              type: number
            description: Existing measurement id
          - in: query
            required: false
            name: includeAnnotations
            schema:
              type: boolean
            description: Whether or not to include annotations on the image
          - in: query
            required: false
            name: includeValidatedAnnotations
            schema:
              type: boolean
            description: Whether or not to include validated annotations on the image
          - in: query
            required: false
            name: includeOriginalAnnotations
            schema:
              type: boolean
            description: Whether or not to include original annotations on the image
        responses:
          200:
            description: image file content for the provided measurement id with optional annotations
            schema:
              id: MeasurementImageFileContent
              type: string
              format: binary
        """

        include_annotations_str = request.args.get('includeAnnotations')
        if include_annotations_str is not None:
            include_annotations_str = include_annotations_str.capitalize()
        logger.info(f'include_annotations_str: {include_annotations_str}')
        include_annotations = False
        if include_annotations_str == 'True':
            include_annotations = True

        include_validated_annotations_str = request.args.get(
            'includeValidatedAnnotations')
        if include_validated_annotations_str is not None:
            include_validated_annotations_str = include_validated_annotations_str.capitalize()
        logger.info(f'include_validated_annotations_str: {include_validated_annotations_str}')
        include_validated_annotations = False
        if include_validated_annotations_str == 'True':
            include_validated_annotations = True

        include_original_annotations_str = request.args.get(
            'includeOriginalAnnotations')
        if include_original_annotations_str is not None:
            include_original_annotations_str = include_original_annotations_str.capitalize()
        logger.info(f'include_origional_annotations_str: {include_original_annotations_str}')
        include_original_annotations = False
        if include_original_annotations_str == 'True':
            include_original_annotations = True

        logger.info(f'Reading image_file content for measurement.id {id}')
        logger.info(f'include_annotations: {include_annotations}')
        logger.info(f'include_validated_annotations: {include_validated_annotations}')

        with Session(db) as session:
            measurement_image_results = session.scalars(
                select(MeasurementImageFile).where(MeasurementImageFile.measurement_id == id)).all()

            if measurement_image_results is not None and len(measurement_image_results) > 0:
                image_file = measurement_image_results[0]

                file_content = read_file_record_content(image_file)
                if file_content is not None:

                    if include_annotations == True:
                        logging.info('including annotations...')
                        annotation_file_results = session.scalars(
                            select(MeasurementAnnotationFile).where(MeasurementAnnotationFile.measurement_id == id)).all()
                        if annotation_file_results is not None and len(annotation_file_results) > 0:
                            annotation_file = annotation_file_results[0]
                            annotation_str = annotation_file.content
                            annotation_json = json.loads(annotation_str)

                            # overwrite the image file content with the new one including annotations
                            file_content = draw_annotations_on_image(
                                file_content, annotation_json)

                    if include_validated_annotations == True:
                        logging.info('including validated annotations...')
                        validated_annotation_file_results = session.scalars(
                            select(ValidatedMeasurementAnnotationFile).where(ValidatedMeasurementAnnotationFile.measurement_id == id)).all()
                        if validated_annotation_file_results is not None and len(validated_annotation_file_results) > 0:
                            validated_annotation_file = validated_annotation_file_results[0]
                            validated_annotation_str = validated_annotation_file.content
                            validated_annotation_json = json.loads(
                                validated_annotation_str)

                            # overwrite the image file content with the new one including annotations
                            file_content = draw_annotations_on_image(
                                file_content, validated_annotation_json)

                    if include_original_annotations == True:
                        logging.info('including original annotations...')
                        original_annotation_file_results = session.scalars(
                            select(OriginalMeasurementAnnotationFile).where(OriginalMeasurementAnnotationFile.measurement_id == id)).all()
                        if original_annotation_file_results is not None and len(original_annotation_file_results) > 0:
                            original_annotation_file = original_annotation_file_results[0]
                            original_annotation_str = original_annotation_file.content
                            original_annotation_json = json.loads(
                                original_annotation_str)

                            # overwrite the image file content with the new one including annotations
                            file_content = draw_annotations_on_image(
                                file_content, original_annotation_json)

                    return send_file_content_as_jpg(image_file.filename, file_content)
            else:
                logging.info(f'measurement_image_results for measurement_id {id} not found.')

        return {'message': f'Image File for Measurement id: {id} not found.'}


def get_content_mime_type(file_content):
    file_type_obj = filetype.guess(file_content)
    mime_type = file_type_obj.mime
    return mime_type


def send_file_content_as_jpg(filename, file_content):
    file_type_obj = filetype.guess(file_content)
    mime_type = file_type_obj.mime
    logging.info(f'content mime_type: {mime_type}')
    if 'png' in mime_type:
        return send_file(
            BytesIO(convert_png_content_to_jpg(file_content)),
            mimetype='image/jpeg',
            as_attachment=True,
            download_name=filename.replace('.png', '.jpg'))
    else:
        return send_file(
            BytesIO(file_content),
            mimetype='image/jpeg',
            as_attachment=True,
            download_name=filename)


def draw_annotations_on_image(file_content, annotation_json, line_thickness=3, include_text=True):
    logging.info(f'draw_annotations_on_image() called')
    #logging.info(f'draw_annotations_on_image() called with annotation_json: {annotation_json}')
    img_data = PImage.open(BytesIO(file_content))
    overlay_img = overlay_polygons(
        img_data, annotation_json, thickness=line_thickness, draw_text=include_text)

    is_success, im_buf_arr = cv2.imencode(".jpg", overlay_img)
    new_file_content = im_buf_arr.tobytes()
    return new_file_content


def draw_text_below_polygon(image, points, text, font=cv2.FONT_HERSHEY_SIMPLEX,
                            font_scale=1.5, font_color=(255, 255, 255),
                            font_thickness=2, offset=30):
    """Draw text below the polygon defined by points."""
    # Find the bottom center point of the polygon
    x_coords, y_coords = zip(*points)
    bottom_center = (min(x_coords) + max(x_coords)
                     ) // 2, max(y_coords) + offset
    # Put text on the image
    cv2.putText(image, text, bottom_center, font,
                font_scale, font_color, font_thickness)


def apply_scale(x_factor, y_factor, points):
    if points == None or len(points) == 0:
        return []
    new_points = []
    for point in points:
        curr_x = point[0]
        curr_y = point[1]
        new_x = round(int(curr_x)*x_factor)
        new_y = round(int(curr_y)*y_factor)
        # logging.info(f'apply_scale: ({curr_x}, {curr_y}) -> ({new_x}, {new_y})')
        new_points.append([new_x, new_y])
    return new_points


def overlay_polygons(image, json_data, thickness=3, draw_text=True):
    logging.info('overlay_polygons() called')

    is_ai = False
    if json_data is not None and 'AI' in json_data:
        logging.info(f'AI prop value: {json_data.get("AI")}')
        is_ai = json_data.get('AI')

    # Create a drawing context
    image = cv2.cvtColor(np.array(image, dtype=np.uint8)
                         [:, :, :3], cv2.COLOR_RGB2BGR)
    image_h, image_w, image_c = image.shape
    # logging.info(f'img width: {image_w}')
    # logging.info(f'img height: {image_h}')
    # logging.info(f'img channel: {image_c}')

    # logging.info(f'json_data: {json_data}')
    if json_data.get('imageWidth') is None or json_data.get('imageHeight') is None or json_data.get('shapes') is None:
        logging.info(
            'Incorrect annotation file format. Will not overlay_polygons.')
        return image

    # logging.info(f'json_data: {json_data}')
    anno_w = json_data['imageWidth']
    anno_h = json_data['imageHeight']
    # logging.info(f'anno_width: {anno_w}')
    # logging.info(f'anno_height: {anno_h}')

    x_factor = image_w/anno_w if anno_w else 1
    y_factor = image_h/anno_h if anno_h else 1
    # logging.info(f'x_factor: {x_factor}')
    # logging.info(f'y_factor: {y_factor}')

    im_overlay = image
    for shape in json_data['shapes']:
        #logging.info(f'shape: {shape}')
        # Assuming annotations are polygons
        if 'points' in shape and len(shape["points"]) > 0:
            #logging.info(f'shape[points]: {shape["points"]} ')
            points = apply_scale(x_factor, y_factor, shape['points'])
            points = np.array(points, dtype=int)
            color = (0, 234, 255)
            if 'label' in shape:
                color = get_defect_bgr_color(shape['label'], is_ai)

            if shape['shape_type'] == 'polygon':
                np_points = points.reshape((-1, 1, 2))
                im_overlay = cv2.polylines(np.array(image, dtype=np.uint8), [np_points], isClosed=True,
                                           color=color, thickness=thickness)
            elif shape['shape_type'] == 'rectangle':
                im_overlay = cv2.rectangle(np.array(image, dtype=np.uint8), points[0], points[1],
                                           color=color, thickness=thickness)
            elif shape['shape_type'] == 'linestrip':
                im_overlay = cv2.polylines(np.array(image, dtype=np.uint8), points[0], points[1],
                                           color=color, thickness=thickness)
            # draw text over image
            if draw_text and 'label' in shape:
                label_text = shape['label']
                if is_ai:
                    label_text = shape['label'] + " (AI)"

                draw_text_below_polygon(
                    im_overlay, points.tolist(), label_text, font_color=color)

            image = im_overlay

    return im_overlay


def is_valid_date(date_text):
    is_valid = True
    try:
        parse(date_text)
    except ValueError:
        is_valid = False
    logging.info(f'is_valid_date() called with: {date_text} -> {is_valid}')
    return is_valid


def get_thumbnail_content(image_file_path):
    
    thumb_bin_content = None
    if '.png' in image_file_path or '.jpg' in image_file_path:
        tmp_thumbnail_file_path = image_file_path.replace( r'\.(jpg|png)', '_thumbnail.jpg')
        p_image = PImage.open(image_file_path)
        p_image = p_image.convert("RGB")
        width, height = p_image.size
        ratio = width / height
        new_height = 100
        new_width = int(width * ratio)
        p_image.thumbnail((new_width, new_height))
        p_image.save(tmp_thumbnail_file_path)
        if os.path.isfile(tmp_thumbnail_file_path):
            with open(tmp_thumbnail_file_path, mode="rb") as f:
                thumb_bin_content = f.read()
            os.remove(tmp_thumbnail_file_path)
    return thumb_bin_content


class MeasurementAnnotationFileAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):
        super(MeasurementAnnotationFileAPI, self).__init__()

    def allowed_file(self, filename):
        return '.' in filename and \
            filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

    def post(self, id):
        """
        Update an existing measurement 2d image annotation file for a measurement id or create a new one
        ---
        consumes:
          - multipart/form-data
        parameters:
          - in: path
            required: true
            name: id
            schema:
              type: number
            description: Existing measurement id
          - in: formData
            required: true
            name: annotation_file
            type: file
            description: new annotation file
          - in: formData
            required: false
            name: compute_measurements
            type: boolean
            description: whether to force computing of measurements when measurements are already present in the file
        responses:
          200:
            description: Success or failure message
            schema:
              properties:
                message:
                  type: string

        """
        logging.info('MeasurementAnnotationFileAPI.post()')

        force_compute_measurements = request.form.get(
            'compute_measurements', False)
        logging.info(f'compute_measurements: {force_compute_measurements}')

        # formData annotation_file
        annotation_file = request.files.get("annotation_file")
        is_file_type_allowed = self.allowed_file(annotation_file.filename)
        logging.info(f'is_file_type_allowed: {is_file_type_allowed}')
        if annotation_file is None or not is_file_type_allowed:
            msg = 'Error. Annotation file not found, or no annotation_file provided.'
            return {'message': msg}

        annotation_filename = secure_filename(annotation_file.filename)
        unique_annotation_filename = str(uuid.uuid4())+'_'+annotation_filename
        
        with tempfile.TemporaryDirectory(dir=UPLOAD_FOLDER, prefix='temp-', suffix='-folder') as tempdirname:
            annotation_file_path = os.path.join(tempdirname, unique_annotation_filename)
            annotation_file.save(annotation_file_path)

            with Session(db) as session:

                # Check if measurement record exists
                measurement_rec = session.query(Measurement).get(id)
                if measurement_rec is None:
                    msg = f'Measurement record id: {id} not found. You must create a measurement record first.'
                    return {'message': msg}

                logging.info('Reading annotation_file content...')
                annotation_json_content = None
                if os.path.isfile(annotation_file_path):
                    annotation_json_content = parse_annotation_file_and_save_props(
                        annotation_file_path, measurement_rec, session)
                    # since it is not a new measurement record, we have the coordinates already in the measurement rec
                    annotation_json_content_coord = None
                    if annotation_json_content is not None:
                        annotation_json_content_coord = add_snapshot_coordinates(
                            annotation_json_content, measurement_rec)
                    
                    annotation_json_content_measurements = None
                    if annotation_json_content_coord is not None:
                        annotation_json_content_measurements = measure_annotation_file_and_update_json_and_measurement_rec(annotation_json_content_coord, measurement_rec, session, force_compute_measurements)
                    
                    # ping_resp = ping_measurement_httpx()
                    # print(f'ping measurement resp: {ping_resp}')
                    # annotation_json_content_measurements = annotation_json_content_coord
                    
                    #os.remove(annotation_file_path)

                str_annotation_json_content = None
                if annotation_json_content_measurements is not None:
                    str_annotation_json_content = json.dumps(annotation_json_content_measurements)
                logging.info(
                    f'annotation_file_content str: {str_annotation_json_content}')

                logger.info(
                    f'Reading measurement_annotation_file content for measurement.id {id}')
                measurement_annotation_file_list = session.scalars(
                    select(MeasurementAnnotationFile).where(
                        MeasurementAnnotationFile.measurement_id == id)
                ).all()
                measurement_annotation_file_rec = None

                # no existing annotation files, create new one...
                if len(measurement_annotation_file_list) == 0:
                    logger.info('Creating new measurement_annotation_file record')
                    measurement_annotation_file_rec = MeasurementAnnotationFile(
                        image_id=measurement_rec.image_id,
                        measurement_id=id,
                        filename=annotation_filename,
                        s3key=None,
                        content=str_annotation_json_content,

                    )
                    session.add(measurement_annotation_file_rec)
                else:
                    logger.info('Update existing measurement_annotation_file record')
                    measurement_annotation_file_rec = measurement_annotation_file_list[0]
                    measurement_annotation_file_rec.filename = annotation_filename
                    measurement_annotation_file_rec.content = str_annotation_json_content

                session.add(measurement_annotation_file_rec)
                session.commit()
                session.refresh(measurement_annotation_file_rec)

                logging.info(
                    f'measurement_annotation_file_rec: {measurement_annotation_file_rec}')
            
                # measurement_annotation_file_rec_json = measurement_annotation_file_rec.toJson()
                # logging.info(
                #     f'measurement_annotation_file_json: {measurement_annotation_file_rec_json}')

                json_defect_list = parse_annotation_file_defects(measurement_rec, measurement_annotation_file_rec)

                image = session.query(Image).get(measurement_rec.image_id)
                inspection_id = image.inspection_id
                invalidate_inspection_cache(inspection_id)

                resp = {
                    'message': f' Measurement Annotation file {annotation_file.filename} updated successfully.',
                    'content': json.loads(measurement_annotation_file_rec.content)
                    }
                return resp

        return {'message': 'could not create temp folder'}
    

    def get(self, id):
        """
        Read an existing measurement annotation file record, representing annotations on 2d still image
        ---
        parameters:
          - in: path
            required: true
            name: id
            schema:
              type: number
            description: Existing measurement id

        responses:
          200:
            description: annotation file content for the provided measurement id
            schema:
              id: MeasurementAnnotationFileContent
              type: string
              format: json
        """
        logger.info(f'Reading annotation_file content for measurement.id {id}')
        with Session(db) as session:
            measurement_annotation_results = session.scalars(
                select(MeasurementAnnotationFile).where(MeasurementAnnotationFile.measurement_id == id)).all()
            logger.info(
                f'measurement_annotation_results: {measurement_annotation_results}')

            if measurement_annotation_results is not None and len(measurement_annotation_results) > 0:
                annotation_file = measurement_annotation_results[0]
                file_content = annotation_file.content
                if file_content is not None:
                    ret = json.loads(file_content)
                    logger.info(f'fihzor2 json.loads(file_content) = {ret}')
                    return ret

        return {'message': f'Annotation File content for Measurement id: {id} not found.'}

    def delete(self, id):
        """
        Deletes an existing measurement annotation file
        ---
        parameters:
          - in: path
            required: true
            name: id
            schema:
              type: number
            description: Existing measurement id
        responses:
          200:
            description: Success/failure message
            schema:
              message:
                type: string
        """
        logger.info(f'DELETE annotation for measurement id: {id}')
        with Session(db) as session:
            measurement_annotation_file_list = session.scalars(
                select(MeasurementAnnotationFile).where(
                    MeasurementAnnotationFile.measurement_id == id)
            ).all()

            if measurement_annotation_file_list is None or len(measurement_annotation_file_list) == 0:
                return {'message': f'annotation for measurement record id {id} not found'}

            # logging.info(f'del measurement_annotation_file_list: {measurement_annotation_file_list}')
            for measurement_annotation in measurement_annotation_file_list:
                logging.info(
                    f'del measurement_annotation id {measurement_annotation.id}')
                session.delete(measurement_annotation)
                session.commit()

            return {'message': f'annotation file for measurement id: {id} successfully deleted.'}

# --------------------------------- Defect API ------------------------------------------

# The defect and measurement use the same file. This API plots a single defect annotation on that file.


class DefectImageFileAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):
        super(DefectImageFileAPI, self).__init__()

    def allowed_file(self, filename):
        return '.' in filename and \
            filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

    def get(self, id):
        """
        Read an existing measurement image file of a defect with optional annotations layer
        ---
        parameters:
          - in: path
            required: true
            name: id
            schema:
              type: number
            description: Existing defect id
          - in: query
            required: false
            name: includeAnnotations
            schema:
              type: boolean
            description: Whether or not to include annotations on the image

        responses:
          200:
            description: measurement image file content for the provided defect id
            schema:
              id: MeasurementImageFileContent
              type: string
              format: binary
        """

        include_annotations_str = request.args.get('includeAnnotations')
        if include_annotations_str is not None:
            include_annotations_str = include_annotations_str.capitalize()
        logger.info(f'include_annotations_str: {include_annotations_str}')
        include_annotations = False
        if include_annotations_str == 'True':
            include_annotations = True

        # include_validated_annotations_str = request.args.get(
        #     'includeValidatedAnnotations')
        # if include_validated_annotations_str is not None:
        #     include_validated_annotations_str = include_validated_annotations_str.capitalize()
        # logger.info(f'include_validated_annotations_str: {include_validated_annotations_str}')
        # include_validated_annotations = False
        # if include_validated_annotations_str == 'True':
        #     include_validated_annotations = True

        logger.info(f'Reading image_file content for defect id# {id}')
        logger.info(f'include_annotations: {include_annotations}')
        # logger.info(f'include_validated_annotations: {include_validated_annotations}')

        with Session(db) as session:

            defect = session.query(Defect).get(id)
            if defect is None:
                return {"message": f"Defect record with id# {id} not found."}

            measurement_id = defect.measurement_id
            measurement_image_results = session.scalars(
                select(MeasurementImageFile).where(MeasurementImageFile.measurement_id == measurement_id)).all()

            if measurement_image_results is not None and len(measurement_image_results) > 0:
                image_file = measurement_image_results[0]

                file_content = read_file_record_content(image_file)
                if file_content is not None:

                    if include_annotations == True:
                        logging.info('including annotations...')
                        annotation_file_results = session.scalars(
                            select(DefectAnnotationFragment).where(DefectAnnotationFragment.defect_id == id)).all()
                        if annotation_file_results is not None and len(annotation_file_results) > 0:
                            annotation_file = annotation_file_results[0]
                            annotation_str = annotation_file.content
                            annotation_json = json.loads(annotation_str)

                            # overwrite the original content with the new one including annotations
                            file_content = draw_annotations_on_image(
                                file_content, annotation_json)

                    # if include_validated_annotations == True:
                    #     logging.info('including validated annotations...')
                    #     validated_annotation_file_results = session.scalars(
                    #         select(ValidatedMeasurementAnnotationFile).where(ValidatedMeasurementAnnotationFile.measurement_id == id)).all()
                    #     if validated_annotation_file_results is not None and len(validated_annotation_file_results) > 0:
                    #         validated_annotation_file = validated_annotation_file_results[0]
                    #         validated_annotation_str = validated_annotation_file.content
                    #         validated_annotation_json = json.loads(
                    #             validated_annotation_str)

                    #         # overwrite the original content with the new one including annotations
                    #         file_content = draw_annotations_on_image(
                    #             file_content, validated_annotation_json)

                    download_filename_ext = image_file.filename.split('.')
                    donwload_filename = download_filename_ext[0]+f'-did{defect.id}'+'.'+download_filename_ext[1]
                    return send_file_content_as_jpg(donwload_filename, file_content)
            else:
                logging.info(f'measurement_image_results for measurement_id {id} not found.')

        return {'message': f'Image File for Measurement id: {id} not found.'}


class DefectImageCommentsAPI(Resource):

    def __init__(self):
        super(DefectImageCommentsAPI, self).__init__()

    def get(self, id):
        """
        Read an existing measurement image file comments of a defect
        ---
        parameters:
          - in: path
            required: true
            name: id
            schema:
              type: number
            description: Existing defect id

        responses:
          200:
            description: measurement image file comments for the provided defect id
            schema:
              id: MeasurementImageFileComments
              type: string
              format: binary
        """
        try:
            with Session(db) as session:
                defect = session.query(Defect).get(id)
                if defect is None:
                    return {"message": f"Defect record with id# {id} not found."}

                measurement_image = session.query(MeasurementImageFile).filter_by(
                    measurement_id=defect.measurement_id).first()
                if measurement_image is None:
                    return {"message": f"Image File for Measurement id: {id} not found."}

                return measurement_image.comments

        except Exception as e:
            return {"message": f"Error occurred while getting measurement image file comments for defect id: {id}. Error: {e}"}

    def post(self, id):
        """
        Save measurement image file comments of a defect
        ---
        parameters:
          - in: path
            required: true
            name: id
            schema:
              type: number
            description: Existing defect id
          - in: body
            required: true
            name: comments
            schema:
              type: string
            description: Measurement image file comments for the provided defect id

        responses:
          200:
            description: Successfully saved measurement image file comments for the provided defect id
            schema:
              id: MeasurementImageFileComments
              type: string
              format: binary
        """
        try:
            with Session(db) as session:
                defect = session.query(Defect).get(id)
                if defect is None:
                    return {"message": f"Defect record with id# {id} not found."}

                measurement_image = session.query(MeasurementImageFile).filter_by(
                    measurement_id=defect.measurement_id).first()
                if measurement_image is None:
                    return {"message": f"Image File for Measurement id: {id} not found."}

                measurement_image.comments = request.json['comments']

                session.add(measurement_image)
                session.commit()

                return measurement_image.comments

        except Exception as e:
            return {"message": f"Error occurred while saving measurement image file comments for defect id: {id}. Error: {e}"}

class DefectRepairEvidenceCommentsAPI(Resource):

    def __init__(self):
        super(DefectRepairEvidenceCommentsAPI, self).__init__()

    def get(self, id):
        """
        Read an existing defect repair evidence comments of a defect
        ---
        parameters:
          - in: path
            required: true
            name: id
            schema:
              type: number
            description: Existing defect id

        responses:
          200:
            description: defect repair evidence comments for the provided defect id
            schema:
              id: Defect Repair Evidence image id
              type: string
              format: binary
        """
        try:
            with Session(db) as session:
                repair_evidence_file_list = session.scalars(
                    select(RepairEvidenceFile).where(RepairEvidenceFile.defect_id == id)
                ).all()

                if not repair_evidence_file_list:
                    return {"message": f"No repair evidence files found for defect id: {id}"}

            return repair_evidence_file_list[0].comments

        except Exception as e:
            return {"message": f"Unexpected error occurred while getting defect repair evidence comments for defect id: {id}. Error: {str(e)}"}


    def post(self, id):
        """
        Save defect repair evidence comments of a defect
        ---
        parameters:
          - in: path
            required: true
            name: id
            schema:
              type: number
            description: Existing defect id
          - in: body
            required: true
            name: comments
            schema:
              type: string
            description: defect repair evidence comments for the provided defect id

        responses:
          200:
            description: Successfully saved defect repair evidence comments for the provided defect id
            schema:
              type: string
              format: binary
        """
        try:
            with Session(db) as session:
                repair_evidence_files = session.scalars(
                    select(RepairEvidenceFile).where(RepairEvidenceFile.defect_id == id)
                ).all()

                if not repair_evidence_files:
                    return {"message": f"No repair evidence files found for defect id: {id}"}

                for repair_evidence_file in repair_evidence_files:
                    repair_evidence_file.comments = request.json['comments']

                session.commit()

                return {"message": "Comments successfully updated for all repair evidence files."}

        except Exception as e:
            return {"message": f"Unexpected error occurred while saving comments for defect id: {id}. Error: {str(e)}"}


class DefectImageThumbnailAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):
        super(DefectImageThumbnailAPI, self).__init__()

    def get(self, id):
        """
        Read an existing measurement image file thumbnail for a defect with or without annotations
        ---
        parameters:
          - in: path
            required: true
            name: id
            schema:
              type: number
            description: Existing defect id
          - in: query
            required: false
            name: includeAnnotations
            schema:
              type: boolean
            description: Whether or not to include annotations on the image

        responses:
          200:
            description: measurement image file content for the provided defect id
            schema:
              id: MeasurementImageThumbnailContent
              type: string
              format: binary
        """
        include_annotations_str = request.args.get('includeAnnotations')
        if include_annotations_str is not None:
            include_annotations_str = include_annotations_str.capitalize()
        logger.info(f'include_annotations_str: {include_annotations_str}')
        include_annotations = False
        if include_annotations_str == 'True':
            include_annotations = True

        logger.info(f'Reading image_file content for defect id# {id}')
        logger.info(f'include_annotations: {include_annotations}')
        with Session(db) as session:

            defect = session.query(Defect).get(id)
            if defect is None:
                return {"message": f"Defect record with id# {id} not found."}
            measurement_id = defect.measurement_id

            image_results = session.scalars(
                select(MeasurementImageFile).where(MeasurementImageFile.measurement_id == measurement_id)).all()
            if image_results is not None and len(image_results) > 0:
                image_file = image_results[0]
                thumbnail_content = image_file.thumbnail
                if thumbnail_content is not None:

                    if include_annotations == True:
                        logging.info('including annotations...')
                        annotation_file_results = session.scalars(
                            select(DefectAnnotationFragment).where(DefectAnnotationFragment.defect_id == id)).all()
                        if annotation_file_results is not None and len(annotation_file_results) > 0:
                            annotation_file = annotation_file_results[0]
                            annotation_str = annotation_file.content  # json annotaiton content
                            annotation_json = json.loads(annotation_str)

                            # overwrite the original content with the new one including annotations
                            thumbnail_content = draw_annotations_on_image(
                                thumbnail_content, annotation_json, 1, False)

                    download_filename_ext = image_file.filename.split('.')
                    donwload_filename = download_filename_ext[0]+f'-did{defect.id}_thumb'+'.'+download_filename_ext[1]
                    return send_file_content_as_jpg(donwload_filename, thumbnail_content)

        return {'message': f'Thumbnail for Defect id: {id} not found.'}


class ComputeMeasurementsForMeasurementAnnotationFileAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):
        super(ComputeMeasurementsForMeasurementAnnotationFileAPI, self).__init__()

    def get(self, id):
        """
        Recomputes the measuremetns for a measurement annotation and updates the DB
        ---
        parameters:
          - in: path
            required: true
            name: id
            schema:
              type: number
            description: Existing measurement id

        responses:
          200:
            description: updated annotation file content for the provided measurement id
            schema:
              id: MeasurementAnnotationFileContent
              type: string
              format: json
        """
        logger.info(f'Reading annotation_file content for measurement.id {id}')
        with Session(db) as session:

            measurement_rec = session.query(Measurement).get(id)
            if measurement_rec is None:
                return {"message": f"Measurement {id} not found."}
            logger.info(f'measurement found: {measurement_rec}')

            measurement_annotation_results = session.scalars(
                select(MeasurementAnnotationFile).where(MeasurementAnnotationFile.measurement_id == id)).all()
            logger.info(
                f'measurement_annotation_results: {measurement_annotation_results}')

            if measurement_annotation_results is None or len(measurement_annotation_results) == 0:
                return {'message': f'Nothing to do. Annotation File for Measurement id: {id} not found.'}

            measurement_annotation_file_rec = measurement_annotation_results[0]
            annotation_file_content = measurement_annotation_file_rec.content # string serialized json
            if annotation_file_content is None:
                return {'message': f'Nothing to do. Empty annotaiton for Measurement id: {id}.'}

            measurement_annotation_file_content_json = json.loads(annotation_file_content)
            updated_annotation_json_content_measurements = measure_annotation_file_and_update_json_and_measurement_rec(measurement_annotation_file_content_json, measurement_rec, session, True)
            
            str_annotation_json_content = json.dumps(updated_annotation_json_content_measurements)
            measurement_annotation_file_rec.content = str_annotation_json_content
            session.add(measurement_annotation_file_rec)
            session.commit()
            session.refresh(measurement_annotation_file_rec)

            # propagate changes to the defects
            defect_list = parse_annotation_file_defects(measurement_rec, measurement_annotation_file_rec)

            image = session.query(Image).get(measurement_rec.image_id)
            inspection_id = image.inspection_id
            invalidate_inspection_cache(inspection_id)

            return jsonify(updated_annotation_json_content_measurements)


class ParseDefectsMeasurementAnnotationFileAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):
        super(ParseDefectsMeasurementAnnotationFileAPI, self).__init__()

    def allowed_file(self, filename):
        return '.' in filename and \
            filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

    def get(self, id):
        """
        Parses an existing measurement annotation file into a list of new defect records. Deletes pre-existing records.
        ---
        parameters:
          - in: path
            required: true
            name: id
            schema:
              type: number
            description: Existing measurement id

        responses:
          200:
            description: annotation file content for the provided measurement id
            schema:
              id: MeasurementAnnotationFileContent
              type: string
              format: json
        """
        logger.info(f'Reading annotation_file content for measurement.id {id}')
        with Session(db) as session:

            measurement_list = session.scalars(
                select(Measurement).where(Measurement.id == id)).all()
            if measurement_list is None or len(measurement_list) == 0:
                return {"message": f"Measurement {id} not found."}
            measurement = measurement_list[0]
            logger.info(f'measurement found: {measurement}')

            measurement_annotation_results = session.scalars(
                select(MeasurementAnnotationFile).where(MeasurementAnnotationFile.measurement_id == id)).all()
            logger.info(
                f'measurement_annotation_results: {measurement_annotation_results}')

            if measurement_annotation_results is None or len(measurement_annotation_results) == 0:
                return {'message': f'Nothing to do. Annotation File for Measurement id: {id} not found.'}

            measurement_annotation_file = measurement_annotation_results[0]
            file_content = measurement_annotation_file.content
            if file_content is None:
                return {'message': f'Nothing to do. Emtpy annotaiton for Measurement id: {id}.'}

            # TODO: instead of deleting and re-creating we may need to update annotations accordinly,
            # using their unique id which is the combination of finding_type and measurement_id

            # delete existing defects and annotation fragments
            # delete_parsed_defects(measurement.id)

            # create new defects (or reuse existing) and update or create new annotation fragments
            defect_list = parse_annotation_file_defects(
                measurement, measurement_annotation_file)

            return jsonify(defect_list)


# parse the annotation file into a set of defect records. Reuse existing records if measurement_id and finding_type match
# delete stale defects. returned the parsed defect list of json objects.
def parse_annotation_file_defects(measurement_rec, measurement_annotation_file):
    logging.info(f'parse_annotation_file_defects() called')
    defect_list = []
    file_content = measurement_annotation_file.content
    json_content = json.loads(file_content)
    #logging.info(f'parse_annotation_file_defects() called with json: {json_content}')
    shapes = json_content['shapes']

    if shapes is None or len(shapes) == 0:
        logging.info('Annotation file has no shapes to parse')
        logging.info(f'returning defect_list: {defect_list}')
        return defect_list
    else:
        logging.info(f'found {len(shapes)} shapes in file')

    # maps a defect type to a list of shapes
    label_shapes_dict = {}
    for el in shapes:
        #logging.info(f'element: {el}')
        finding_label = 'Other'
        if 'label' in el:
            finding_label = el['label']  # finding_type
        logging.info(f'defect label: {finding_label}')
        key = finding_label.replace(' ', '_')
        if not key in label_shapes_dict:
            label_shapes_dict[key] = []
        label_shapes_dict[key].append(el)
    #logging.info(f'label_shapes_dict: {label_shapes_dict}')

    with Session(db) as session:

        # The user might have deleted annotations of a type, so we delete them in the DB
        curr_defect_list = session.scalars(
            select(Defect)
            .where(Defect.measurement_id == measurement_rec.id)
        ).all()
        for defect in curr_defect_list:
            #defect_key = defect.finding_type.replace(' ', '_')
            delete_defect_and_annotation_fragment(defect.id)
            # if not defect_key in label_shapes_dict:
            #     logging.info(f'deleting existing defect record for finding_type: {defect.finding_type}')
            #     delete_defect_and_annotation(defect.id)
            # else:
            #     logging.info(f'keeping existing defect record for finding_type: {defect.finding_type}')

        for key in label_shapes_dict:
            # The fragment is the annotation file with only one shape inside
            json_fragment = copy.deepcopy(json_content)
            if 'polygons' in json_fragment:
                del json_fragment['polygons']

            json_fragment['shapes'] = []
            json_fragment['shapes'] = label_shapes_dict[key]
            finding_label = 'Other'
            if 'label' in label_shapes_dict[key][0]:
                finding_label = label_shapes_dict[key][0]['label']


            logger.info(f'Creating a new defect')
            new_defect = Defect(
                image_id=measurement_rec.image_id,
                measurement_id=measurement_rec.id,

                date=measurement_rec.date,
                location=measurement_rec.location,

                root_face_distance=measurement_rec.root_face_distance,
                span_wise_length=measurement_rec.span_wise_length,
                chord_wise_width=measurement_rec.chord_wise_width,
                depth=measurement_rec.depth,
                height=measurement_rec.height,
                width=measurement_rec.width,
                length=measurement_rec.length,
                aspect_ratio=measurement_rec.aspect_ratio,
                area=measurement_rec.area,
                percent_area=measurement_rec.percent_area,

                finding_type=finding_label,
                ge_disposition=measurement_rec.ge_disposition,

                is_priority=measurement_rec.is_priority,
                description=measurement_rec.description,

                image_pitch=measurement_rec.image_pitch,
                image_yaw=measurement_rec.image_yaw,
                image_hfov=measurement_rec.image_hfov,

                design_tolerance=measurement_rec.design_tolerance,
                disposition_provided_by=measurement_rec.disposition_provided_by,
                status=measurement_rec.status,
                repair_date=measurement_rec.repair_date,
                repair_report_id=measurement_rec.repair_report_id,
                repair_approved_by=measurement_rec.repair_approved_by,
                is_manual=measurement_rec.is_manual,
                sso=measurement_rec.sso,
            )

            # apply the read measurements to the new_defect object
            update_defect_measurements(new_defect, json_fragment)
            
            #TODO: implement the auto close feature using tolerances and rules TBD

            # if is_defect_within_tolerance(new_defect):
            #     logging.info(f'Defect {new_defect.id} is within tolerance. It will be auto closed')
            #     new_defect.status = "Closed"
            #     new_defect.ge_disposition = "Within Tolerance - No Repair Needed"
            #     new_defect.sso = "AUTO_CLOSE"

            session.add(new_defect)
            session.commit()
            session.refresh(new_defect)  # to read the id
            logger.info(f'Created defect record id# {str(new_defect.id)}')

            new_defect_annotation_fragment = DefectAnnotationFragment(
                image_id=measurement_rec.image_id,
                measurement_id=measurement_rec.id,
                defect_id=new_defect.id,
                content=json.dumps(json_fragment)
            )
            session.add(new_defect_annotation_fragment)
            session.commit()
            # to read the id
            session.refresh(new_defect_annotation_fragment)
            logging.info(f'Created a new defect fratment id# {new_defect_annotation_fragment.id}')
            
            defect_list.append(Defect.serialize(new_defect))

            
    logging.info(f'returning json defect_list: {defect_list}')
    return defect_list # list of json defects updated by this function


def is_defect_within_tolerance(defect):
    if not defect.is_manual and defect.finding_type == 'CoreGap' and ( (defect.width > 0 and defect.width < 0.010) or defect.root_face_distance > 30):
        return True
    return False


# update the record but does not save it.
def update_defect_measurements(defect, json_fragment):
    logging.info(f'update_defect_measurements() called.')
    #logging.info(f'update_defect_measurements() called with json_fragment: {json_fragment}')
    if json_fragment is not None:
        if 'shapes' in json_fragment:
            shapes = json_fragment['shapes']
            total_area = 0.0
            max_width = 0.0
            max_length = 0.0
            for shape in shapes:
                
                if 'width_blob' in shape:
                    max_width = max(max_width, shape['width_blob'])

                if 'area' in shape and shape['area'] > 0:
                    total_area += shape['area']
                if 'width' in shape:
                    max_width = max(max_width, shape['width'])
                if 'length' in shape:
                    max_length = max(max_length, shape['length'])
            logging.info(f'update defect area to: {total_area}')
            defect.area = total_area
            logging.info(f'update defect length to: {max_length}')
            defect.length = max_length
            logging.info(f'update defect width to: {max_width}')
            defect.width = max_width


def delete_parsed_defects(measurement_id):
    logging.info(f'delete_parsed_defects() called for measurement id# {measurement_id}')
    with Session(db) as session:
        defect_list = session.scalars(
            select(Defect).where(
                Defect.measurement_id == measurement_id)
        ).all()
        for defect in defect_list:
            id = defect.id
            # Delete DefectAnnotationFragment if any
            defect_annotation_fragment_list = session.scalars(
                select(DefectAnnotationFragment).where(
                    DefectAnnotationFragment.defect_id == id)
            ).all()

            if defect_annotation_fragment_list is not None and len(defect_annotation_fragment_list) > 0:
                defect_annotation_fragment_rec = defect_annotation_fragment_list[0]
                session.delete(defect_annotation_fragment_rec)
                session.commit()
            else:
                logging.info("No defect annotation fragment found.")

            # Then finally delete Defect (it can only be deelted if there is no back references via defect_id)
            session.delete(defect)
            session.commit()
            logging.info(f'Successfully deleted defect id# {defect.id}')


# defect records have companion defect_annotation_fragment records
def delete_defect_and_annotation_fragment(defect_id):
    logging.info(f'delete_defect_and_annotation() called for defect_id: {defect_id}')
    with Session(db) as session:
        defect_list = session.scalars(
            select(Defect).where(
                Defect.id == defect_id)
        ).all()
        for defect in defect_list:
            id = defect.id
            # Delete DefectAnnotationFragment if any
            defect_annotation_fragment_list = session.scalars(
                select(DefectAnnotationFragment).where(
                    DefectAnnotationFragment.defect_id == id)
            ).all()

            if defect_annotation_fragment_list is not None and len(defect_annotation_fragment_list) > 0:
                defect_annotation_fragment_rec = defect_annotation_fragment_list[0]
                session.delete(defect_annotation_fragment_rec)
                session.commit()
            else:
                logging.info("No defect annotation fragment found.")

            # Then finally delete Defect (it can only be deelted if there is no back references via defect_id)
            session.delete(defect)
            session.commit()
            logging.info(f'Successfully deleted defect id# {defect.id}')


class MeasurementDefectListAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):
        super(MeasurementDefectListAPI, self).__init__()

    def get(self, id):
        """
        Read the list of measurement defects for a provided measurement id
        ---
        parameters:
          - in: path
            required: false
            name: id
            schema:
              type: number
            description: Measurement id

        responses:
          200:
            description: List of measuremtn defect records parsed out of the measurement annotation file
            schema:
              id: DefectList
              type: array
              items:
                schema:
                  id: DefectObject

        """
        logging.info('MeasurementDefectListAPI')
        with Session(db) as session:
            measurement = session.query(Measurement).get(id)
            if measurement is None:
                return {"message": f"Measurement {id} not found."}

            defect_list = session.scalars(
                select(Defect).where(Defect.measurement_id == id)
            ).all()

            resp_list = Defect.serialize_list(defect_list)

            # logging.info(f'resp_list: {resp_list}')
            return jsonify(resp_list)

# --------------------------------------------- Defect -----------------------------------------------


class DefectAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):
        super(DefectAPI, self).__init__()

    def delete(self, id):
        """
        Deletes an existing defect record with its companion defect_annotation_fragment
        ---
        parameters:
          - in: path
            required: true
            name: id
            schema:
              type: number
            description: Existing defect id

        responses:
          200:
            description: Success or failure message
            schema:
              properties:
                message:
                  type: string

        """
        with Session(db) as session:

            defect = session.query(Defect).get(id)
            if defect is None:
                return {'message': f'Defect record id #{id} not found.'}

            image = session.query(Image).get(defect.image_id)
            inspection_id = image.inspection_id

            # Delete DefectAnnotationFragment if any
            defect_annotation_fragment_list = session.scalars(
                select(DefectAnnotationFragment).where(
                    DefectAnnotationFragment.defect_id == id)
            ).all()

            if defect_annotation_fragment_list is not None and len(defect_annotation_fragment_list) > 0:
                defect_annotation_fragment_rec = defect_annotation_fragment_list[0]
                session.delete(defect_annotation_fragment_rec)
            else:
                logging.info("No defect annotation fragment found.")

            # Then finally delete Defect (it can only be deelted if there is no back references via defect_id)
            session.delete(defect)
            session.commit()
            logging.info(f'Successfully deleted defect id# {defect.id}')

            invalidate_inspection_cache(inspection_id)

            return {'message': f'Successfully deleted defect id# {id}'}

    def get(self, id):
        """
        Read an existing defect meta-data
        ---
        parameters:
          - in: path
            required: true
            name: id
            schema:
              type: number
            description: Existing defect id

        responses:
          200:
            description: Existing defect metadata
            schema:
              id: DefectObject

        """

        logging.info(f'Look for defect record id {id}')
        with Session(db) as session:
            defect = session.query(Defect).get(id)
            if defect is None:
                return {'message': f'Defect {id} not found.'}
            jsonResp = defect.toJson()
            return jsonResp


class CreateDefectAPI(Resource):

    def __init__(self):
        self.reqparse = reqparse.RequestParser()

        self.reqparse.add_argument('image_id', type=int, location='json')
        self.reqparse.add_argument('measurement_id', type=int, location='json')

        self.reqparse.add_argument('date', type=str, location='json')
        self.reqparse.add_argument('location', type=str, location='json')

        self.reqparse.add_argument(
            'root_face_distance', type=float, location='json')
        self.reqparse.add_argument(
            'span_wise_length', type=float, location='json')
        self.reqparse.add_argument(
            'chord_wise_width', type=float, location='json')
        self.reqparse.add_argument('depth', type=float, location='json')
        self.reqparse.add_argument('height', type=float, location='json')
        self.reqparse.add_argument('width', type=float, location='json')
        self.reqparse.add_argument('length', type=float, location='json')
        self.reqparse.add_argument('aspect_ratio', type=float, location='json')
        self.reqparse.add_argument('area', type=float, location='json')
        self.reqparse.add_argument('percent_area', type=float, location='json')

        self.reqparse.add_argument('finding_type', type=str, location='json')
        self.reqparse.add_argument('ge_disposition', type=str, location='json')

        self.reqparse.add_argument('is_priority', type=bool, location='json')
        self.reqparse.add_argument('description', type=str, location='json')

        self.reqparse.add_argument('image_pitch', type=float, location='json')
        self.reqparse.add_argument('image_yaw', type=float, location='json')
        self.reqparse.add_argument('image_hfov', type=float, location='json')

        self.reqparse.add_argument(
            'design_tolerance', type=str, location='json')
        self.reqparse.add_argument(
            'disposition_provided_by', type=str, location='json')
        self.reqparse.add_argument('status', type=str, location='json')
        self.reqparse.add_argument('repair_date', type=str, location='json')
        self.reqparse.add_argument(
            'repair_report_id', type=str, location='json')
        self.reqparse.add_argument(
            'repair_approved_by', type=str, location='json')
        self.reqparse.add_argument('is_manual', type=bool, location='json')
        self.reqparse.add_argument(
            'sso', type=str, location='json')

        super(CreateDefectAPI, self).__init__()

    def post(self):
        """
        Create a new defect record represneting a measurement shape meta-data
        ---
        parameters:
          - name: CreateDefectRequest
            in: body
            required: true
            schema:
              id: CreateDefectBody
              properties:
                image_id:
                  type: number
                  example: 1
                measurement_id:
                  type: number
                  example: 1
                date:
                  type: string
                  example: "2023-11-01"
                location:
                  type: string
                root_face_distance:
                  type: number
                span_wise_length:
                  type: number
                chord_wise_width:
                  type: number
                depth:
                  type: number
                height:
                  type: number
                width:
                  type: number
                length:
                  type: number
                aspect_ratio:
                  type: number
                area:
                  type: number
                percent_area:
                  type: number
                finding_type:
                  type: string
                ge_disposition:
                  type: string
                is_priority:
                  type: boolean
                description:
                  type: string
                image_pitch:
                  type: number
                image_yaw:
                  type: number
                image_hfov:
                  type: number
                design_tolerance:
                  type: string
                disposition_provided_by:
                  type: string
                status:
                  type: string
                repair_date:
                  type: string
                  example: "2023-11-01"
                repair_report_id:
                  type: string
                repair_approved_by:
                  type: string
                is_manual:
                  type: boolean
                sso:
                  type: string
        responses:
          200:
            description: Created defect
            schema:
              id: DefectObject
              properties:
                id:
                  type: number

                image_id:
                  type: number
                measurement_id:
                  type: number

                date:
                  type: string
                location:
                  type: string

                root_face_distance:
                  type: number
                span_wise_length:
                  type: number
                chord_wise_width:
                  type: number
                depth:
                  type: number
                height:
                  type: number
                width:
                  type: number
                length:
                  type: number
                aspect_ratio:
                  type: number
                area:
                  type: number
                percent_area:
                  type: number

                finding_type:
                  type: string

                ge_disposition:
                  type: string

                is_priority:
                  type: boolean
                description:
                  type: string

                image_pitch:
                  type: number
                image_yaw:
                  type: number
                image_hfov:
                  type: number

                design_tolerance:
                  type: string
                disposition_provided_by:
                  type: string
                status:
                  type: string
                repair_date:
                  type: string
                repair_report_id:
                  type: string
                repair_approved_by:
                  type: string
                is_manual:
                  type: boolean
                sso:
                  type: boolean


        """
        args = self.reqparse.parse_args()
        logger.info(args.items())

        if args['date'] == '':
            # current_date = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            current_date = datetime.now().isoformat()
            args['date'] = current_date

        image_id = args['image_id']
        inspection_id = -1
        with Session(db) as session:
            image = session.query(Image).get(image_id)
            if image is None:
                return {'message': f'Could not create defect. Image record for image_id: {image_id} not found.'}
            inspection_id = image.inspection_id # save to invalidate cache later

        measurement_id = args['measurement_id']
        with Session(db) as session:
            measurement = session.query(Measurement).get(measurement_id)
            if measurement is None:
                return {'message': f'Could not create defect. Measurement record for measurement_id: {measurement_id} not found.'}

        defect = Defect(
            image_id=args['image_id'],
            measurement_id=args['measurement_id'],

            date=args['date'],
            location=args['location'],

            root_face_distance=args['root_face_distance'],
            span_wise_length=args['span_wise_length'],
            chord_wise_width=args['chord_wise_width'],
            depth=args['depth'],
            height=args['height'],
            width=args['width'],
            length=args['length'],
            aspect_ratio=args['aspect_ratio'],
            area=args['area'],
            percent_area=args['percent_area'],

            finding_type=args['finding_type'],
            ge_disposition=args['ge_disposition'],

            is_priority=args['is_priority'],
            description=args['description'],

            image_pitch=args['image_pitch'],
            image_yaw=args['image_yaw'],
            image_hfov=args['image_hfov'],

            design_tolerance=args['design_tolerance'],
            disposition_provided_by=args['disposition_provided_by'],
            status=args['status'],
            repair_date=args['repair_date'],
            repair_report_id=args['repair_report_id'],
            repair_approved_by=args['repair_approved_by'],
            is_manual=args['is_manual'],
            sso=args['sso'],
        )

        with Session(db) as session:
            session.add(defect)
            session.commit()
            session.refresh(defect)  # to read the id

            logger.info(f'Created defect: {str(defect)}')

            invalidate_inspection_cache(inspection_id)

        # resp = jsonify(inspection.serialize())
        resp = defect.toJson()
        logging.info(f'resp: {resp}')
        return resp


class SearchImageAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):
        super(SearchImageAPI, self).__init__()

    def get(self):
        """
        Read an existing image list meta-data according to filter parameters
        ---
        parameters:
          - in: query
            required: false
            name: esn
            schema:
              type: string
          - in: query
            required: false
            name: min_distance
            schema:
              type: number
          - in: query
            required: false
            name: max_distance
            schema:
              type: number

        responses:
          200:
            description: Existing image records for query
            schema:
              id: ImageSearchResponse
              type: array
              items:
                type: ImageObject

        """

        # Inspection filter props
        esn = request.args.get('esn', '%')
        min_distance = request.args.get('min_distance', 0)
        max_distance = request.args.get('max_distance', 200)

        logging.info(f'esn: {esn}')
        logging.info(f'min_distance: {min_distance}')
        logging.info(f'max_distance: {max_distance}')

        cache_key = 'SearchImageAPI.get_'+esn+'_'+str(min_distance)+'_'+str(max_distance)
        if cache_key in cache:
            logging.info(f'Cache hit: {cache_key}')
            resp_list = cache[cache_key]
            return jsonify(resp_list)
        else:
            logging.info('Cache miss')

        logging.info(f'Search for image records...')
       
        with Session(db) as session:
            image_list = []
            if esn != '%':
                image_list = session.query(
                    Inspection.id.label('inspection_id'),
                    Inspection.esn.label('inspection_esn'),
                    Inspection.sect.label('inspection_sect'),
                    Inspection.blade_type.label('inspection_blade_type'),
                    Inspection.upload_date.label('inspection_upload_date'),
                    Image.id.label('image_id'),
                    Image.distance.label('distance')
                ).join(Image, Image.inspection_id == Inspection.id,  isouter=False
                ).filter(and_(Inspection.esn.like(esn),Image.distance >= min_distance, Image.distance <= max_distance)
                ).distinct('image_id')
            else:
                image_list = session.query(
                    Inspection.id.label('inspection_id'),
                    Inspection.esn.label('inspection_esn'),
                    Inspection.sect.label('inspection_sect'),
                    Inspection.blade_type.label('inspection_blade_type'),
                    Inspection.upload_date.label('inspection_upload_date'),
                    Image.id.label('image_id'),
                    Image.distance.label('distance')
                ).join(Image, Image.inspection_id == Inspection.id,  isouter=False
                ).filter(and_(Image.distance >= min_distance, Image.distance <= max_distance))

            resp_list = []
            for el in image_list:
                resp_list.append({
                    'inspection_id':el[0],
                    'esn':el[1],
                    'sect':el[2],
                    'blade_type':el[3],
                    'upload_date':el[4],
                    'id':el[5],
                    'distance':el[6]
                })                               
                  
            logging.info(f'returning {len(resp_list)} images.')
            
            # save resutls for 1/2 hour
            cache.set(cache_key, resp_list, expire=CACHE_TTL)
            return jsonify(resp_list)


class SearchDefectAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):
        super(SearchDefectAPI, self).__init__()

    def get(self):
        """
        Read an existing defect list meta-data according to filter parameters
        ---
        parameters:
          - in: query
            required: false
            name: esn
            schema:
              type: string

        responses:
          200:
            description: Existing defect metadata and content
            schema:
              id: DefectSearchResponse
              type: array
              items:
                type: DefectObject

        """

        # Inspection filter props
        esn = request.args.get('esn', '%')

        logging.info(f'esn: {esn}')

        cache_key = 'SearchDefectAPI.get_'+esn
        if cache_key in cache:
            logging.info(f'Cache hit: {cache_key}')
            resp_list = cache[cache_key]
            return jsonify(resp_list)
        else:
            logging.info('Cache miss')

        logging.info(f'Search for defect inspection records...')
        with Session(db) as session:
            image_list = []
            if esn != '%':
                image_list = session.query(
                    Inspection.id.label('inspection_id'),
                    Inspection.esn.label('inspection_esn'),
                    Inspection.sect.label('inspection_sect'),
                    Inspection.blade_type.label('inspection_blade_type'),
                    Inspection.upload_date.label('inspection_upload_date'),
                    Image.id.label('image_id'),
                ).join(Image, Image.inspection_id == Inspection.id,  isouter=False
                    ).filter(Inspection.esn.like(esn)
                                ).distinct('image_id')
            else:
                image_list = session.query(
                    Inspection.id.label('inspection_id'),
                    Inspection.esn.label('inspection_esn'),
                    Inspection.sect.label('inspection_sect'),
                    Inspection.blade_type.label('inspection_blade_type'),
                    Inspection.upload_date.label('inspection_upload_date'),
                    Image.id.label('image_id'),
                ).join(Image, Image.inspection_id == Inspection.id,  isouter=False)

            image_inspection_map = {}
            image_id_set = set()

            for el in image_list:
                image_id = el[5]
                image_inspection_map[str(image_id)] = {
                    "id":el[0],
                    "esn":el[1],
                    "sect":el[2],
                    "blade_type":el[3],
                    "upload_date":el[4],
                }                               
                    
                image_id_set.add(image_id)

            image_id_list = list(image_id_set)
            logging.info(f'image_id_list length: {len(image_id_list)} for esn: {esn}')
            #logging.info(f'image_inspection_map: {image_inspection_map}')

            # defect_list = session.query(
            #     Defect, 
            # ).filter(Defect.image_id.in_(image_id_list)).order_by(Defect.root_face_distance, Defect.id).all()
            # resp_list = Defect.serialize_list(defect_list)

            logging.info(f'Search for defect records...')
            defect_list = []
            if esn != '%':
                defect_list = session.query(
                    Defect
                ).join(Image, Defect.image_id == Image.id,  isouter=False
                ).join(Inspection, Image.inspection_id == Inspection.id,  isouter=False            
                ).filter(Inspection.esn.like(esn)).order_by(Defect.root_face_distance, Defect.id).all()
            else:
                defect_list = session.query(
                    Defect
                ).order_by(Defect.root_face_distance, Defect.id).all()

            resp_list = Defect.serialize_list(defect_list)

            # function to push inspection data after json serialization
            def add_inspection_data(defect):
                image_id = defect['image_id']
                inspection = image_inspection_map[str(image_id)]
                #logging.info(f'inspection: {inspection}')

                defect["inspection_id"] = inspection['id']
                defect["inspection_esn"] = inspection['esn']
                defect["inspection_sect"] = inspection['sect']
                defect["inspection_blade_type"] = inspection['blade_type']
                defect["inspection_upload_date"] = inspection['upload_date']

                #logging.info(f'inspection for {image_id} is: {inspection}')

            # logging.info('executing in parallel...')
            # futures = []
            # with ThreadPoolExecutor(THREAD_POOL_SIZE) as executor:
            #     for resp in resp_list:
            #         futures.append(executor.submit(add_inspection_data, resp))
            #     wait(futures)

            logging.info('Executing in series...')
            for resp in resp_list:
                add_inspection_data(resp)     


            # defect_list = session.query(
            #     Defect,
            #     Inspection.id.label('inspection_id'),
            #     Inspection.esn.label('inspection_esn'),
            #     Inspection.sect.label('inspection_sect'),
            #     Inspection.blade_type.label('inspection_blade_type'),
            #     Inspection.upload_date.label('inspection_upload_date')
            # ).filter(Defect.image_id.in_(image_id_list)).order_by(Defect.root_face_distance, Defect.id).all()

            # resp_list = []
            # for defect in defect_list:
            #     obj = {}
            #     obj["area"] = defect.area
            #     obj["aspect_ratio"] = defect.aspect_ration
            #     obj["chord_wise_width"] = defect.chord_wise_width,
            #     obj["date"] = defect.date
            #     obj["depth"] = defect.depth
            #     obj["description"]=defect.description
            #     obj["design_tolerance"] = defect.design_tolerance
            #     obj["disposition_provided_by"] = defect.disposition_provided_by
            #     obj["finding_type"] = defect.finding_type
            #     obj["ge_disposition"]= defect.ge_disposition
            #     obj["height"] = defect.height
            #     obj["id"] = defect.id
            #     obj["image_hfov"] = defect.image_hov
            #     obj["image_id"] = defect.image_id
            #     obj["image_pitch"] = defect.image_pitch
            #     obj["image_yaw"] = defect.image_yaw
            #     obj["is_manual"] = defect.is_manual
            #     obj["is_priority"] = defect.is_priority
            #     obj["length"] = defect.length
            #     obj["location"] = defect.location
            #     obj["measurement_id"] = defect.measurement_id
            #     obj["percent_area"] = defect.percent_area
            #     obj["repair_approved_by"] = defect.approved_by
            #     obj["repair_date"] = defect.repair_date
            #     obj["repair_report_id"] = defect.repair_report_id
            #     obj["root_face_distance"] = defect.root_face_distance,
            #     obj["span_wise_length"] = defect.span_wise_legth
            #     obj["sso"] = defect.sso
            #     obj["status"] = defect.status
            #     obj["width"] = defect.width

            #     obj['inspection_id'] = defect.inspection_id
            #     obj['inspection_esn'] = defect.inspection_esn
            #     obj['inspection_sect'] = defect.inspection_sect
            #     obj['inspection_blade_type'] = defect.inspection_blade_type
            #     obj['inspection_upload_date'] = defect.inspection_upload_date

            #     resp_list.append(obj)
            

            logging.info(f'returning {len(resp_list)} defects.')
            
            # save resutls for 1/2 hour
            cache.set(cache_key, resp_list, expire=CACHE_TTL)
            return jsonify(resp_list)


class SearchDefectCsvAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):
        super(SearchDefectCsvAPI, self).__init__()

    def get(self):
        """
        Read a csv file with existing defect list meta-data according to filter parameters
        ---
        parameters:
          - in: query
            required: false
            name: esn
            schema:
              type: string

        responses:
          200:
            description: defect list csv file
            schema:
              id: DefectListCSVFile
              type: string
              format: binary

        """

        # Inspection filter props
        esn = request.args.get('esn', '%')

        logging.info(f'esn: {esn}')
        resp_list = []

        cache_key = 'SearchDefectCsvAPI.get_'+esn
        if cache_key in cache:
            logging.info(f'Cache hit: {cache_key}')
            resp_list = cache[cache_key]

        else:
            logging.info('Cache miss')

        if len(resp_list) == 0:
            logging.info(f'Search for defect inspection records...')
            with Session(db) as session:
                image_list = []
                if esn != '%':
                    image_list = session.query(
                        Inspection.id.label('inspection_id'),
                        Inspection.esn.label('inspection_esn'),
                        Inspection.sect.label('inspection_sect'),
                        Inspection.blade_type.label('inspection_blade_type'),
                        Inspection.upload_date.label('inspection_upload_date'),
                        Image.id.label('image_id'),
                    ).join(Image, Image.inspection_id == Inspection.id,  isouter=False
                        ).filter(Inspection.esn.like(esn)
                                    ).distinct('image_id')
                else:
                    image_list = session.query(
                        Inspection.id.label('inspection_id'),
                        Inspection.esn.label('inspection_esn'),
                        Inspection.sect.label('inspection_sect'),
                        Inspection.blade_type.label('inspection_blade_type'),
                        Inspection.upload_date.label('inspection_upload_date'),
                        Image.id.label('image_id'),
                    ).join(Image, Image.inspection_id == Inspection.id,  isouter=False)

                image_inspection_map = {}
                image_id_set = set()

                for el in image_list:
                    image_id = el[5]
                    image_inspection_map[str(image_id)] = {
                        "id":el[0],
                        "esn":el[1],
                        "sect":el[2],
                        "blade_type":el[3],
                        "upload_date":el[4],
                    }                               
                        
                    image_id_set.add(image_id)

                image_id_list = list(image_id_set)
                logging.info(f'image_id_list length: {len(image_id_list)} for esn: {esn}')
                #logging.info(f'image_inspection_map: {image_inspection_map}')

                logging.info(f'Search for defect records...')
                defect_list = []
                if esn != '%':
                    defect_list = session.query(
                        Defect
                    ).join(Image, Defect.image_id == Image.id,  isouter=False
                    ).join(Inspection, Image.inspection_id == Inspection.id,  isouter=False            
                    ).filter(Inspection.esn.like(esn)).order_by(Defect.root_face_distance, Defect.id).all()
                else:
                    defect_list = session.query(
                        Defect
                    ).order_by(Defect.root_face_distance, Defect.id).all()

                resp_list = Defect.serialize_list(defect_list)

                # function to push inspection data after json serialization
                def add_inspection_data(defect):
                    image_id = defect['image_id']
                    inspection = image_inspection_map[str(image_id)]
                    #logging.info(f'inspection: {inspection}')

                    defect["inspection_id"] = inspection['id']
                    defect["inspection_esn"] = inspection['esn']
                    defect["inspection_sect"] = inspection['sect']
                    defect["inspection_blade_type"] = inspection['blade_type']
                    defect["inspection_upload_date"] = inspection['upload_date']

                logging.info('Executing in series...')
                for resp in resp_list:
                    add_inspection_data(resp)     

                logging.info(f'returning {len(resp_list)} defects.')
                
                # save resutls for 1/2 hour
                cache.set(cache_key, resp_list, expire=CACHE_TTL)


        with tempfile.TemporaryDirectory(dir=UPLOAD_FOLDER, prefix='temp-', suffix='-folder') as tempdirname:
            output_filename = 'defect_list.csv'
            csv_path = os.path.join(tempdirname, output_filename)
            
            with open(csv_path, 'w') as csvfile:
                writer = csv.writer(csvfile)

                # Write the header row (if applicable)
                if isinstance(resp_list, list) and len(resp_list) > 0:
                    writer.writerow(resp_list[0].keys())
                
                # Write the data rows
                for row in resp_list:
                    writer.writerow(row.values())

            return send_file(os.path.abspath(csv_path),
                            as_attachment=True,
                            download_name="defect_list.csv")


        return jsonify(resp_list)


class UpdateDefectAPI(Resource):

    def __init__(self):
        self.reqparse = reqparse.RequestParser()

        # self.reqparse.add_argument('image_id', type=int, location='json')

        self.reqparse.add_argument('date', type=str, location='json')

        self.reqparse.add_argument('location', type=str, location='json')

        self.reqparse.add_argument(
            'root_face_distance', type=float, location='json')

        self.reqparse.add_argument(
            'span_wise_length', type=float, location='json')
        self.reqparse.add_argument(
            'chord_wise_width', type=float, location='json')
        self.reqparse.add_argument('depth', type=float, location='json')
        self.reqparse.add_argument('height', type=float, location='json')
        self.reqparse.add_argument('width', type=float, location='json')
        self.reqparse.add_argument('length', type=float, location='json')
        self.reqparse.add_argument('aspect_ratio', type=float, location='json')
        self.reqparse.add_argument('area', type=float, location='json')
        self.reqparse.add_argument('percent_area', type=float, location='json')

        self.reqparse.add_argument('finding_type', type=str, location='json')

        self.reqparse.add_argument(
            'finding_reference', type=str, location='json')
        self.reqparse.add_argument('ge_disposition', type=str, location='json')

        self.reqparse.add_argument('is_priority', type=bool, location='json')
        self.reqparse.add_argument('description', type=str, location='json')

        self.reqparse.add_argument('image_pitch', type=float, location='json')
        self.reqparse.add_argument('image_yaw', type=float, location='json')
        self.reqparse.add_argument('image_hfov', type=float, location='json')

        self.reqparse.add_argument(
            'design_tolerance', type=str, location='json')
        self.reqparse.add_argument(
            'disposition_provided_by', type=str, location='json')
        self.reqparse.add_argument('status', type=str, location='json')
        self.reqparse.add_argument('repair_date', type=str, location='json')
        self.reqparse.add_argument(
            'repair_report_id', type=str, location='json')
        self.reqparse.add_argument(
            'repair_approved_by', type=str, location='json')
        self.reqparse.add_argument('is_manual', type=bool, location='json')
        self.reqparse.add_argument('sso', type=str, location='json')

        super(UpdateDefectAPI, self).__init__()

    def post(self, id):
        """
        Update an existing defect record meta-data
        ---
        parameters:
          - in: path
            required: true
            name: id
            schema:
              type: number
            description: Existing defect id
          - name: UpdateDefectRequest
            in: body
            required: true
            schema:
              id: CreateDefectBody

        responses:
          200:
            description: Updated defect record
            schema:
              id: DefectObject

        """
        args = self.reqparse.parse_args()
        logger.info(f'body: {args.items()}')

        # TODO: validate date instead
        if args['date'] is None or args['date'] == '':
            # current_date = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            current_date = datetime.now().isoformat()
            args['date'] = current_date

        with Session(db) as session:
            defect = session.query(Defect).get(id)
            if defect is None:
                return {'message': f'Could not find defect id#: {id}.'}

            # Propagate changes to the measuremnet if the measurement has only one defect type
            measurement_id = defect.measurement_id
            defect_list = session.query(Defect).filter(Defect.measurement_id == measurement_id).all()
            measurement = None
            if len(defect_list) == 1:
                measurement = session.query(Measurement).get(measurement_id)

            if args['date'] is not None:
                defect.date = args['date']
                if measurement is not None: measurement.date = args['date']

            if args['location'] is not None:
                defect.location = args['location']
                if measurement is not None: measurement.location = args['location']

            if args['root_face_distance'] is not None:
                defect.root_face_distance = args['root_face_distance']
                if measurement is not None: measurement.root_face_distance = args['root_face_distance']

            if args['span_wise_length'] is not None:
                defect.span_wise_length = args['span_wise_length']
                if measurement is not None: measurement.span_wise_length = args['span_wise_length']

            if args['chord_wise_width'] is not None:
                defect.chord_wise_width = args['chord_wise_width']
                if measurement is not None: measurement.chord_wise_width = args['chord_wise_width']

            if args['depth'] is not None:
                defect.depth = args['depth']
                if measurement is not None: measurement.depth = args['depth']

            if args['height'] is not None:
                defect.height = args['height']
                if measurement is not None: measurement.height = args['height']

            if args['width'] is not None:
                defect.width = args['width']
                if measurement is not None: measurement.width = args['width']

            if args['length'] is not None:
                defect.length = args['length']
                if measurement is not None: measurement.length = args['length']

            if args['aspect_ratio'] is not None:
                defect.aspect_ratio = args['aspect_ratio']
                if measurement is not None: measurement.aspect_ratio = args['aspect_ratio']

            if args['area'] is not None:
                defect.area = args['area']
                if measurement is not None: measurement.area = args['area']

            if args['percent_area'] is not None:
                defect.percent_area = args['percent_area']
                if measurement is not None: measurement.percent_area = args['percent_area']

            if args['finding_type'] is not None:
                defect.finding_type = args['finding_type']
                if measurement is not None: measurement.finding_type = args['finding_type']

            if args['ge_disposition'] is not None:
                defect.ge_disposition = args['ge_disposition']
                if measurement is not None: measurement.ge_disposition = args['ge_disposition']

            if args['is_priority'] is not None:
                defect.is_priority = args['is_priority']
                if measurement is not None: measurement.is_priority = args['is_priority']

            if args['description'] is not None:
                defect.description = args['description']
                if measurement is not None: measurement.description = args['description']

            if args['image_pitch'] is not None:
                defect.image_pitch = args['image_pitch']
                if measurement is not None: measurement.image_pitch = args['image_pitch']
            if args['image_yaw'] is not None:
                defect.image_yaw = args['image_yaw']
                if measurement is not None: measurement.image_yaw = args['image_yaw']
            if args['image_hfov'] is not None:
                defect.image_hfov = args['image_hfov']
                if measurement is not None: measurement.image_hfov = args['image_hfov']

            if args['design_tolerance'] is not None:
                defect.design_tolerance = args['design_tolerance']
                if measurement is not None: measurement.design_tolerance = args['design_tolerance']
            if args['disposition_provided_by'] is not None:
                defect.disposition_provided_by = args['disposition_provided_by']
                if measurement is not None: measurement.disposition_provided_by = args['disposition_provided_by']
            if args['status'] is not None:
                defect.status = args['status']
                if measurement is not None: measurement.status = args['status']
            if args['repair_date'] is not None:
                # if one explicitly sets the date to empty, we remove the date
                if args['repair_date'].strip() == '':
                    logging.info('Setting repair_date to None')
                    defect.repair_date = None
                    if measurement is not None: measurement.repair_date = None
                else:
                    defect.repair_date = args['repair_date']
                    if measurement is not None: measurement.repair_date = args['repair_date']
                    
            if args['repair_report_id'] is not None:
                defect.repair_report_id = args['repair_report_id']
                if measurement is not None: measurement.repair_report_id = args['repair_report_id']
            if args['repair_approved_by'] is not None:
                defect.repair_approved_by = args['repair_approved_by']
                if measurement is not None: measurement.repair_approved_by = args['repair_approved_by']
            if args['is_manual'] is not None:
                defect.is_manual = args['is_manual']
                if measurement is not None: measurement.is_manual = args['is_manual']
            if args['sso'] is not None:
                defect.sso = args['sso']
                if measurement is not None: measurement.sso = args['sso']

            # logger.info(f'Measurement before update: {str(measurement)}')
            if measurement is not None:
                logging.info(f'Will also update measurement id# {measurement_id}') 
                session.add(measurement)
            session.add(defect)
            session.commit()
            session.refresh(defect)

            logger.info(f'Updated defect: {str(defect)}')

            image = session.query(Image).get(defect.image_id)
            inspection_id = image.inspection_id
            invalidate_inspection_cache(inspection_id)

            if measurement is not None:
                logger.info(f'Also updated measurement: {str(measurement)}')

        # resp = jsonify(inspection.serialize())
        resp = defect.toJson()
        logging.info(f'resp: {resp}')
        return resp


def invalidate_inspection_cache(inspection_id):
    if USE_INSPECTION_CACHE:
        cache_defect_list_key = 'InspectionDefectListAPI.get_'+str(inspection_id)+'_True'
        if cache_defect_list_key in cache:
          logging.info(f'delete cache entry: {cache_defect_list_key}')  
          cache.delete(cache_defect_list_key)
        cache_defect_list_key = 'InspectionDefectListAPI.get_'+str(inspection_id)+'_False'
        if cache_defect_list_key in cache:
          logging.info(f'delete cache entry: {cache_defect_list_key}')  
          cache.delete(cache_defect_list_key)


class UpdateDefectListAPI(Resource):

    def __init__(self):
        self.reqparse = reqparse.RequestParser()

        # self.reqparse.add_argument('image_id', type=int, location='json')

        self.reqparse.add_argument('id_list', type=int, action='append', help='List of defect IDs', location='json')
        self.reqparse.add_argument('date', type=str, location='json')

        self.reqparse.add_argument('location', type=str, location='json')

        self.reqparse.add_argument(
            'root_face_distance', type=float, location='json')

        self.reqparse.add_argument(
            'span_wise_length', type=float, location='json')
        self.reqparse.add_argument(
            'chord_wise_width', type=float, location='json')
        self.reqparse.add_argument('depth', type=float, location='json')
        self.reqparse.add_argument('height', type=float, location='json')
        self.reqparse.add_argument('width', type=float, location='json')
        self.reqparse.add_argument('length', type=float, location='json')
        self.reqparse.add_argument('aspect_ratio', type=float, location='json')
        self.reqparse.add_argument('area', type=float, location='json')
        self.reqparse.add_argument('percent_area', type=float, location='json')

        self.reqparse.add_argument('finding_type', type=str, location='json')

        self.reqparse.add_argument(
            'finding_reference', type=str, location='json')
        self.reqparse.add_argument('ge_disposition', type=str, location='json')

        self.reqparse.add_argument('is_priority', type=bool, location='json')
        self.reqparse.add_argument('description', type=str, location='json')

        self.reqparse.add_argument('image_pitch', type=float, location='json')
        self.reqparse.add_argument('image_yaw', type=float, location='json')
        self.reqparse.add_argument('image_hfov', type=float, location='json')

        self.reqparse.add_argument(
            'design_tolerance', type=str, location='json')
        self.reqparse.add_argument(
            'disposition_provided_by', type=str, location='json')
        self.reqparse.add_argument('status', type=str, location='json')
        self.reqparse.add_argument('repair_date', type=str, location='json')
        self.reqparse.add_argument(
            'repair_report_id', type=str, location='json')
        self.reqparse.add_argument(
            'repair_approved_by', type=str, location='json')
        self.reqparse.add_argument('is_manual', type=bool, location='json')
        self.reqparse.add_argument('sso', type=str, location='json')

        super(UpdateDefectListAPI, self).__init__()

    def post(self):
        """
        Update an existing list of defect records, identified by the id_list, with the provided properties 
        ---
        parameters:
          - name: UpdateDefectListRequest
            in: body
            required: true
            schema:
              id: UpdateDefectListBody
              properties:
                id_list:
                  type: array
                  items:
                    type: number
                  example: [1, 2, 3, 4]
                date:
                  type: string
                  example: "2023-11-01"
                location:
                  type: string
                root_face_distance:
                  type: number
                span_wise_length:
                  type: number
                chord_wise_width:
                  type: number
                depth:
                  type: number
                height:
                  type: number
                width:
                  type: number
                length:
                  type: number
                aspect_ratio:
                  type: number
                area:
                  type: number
                percent_area:
                  type: number
                finding_type:
                  type: string
                ge_disposition:
                  type: string
                is_priority:
                  type: boolean
                description:
                  type: string
                image_pitch:
                  type: number
                image_yaw:
                  type: number
                image_hfov:
                  type: number
                design_tolerance:
                  type: string
                disposition_provided_by:
                  type: string
                status:
                  type: string
                repair_date:
                  type: string
                  example: "2023-11-01"
                repair_report_id:
                  type: string
                repair_approved_by:
                  type: string
                is_manual:
                  type: boolean
                sso:
                  type: string

        responses:
          200:
            description: Message indicating the number of updated records
            schema:
              message: string

        """
        args = self.reqparse.parse_args()
        logger.info(f'body: {args.items()}')

        # TODO: validate date instead
        if args['id_list'] is None or args['id_list'] == '':
            return {'message': f'id_list is required'}
    
        id_list = args['id_list']

        updated_defect_list = []
        with Session(db) as session:

            defect_list = session.query(Defect).filter(Defect.id.in_(id_list)).all()
            if defect_list is None:
                return {'message': f'Could not find defects for id_list: {id_list}.'}

            # walk through the existing list of defects from the DB, updating props if they are provided in the rest call body
            for defect in defect_list:
                # Propagate changes to the measuremnet if the measurement has only one parsed defect type
                measurement_id = defect.measurement_id
                parsed_defects = session.query(Defect).filter(Defect.measurement_id == measurement_id).all()
                measurement = None
                if len(parsed_defects) == 1:
                    measurement = session.query(Measurement).get(measurement_id)

                if args['date'] is not None:
                    defect.date = args['date']
                    if measurement is not None: measurement.date = args['date']

                if args['location'] is not None:
                    defect.location = args['location']
                    if measurement is not None: measurement.location = args['location']

                if args['root_face_distance'] is not None:
                    defect.root_face_distance = args['root_face_distance']
                    if measurement is not None: measurement.root_face_distance = args['root_face_distance']

                if args['span_wise_length'] is not None:
                    defect.span_wise_length = args['span_wise_length']
                    if measurement is not None: measurement.span_wise_length = args['span_wise_length']

                if args['chord_wise_width'] is not None:
                    defect.chord_wise_width = args['chord_wise_width']
                    if measurement is not None: measurement.chord_wise_width = args['chord_wise_width']

                if args['depth'] is not None:
                    defect.depth = args['depth']
                    if measurement is not None: measurement.depth = args['depth']

                if args['height'] is not None:
                    defect.height = args['height']
                    if measurement is not None: measurement.height = args['height']

                if args['width'] is not None:
                    defect.width = args['width']
                    if measurement is not None: measurement.width = args['width']

                if args['length'] is not None:
                    defect.length = args['length']
                    if measurement is not None: measurement.length = args['length']

                if args['aspect_ratio'] is not None:
                    defect.aspect_ratio = args['aspect_ratio']
                    if measurement is not None: measurement.aspect_ratio = args['aspect_ratio']

                if args['area'] is not None:
                    defect.area = args['area']
                    if measurement is not None: measurement.area = args['area']

                if args['percent_area'] is not None:
                    defect.percent_area = args['percent_area']
                    if measurement is not None: measurement.percent_area = args['percent_area']

                if args['finding_type'] is not None:
                    defect.finding_type = args['finding_type']
                    if measurement is not None: measurement.finding_type = args['finding_type']

                if args['ge_disposition'] is not None:
                    defect.ge_disposition = args['ge_disposition']
                    if measurement is not None: measurement.ge_disposition = args['ge_disposition']

                if args['is_priority'] is not None:
                    defect.is_priority = args['is_priority']
                    if measurement is not None: measurement.is_priority = args['is_priority']

                if args['description'] is not None:
                    defect.description = args['description']
                    if measurement is not None: measurement.description = args['description']

                if args['image_pitch'] is not None:
                    defect.image_pitch = args['image_pitch']
                    if measurement is not None: measurement.image_pitch = args['image_pitch']
                if args['image_yaw'] is not None:
                    defect.image_yaw = args['image_yaw']
                    if measurement is not None: measurement.image_yaw = args['image_yaw']
                if args['image_hfov'] is not None:
                    defect.image_hfov = args['image_hfov']
                    if measurement is not None: measurement.image_hfov = args['image_hfov']

                if args['design_tolerance'] is not None:
                    defect.design_tolerance = args['design_tolerance']
                    if measurement is not None: measurement.design_tolerance = args['design_tolerance']
                if args['disposition_provided_by'] is not None:
                    defect.disposition_provided_by = args['disposition_provided_by']
                    if measurement is not None: measurement.disposition_provided_by = args['disposition_provided_by']
                if args['status'] is not None:
                    defect.status = args['status']
                    if measurement is not None: measurement.status = args['status']
                if args['repair_date'] is not None:
                    # if one explicitly sets the date to empty, we remove the date
                    if args['repair_date'].strip() == '':
                        logging.info('Setting repair_date to None')
                        defect.repair_date = None
                        if measurement is not None: measurement.repair_date = None
                    else:
                        defect.repair_date = args['repair_date']
                        if measurement is not None: measurement.repair_date = args['repair_date']
                        
                if args['repair_report_id'] is not None:
                    defect.repair_report_id = args['repair_report_id']
                    if measurement is not None: measurement.repair_report_id = args['repair_report_id']
                if args['repair_approved_by'] is not None:
                    defect.repair_approved_by = args['repair_approved_by']
                    if measurement is not None: measurement.repair_approved_by = args['repair_approved_by']
                if args['is_manual'] is not None:
                    defect.is_manual = args['is_manual']
                    if measurement is not None: measurement.is_manual = args['is_manual']
                if args['sso'] is not None:
                    defect.sso = args['sso']
                    if measurement is not None: measurement.sso = args['sso']

                # logger.info(f'Measurement before update: {str(measurement)}')
                if measurement is not None:
                    logging.info(f'Will also update measurement id# {measurement_id}') 
                    session.add(measurement)

                session.add(defect)
                session.commit()
                session.refresh(defect)
                updated_defect_list.append(defect)

                logger.info(f'Updated defect: {str(defect)}')

                image = session.query(Image).get(defect.image_id)
                inspection_id = image.inspection_id
                invalidate_inspection_cache(inspection_id)

                if measurement is not None:
                    logger.info(f'Also updated measurement: {str(measurement)}')

            updated_id_list = []
            for defect in updated_defect_list:
                updated_id_list.append(defect.id)

        # resp = jsonify(inspection.serialize())
        resp = {'message': f'updated {len(updated_defect_list)} defect record(s)',
                'updated_defect_id_list': updated_id_list
               }
        logging.info(f'resp: {resp}')
        return resp


class DefectFrameAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):
        super(DefectFrameAPI, self).__init__()

    def allowed_file(self, filename):
        return '.' in filename and \
            filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

    def get(self, id):
        """
        Read an existing defect image frame json record
        ---
        parameters:
          - in: path
            required: true
            name: id
            schema:
              type: number
            description: Existing defect id

        responses:
          200:
            description: image frame for the provided defect id
            schema:
              id: FrameContent
              type: string
              format: json
        """
        logger.info(f'Reading fragment_file content for defect.id {id}')
        with Session(db) as session:
            
            defect_rec =  session.query(Defect).get(id)
            if defect_rec is None:
                return {'message': f'No defect id# {id} found in DB.'}
            image_id = defect_rec.image_id

            image =  session.query(Image).get(image_id)
            if image is None:
                return {'message': f'No image id# {image_id} associated to defect id# {id} found in DB.'}
            
            frame = image.frame
            if frame is None:
                return {'message': f'No frame for defect id# {id} found in DB.'}

            return json.loads(frame)


class DefectFragmentFileAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):
        super(DefectFragmentFileAPI, self).__init__()

    def allowed_file(self, filename):
        return '.' in filename and \
            filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

    def get(self, id):
        """
        Read an existing defect fragment file record, representing annotations on 2d still image
        ---
        parameters:
          - in: path
            required: true
            name: id
            schema:
              type: number
            description: Existing defect id

        responses:
          200:
            description: annotation file content for the provided defect id
            schema:
              id: DefectAnnotationFragmentContent
              type: string
              format: json
        """
        logger.info(f'Reading fragment_file content for defect.id {id}')
        with Session(db) as session:
            defect_annotation_fragment_results = session.scalars(
                select(DefectAnnotationFragment).where(DefectAnnotationFragment.defect_id == id)).all()
            logger.info(
                f'defect_annotation_fragment_results: {defect_annotation_fragment_results}')

            if defect_annotation_fragment_results is not None and len(defect_annotation_fragment_results) > 0:
                annotation_fragment = defect_annotation_fragment_results[0]
                fragment_content = annotation_fragment.content
                if fragment_content is not None:
                    return json.loads(fragment_content)

        return {'message': f'Annotation Fragment content for Defect id: {id} not found.'}


# ------------------------------------ Defect Repair Evidence --------------------------

class DefectRepairEvidenceFileAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):
        super(DefectRepairEvidenceFileAPI, self).__init__()

    def allowed_file(self, filename):
        return '.' in filename and \
            filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

    def get(self, id):
        """
        Read an existing defect repair evidence file by its id
        ---
        parameters:
          - in: path
            required: true
            name: id
            schema:
              type: number
            description: Existing defect repair evidence file id

        responses:
          200:
            description: defect repair evidence file content of the evidence file id
            schema:
              id: DefectEvidenceFileContent
              type: string
              format: binary
        """

        logger.info(f'Reading repair_evidence_file content id# {id}')
       
        with Session(db) as session:

            repair_efidence_file_rec = session.query(RepairEvidenceFile).get(id)
            if repair_efidence_file_rec is None:
                return {"message": f"Defect repair evidence file id# {id} not found."}

            file_content = read_file_record_content(repair_efidence_file_rec)
            if file_content is not None:

                download_filename_ext = repair_efidence_file_rec.filename.split('.')
                donwload_filename = download_filename_ext[0]+f'-did{repair_efidence_file_rec.id}'+'.'+download_filename_ext[1]
                
                mime_type = get_content_mime_type(file_content)

                return send_file(
                    BytesIO(file_content),
                    mimetype=mime_type,
                    as_attachment=True,
                    download_name=donwload_filename)
            else:
                logging.info(f'measurement_image_results for measurement_id {id} not found.')

        return {'message': f'Image File for Measurement id: {id} not found.'}

    def delete(self, id):
        """
        Deletes an existing repair evidence file
        ---
        parameters:
        - in: path
          required: true
          name: id
          schema:
            type: number
          description: Existing repair evidence file id
        responses:
          200:
            description: Success/failure message
            schema:
              message:
                type: string
        """
        logger.info(f'DELETE repaie evidence file id: {id}')
        with Session(db) as session:
            repair_evidence_file = session.query(RepairEvidenceFile).get(id)
            
            if repair_evidence_file is None:
                return {'message': f'repair evidence file record id {id} not found'}

            session.delete(repair_evidence_file)
            session.commit()

            return {'message': f'Repair evidence file id: {id} successfully deleted.'}


class DefectRepairEvidenceFileListAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):
        super(DefectRepairEvidenceFileListAPI, self).__init__()

    def get(self, id):
        """
        Read list of defect repair evidence files uploaded for the provided defect id
        ---
        parameters:
          - in: path
            required: true
            name: id
            schema:
              type: number
            description: Existing defect id

        responses:
          200:
            description: defect repair evidence file list
            schema:
              id: DefectEvidenceFileContent
              type: string
              format: binary
        """

        logger.info(f'Reading repair_evidence_file list for defect id# {id}')
       
        with Session(db) as session:
            resp_list = []
            repair_evidence_file_list = session.scalars(
                        select(RepairEvidenceFile).where(RepairEvidenceFile.defect_id == id)
                    ).all()
            for repair_file in repair_evidence_file_list:
                resp_list.append({'id':repair_file.id,
                                  'filename':repair_file.filename,
                                  'mime_type':repair_file.mime_type,
                                  'repairEvidenceComments':repair_file.comments,
                                 })

            return jsonify(resp_list)
        

class CreateDefectRepairEvidenceFileAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):
        super(CreateDefectRepairEvidenceFileAPI, self).__init__()

    def allowed_file(self, filename):
        return '.' in filename and \
            filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

    def post(self, id):
        """
        Upload a new defect repair evidence file
        ---
        consumes:
          - multipart/form-data
        parameters:
          - in: path
            required: true
            name: id
            schema:
              type: number
            description: Existing defect record id
          - in: formData
            required: true
            name: evidence_file
            type: file
            description: new evidence file
        responses:
          200:
            description: Success or failure message
            schema:
              properties:
                message:
                  type: string
                id:
                  type: number

        """
        logging.info(f'CreateDefectRepairEvidenceFileAPI.POST called with {request.files}')

        # formData image_file
        evidence_file = request.files.get("evidence_file")
        repair_evidence_comments = request.form.get("repair_evidence_comments")
        if evidence_file is None or not self.allowed_file(evidence_file.filename):
            msg = 'Evidence file content type not allowed, or no evidence_file provided.'
            return {'message': msg}

        evidence_file_filename = secure_filename(
            evidence_file.filename)
        unique_evidence_filename = str(uuid.uuid4()) + \
        '_'+evidence_file_filename
        evidence_file_path = os.path.join(
            app.config['UPLOAD_FOLDER'], unique_evidence_filename)
        evidence_file.save(evidence_file_path)

        logging.info('Reading evidence_file formData content...')
        bin_content = None
        with open(evidence_file_path, mode="rb") as f:
            bin_content = f.read()
        
        mime_type = get_content_mime_type(bin_content)

        logging.info('Generating evidence_file formData thumbnail...')
        thumb_bin_content = get_thumbnail_content(evidence_file_path)
        if os.path.isfile(evidence_file_path):
            os.remove(evidence_file_path)

        with Session(db) as session:
            defect_rec = session.query(Defect).get(id)
            if defect_rec is None:
                msg = f'Defect record id: {id} not found.'
                return {'message': msg}

            repair_evidence_file_rec = RepairEvidenceFile(
                defect_id=defect_rec.id,
                filename=evidence_file_filename,
                mime_type=mime_type,
                s3key=None,
                content=bin_content,
                thumbnail=thumb_bin_content,
                comments=repair_evidence_comments
            )
            session.add(repair_evidence_file_rec)
    

            if USE_S3:
                # Obtain image, image_file and inspection for the measurement
                image_rec = session.query(Image).get(defect_rec.image_id)
                image_file_rec = session.scalars(
                    select(ImageFile).filter(
                        ImageFile.image_id == defect_rec.image_id)
                ).all()[0]
                inspection_rec = session.query(
                    Inspection).get(image_rec.inspection_id)
                # name of image where this measurement came from
                pan_image_filename = image_file_rec.filename
                measurements_folder_name = pan_image_filename.split(r'\.(png|jpg)')[0]+'_measurements'
                repair_evidence_folder_name = 'evidence_defect_id{defect_rec.id}'
                s3key = get_inspection_s3key(inspection_rec)+f'/images/{measurements_folder_name}/{repair_evidence_folder_name}/{evidence_file_filename}'
                response = upload_content_to_s3(bin_content, s3key)
                logging.info(f'upload to s3 resp: {response}')
                if response.status_code != 200:
                    return {'message': f'Error uploading file: {evidence_file_filename} to S3'}
                repair_evidence_file_rec.s3key = s3key
                repair_evidence_file_rec.content = None

            session.commit()
            session.refresh(repair_evidence_file_rec)

            repair_evidence_file_rec.content = None
            repair_evidence_file_rec.thumbnail = None
            logging.info(
                f'new_repair_evidence_file_rec: {repair_evidence_file_rec}')
            repair_evidence_file_rec_json = repair_evidence_file_rec.toJson()
            logging.info(
                f'repair_evidence_file_rec_json: {repair_evidence_file_rec_json}')

            resp = {
                'message': f' Repair evidence file {evidence_file.filename} uploaded successfully.',
                'id': repair_evidence_file_rec.id 
            }
            return resp


class DefectRepairEvidenceImageAPI(Resource):

    def __init__(self):
        super(DefectRepairEvidenceImageAPI, self).__init__()
    
    def get(self, id):
        """
        Read a repaired image file of a defect 
        ---
        parameters:
          - in: path
            required: true
            name: id
            schema:
              type: number
            description: Existing defect repaired image id

        responses:
          200:
            description: repaired image file content for the provided defectrepair id
            schema:
              id: RepairEvidenceFile Id
              type: string
              format: binary
        """

        logger.info(f'Reading defect_repaired_image_file content for id# {id}')

        with Session(db) as session:

            RepairEvidenceFileImage = session.query(RepairEvidenceFile).get(id)
            if RepairEvidenceFileImage is None:
                return {"message": f"Defect record with id# {id} not found."}

            file_content = read_file_record_content(RepairEvidenceFileImage)
            if file_content is not None:
                    download_filename_ext = RepairEvidenceFileImage.filename.split('.')
                    donwload_filename = download_filename_ext[0]+f'-did{id}'+'.'+download_filename_ext[1]
                    return send_file_content_as_jpg(donwload_filename, file_content)
            else:
                logging.info(f'repaired_image_results for defect {id} not found.')
                return {'failure': True, 'message': f'Image File for Defect id: {id} not found.'}, 404


# ---------------------------------- OrigionalMeasurementAnnotationFile ----------------
# When expertes review MeasurementAnnotationFiles, they may choose to overwite that annotaiton file.
# A copy of the original file is then saved using this API
class OriginalMeasurementAnnotationFileMetadataAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):
        super(OriginalMeasurementAnnotationFileMetadataAPI, self).__init__()

    def get(self, id):
        """
        Read the original annotaiton file meta-data including replaced_by and replaced_timestamp
        ---
        parameters:
          - in: path
            required: true
            name: id
            schema:
              type: number
            description: Existing measurement id

        responses:
          200:
            schema:
              replaced_by:
                type: string
              replaced_timestamp:
                type: string
        """
        logger.info(
            f'Reading original_meassurement_annotation_file content for measurement.id {id}')

        image_id = None
        label_list = []
        with Session(db) as session:
            measurement = session.query(Measurement).get(id)
            if measurement is None:
                return {"message": f"Measurement {id} not found."}
            image_id = measurement.image_id
            original_measurement_annotation_results = session.scalars(
                select(OriginalMeasurementAnnotationFile).where(OriginalMeasurementAnnotationFile.measurement_id == id)).all()
            logger.info(
                f'origional_annotation_validation_status_results:  {original_measurement_annotation_results}')

            if original_measurement_annotation_results is not None and len(original_measurement_annotation_results) > 0:
                original_measurement_annotation_file_rec = original_measurement_annotation_results[
                    0]
                replaced_by = original_measurement_annotation_file_rec.replaced_by
                replaced_timestamp = original_measurement_annotation_file_rec.replaced_timestamp
                measurement_id = id

                label_list = get_annotation_content_label_list(
                    original_measurement_annotation_file_rec.content)

                return {
                    'id': original_measurement_annotation_file_rec.id,
                    'replaced_by': replaced_by,
                        'replaced_timestamp': str(replaced_timestamp),
                        'measurement_id': measurement_id,
                        'image_id': image_id,
                        'label_list': label_list}
            else:
                # check the annotation file, if any, for the list of labels and finding_type
                measurement_annotation_file_results = session.scalars(
                    select(MeasurementAnnotationFile).where(MeasurementAnnotationFile.measurement_id == id)).all()
                if measurement_annotation_file_results is not None and len(measurement_annotation_file_results) > 0:
                    measurement_annotation_file_rec = measurement_annotation_file_results[0]
                    label_list = get_annotation_content_label_list(
                        measurement_annotation_file_rec.content)

        return {
            'message': f'No original record found for measurement_id {id}',
            'id': -1,
            'replaced_by': None,
            'replaced_timestamp': None,
            'measurement_id': id,
            'image_id': image_id,
            'label_list': label_list}


class OriginalMeasurementAnnotationFileAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):
        super(OriginalMeasurementAnnotationFileAPI, self).__init__()

    def allowed_file(self, filename):
        return '.' in filename and \
            filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

    def post(self, id):
        """
        Creates an origional measurement 2d image annotation file record for a measurement or updates it
        ---
        consumes:
          - multipart/form-data
        parameters:
          - in: path
            required: true
            name: id
            schema:
              type: number
            description: Existing measurement id
          - in: formData
            required: false
            name: replaced_by
            type: string
            description: sso or name of the person who replaced
          - in: formData
            required: false
            name: replaced_timestamp
            type: string
            description: ISO date indicating when it was replaced. Now() if blank
          - in: formData
            required: true
            name: annotation_file
            type: file
            description: new annotation file object
        responses:
          200:
            description: Success or failure message
            schema:
              properties:
                message:
                  type: string

        """

        replaced_by = request.form.get('replaced_by', None)
        replaced_timestamp_str = request.form.get(
            'replaced_timestamp', None)
        label = None
        if replaced_timestamp_str == None:
            replaced_timestamp_str = str(datetime.now())

        logging.info(f'replaced_by: {replaced_by}')
        logging.info(f'replaced_timestamp_str: {replaced_timestamp_str}')

        if replaced_by is None:
            msg = 'No replaced_by parameter provided.'
            return {'message': msg}

        if replaced_timestamp_str is None:
            msg = 'No replaced_timestamp parameter provided.'
            return {'message': msg}

        # formData annotation_file
        annotation_file = request.files.get("annotation_file")
        if annotation_file is None or not self.allowed_file(annotation_file.filename):
            msg = 'Annotation file not found, or no annotation_file attachment provided.'
            return {'message': msg}

        annotation_filename = secure_filename(annotation_file.filename)
        unique_annotation_filename = str(uuid.uuid4())+'_'+annotation_filename
        
        with Session(db) as session:
            with tempfile.TemporaryDirectory(dir=UPLOAD_FOLDER, prefix='temp-', suffix='-folder') as tempdirname:
                annotation_file_path = os.path.join(tempdirname, unique_annotation_filename)
                annotation_file.save(annotation_file_path)

                # Check if measurement record exists
                measurement_rec = session.query(Measurement).get(id)
                if measurement_rec is None:
                    msg = f'Measurement record id: {id} not found.'
                    return {'message': msg}

                logging.info(
                    'Reading annotation_file attachment content and update measurement_rec...')
                annotation_json_content = None
                if os.path.isfile(annotation_file_path):
                    # We should not update the hov, pitch and yaw based on validated file...
                    # annotation_json_content = parse_annotation_file_and_save_props(
                    #     annotation_file_path, measurement_rec, session)
                    with open(annotation_file_path) as f:
                        annotation_json_content = json.load(f)
                    #os.remove(annotation_file_path)

                str_annotation_json_content = None
                str_annotation_json_content = json.dumps(annotation_json_content)
                logging.info(
                    f'original_annotation_file_content str: {str_annotation_json_content}')

                logger.info(
                    f'Reading original_measurement_annotation_file content for measurement.id {id}')
                original_measurement_annotation_file_list = session.scalars(
                    select(OriginalMeasurementAnnotationFile).where(
                        OriginalMeasurementAnnotationFile.measurement_id == id)
                ).all()
                original_measurement_annotation_file_rec = None

                # no existing annotation files, create new one...
                if len(original_measurement_annotation_file_list) == 0:
                    original_measurement_annotation_file_rec = OriginalMeasurementAnnotationFile(
                        image_id=measurement_rec.image_id,
                        measurement_id=id,
                        filename=annotation_filename,
                        s3key=None,
                        content=str_annotation_json_content,
                        replaced_by=replaced_by,
                        replaced_timestamp=replaced_timestamp_str
                    )
                    session.add(original_measurement_annotation_file_rec)
                else:  # update existing record
                    original_measurement_annotation_file_rec = original_measurement_annotation_file_list[
                        0]
                    original_measurement_annotation_file_rec.filename = annotation_filename
                    original_measurement_annotation_file_rec.content = str_annotation_json_content

                    if replaced_by is not None:
                        original_measurement_annotation_file_rec.replaced_by = replaced_by
                    if replaced_timestamp_str is not None:
                        original_measurement_annotation_file_rec.replaced_timestamp = replaced_timestamp_str
                session.commit()
                session.refresh(original_measurement_annotation_file_rec)

                logging.info(
                    f'original_measurement_annotation_file_rec: {original_measurement_annotation_file_rec}')
                measurement_annotation_file_rec_json = original_measurement_annotation_file_rec.toJson()
                logging.info(
                    f'original_measurement_annotation_file_json: {measurement_annotation_file_rec_json}')
                return measurement_annotation_file_rec_json

    # returns the file content as opposed to meta-data

    def get(self, id):
        """
        Read an existing original measurement annotation file record, representing annotations on 2d still image
        ---
        parameters:
          - in: path
            required: true
            name: id
            schema:
              type: number
            description: Existing measurement id

        responses:
          200:
            description: original measurement annotation file content for the provided measurement id
            schema:
              id: OriginalMeasurementAnnotationFileContent
              type: string
              format: json
        """
        logger.info(
            f'Reading validated_measurement_annotation_file content for measurement.id {id}')
        with Session(db) as session:
            original_measurement_annotation_results = session.scalars(
                select(OriginalMeasurementAnnotationFile).where(
                    OriginalMeasurementAnnotationFile.measurement_id == id)).all()
            logger.info(
                f'original_measurement_annotation_results: {original_measurement_annotation_results}')

            if original_measurement_annotation_results is not None and len(original_measurement_annotation_results) > 0:
                annotation_file = original_measurement_annotation_results[0]
                file_content = annotation_file.content
                if file_content is not None:
                    return json.loads(file_content)

        return {'message': f'Original annotation File content for Measurement id: {id} not found.'}

    def delete(self, id):
        """
        Deletes an existing original measurement annotation file
        ---
        parameters:
          - in: path
            required: true
            name: id
            schema:
              type: number
            description: Existing measurement id
        responses:
          200:
            description: Success/failure message
            schema:
              message:
                type: string
        """
        logger.info(f'DELETE validated annotation for measurement id: {id}')
        with Session(db) as session:
            original_measurement_annotation_file_list = session.scalars(
                select(OriginalMeasurementAnnotationFile).where(
                    OriginalMeasurementAnnotationFile.measurement_id == id)
            ).all()

            if original_measurement_annotation_file_list is None or len(original_measurement_annotation_file_list) == 0:
                return {'message': f'annotation for measurement record id {id} not found'}

            # logging.info(f'del measurement_annotation_file_list: {measurement_annotation_file_list}')
            for original_measurement_annotation in original_measurement_annotation_file_list:
                logging.info(
                    f'del original_measurement_annotation id {original_measurement_annotation.id}')
                session.delete(original_measurement_annotation)
                session.commit()

            return {'message': f'Original annotation file for measurement id: {id} successfully deleted.'}


# ---------------------------------- ValidatedMeasurementAnnotationFile ----------------
# When expertes review MeasurementAnnotationFiles, they save them to the ValidatedMeasurementAnnotationFiles record.
class ValidatedMeasurementAnnotationFileMetadataAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):
        super(ValidatedMeasurementAnnotationFileMetadataAPI, self).__init__()

    def get(self, id):
        """
        Read the meta-data including validation status of the validated measurement annotation
        ---
        parameters:
          - in: path
            required: true
            name: id
            schema:
              type: number
            description: Existing measurement id

        responses:
          200:
            description: the validation status of the measurement annotation if any
            schema:
              validation_status:
                type: string
              validated_by:
                type: string
              validation_timestamp:
                type: string
        """
        logger.info(
            f'Reading validated_meassurement_annotation_file content for measurement.id {id}')

        image_id = None
        label_list = []
        validation_status = 'no_annotation'
        finding_type = 'unknown'
        with Session(db) as session:
            measurement = session.query(Measurement).get(id)
            if measurement is None:
                return {"message": f"Measurement {id} not found."}
            image_id = measurement.image_id
            validated_measurement_annotation_results = session.scalars(
                select(ValidatedMeasurementAnnotationFile).where(ValidatedMeasurementAnnotationFile.measurement_id == id)).all()
            logger.info(
                f'measurement_annotation_validation_status_results: {validated_measurement_annotation_results}')

            if validated_measurement_annotation_results is not None and len(validated_measurement_annotation_results) > 0:
                validated_measurement_annotation_file_rec = validated_measurement_annotation_results[
                    0]
                validation_status = validated_measurement_annotation_file_rec.validation_status
                validated_by = validated_measurement_annotation_file_rec.validated_by
                validation_timestamp = validated_measurement_annotation_file_rec.validation_timestamp
                finding_type = validated_measurement_annotation_file_rec.finding_type
                measurement_id = id

                label_list = get_annotation_content_label_list(
                    validated_measurement_annotation_file_rec.content)

                return {
                    'id': validated_measurement_annotation_file_rec.id,
                    'validation_status': validation_status,
                        'validated_by': validated_by,
                        'validation_timestamp': str(validation_timestamp),
                        'measurement_id': measurement_id,
                        'image_id': image_id,
                        'finding_type': finding_type,
                        'label_list': label_list}
            else:
                # check the annotation file, if any, for the list of labels and finding_type
                measurement_annotation_results = session.scalars(
                    select(MeasurementAnnotationFile).where(MeasurementAnnotationFile.measurement_id == id)).all()
                if measurement_annotation_results is not None and len(measurement_annotation_results) > 0:
                    measurement_annotation_file_rec = measurement_annotation_results[0]
                    label_list = get_annotation_content_label_list(
                        measurement_annotation_file_rec.content)
                    finding_type = measurement_annotation_file_rec.finding_type
                    if (len(label_list) > 0):
                        validation_status = 'pending'

        # return a templace record so it can show up in the UI.
        return {
            'id': -1,
            'validation_status': validation_status,
            'validated_by': None,
            'validation_timestamp': None,
            'measurement_id': id,
            'image_id': image_id,
            'finding_type': finding_type,
            'label_list': label_list}


class ValidatedMeasurementAnnotationFileAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):
        super(ValidatedMeasurementAnnotationFileAPI, self).__init__()

    def allowed_file(self, filename):
        return '.' in filename and \
            filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

    def post(self, id):
        """
        Update an existing validated measurement 2d image annotation file for a measurement id or create a new one
        ---
        consumes:
          - multipart/form-data
        parameters:
          - in: path
            required: true
            name: id
            schema:
              type: number
            description: Existing measurement id
          - in: formData
            required: false
            name: validation_status
            type: string
            description: valid, invalid, corrected
          - in: formData
            required: false
            name: validated_by
            type: string
            description: sso or name of the person who validated
          - in: formData
            required: false
            name: validation_timestamp
            type: string
            description: ISO date indicating when it was validated. Now() if blank
          - in: formData
            required: true
            name: annotation_file
            type: file
            description: new annotation file object
        responses:
          200:
            description: Success or failure message
            schema:
              properties:
                message:
                  type: string

        """

        validation_status = request.form.get('validation_status', None)
        validated_by = request.form.get('validated_by', None)
        validation_timestamp_str = request.form.get(
            'validation_timestamp', None)
        label = None
        if validation_timestamp_str == None:
            validation_timestamp_str = str(datetime.now())

        logging.info(f'validation_statuss: {validation_status}')
        logging.info(f'validated_by: {validated_by}')
        logging.info(f'validation_timestamp_str: {validation_timestamp_str}')

        # formData annotation_file
        annotation_file = request.files.get("annotation_file")
        logging.info(f'annotation_file Hobbo = {annotation_file}')
        if annotation_file is None or not self.allowed_file(annotation_file.filename):
            msg = 'Annotation file not found, or no annotation_file provided.'
            logging.info(f'bummer2 : {msg}')
            return {'message': msg}

        logging.info(f'I passed by here after bummer2')
        annotation_filename = secure_filename(annotation_file.filename)
        unique_annotation_filename = str(uuid.uuid4())+'_'+annotation_filename
        annotation_file_path = os.path.join(
            app.config['UPLOAD_FOLDER'], unique_annotation_filename)
        annotation_file.save(annotation_file_path)
        logging.info(f'I passed by here after bummer3')

        with Session(db) as session:

            # Check if measurement record exists
            measurement_rec = session.query(Measurement).get(id)
            logging.info(f'I passed by here after bummer4')
            if measurement_rec is None:
                msg = f'Measurement record id: {id} not found.'
                return {'message': msg}

            logging.info(
                'Reading annotation_file content and update measurement_rec...')
            annotation_json_content = None
            logging.info(f'I passed by here after bummer5')
            if os.path.isfile(annotation_file_path):
                # We should not update the hov, pitch and yaw based on validated file...
                # annotation_json_content = parse_annotation_file_and_save_props(
                #     annotation_file_path, measurement_rec, session)
                with open(annotation_file_path) as f:
                    annotation_json_content = json.load(f)
                os.remove(annotation_file_path)
            logging.info(f'I passed by here after bummer666 = annotation_json_content = {annotation_json_content}')

            # if the user does not provide the parameners in the rest call, we look for them in the annotation file
            validation_props = read_validated_annotation_props( annotation_json_content)
            logging.info(f'I passed by here after bummer777')
            logging.info(f'validation_props: {validation_props}')
            if validation_status is None:
                validation_status = validation_props.get(
                    'validationStatus', None)
                logging.info(f'using annotation file validation_status: {validation_status}')
            if validated_by is None:
                validated_by = validation_props.get('validatedBy', None)
                logging.info(f'using annotation file validated_by: {validated_by}')
            if validation_timestamp_str is None:
                validation_timestamp_str = validation_props.get(
                    'validationTimestamp', None)
                logging.info(f'using annotation file validation_timestamp_str: {validation_timestamp_str}')
            label = validation_props.get(
                'label', None)

            str_annotation_json_content = None
            str_annotation_json_content = json.dumps(annotation_json_content)
            logging.info(
                f'validated_annotation_file_content str: {str_annotation_json_content}')

            logger.info(
                f'Reading validated_measurement_annotation_file content for measurement.id {id}')
            validated_measurement_annotation_file_list = session.scalars(
                select(ValidatedMeasurementAnnotationFile).where(
                    ValidatedMeasurementAnnotationFile.measurement_id == id)
            ).all()
            validated_measurement_annotation_file_rec = None

            # no existing annotation files, create new one...
            if len(validated_measurement_annotation_file_list) == 0:
                validated_measurement_annotation_file_rec = ValidatedMeasurementAnnotationFile(
                    image_id=measurement_rec.image_id,
                    measurement_id=id,
                    filename=annotation_filename,
                    s3key=None,
                    content=str_annotation_json_content,
                    validation_status=validation_status,
                    validated_by=validated_by,
                    validation_timestamp=validation_timestamp_str,
                    finding_type=label
                )
                session.add(validated_measurement_annotation_file_rec)
            else:  # update existing record
                validated_measurement_annotation_file_rec = validated_measurement_annotation_file_list[
                    0]
                validated_measurement_annotation_file_rec.filename = annotation_filename
                validated_measurement_annotation_file_rec.content = str_annotation_json_content
                if validation_status is not None:
                    validated_measurement_annotation_file_rec.validation_status = validation_status
                if validated_by is not None:
                    validated_measurement_annotation_file_rec.validated_by = validated_by
                if validation_timestamp_str is not None:
                    validated_measurement_annotation_file_rec.validation_timestamp = validation_timestamp_str
                if label is not None:
                    validated_measurement_annotation_file_rec.finding_type = label
            session.commit()
            session.refresh(validated_measurement_annotation_file_rec)

            logging.info(
                f'validated_measurement_annotation_file_rec: {validated_measurement_annotation_file_rec}')
            measurement_annotation_file_json = validated_measurement_annotation_file_rec.toJson()
            logging.info(
                f'validated_measurement_annotation_file_json: {measurement_annotation_file_json}')
            return measurement_annotation_file_json

    def get(self, id):
        """
        Read an existing validated measurement annotation file record, representing annotations on 2d still image
        ---
        parameters:
          - in: path
            required: true
            name: id
            schema:
              type: number
            description: Existing measurement id

        responses:
          200:
            description: validated measurement annotation file content for the provided measurement id
            schema:
              id: ValidatedMeasurementAnnotationFileContent
              type: string
              format: json
        """
        logger.info(
            f'Reading validated_measurement_annotation_file content for measurement.id {id}')
        with Session(db) as session:
            validated_measurement_annotation_results = session.scalars(
                select(ValidatedMeasurementAnnotationFile).where(
                    ValidatedMeasurementAnnotationFile.measurement_id == id)).all()
            logger.info(
                f'validated_measurement_annotation_results: {validated_measurement_annotation_results}')

            if validated_measurement_annotation_results is not None and len(validated_measurement_annotation_results) > 0:
                annotation_file = validated_measurement_annotation_results[0]
                file_content = annotation_file.content
                if file_content is not None:
                    return json.loads(file_content)

        return {'message': f'Validated annotation File content for Measurement id: {id} not found.'}

    def delete(self, id):
        """
        Deletes an existing validated measurement annotation file
        ---
        parameters:
          - in: path
            required: true
            name: id
            schema:
              type: number
            description: Existing measurement id
        responses:
          200:
            description: Success/failure message
            schema:
              message:
                type: string
        """
        logger.info(f'DELETE validated annotation for measurement id: {id}')
        with Session(db) as session:
            validated_measurement_annotation_file_list = session.scalars(
                select(ValidatedMeasurementAnnotationFile).where(
                    ValidatedMeasurementAnnotationFile.measurement_id == id)
            ).all()

            if validated_measurement_annotation_file_list is None or len(validated_measurement_annotation_file_list) == 0:
                return {'message': f'annotation for measurement record id {id} not found'}

            # logging.info(f'del measurement_annotation_file_list: {measurement_annotation_file_list}')
            for validated_measurement_annotation in validated_measurement_annotation_file_list:
                logging.info(
                    f'del validated_measurement_annotation id {validated_measurement_annotation.id}')
                session.delete(validated_measurement_annotation)
                session.commit()

            return {'message': f'annotation file for measurement id: {id} successfully deleted.'}


# -------------------------------------------- Blade ----------------------------------

class BladeListAPI(Resource):
    def __init__(self):
        self.reqparse = reqparse.RequestParser()

        self.reqparse.add_argument('offset', type=str, location='args')
        self.reqparse.add_argument('limit', type=str, location='args')

        super(BladeListAPI, self).__init__()

    def get(self):
        """
        Read an existing blades list
        ---
        parameters:
          - in: query
            required: false
            name: limit
            schema:
              type: number
            description: limit number of returned values from offset position
            default: 100000
          - in: query
            required: false
            name: offset
            schema:
              type: number
            description: initial index position for return values
            default: 0

        responses:
          200:
            description: Existing blade metadata
            schema:
              id: BladeObjectList
              type: array
              items:
                schema:
                  id: BladeObject

        """
        args = self.reqparse.parse_args()
        logger.info(args.items())

        offset = args['offset']
        limit = args['limit']

        if offset is None:
            offset = 0
        if limit is None:
            limit = 100000

        logger.info(f'offset: {offset}, limit: {limit}')
        with Session(db) as session:
            blade_list = session.scalars(
                select(Blade)
            ).all()
            logging.info(f'blade_list: {blade_list}')
            return Blade.toJsonList(blade_list)

        return {'message': 'No blade records in DB.'}


class BladeAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):
        super(BladeAPI, self).__init__()

    def get(self, id):
        """
        Read an existing blade
        ---
        parameters:
          - in: path
            required: true
            name: id
            schema:
              type: number
            description: Existing blade id

        responses:
          200:
            description: Existing blade
            schema:
              id: BladeObject

        """
        with Session(db) as session:
            inspection = session.query(Blade).get(id)
            if inspection is not None:
                return inspection.toJson()

        return {'message': f'Blade {id} not found.'}

    def delete(self, id):
        """
        Deletes an existing blade
        ---
        parameters:
          - in: path
            required: true
            name: id
            schema:
              type: number
            description: Existing blade id

        responses:
          200:
            description: Existing blade
            schema:
              id: BladeObject

        """
        with Session(db) as session:
            blade = session.query(Blade).get(id)
            if blade is not None:
                session.delete(blade)
                session.commit()
                # TODO: cascade delete all dependence records
                return blade.toJson()

        return {'message': f'Blade {id} not found.'}


class CreateBladeAPI(Resource):

    def __init__(self):
        self.reqparse = reqparse.RequestParser()
        self.reqparse.add_argument('serial_number', type=str, location='json')
        self.reqparse.add_argument('model', type=str, location='json')
        self.reqparse.add_argument('manufacturer', type=str, location='json')
        self.reqparse.add_argument('tpi_ncr_number', type=str, location='json')
        self.reqparse.add_argument('number', type=str, location='json')
        self.reqparse.add_argument('set_number', type=str, location='json')
        self.reqparse.add_argument('length', type=float, location='json')

        super(CreateBladeAPI, self).__init__()

    def post(self):
        """
        Create a blade
        ---
        parameters:
          - name: blade
            in: body
            required: true
            schema:
              id: BladeBody
              required:
                - name
                - date
                - part_number
              properties:
                serial_number:
                  type: string
                  example: "TPI-51179"
                model:
                  type: string
                  example: "GE 68.7 Root"
                manufacturer:
                  type: string
                  example: "TPI MX3"
                tpi_ncr_number:
                  type: string
                  example: "1234"
                number:
                  type: string
                  example: "1"
                  description: "Number of the blade e.g. 1, 2, 3, etc."
                set_number:
                  type: string
                  example: "378"
                length:
                  type: number
                  example: 12.5
        responses:
          200:
            description: Created blade object
            schema:
              id: BladeObject
              properties:
                id:
                  type: number
                number:
                  type: string
                set_number:
                  type: string
                serial_number:
                  type: string
                model:
                  type: string
                manufacturer:
                  type: string
                tpi_ncr_number:
                  type: string
                length:
                  type: number


        """
        args = self.reqparse.parse_args()
        logger.info(args.items())

        blade = Blade(
            serial_number=args['serial_number'],
            model=args['model'],
            manufacturer=args['manufacturer'],
            tpi_ncr_number=args['tpi_ncr_number'],
            number=args['number'],
            set_number=args['set_number'])

        with Session(db) as session:
            session.add(blade)
            session.commit()
            session.refresh(blade)  # to read the id

        logger.info(f'Created blade: {str(blade)}')

        # resp = jsonify(inspection.serialize())
        resp = blade.toJson()
        logging.info(f'resp: {resp}')
        return resp


# --------------------------------- Inspection ----------------------------------

# Note that defects are sub-groups of shapes from a measurement that have a common label.
class InspectionDefectListAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):

        self.reqparse = reqparse.RequestParser()
        self.reqparse.add_argument('includeAI', type=bool, location='args')

        super(InspectionDefectListAPI, self).__init__()

    def get(self, id):
        """
        Read the list of measurement defects for a provided inspection id
        ---
        parameters:
          - in: path
            required: true
            name: id
            schema:
              type: number
            description: Inspection id
          - in: query
            required: false
            name: includeAI
            schema:
              type: boolean
            description: Whether to include AI annotations or not

        responses:
          200:
            description: List of measurement defects
            schema:
              id: InspectionDefectList
              type: array
              items:
                schema:
                  id: DefectObject

        """

        logging.info('InspectionDefectListAPI')

        include_ai_str = request.args.get('includeAI')
        if include_ai_str is not None:
            include_ai_str = include_ai_str.capitalize()
        logger.info(f'include_ai_str: {include_ai_str}')
        include_ai = False
        if include_ai_str == 'True':
            include_ai = True

        cache_key = 'InspectionDefectListAPI.get_'+str(id)+'_'+include_ai_str
        if USE_INSPECTION_CACHE:
            if cache_key in cache:
                logging.info(f'Cache hit: {cache_key}')
                resp = cache[cache_key]
                return jsonify(resp)
            else:
                logging.info('Cache miss')

        with Session(db) as session:
            image_list = session.scalars(
                select(Image).where(Image.inspection_id ==
                                    id).order_by(Image.distance)
            ).all()

            resp = []
            if image_list is not None and len(image_list) > 0:
                for img in image_list:

                    defect_list = []
                    if include_ai == True:
                        defect_list = session.scalars(
                            select(Defect).where(
                                Defect.image_id == img.id)
                        ).all()
                    else:
                        defect_list = session.scalars(
                            select(Defect).where(
                                and_(Defect.image_id == img.id, Defect.is_manual == True))
                        ).all()

                    for defect in defect_list:
                        if defect.root_face_distance == 0 and defect.root_face_distance != img.distance:
                            # logging.info(f'Measurement id: {measurement.id} is missing root face distance. Using image distance from image id {img.id}: {img.distance}')
                            defect.root_face_distance = img.distance

                    resp.extend(
                        Defect.serialize_list(defect_list))

                # logging.info(f'resp_list: {resp_list}')
                
                if USE_INSPECTION_CACHE:
                    cache.set(cache_key, resp, expire=CACHE_TTL)
                
                return jsonify(resp)

        resp = {
            'message': f'No defect records for inspection id {id} found in DB.'}
        
        return make_response(jsonify(resp), 200)


class InspectionMeasurementListAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):
        super(InspectionMeasurementListAPI, self).__init__()

    def get(self, id):
        """
        Read the list of measurements for a provided inspection id
        ---
        parameters:
          - in: path
            required: false
            name: id
            schema:
              type: number
            description: Inspection id

        responses:
          200:
            description: List of measurements
            schema:
              id: InspectionMeasurementList
              type: array
              items:
                schema:
                  id: MeasurementObject

        """
        logging.info('InspectionMeasurementListAPI')
        with Session(db) as session:
            image_list = session.scalars(
                select(Image).where(Image.inspection_id ==
                                    id).order_by(Image.distance)
            ).all()

            resp_list = []
            if image_list is not None and len(image_list) > 0:
                for img in image_list:

                    measurement_list = session.scalars(
                        select(Measurement).where(
                            Measurement.image_id == img.id)
                    ).all()

                    for measurement in measurement_list:
                        if measurement.root_face_distance == 0 and measurement.root_face_distance != img.distance:
                            # logging.info(f'Measurement id: {measurement.id} is missing root face distance. Using image distance from image id {img.id}: {img.distance}')
                            measurement.root_face_distance = img.distance

                    resp_list.extend(
                        Measurement.serialize_list(measurement_list))

                # logging.info(f'resp_list: {resp_list}')
                return jsonify(resp_list)

        resp = {
            'message': f'No measurement records for inspection id {id} found in DB.'}
        return make_response(jsonify(resp), 200)


class InspectionImageFindAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):
        self.reqparse = reqparse.RequestParser()
        self.reqparse.add_argument('distance', type=float, location='args')

        super(InspectionImageFindAPI, self).__init__()

    def get(self, id):
        """
        Find existing 360 image for the provided inspection id and distance
        ---
        parameters:
          - in: path
            required: true
            name: id
            schema:
              type: number
            description: Inspection id
          - in: query
            required: true
            name: distance
            schema:
              type: number
            description: Distance from root

        responses:
          200:
            description: List of 360 images metadata
            schema:
              id: InspectionImageList
              type: array
              items:
                schema:
                  id: ImageObject

        """
        logging.info('InspectionImageFindAPI')

        args = self.reqparse.parse_args()
        logger.info(args.items())

        distance = args['distance']

        with Session(db) as session:
            image_list = session.scalars(
                select(Image).where(Image.inspection_id ==
                                    id).order_by(Image.distance)
            ).all()

            if image_list is not None and len(image_list) > 0:
                resp_list = []
                for img in image_list:
                    img_dist = img.distance
                    # logging.info(f'img {img.id} distance: {img_dist}')
                    if (abs(distance - img_dist) < 0.5):
                        resp_list.append(img)
                # logging.info(f'resp_list: {resp_list}')

                return jsonify(Image.serialize_list(resp_list))

        resp = {'message': f'No image records for inspection id {id} found in DB.'}
        return make_response(jsonify(resp), 200)


class InspectionImageListAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):
        super(InspectionImageListAPI, self).__init__()

    def get(self, id):
        """
        Read existing 360 images list for the provided inspection id
        ---
        parameters:
          - in: path
            required: false
            name: id
            schema:
              type: number
            description: Inspection id

        responses:
          200:
            description: List of 360 images metadata
            schema:
              id: InspectionImageList
              type: array
              items:
                schema:
                  id: ImageObject

        """
        logging.info('InspectionImageListAPI')
        with Session(db) as session:
            image_list = session.scalars(
                select(Image).where(Image.inspection_id ==
                                    id).order_by(Image.distance)
            ).all()

            if image_list is not None and len(image_list) > 0:
                resp_list = Image.serialize_list(image_list)
                # logging.info(f'resp_list: {resp_list}')
                for img in resp_list:
                    m_count = session.query(Measurement.id).filter(
                        Measurement.image_id == img['id']).count()
                    if m_count > 0:
                        img['measurement_count'] = m_count
                    else:
                        img['measurement_count'] = 0

                    vts_count = session.query(VTShot.id).filter(
                        VTShot.image_id == img['id']).count()
                    if vts_count > 0:
                        img['vtshot_count'] = vts_count
                    else:
                        img['vtshot_count'] = 0
                # logging.info(f'resp_list: {resp_list}')

                return jsonify(resp_list)

        resp = {'message': f'No image records for inspection id {id} found in DB.'}
        return make_response(jsonify(resp), 200)


class InspectionImageDistancesAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):
        super(InspectionImageDistancesAPI, self).__init__()

    def get(self, id):
        """
        Read list of valid 360 images distances for the provided inspection id
        ---
        parameters:
          - in: path
            required: false
            name: id
            schema:
              type: number
            description: Inspection id

        responses:
          200:
            description: List of 360 images metadata
            schema:
              type: array
              items:
                type: number
              example: [1, 2, 3, 4]

        """
        logging.info('InspectionImageDistancesAPI')
        with Session(db) as session:
            image_list = session.scalars(
                select(Image).where(Image.inspection_id ==
                                    id).order_by(Image.distance)
            ).all()

            if image_list is not None and len(image_list) > 0:
                img_list = Image.serialize_list(image_list)
                resp_list = []
                # logging.info(f'image_list: {resp_list}')
                for img in img_list:
                    # logging.info(f'img: {img}')
                    dist = round(img.get('distance'), 1)
                    if dist not in resp_list:
                        resp_list.append(dist)
                logging.info(f'resp_list: {resp_list}')

                return jsonify(resp_list)

        resp = {'message': f'No image records for inspection id {id} found in DB.'}
        return make_response(jsonify(resp), 200)


# split the key into two sub-categories: AI and MANUAL
def build_total_template_manual_and_ai(template):
    count_dict = {}
    for key in template.keys():
        count_dict[key+' (M)'] = 0
        count_dict[key+' (AI)'] = 0
    return count_dict


# split the key into two sub-categories: AI and MANUAL
def build_total_template_manual_ai_disposition(defect_template, disposition_list):
    count_dict = {}
    for defect_key in defect_template.keys():
        manual_key = defect_key+' (M)'
        ai_key = defect_key+' (AI)'
        #count_dict[manual_key] = 0
        #count_dict[ai_key] = 0
        for disposition_key in disposition_list:
            count_dict[manual_key+' - '+disposition_key] = 0
            count_dict[ai_key+' - '+disposition_key] = 0
    return count_dict


# return one of: ['Open_No_Disposition', 'False_Positive', 'Within_Tolerance', 'Out_of_Tolerance', 'Repaired', 'Duplicate']
def get_disposition_key(ge_disposition):
    if 'False' in ge_disposition:
        return 'False_Positive'
    elif 'Within' in ge_disposition:
        return 'Within_Tolerance'
    elif 'Out' in ge_disposition:
        return 'Out_of_Tolerance'
    elif 'Repaired' in ge_disposition:
        return 'Repaired'
    elif 'Duplicate' in ge_disposition:
        return 'Duplicate'
    return 'Open_No_Disposition'
    


def build_defect_stats_template_manual_and_ai(template):
    stats_dict = {}

    for key in template.keys():
        stats_dict[key+' (M)'] = DISPOSITION_TOTAL_TEMPLATE.copy()
        stats_dict[key+' (AI)'] = DISPOSITION_TOTAL_TEMPLATE.copy()
    return stats_dict


# total defect counts by category
DEFECT_TOTAL_TEMPLATE = {
    'Adhesive Cracks': 0,
    'Adhesive Voids': 0,
    'CoreGap': 0,
    'Cuts in Web Flange': 0,
    'Damaged Glass': 0,
    'Dust & dirt': 0,
    'Entrained air': 0,
    'Exposed core / Missing laminate': 0,
    'Foreign Objects':  0,
    'Laminate Loose Edge': 0,
    'Layer end': 0,
    'Layer misplacement': 0,
    'Layers Overlap': 0,
    'LPS Cable Damage': 0,
    'Main SW Web Foot Lam': 0,
    'Metal Shavings': 0,
    'RCO Bracket bond': 0,
    'RCO Seal': 0,
    'Repairs incorrect staggering': 0,
    'Shearclips missing': 0,
    'TE SW Web Foot Lam': 0,
    'TEBC Overlam Overlap': 0,
    'TEBC Paste Thickness': 0,
    'TEBC Wave': 0,
    'Uncured laminate': 0,
    'Other': 0,
    'Voids Overlaminate': 0,
    'Waves Laminate': 0,
    "Core Offset": 0,
    "Semi Dry Glass / LFR": 0,
    "Delamination / LDL": 0,
    "Core Misplacement / LCM": 0,
    "LPS Loose Cable": 0,
    "Laminate Roving Misplacement / LRM": 0,
    "Excess Adhesive": 0,
}

DISPOSITION_LIST = ['Open_No_Disposition', 'False_Positive', 'Within_Tolerance', 'Out_of_Tolerance', 'Repaired', 'Duplicate']

# builds a dict where we have (AI) and (M) columns for each defect type
DEFECT_TOTAL_TEMPLATE_MANUAL_AI = build_total_template_manual_and_ai(
    DEFECT_TOTAL_TEMPLATE)

DEFECT_TOTAL_TEMPLATE_MANUAL_AI_DISPOSITION = build_total_template_manual_ai_disposition(DEFECT_TOTAL_TEMPLATE, DISPOSITION_LIST)

DISPOSITION_TOTAL_TEMPLATE = {
    'Total_Open_No_Disposition': 0,
    'Total_False_Positive': 0,
    'Total_Within_Tolerance': 0,
    'Total_Out_of_Tolerance': 0,
    'Total_Repaired': 0,
    'Total_Duplicate': 0,
}

# builds a dict with (AI) and (M) counts for each type of disposition
DISPOSITION_TOTAL_TEMPLATE_MANUAL_AI = build_total_template_manual_and_ai(
    DISPOSITION_TOTAL_TEMPLATE)

DEFECT_STATS_TEMPLATE_MANUAL_AI = build_defect_stats_template_manual_and_ai(
    DEFECT_TOTAL_TEMPLATE)


# returns a dictionalry with number of defects in each defect category
def get_annotation_content_label_list(file_content):
    label_list = []

    # logging.info(f'file_content: {file_content}')
    if (file_content is not None and len(file_content) > 0):
        json_content = json.loads(file_content)
        if 'shapes' in json_content:
            shapes = json_content['shapes']
            if shapes is not None:
                for el in shapes:
                    # logging.info(f'element: {el}')
                    finding_label = el['label']
                    label_list.append(finding_label)
    return label_list


# totalizes defects by defect type for a 360 image
def image_defect_count(img_id, count_dict, session):
    # count_dict = DEFECT_TOTAL_TEMPLATE.copy()
    logging.info(f'Counting defects by type for image id# {img_id}')
    defect_list = session.scalars(
        select(Defect).where(
            Defect.image_id == img_id)
    ).all()
    for defect in defect_list:
        finding_label = defect.finding_type
        if finding_label == '' or finding_label is None:
            finding_label = 'Other'

        if finding_label in count_dict:
            count_dict[finding_label] += 1
            logging.info(f'Counting finding_type: {finding_label} for 360 image id# {img_id}')
        else:
            logging.info(f'Ignoring finding_type: {finding_label}')


# Adopt standard names or group names for variations in labels
def normalize_finding_label(label):
    finding_label = label
    if finding_label == '' or finding_label is None:
        #logging.info(f'Counting finding_type: {finding_label} as "Other" ')
        finding_label = 'Other'
    elif finding_label == 'Core Gap' or finding_label == 'core-gap' or finding_label == 'Core gaps between panels':
        #logging.info(f'Counting finding_type: {finding_label} as "CoreGap" ')
        finding_label = 'CoreGap'
    elif finding_label == 'Adhesive Void':
        finding_label = 'Adhesive Voids'
    elif finding_label == 'Delamination':
        #logging.info(f'Counting finding_type: {finding_label} as "Delamination / LDL" ')
        finding_label = "Delamination / LDL"
    elif 'Foreign Objects' in finding_label:
        finding_label = 'Foreign Objects'
    return finding_label


def inspection_defect_stats_count(inspection_id, count_dict, session):
    # count_dict = DEFECT_TOTAL_TEMPLATE.copy()
    logging.info(f'Gathering defect stats for inspection id# {inspection_id}')
    logging.info(f'Counting defects using count_dict: {count_dict}')
    defect_list = session.query(
        Defect.id, Defect.image_id, Defect.finding_type, Defect.is_manual, Defect.ge_disposition
    ).join(Image, Image.id == Defect.image_id
           ).filter(Image.inspection_id == inspection_id
                        ).all()
    for defect in defect_list:

        # normalize labels
        finding_label = normalize_finding_label(defect.finding_type)

        label_suffix = ""  # separate manual and AI accounts
        if defect.is_manual:
            label_suffix = " (M)"
        else:
            label_suffix = " (AI)"
        finding_label = finding_label.strip() + label_suffix

        if finding_label in count_dict:
            stat = count_dict[finding_label]
            if defect.status == 'Open' and defect.ge_disposition == '':
                stat['Total_Open_No_Disposition'] += 1
            if 'False' in defect.ge_disposition:
                stat['Total_False_Positive'] += 1
            if 'Within' in defect.ge_disposition:
                stat['Total_Within_Tolerance'] += 1
            if 'Out' in defect.ge_disposition:
                stat['Total_Out_of_Tolerance'] += 1
            if defect.status == 'Closed' and 'Out' in defect.ge_disposition:
                stat['Total_Repaired'] += 1
        else:
            logging.info(f'Could not match finding_type: {finding_label}')


# Count defect categories in two groups AI and manual
def inspection_defect_count(inspection_id, count_dict, session):
    # count_dict = DEFECT_TOTAL_TEMPLATE.copy()
    logging.info(f'Counting defects by type for inspection id# {inspection_id}')
    #logging.info(f'Counting defects using count_dict {count_dict}')
    defect_list = session.query(
        Defect.id, Defect.image_id, Defect.finding_type, Defect.is_manual, Defect.ge_disposition, Defect.status
    ).join(Image, Image.id == Defect.image_id
           ).filter(Image.inspection_id == inspection_id
                        ).all()
    for defect in defect_list:

        # First we normalize the finding labels
        finding_label = normalize_finding_label(defect.finding_type)

        # when we find out whether the finding is AI or manual
        label_suffix = ""  
        if defect.is_manual:
            label_suffix = " (M)"
        else:
            label_suffix = " (AI)"
        finding_label = finding_label + label_suffix

        # Increment individual defect type counts e.g. "CoreGap (AI)", "CoreGap (M)", etc.
        # if finding_label in count_dict:
        #     count_dict[finding_label] += 1
        # else:
        #     logging.info(f'Could not match plain finding_type: {finding_label}')

        disposition = get_disposition_key(defect.ge_disposition)
        disposition_finding_label = finding_label + ' - '+disposition
        
        if disposition_finding_label in count_dict:
            count_dict[disposition_finding_label] += 1
            #logging.info(f'Counting disposition_finding_label: {disposition_finding_label}')
        else:
            logging.info(f'Could not match disposition_finding_label: {disposition_finding_label}')


# count the defects of an inspection by disposition type, increment count_dict stats
def inspection_disposition_count(inspection_id, count_dict, session):
    # count_dict = DEFECT_TOTAL_TEMPLATE.copy()
    logging.info(f'Counting defect status by ge_disposition for inspection id# {inspection_id}')
    defect_list = session.query(
        Defect
    ).join(Image, Image.id == Defect.image_id
           ).filter(Image.inspection_id == inspection_id
                        ).all()

    for defect in defect_list:
        disposition_label = defect.ge_disposition
        if disposition_label == '' or disposition_label is None:
            #logging.info(f'Counting ge_disposition: {disposition_label} as "Open" for inspection id# {inspection_id}')
            disposition_label = 'Total_Open_No_Disposition'
        elif 'Out of' in disposition_label:
            if defect.status == 'Closed':
                disposition_label = 'Total_Repaired'
            else:
                disposition_label = 'Total_Out_of_Tolerance'
        elif 'False' in disposition_label:
            disposition_label = 'Total_False_Positive'
        elif 'Within' in disposition_label:
            disposition_label = 'Total_Within_Tolerance'
        elif 'Duplicate' in disposition_label:
            disposition_label = 'Total_Duplicate'

        label_suffix = ""  # separate manual and AI accounts
        if defect.is_manual:
            label_suffix = " (M)"
        else:
            label_suffix = " (AI)"
        disposition_label = disposition_label + label_suffix

        if disposition_label in count_dict:
            count_dict[disposition_label] += 1
        else:
            logging.info(f'Could not match ge_disposition: {disposition_label}')


class DefectStatsCsvAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):
        # self.TASK_EXECUTOR = ThreadPoolExecutor(THREAD_POOL_SIZE)

        self.reqparse = reqparse.RequestParser()

        self.reqparse.add_argument('esn', type=str, location='args')
        self.reqparse.add_argument('offset', type=str, location='args')
        self.reqparse.add_argument('limit', type=str, location='args')
        self.reqparse.add_argument('async', type=bool, location='args')

        super(DefectStatsCsvAPI, self).__init__()

    def __gather_inspection_defect_stats(self, inspection, stats_dict):
        with Session(db) as session:
            # defects within all measurements of a 360 image
            inspection_defect_stats_count(inspection.id, stats_dict, session)

    def __run_generate_defect_csv_report(self, current_app, environ, task_id, is_async, inspection_list):
        with current_app.request_context(environ):
            with tempfile.TemporaryDirectory(dir=UPLOAD_FOLDER, prefix='temp-', suffix='-folder') as tempdirname:

                output_filename = 'defect_stats.csv'

                # we cannot delete the file on exit for async calls,
                # so we save it elsewhere with a unique name
                if is_async:
                    tempdirname = TEMP_DOWNLOAD_FILES_DIR
                    output_filename = task_id+'_'+output_filename

                csv_path = os.path.join(tempdirname, output_filename)

                stats_dict = DEFECT_STATS_TEMPLATE_MANUAL_AI.copy()
                with open(csv_path, 'w') as csvfile:
                    writer = csv.writer(csvfile)
                    defect_headers = ['Defect Type']
                    disposition_headers = list(
                        DISPOSITION_TOTAL_TEMPLATE.keys())
                    headers_with_defects = defect_headers + disposition_headers
                    # writer.writerow(headers_with_defects)

                    futures = []
                    with ThreadPoolExecutor(THREAD_POOL_SIZE) as executor:
                        for inspection in inspection_list:
                            futures.append(executor.submit(
                                self.__gather_inspection_defect_stats, inspection, stats_dict))
                        wait(futures)

                    writer.writerow(headers_with_defects)
                    logging.info('writing csv file rows...')
                    for defect_type in stats_dict.keys():
                        defect_stat = stats_dict[defect_type]
                        row = []
                        row.append(defect_type)
                        for stat_key in DISPOSITION_TOTAL_TEMPLATE:
                            row.append(defect_stat[stat_key])
                        writer.writerow(row)

                if is_async:
                    TASK_STATUS_REGISTRY[task_id] = {
                        'id': task_id,
                        'timestamp': datetime.now().timestamp(),
                        'status': 'COMPLETE',
                        'filename': output_filename,
                        'path': os.path.abspath(csv_path)
                    }
                    logging.info(f'Set TASK_STATUS: for task_id: {task_id} to: {TASK_STATUS_REGISTRY[task_id]}')
                else:
                    logging.info(f'wrote file: {csv_path}')
                    return send_file(os.path.abspath(csv_path),
                                     as_attachment=True,
                                     download_name="defect_stats.csv")

    def get(self):
        """
        Read stats for DB inspections as a CSV file
        ---
        parameters:
          - in: query
            required: false
            name: esn
            schema:
              type: string
            description: filter results by blade serial number
          - in: query
            required: false
            name: limit
            schema:
              type: number
            description: limit number of returned values from offset position
            default: 100000
          - in: query
            required: false
            name: offset
            schema:
              type: number
            description: initial index position for return values
            default: 0
          - in: query
            required: false
            name: async
            schema:
              type: boolean
            description: use asynchronous strategy with background tasks
            default: false

        responses:
          200:
            description: defect stats csv file
            schema:
              id: DefectStatsCSVFile
              type: string
              format: binary

        """
        args = self.reqparse.parse_args()
        logger.info(args.items())

        offset = args['offset']
        limit = args['limit']
        esn = args['esn']

        if offset is None:
            offset = 0
        if limit is None:
            limit = 100000

        # booleans are handled differently in python, so we use string inputs
        async_str = request.args.get('async')
        if async_str is not None:
            async_str = async_str.capitalize()
        logger.info(f'asyncStr: {async_str}')
        is_async = False
        if async_str == 'True':
            is_async = True
        logger.info(f'is_async: {is_async}')

        logger.info(f'offset: {offset}, limit: {limit}, esn: {esn}, async: {is_async}')

        with Session(db) as session:
            inspection_list = []
            if esn is None:
                inspection_list = session.scalars(
                    select(Inspection).offset(offset).limit(limit)
                ).all()
            else:
                inspection_list = session.scalars(
                    select(Inspection).where(
                        Inspection.esn == esn).offset(offset).limit(limit)
                ).all()

            logging.info(f'number of inspections: {len(inspection_list)}')

            task_id = uuid.uuid4().hex
            if is_async:
                TASK_STATUS_REGISTRY[task_id] = {
                    'id': task_id,
                    'timestamp': datetime.now().timestamp(),
                      'status': 'RUNNING',
                      'filename': None,
                }
                TASK_EXECUTOR.submit(self.__run_generate_defect_csv_report, app,
                                     request.environ, task_id, is_async, inspection_list)
                resp =  {'id': task_id,
                        'status': 'RUNNING'}
                logging.info(f'returning: {resp}')
                return resp, 202

            else:
                logging.info('generate and return report file')
                return self.__run_generate_defect_csv_report(app,
                                                             request.environ, task_id, is_async, inspection_list)


def make_blank_row(headers):
    logging.info(f'make_blank_row() called with: {headers}')
    row = len(headers) * ['']
    return row


def add_total(total_row, row):
    logging.info(f'add_todal(): {row}')
    for idx, col in enumerate(row):
        col_str = str(col).strip()
        #logging.info(f'processing col: {col_str}')
        if col_str.isnumeric():
            if total_row[idx] == '':
                total_row[idx] = '0'
            total_row[idx] = str(int(total_row[idx]) + int(col_str))
            #logging.info(f'total col value: {total_row[idx]}')
        else:
            logging.info(f'add_total() ignoring non-numeric col: {col_str}')


class InspectionStatsCsvAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):
        # self.TASK_EXECUTOR = ThreadPoolExecutor(THREAD_POOL_SIZE)

        self.reqparse = reqparse.RequestParser()

        self.reqparse.add_argument('esn', type=str, location='args')
        self.reqparse.add_argument('manufacture_stage', type=str, location='args')
        self.reqparse.add_argument('offset', type=str, location='args')
        self.reqparse.add_argument('limit', type=str, location='args')
        self.reqparse.add_argument('async', type=bool, location='args')

        super(InspectionStatsCsvAPI, self).__init__()

    def __build_row(self, inspection):
        with Session(db) as session:
            img_2d_shots_total = 0
            image_list = session.scalars(
                select(Image).where(
                    Image.inspection_id == inspection.id)
            ).all()

            # empty count dictionaries to be populated for each inspection
            #count_defects_dict = DEFECT_TOTAL_TEMPLATE_MANUAL_AI.copy()
            count_defects_dict = DEFECT_TOTAL_TEMPLATE_MANUAL_AI_DISPOSITION.copy()
            count_disposition_dict = DISPOSITION_TOTAL_TEMPLATE_MANUAL_AI.copy() # manual and AI
            if image_list is not None and len(image_list) > 0:
                for img in image_list:
                    # measurements represent 2d shots. measurement_annotation_file represent findings
                    # defects are parsed out of measuremetn_annotation_file's
                    measurement_results = session.scalars(
                        select(MeasurementImageFile.id).where(
                            MeasurementImageFile.image_id == img.id)
                    ).all()
                    img_2d_shots_total += len(measurement_results)
                    # image_defect_count(img.id, count_dict, session) # defects within all measurements of a 360 image

            # defects within all measurements of a 360 image
            inspection_defect_count(inspection.id, count_defects_dict, session)
            # defects within all measurements of a 360 image
            inspection_disposition_count(inspection.id, count_disposition_dict, session)

            row = [inspection.esn, inspection.blade_type, inspection.factory_name, get_section_name(
                inspection.sect), inspection.manufacture_stage, inspection.upload_date, inspection.date,
                inspection.sso, inspection.status, img_2d_shots_total]
            for key in count_defects_dict.keys():
                row.append(count_defects_dict[key])
            for key in count_disposition_dict.keys():
                row.append(count_disposition_dict[key])
            return row

    def __run_generate_inspection_csv_report(self, current_app, environ, task_id, is_async, inspection_list, esn, manufacture_stage):
        with current_app.request_context(environ):
            with tempfile.TemporaryDirectory(dir=UPLOAD_FOLDER, prefix='temp-', suffix='-folder') as tempdirname:

                output_filename = 'inspection_stats.csv'
                output_cache_filename = str(esn)+'_'+manufacture_stage.replace(' ','_')+'_'+output_filename

                # we cannot delete the file on exit for async calls,
                # so we save it elsewhere with a unique name
                if is_async:
                    tempdirname = TEMP_DOWNLOAD_FILES_DIR
                    output_filename = task_id+'_'+output_filename

                csv_path = os.path.join(tempdirname, output_filename)
                using_cached_result = False

                cache_file_path = os.path.abspath(FILE_CACHE_FOLDER+'/'+output_cache_filename)
                if pathlib.Path(cache_file_path).is_file():
                    cached_file_ts = os.path.getmtime(cache_file_path)
                    logging.info(f'cached_file_ts: {cached_file_ts}s')
                    current_time = time.time()
                    logging.info(f'current time: {current_time}s') 
                    file_age = (current_time - cached_file_ts) 
                    logging.info(f'cache age: {file_age}s')

                    if file_age < FILE_CACHE_TIMEOUT:
                        logging.info('Returing cached file content')
                        logging.info(f'copying: {cache_file_path} to: {os.path.abspath(csv_path)}')
                        shutil.copy(cache_file_path, os.path.abspath(csv_path))
                        using_cached_result = True
                    else:
                        logging.info('Cached file is too old, will re-generate data')
                else:
                    logging.info(f'{cache_file_path} not found')

                if not using_cached_result:
                    logging.info(f'creating csv file: {csv_path}')
                    with open(csv_path, 'w') as csvfile:
                        writer = csv.writer(csvfile)
                        #defect_headers = list(DEFECT_TOTAL_TEMPLATE_MANUAL_AI.keys())
                        defect_headers = list(DEFECT_TOTAL_TEMPLATE_MANUAL_AI_DISPOSITION.keys())
                        disposition_headers = list(
                            DISPOSITION_TOTAL_TEMPLATE_MANUAL_AI.keys())
                        headers = ['Blade ID', 'Blade Type', 'Factory Name', 'Blade Cavity', 'Manufacture Stage',
                                'Upload Date', 'Inspection Date',  'Inspector SSO', 'Annotation Status', '2D Images Count']
                        headers_with_defects = headers + defect_headers + disposition_headers
                        writer.writerow(headers_with_defects)

                        total_row = make_blank_row(headers_with_defects)

                        futures = []
                        with ThreadPoolExecutor(THREAD_POOL_SIZE) as executor:
                            for inspection in inspection_list:
                                futures.append(executor.submit(
                                    self.__build_row, inspection))
                            wait(futures)
                        for future in futures:
                            row = future.result()
                            add_total(total_row, row)
                            logging.info(f'write row: {row}')
                            writer.writerow(row)

                        blank = make_blank_row(headers_with_defects)
                        writer.writerow(blank)
                        logging.info(f'blank_row: {blank}')

                        total_row[0] = "Total"
                        logging.info(f'total_row: {total_row}')
                        writer.writerow(total_row)

                    logging.info(f'Done writing csv file: {csv_path}')
                    # cache results for later use.
                    logging.info(f'creating a cache copy of {os.path.abspath(csv_path)} at: {cache_file_path}')
                    shutil.copy(os.path.abspath(csv_path), cache_file_path)
        

                if is_async:
                    TASK_STATUS_REGISTRY[task_id] = {
                        'id': task_id,
                        'timestamp': datetime.now().timestamp(),
                        'status': 'COMPLETE',
                        'filename': output_filename,
                        'path': os.path.abspath(csv_path)
                    }
                    logging.info(f'Set TASK_STATUS: for task_id: {task_id} to: {TASK_STATUS_REGISTRY[task_id]}')
                else:
                    logging.info(f'wrote file: {csv_path}')
                    return send_file(os.path.abspath(csv_path),
                                     as_attachment=True,
                                     download_name="inspection_stats.csv")
                


    def get(self):
        """
        Read stats for DB inspections as a CSV file
        ---
        parameters:
          - in: query
            required: false
            name: esn
            schema:
              type: string
            description: filter results by blade serial number
          - in: query
            required: false
            name: manufacture_stage
            schema:
              type: string
            description: filter results by blade manufacture stage
          - in: query
            required: false
            name: limit
            schema:
              type: number
            description: limit number of returned values from offset position
            default: 100000
          - in: query
            required: false
            name: offset
            schema:
              type: number
            description: initial index position for return values
            default: 0
          - in: query
            required: false
            name: async
            schema:
              type: boolean
            description: use asynchronous strategy with background tasks
            default: false

        responses:
          200:
            description: image file content for the provided measurement id
            schema:
              id: InspectionListCSVFile
              type: string
              format: binary

        """
        args = self.reqparse.parse_args()
        logger.info(args.items())

        offset = args['offset']
        limit = args['limit']
        esn = args['esn']
        manufacture_stage = args['manufacture_stage']

        if offset is None:
            offset = 0
        if limit is None:
            limit = 100000

        # booleans are handled differently in python, so we use string inputs
        async_str = request.args.get('async')
        if async_str is not None:
            async_str = async_str.capitalize()
        logger.info(f'asyncStr: {async_str}')
        is_async = False
        if async_str == 'True':
            is_async = True
        logger.info(f'is_async: {is_async}')

        if manufacture_stage == None:
            manufacture_stage = '%'
        if esn == None:
            esn = '%'

        logger.info(f'offset: {offset}, limit: {limit}, esn: {esn}, manufactue_stage: {manufacture_stage}, async: {is_async}')


        with Session(db) as session:
            inspection_list = []
            if esn is None:
                inspection_list = session.scalars(
                    select(Inspection).offset(offset).limit(limit)
                ).all()
            else:
                inspection_list = session.scalars(
                    select(Inspection).where(
                        and_(Inspection.esn.ilike(esn), Inspection.manufacture_stage.ilike(manufacture_stage))).offset(offset).limit(limit)
                ).all()

            logging.info(f'number of inspections: {len(inspection_list)}')

            task_id = uuid.uuid4().hex
            if is_async:
                TASK_STATUS_REGISTRY[task_id] = {
                    'id': task_id,
                    'timestamp': datetime.now().timestamp(),
                      'status': 'RUNNING',
                      'filename': None,
                }
                TASK_EXECUTOR.submit(self.__run_generate_inspection_csv_report, app,
                                     request.environ, task_id, is_async, inspection_list, esn, manufacture_stage)
                return {'id': task_id,
                        'status': 'RUNNING'}, 202

            else:
                return self.__run_generate_inspection_csv_report(app,
                                                                 request.environ, task_id, is_async, inspection_list, esn, manufacture_stage)


INSPECTION_ENTRY_TEMPLATE = {
    "id": 10,
    "esn": "",
    "app_type": "",
    "blade_type": "",
    "certificate_id": 0,
    "certification_status": "",
    "customer_name": "",
    "d3_date": "",
    "date": "",
    "disp": "",
    "engine_type": "",
    "factory_name": "",
    "inspector_name": "",
    "location": "",
    "manufacture_date": "",
    "manufacture_stage": "",
    "misc": "",
    "post_molding_date": "",
    "upload_date": "",
    "sect": "",
    "sso": "",
    "status": ""
}


class InspectionListCsvAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):
        self.reqparse = reqparse.RequestParser()

        self.reqparse.add_argument('esn', type=str, location='args')
        self.reqparse.add_argument('offset', type=str, location='args')
        self.reqparse.add_argument('limit', type=str, location='args')
        self.reqparse.add_argument('async', type=bool, location='args')

        super(InspectionListCsvAPI, self).__init__()

    def __run_generate_inspection_csv_report(self, current_app, environ, task_id, is_async, inspection_list):
        with current_app.request_context(environ):

            with tempfile.TemporaryDirectory(dir=UPLOAD_FOLDER, prefix='temp-', suffix='-folder') as tempdirname:
                
                output_filename = 'inspection_list.csv'
            
                # we cannot delete the file on exit for async calls,
                # so we save it elsewhere with a unique name
                if is_async:
                    tempdirname = TEMP_DOWNLOAD_FILES_DIR
                    output_filename = task_id+'_'+output_filename

                csv_path = os.path.join(tempdirname, output_filename)

                with open(csv_path, 'w') as csvfile:
                    writer = csv.writer(csvfile)
                    field_names = list(INSPECTION_ENTRY_TEMPLATE.keys())
                    # added by Habib to replace esn by Blade_Serial_Number
                    field_names_plus_total = field_names.copy()
                    headers = [
                        'Blade_Serial_Number' if field == 'esn'
                        else 'blade_type' if field == 'engine_type'
                        else 'crawler_scan_date' if field == 'post_molding_date'
                        else field
                        for field in field_names]

                    headers.append('number_defects')
                    writer.writerow(headers)

                    with Session(db) as session:
                        for inspection in inspection_list:
                            total_defects = 0
                            defect_list = session.query(
                                Defect
                            ).join(Image, Image.id == Defect.image_id
                                   ).filter(Image.inspection_id == inspection.id
                                                ).all()
                            total_defects += len(defect_list)

                            # image_list = session.scalars(
                            #     select(Image).where(
                            #         Image.inspection_id == inspection.id)
                            # ).all()

                            # if image_list is not None and len(image_list) > 0:
                            #     for img in image_list:
                            #         defect_list_results = session.scalars(
                            #             select(Defect.id).where(
                            #                 Defect.image_id == img.id)
                            #         ).all()
                            #         total_defects += len(defect_list_results)

                            logging.info(f'field_names: {field_names}')

                            values_list = []
                            for field in field_names:
                                values_list.append(getattr(inspection, field))
                            values_list.append(total_defects)
                            writer.writerow(values_list)
                
                if is_async:
                    TASK_STATUS_REGISTRY[task_id] = {
                        'id': task_id,
                        'timestamp': datetime.now().timestamp(),
                        'status': 'COMPLETE',
                        'filename': output_filename,
                        'path': os.path.abspath(csv_path)
                    }
                    logging.info(f'Set TASK_STATUS: for task_id: {task_id} to: {TASK_STATUS_REGISTRY[task_id]}')
                else:
                    logging.info(f'wrote file: {csv_path}')
                    return send_file(os.path.abspath(csv_path),
                                     as_attachment=True,
                                     download_name="inspection_list.csv")
    

    def get(self):
        """
        Read the full list of DB inspections as a CSV file
        ---
        parameters:
          - in: query
            required: false
            name: esn
            schema:
              type: string
            description: filter results by blade serial number
          - in: query
            required: false
            name: limit
            schema:
              type: number
            description: limit number of returned values from offset position
            default: 100000
          - in: query
            required: false
            name: offset
            schema:
              type: number
            description: initial index position for return values
            default: 0

        responses:
          200:
            description: image file content for the provided measurement id
            schema:
              id: InspectionListCSVFile
              type: string
              format: binary

        """
        args = self.reqparse.parse_args()
        logger.info(args.items())

        offset = args['offset']
        limit = args['limit']
        esn = args['esn']

        if offset is None:
            offset = 0
        if limit is None:
            limit = 100000
        # booleans are handled differently in python, so we use string inputs
        async_str = request.args.get('async')
        if async_str is not None:
            async_str = async_str.capitalize()
        logger.info(f'asyncStr: {async_str}')
        is_async = False
        if async_str == 'True':
            is_async = True
        logger.info(f'is_async: {is_async}')
        logger.info(f'offset: {offset}, limit: {limit}, esn: {esn}, async: {is_async}')

        with Session(db) as session:
            inspection_list = []
            if esn is None:
                inspection_list = session.scalars(
                    select(Inspection).order_by(
                        Inspection.esn).offset(offset).limit(limit)
                ).all()
            else:
                inspection_list = session.scalars(
                    select(Inspection).where(
                        Inspection.esn == esn).offset(offset).limit(limit)
                ).all()
            # logging.info(f'inspection_list: {inspection_list}')

            json_list = Inspection.toJsonList(inspection_list)
            #logging.info(f'json_list: {json_list}')
            logging.info(f'number of inspections: {len(inspection_list)}')

            task_id = uuid.uuid4().hex
            if is_async:
                TASK_STATUS_REGISTRY[task_id] = {
                    'id': task_id,
                    'timestamp': datetime.now().timestamp(),
                      'status': 'RUNNING',
                      'filename': None,
                }
                TASK_EXECUTOR.submit(self.__run_generate_inspection_csv_report, app,
                                     request.environ, task_id, is_async, inspection_list)
                return {'id': task_id,
                        'status': 'RUNNING'}, 202

            else:
                return self.__run_generate_inspection_csv_report(app,
                                                                 request.environ, task_id, is_async, inspection_list)


class InspectionListAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):
        self.reqparse = reqparse.RequestParser()

        self.reqparse.add_argument('esn', type=str, location='args')
        self.reqparse.add_argument(
            'upload_begin_date', type=str, location='args')
        self.reqparse.add_argument(
            'upload_end_date', type=str, location='args')

        self.reqparse.add_argument('offset', type=str, location='args')
        self.reqparse.add_argument('limit', type=str, location='args')

        super(InspectionListAPI, self).__init__()

    def get(self):
        """
        Read an existing inspection list
        ---
        parameters:
          - in: query
            required: false
            name: esn
            schema:
              type: string
            description: filter results by blade serial number
          - in: query
            required: false
            name: upload_begin_date
            schema:
              type: string
            description: ISO date string lower bound for upload_date
          - in: query
            required: false
            name: upload_end_date
            schema:
              type: string
            description: ISO date string upper bound for upload_date
          - in: query
            required: false
            name: limit
            schema:
              type: number
            description: limit number of returned values from offset position
            default: 100000
          - in: query
            required: false
            name: offset
            schema:
              type: number
            description: initial index position for return values
            default: 0

        responses:
          200:
            description: Existing inspection metadata
            schema:
              id: InspectionObjectList
              type: array
              items:
                schema:
                  id: InspectionObject

        """
        args = self.reqparse.parse_args()
        logger.info(args.items())

        offset = args['offset']
        limit = args['limit']
        esn = args['esn']
        upload_begin_date = args['upload_begin_date']
        upload_end_date = args['upload_end_date']

        if offset is None:
            offset = 0
        if limit is None:
            limit = 100000
        if upload_begin_date is None:
            upload_begin_date = '1970-01-01'
        if upload_end_date is None:
            day_after_tomorrow = date.today()+timedelta(days=2)
            upload_end_date = str((day_after_tomorrow).isoformat())

        begin_date = None
        end_date = None
        try:
            begin_date = datetime.fromisoformat(upload_begin_date)
        except ValueError:
            return {'message': f'Wrong upload_begin_date format: {upload_begin_date}. Please provide a valid ISO date'}

        try:
            end_date = datetime.fromisoformat(upload_end_date)
        except ValueError:
            return {'message': f'Wrong upload_end_date format: {upload_end_date}. Please provide a valid ISO date'}

        if upload_begin_date > upload_end_date:
            return {'message': f'Wrong time frame: upload_begin_date: {upload_begin_date} comes after upload_end_date: {upload_end_date}'}

        logger.info(f'upload_begin_date: {begin_date}, upload_end_date: {end_date}')
        logger.info(f'offset: {offset}, limit: {limit}, esn prefix: {IDMGroupAPI.esn_prefix}')

        with Session(db) as session:

            # in masda VM or localhost development, where we have no sso authentication, this value is null.
            # also when using the API outside of the Web app, i.e. via importFolder.js command line script
            # which is used by the post-processing pipeline
            if IDMGroupAPI.esn_prefix is None:
                esn_prefix_set = set()
                for value in session.scalars(select(distinct(Inspection.esn))).all():
                    prefix = value.split('-')[0]
                    esn_prefix_set.add(prefix)
                esn_prefix_list = list(esn_prefix_set)
                logging.info(f'esn_prefix_list: {esn_prefix_list}')
                IDMGroupAPI.esn_prefix = esn_prefix_list

            print("filter with esn_prefix ----> ", IDMGroupAPI.esn_prefix)
            inspection_list = []
            if esn is None:
                inspection_list = session.scalars(
                    select(Inspection)
                    .where(or_(*[Inspection.esn.ilike(esn_pre+'-%') for esn_pre in IDMGroupAPI.esn_prefix]))
                    .filter(or_(Inspection.upload_date.between(begin_date, end_date), Inspection.upload_date == None))
                    .offset(offset).limit(limit)
                ).all()
            else:
                logging.info(f'Searching using esn: {esn}')
                inspection_list = session.scalars(
                    select(Inspection)
                    .where(and_(Inspection.esn == esn, or_(*[Inspection.esn.ilike(esn_pre+'-%') for esn_pre in IDMGroupAPI.esn_prefix])))
                    .filter(or_(Inspection.upload_date.between(begin_date, end_date), Inspection.upload_date == None))
                    .offset(offset).limit(limit)
                ).all()
            # logging.info(f'inspection_list: {inspection_list}')
            return Inspection.toJsonList(inspection_list)

        return {'message': 'No inspection records in DB.'}

# when images have no meta-data .json file, they are not imported.
# we can end up with inspections that have no image records on them.


def find_inspections_with_no_images(session):
    inspection_list = session.query(Inspection.id, func.count(Image.id)).\
        join(Image, Image.inspection_id == Inspection.id, isouter=True).\
        group_by(Inspection.id).\
        having(func.count(Image.id) == 0)

    id_list = []
    for inspection in inspection_list:
        id_list.append(inspection.id)
    return id_list


# measuremes have to have image files attached to them
def find_measurements_with_no_image_files(session):
    measurement_ids_used_by_measurement_image_files = session.query(
        MeasurementImageFile.measurement_id)
    measurement_list = session.query(Measurement.id).\
        filter(Measurement.id.notin_(
            measurement_ids_used_by_measurement_image_files))

    id_list = []
    for measurement in measurement_list:
        id_list.append(measurement.id)
    return id_list

# ======================================= Maintancene API =======================================


class InvalidMeasurementAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):
        super(InvalidMeasurementAPI, self).__init__()

    def get(self):
        """
        Read an list of measurements with no images
        ---
        responses:
          200:
            description: List of measurement ids that have no image
            schema:
              properties:
                id_list:
                  type: array
                  items:
                    type: integer


        """
        with Session(db) as session:
            id_list = find_measurements_with_no_image_files(session)
            return jsonify({'id_list': id_list})

    def delete(self):
        """
        Deletes all measurement that have no images
        ---
        responses:
          200:
            description: Success or failure message
            schema:
              properties:
                message:
                  type: string
                id_list:
                  type: array
                  items:
                    type: integer

        """
        with Session(db) as session:

            id_list = find_measurements_with_no_image_files(session)
            msg = ''
            for id in id_list:
                msg += delete_measurement_dependencies(id)
                measurement = session.query(Measurement).get(id)
                session.delete(measurement)
                session.commit()
                logging.info(f'Successfully deleted measurement id# {measurement.id}')

            if len(id_list) == 0:
                msg = 'Found no measurements with 0 images'

            return {'message': msg, 'id_list': id_list}


class InspectionStatsAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):
        super(InspectionStatsAPI, self).__init__()

    def get(self, id):
        """
        Find inspections with no images
        ---
        parameters:
          - in: path
            required: true
            name: id
            schema:
              type: number
            description: Existing inspection id

        responses:
          200:
            description: List of stats for inspection
            schema:
              properties:
                inspection_id:
                  type: integer
                image_count:
                  type: integer
                vtshot_count:
                  type: integer
                measurement_count:
                  type: integer

        """
        image_count = 0
        vtshot_count = 0
        measurement_count = 0

        with Session(db) as session:
            inspection = session.query(Inspection).get(id)
            if inspection is None:
                return {'message': f'Inspection {id} not found.'}

            image_id_list = session.scalars(
                select(Image.id)
                .where(Image.inspection_id == id)
            ).all()

            for image_id in image_id_list:
                # logging.info(f'image_id: {image_id}')
                image_count += 1

                image_measurement_id_list = session.scalars(
                    select(Measurement.id)
                    .where(Measurement.image_id == image_id)
                ).all()

                image_meas_count = len(image_measurement_id_list)
                # logging.info(f'image_meas_count: {str(image_meas_count)}')

                image_vtshot_id_list = session.scalars(
                    select(VTShot.id)
                    .where(VTShot.image_id == image_id)
                ).all()

                image_vtshot_count = len(image_vtshot_id_list)
                # logging.info(f'image_vtshot_count: {str(image_vtshot_count)}')

                measurement_count += image_meas_count
                vtshot_count += image_vtshot_count

            stats = {
                "inspection_id": id,
                "image_count": image_count,
                "vtshot_count": vtshot_count,
                "measurement_count": measurement_count
            }
            logging.info(f'stats: {stats}')
            return stats


# This is used to detect mal-formed inspections in the DB: those with missin image_file records...
class EmptyInspectionAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):

        # Used by the delete method to log who deleted the record
        self.reqparse = reqparse.RequestParser()
        self.reqparse.add_argument('sso', type=str, location='args')

        super(EmptyInspectionAPI, self).__init__()

    def get(self):
        """
        Find inspections with no images
        ---
        responses:
          200:
            description: List of inspections with image record but missing image_file
            schema:
              properties:
                id_list:
                  type: array
                  items:
                    type: integer


        """
        with Session(db) as session:
            id_list = find_inspections_with_no_images(session)
            return jsonify({'id_list': id_list})

    def delete(self):
        """
        Deletes all inspections that have missing image_file records
        ---
        parameters:
          - in: query
            required: false
            name: sso
            schema:
              type: string
        responses:
          200:
            description: Success or failure message
            schema:
              properties:
                message:
                  type: string
                id_list:
                  type: array
                  items:
                    type: integer

        """

        args = self.reqparse.parse_args()
        logger.info(args.items())
        sso = args['sso']

        with Session(db) as session:

            id_list = find_inspections_with_no_images(session)
            msg = ''
            for id in id_list:
                msg += delete_inspection_with_dependencies(id, sso)

            if len(id_list) == 0:
                msg = 'Found no inspections with 0 images'

            return {'message': msg, 'id_list': id_list}


class MissingImageFilesAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):
        super(MissingImageFilesAPI, self).__init__()

    def get(self, id):
        """
        Search for image records with missing image_file for an inspection
        ---
        parameters:
          - in: path
            required: true
            name: id
            schema:
              type: number
            description: Existing inspection id

        responses:
          200:
            description:  Inspection list
            schema:
              id: ImageObjectList
              type: array
              items:
                schema:
                  id: ImageObject

        """
        with Session(db) as session:

            inspection = session.query(Inspection).get(id)
            if inspection is None:
                return {'message': f'Inspection {id} not found.'}

            image_list = session.scalars(
                select(Image)
                .join(ImageFile, ImageFile.image_id == Image.id, isouter=True)
                .where(and_(ImageFile.id == None, Image.inspection_id == id))
            ).all()

            return jsonify(Image.serialize_list(image_list))


class MoveInspectionImageFilesToS3API(Resource):
    # decorators = [auth.login_required]

    def __init__(self):
        super(MoveInspectionImageFilesToS3API, self).__init__()

    def get(self, id):
        """
        Moves all image files from the DB to S3
        ---
        parameters:
          - in: path
            required: true
            name: id
            schema:
              type: number
            description: Existing inspection id

        responses:
          200:
            description: Success or failure message
            schema:
              properties:
                message:
                  type: string
                moved_image_count:
                  type: number
                moved_vtshot_count:
                  type: number
                moved_measurement_count:
                  type: number

        """
        logging.info(f'Move inspection id# {id} images to s3...')
        with Session(db) as session:

            inspection = session.query(Inspection).get(id)
            if inspection is None:
                return {'message': f'Inspection {id} not found.'}
            logging.info('moving inspection: {inspection} images')

            image_list = session.scalars(
                select(Image)
                .where(Image.inspection_id == id)
            ).all()
            logging.info(f'Found {len(image_list)} images.')

            image_filename_dict = {}

            moved_image = 0
            moved_measurement_image = 0
            moved_vtshot_image = 0
            for image in image_list:
                image_file_list = session.scalars(
                    select(ImageFile)
                    .where(ImageFile.image_id == image.id)
                ).all()
                if image_file_list is not None and len(image_file_list) > 0:
                    image_file = image_file_list[0]
                    image_file_content = image_file.content
                    image_filename = image_file.filename

                    # we need this to build s3 keys for vtshots and measurements
                    image_filename_dict[image_file.image_id] = image_filename

                    if image_file_content is not None:
                        logging.info(f'moving image_file id# {image_file.id} to s3')
                        s3key = get_inspection_s3key(
                            inspection)+f'/images/{image_filename}'
                        response = upload_content_to_s3(
                            image_file_content, s3key)
                        logging.info(f'upload to s3 resp: {response}')
                        if response.status_code != 200:
                            return {'message': f'Error moving image_file id# {image_file.id} to s3'}
                        image_file.s3key = s3key
                        image_file.content = None

                        session.add(image_file)
                        session.commit()
                        moved_image += 1
                    else:
                        logging.info(f'skipping image_file id# {image_file.id} due to empty content')
                else:
                    logging.info(
                        f'Error: no ImageFile for image id# {image.id}')

                # --------------- measurements for a 360 image --------------------
                measurement_image_file_list = session.scalars(
                    select(MeasurementImageFile)
                    .where(MeasurementImageFile.image_id == image.id)
                ).all()

                for measurement_image_file in measurement_image_file_list:
                    measurement_image_file_content = measurement_image_file.content
                    measurement_filename = measurement_image_file.filename
                    image_filename = image_filename_dict[measurement_image_file.image_id]
                    measurements_folder_name = image_filename.split('.')[
                        0]+'_measurements'
                    if measurement_image_file_content is not None:
                        logging.info(f'moving measurement_image_file id# {measurement_image_file.id} to s3')
                        s3key = get_inspection_s3key(
                            inspection)+f'/images/{measurements_folder_name}/{measurement_filename}'
                        response = upload_content_to_s3(
                            measurement_image_file_content, s3key)
                        logging.info(f'upload to s3 resp: {response}')
                        if response.status_code != 200:
                            return {'message': f'Error moving measurement_image_file id# {measurement_image_file.id} to s3'}
                        measurement_image_file.s3key = s3key
                        measurement_image_file.content = None

                        session.add(measurement_image_file)
                        session.commit()
                        moved_measurement_image += 1
                    else:
                        logging.info(f'skipping measurement_image_file id# {measurement_image_file.id} due to empty content')

                # --------------- vtshots for a 360 image --------------------
                vtshot_image_file_list = session.scalars(
                    select(VTShotImageFile)
                    .where(VTShotImageFile.image_id == image.id)
                ).all()

                for vtshot_image_file in vtshot_image_file_list:
                    vtshot_image_file_content = vtshot_image_file.content
                    vtshot_filename = vtshot_image_file.filename
                    image_filename = image_filename_dict[vtshot_image_file.image_id]
                    virtualtour_folder_name = image_filename.split('.')[
                        0]+'_virtualtour'

                    if vtshot_image_file_content is not None:
                        logging.info(f'moving vtshot_image_file id# {vtshot_image_file.id} to s3')

                        s3key = get_inspection_s3key(
                            inspection)+f'/images/{virtualtour_folder_name}/{vtshot_filename}'
                        response = upload_content_to_s3(
                            vtshot_image_file_content, s3key)
                        logging.info(f'upload to s3 resp: {response}')
                        if response.status_code != 200:
                            return {'message': f'Error moving vtshot_image_file id# {vtshot_image_file.id} to s3'}
                        vtshot_image_file.s3key = s3key
                        vtshot_image_file.content = None

                        session.add(vtshot_image_file)
                        session.commit()
                        moved_vtshot_image += 1
                    else:
                        logging.info(f'skipping vtshot_image_file id# {vtshot_image_file.id} due to empty content')

            return jsonify({
                'inspection_id': id,
                'moved_image_count': moved_image,
                            'moved_vtshot_count': moved_vtshot_image,
                            'moved_measurement_count': moved_measurement_image
            })


class CompressInspectionImageFilesAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):
        super(CompressInspectionImageFilesAPI, self).__init__()

    def get(self, id):
        """
        Compresses all .png files from an inspection storing them as .jpeg in the DB
        ---
        parameters:
          - in: path
            required: true
            name: id
            schema:
              type: number
            description: Existing inspection id

        responses:
          200:
            description: Success or failure message
            schema:
              properties:
                message:
                  type: string
                compressed_image_count:
                  type: number
                compressed_vtshot_count:
                  type: number
                compressed_measurement_count:
                  type: number

        """
        with Session(db) as session:

            inspection = session.query(Inspection).get(id)
            if inspection is None:
                return {'message': f'Inspection {id} not found.'}

            image_list = session.scalars(
                select(Image)
                .where(Image.inspection_id == id)
            ).all()

            compressed_image = 0
            compressed_measurement_image = 0
            compressed_vtshot_image = 0
            for image in image_list:
                image_file_list = session.scalars(
                    select(ImageFile)
                    .where(ImageFile.image_id == image.id)
                ).all()
                if image_file_list is not None and len(image_file_list) > 0:
                    image_file = image_file_list[0]
                    mime_type = get_content_mime_type(image_file.content)
                    if ('png' in mime_type):
                        logging.info(f'compressing image_file id# {image_file.id}')
                        jpg_content = convert_png_content_to_jpg(
                            image_file.content)
                        image_file.content = jpg_content
                        image_file.filename = image_file.filename.replace(
                            '.png', '.jpg')
                        session.add(image_file)
                        session.commit()
                        compressed_image += 1
                    else:
                        logging.info(f'skipping image_file id# {image_file.id}')
                else:
                    logging.info(
                        f'Error: no ImageFile for image id# {image.id}')

                # --------------- measurements for a 360 image --------------------
                measurement_image_file_list = session.scalars(
                    select(MeasurementImageFile)
                    .where(MeasurementImageFile.image_id == image.id)
                ).all()

                for measurement_image_file in measurement_image_file_list:
                    meas_mime_type = get_content_mime_type(
                        measurement_image_file.content)
                    if ('png' in meas_mime_type):
                        logging.info(f'compressing measurement_image_file id# {measurement_image_file.id}')
                        jpg_content = convert_png_content_to_jpg(
                            measurement_image_file.content)
                        measurement_image_file.content = jpg_content
                        measurement_image_file.filename = measurement_image_file.filename.replace(
                            '.png', '.jpg')
                        session.add(measurement_image_file)
                        session.commit()
                        compressed_measurement_image += 1
                    else:
                        logging.info(f'skipping measurement_image_file id# {measurement_image_file.id}')

                # --------------- vtshots for a 360 image --------------------
                vtshot_image_file_list = session.scalars(
                    select(VTShotImageFile)
                    .where(VTShotImageFile.image_id == image.id)
                ).all()

                for vtshot_image_file in vtshot_image_file_list:
                    vtshot_mime_type = get_content_mime_type(
                        vtshot_image_file.content)
                    if ('png' in vtshot_mime_type):
                        logging.info(f'compressing vtshot_image_file id# {vtshot_image_file.id}')
                        jpg_content = convert_png_content_to_jpg(
                            vtshot_image_file.content)
                        vtshot_image_file.content = jpg_content
                        vtshot_image_file.filename = vtshot_image_file.filename.replace(
                            '.png', '.jpg')
                        session.add(vtshot_image_file)
                        session.commit()
                        compressed_vtshot_image += 1
                    else:
                        logging.info(f'skipping vtshot_image_file id# {vtshot_image_file.id}')

            return jsonify({'compressed_image_count': compressed_image,
                            'compressed_vtshot_count': compressed_vtshot_image,
                            'compressed_measurement_count': compressed_measurement_image
                            })

            # return jsonify(Image.serialize_list(image_list))

# =================================================================================================


class CertificateAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):
        super(CertificateAPI, self).__init__()

    def get(self, id):
        """
        Read an existing certificate
        ---
        parameters:
          - in: path
            required: true
            name: id
            schema:
              type: number
            description: Existing certificate id

        responses:
          200:
            description: Existing certificate
            schema:
              id: CertificateObject


        """
        with Session(db) as session:
            certificate = session.query(Certificate).get(id)
            if certificate is None:
                return {'message': f'Certificate {id} not found.'}

            certificate_json = certificate.toJson()
            return certificate_json

    def delete(self, id):
        """
        Deletes an existing certificate
        ---
        parameters:
          - in: path
            required: true
            name: id
            schema:
              type: number
            description: Existing certificate id

        responses:
          200:
            description: Success or failure message
            schema:
              properties:
                message:
                  type: string

        """
        logger.info(f'DELETE certificate id: {id}')
        msg = ''
        with Session(db) as session:
            certificate = session.query(Certificate).get(id)
            if certificate is None:
                return {'message': f'certificate id {id} not found'}

            # remove inspection table references first
            inspection_list = session.scalars(
                select(Inspection).filter(Inspection.certificate_id == id)
            ).all()
            id_list = []
            for inspection in inspection_list:
                inspection.certificate_id = None
                id_list.append(inspection.id)
            session.commit()

            session.delete(certificate)
            session.commit()
            msg += f'Certificate {id} db entry deleted. Updated inspection records: {id_list}'
            return {'message': msg}

# =================================================================================================


class CreateInspectionLogsAPI(Resource):

    def __init__(self):
        self.reqparse = reqparse.RequestParser()

        self.reqparse.add_argument('inspection_id', type=int, location='json')
        self.reqparse.add_argument('date', type=str, location='json')
        self.reqparse.add_argument('folder_name', type=str, location='json')
        self.reqparse.add_argument('operation', type=str, location='json')
        self.reqparse.add_argument('status', type=str, location='json')
        self.reqparse.add_argument('message', type=str, location='json')
        self.reqparse.add_argument('sso', type=str, location='json')

        super(CreateInspectionLogsAPI, self).__init__()

    def post(self):
        """
        Create an inpsection log entry
        ---
        parameters:
          - name: InspectionLogsBody
            in: body
            required: true
            schema:
              id: InspectionLogsBody
              properties: 
                inspection_id:
                  type: number
                  example: 123
                date:
                  type: string
                  description: Current date with timezone
                folder_name:
                  type: string
                  example: "central_web"
                operation:
                  type: string
                  example: "CREATE"
                status:
                  type: string
                  example: "SUCCESS"
                message:
                  type: string
                  example: "Imported 300 images"
                sso:
                  type: string
                  example: "212336564"

        responses:  
          200:
            description: Created log entry
            schema:
              id: InspectionLogsObject
              properties:
                id:
                  type: number
                inspection_id:
                  type: number
                  example: 123
                date:
                  type: string
                  description: Current iso date with timezone
                  example: "Wed, 19 Jun 2024 21:27:39 GMT"
                folder_name:
                  type: string
                  example: "central_web"
                operation:
                  type: string
                  example: "CREATE"
                status:
                  type: string
                  example: "SUCCESS"
                message:
                  type: string
                  example: "Imported 300 images"
                sso:
                  type: string
                  example: "212336564"

        """
        args = self.reqparse.parse_args()
        logger.info(args.items())

        with Session(db) as session:
            if args['inspection_id'] is not None:
                inspection_id = args['inspection_id']
                inspection = session.query(Inspection).get(inspection_id)
                if inspection is None:
                    return {'message': f'Inspection {inspection_id} not found. Please provide a valid inspection id or null if import was not successful.'}

        log_entry = add_inspection_log(
            inspection_id=args['inspection_id'],
            date=args['date'],
            folder_name=args['folder_name'],
            operation=args['operation'],
            status=args['status'],
            message=args['message'],
            sso=args['sso'])

        logger.info(f'Created inspection log entry: {str(log_entry)}')

        # resp = jsonify(inspection.serialize())
        resp = log_entry.toJson()
        logging.info(f'resp: {resp}')
        return resp


def add_inspection_log(inspection_id, date, folder_name, operation, status, message, sso):
    logging.info(f'add_inspection_log() called with: {inspection_id}, {date}, {folder_name}, {operation}, {status}, {message}, {sso}')
    # Dates are required for inspection, use current date if none is provided
    if date is None or date == '':
        # current_date = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        current_date = datetime.now().isoformat()
        date = current_date

    if sso == '':
        sso = 'unknonw'

    log_entry = None
    with Session(db) as session:
        log_entry = InspectionLogs(
            inspection_id=inspection_id,
            date=date,
            folder_name=folder_name,
            operation=operation,
            status=status,
            message=message,
            sso=sso
        )
        session.add(log_entry)
        session.commit()
        session.refresh(log_entry)  # to read the id

    return log_entry


class InspectionLogsAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):
        self.reqparse = reqparse.RequestParser()

        self.reqparse.add_argument('esn', type=str, location='args')
        self.reqparse.add_argument('offset', type=str, location='args')
        self.reqparse.add_argument('limit', type=str, location='args')

        super(InspectionLogsAPI, self).__init__()

    def get(self):
        """
        Read an existing inspection log list
        ---
        parameters:
          - in: query
            required: false
            name: esn
            schema:
              type: string
            description: filter results by blade serial number
          - in: query
            required: false
            name: limit
            schema:
              type: number
            description: limit number of returned values from offset position
            default: 100000
          - in: query
            required: false
            name: offset
            schema:
              type: number
            description: initial index position for return values
            default: 0

        responses:
          200:
            description: Existing inspection logs list
            schema:
              id: InspectionLogsObjectList
              type: array
              items:
                schema:
                  id: InspectionLogsObject

        """
        args = self.reqparse.parse_args()
        logger.info(args.items())

        offset = args['offset']
        limit = args['limit']
        esn = args['esn']

        if offset is None:
            offset = 0
        if limit is None:
            limit = 100000

        logger.info(f'offset: {offset}, limit: {limit}')
        with Session(db) as session:
            logs_list = []
            if esn is None:
                logs_list = session.scalars(
                    select(InspectionLogs).offset(offset).limit(limit)
                ).all()
            else:
                logs_list = session.query(InspectionLogs).join(
                    Inspection, Inspection.id == InspectionLogs.inspection_id
                ).filter(Inspection.esn.ilike(esn)
                             ).order_by(InspectionLogs.date).desc().all()
            # logging.info(f'inspection_list: {inspection_list}')
            return InspectionLogs.toJsonList(logs_list)

        return {'message': 'No inspection log records in DB.'}


class SearchInspectionAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):
        super(SearchInspectionAPI, self).__init__()

    def get(self):
        """
        Search inspections matching criteria 
        ---
        parameters:
          - in: query
            required: false
            name: esn
            schema:
              type: string
          - in: query
            required: false
            name: sect
            schema:
              type: string
          - in: query
            required: false
            name: date
            schema:
              type: string
          - in: query
            required: false
            name: manufacture_stage

        responses:
          200:
            description:  Inspection list
            schema:
              id: InspectionObjectList
              type: array
              items:
                schema:
                  id: InspectionObject

        """

        # Inspection filter props
        esn = request.args.get('esn', '%')
        sect = request.args.get('sect', '%')

        # Image filter props
        date = request.args.get('date', None)

        # measurement filter props
        manufacture_stage = request.args.get('manufacture_stage', '%')

        logging.info(f'esn: {esn}')
        logging.info(f'date: {date}')
        logging.info(f'sect: {sect}')
        logging.info(f'manufacture_stage: {manufacture_stage}')

        logging.info(f'Search for inspection records...')

        with Session(db) as session:
            inspection_list = []
            if date is not None:
                inspection_list = session.scalars(
                    select(Inspection)
                    .filter(Inspection.esn.like(esn))
                    .filter(Inspection.sect.like(sect))
                    .filter(Inspection.manufacture_stage.like(manufacture_stage))
                    .filter(Inspection.date == func.date(date))
                ).all()
            else:
                logging.info('ignoring date field.')
                inspection_list = session.scalars(
                    select(Inspection)
                    .filter(Inspection.esn.like(esn))
                    .filter(Inspection.sect.like(sect))
                    .filter(Inspection.manufacture_stage.like(manufacture_stage))
                ).all()

            return Inspection.toJsonList(inspection_list)


class InspectionAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):
        self.reqparse = reqparse.RequestParser()
        self.reqparse.add_argument('sso', type=str, location='args')
        super(InspectionAPI, self).__init__()

    def get(self, id):
        """
        Read an existing inspection
        ---
        parameters:
          - in: path
            required: true
            name: id
            schema:
              type: number
            description: Existing inspection id

        responses:
          200:
            description: Existing inspection
            schema:
              id: InspectionObject


        """
        with Session(db) as session:
            inspection = session.query(Inspection).get(id)
            if inspection is None:
                return {'message': f'Inspection {id} not found.'}

            # print(f'inspection upload_date iso: {inspection.upload_date.isoformat()}')

            inspection_json = inspection.toJson()
            return inspection_json

    def delete(self, id):
        """
        Deletes an existing inspection
        ---
        parameters:
          - in: query
            required: false
            name: sso
            schema:
              type: string
          - in: path
            required: true
            name: id
            schema:
              type: number
            description: Existing inspection id

        responses:
          200:
            description: Success or failure message
            schema:
              properties:
                message:
                  type: string

        """

        args = self.reqparse.parse_args()
        logger.info(args.items())
        sso = args['sso']

        logger.info(f'sso: {sso}')

        logger.info(f'DELETE inspection id: {id}')

        msg = delete_inspection_with_dependencies(id, sso)

        invalidate_inspection_cache(id)

        return {'message': msg}


def delete_inspection_with_dependencies(inspection_id, sso=None):
    msg = ''
    with Session(db) as session:
        inspection = session.query(Inspection).get(inspection_id)
        if inspection is None:
            return f'inspection id {inspection_id} not found'

        images_list = session.scalars(
            select(Image).where(
                Image.inspection_id == inspection_id)
        ).all()

        for image in images_list:
            image_id = image.id
            delete_image_dependencies(image_id)

        # logging.info(f'del images_list: {images_list}')
        msg += f'del {len(images_list)} image entries for inspection id: {inspection_id}. '
        for image in images_list:
            session.delete(image)

        session.delete(inspection)
        msg += f'del inspection id:{inspection_id}. '
        session.commit()

        # will handle this at the UI level
        # add_inspection_log(
        #     inspection_id=inspection_id,
        #     date=None,
        #     input_path=None,
        #     operation='DELETE',
        #     status='SUCCESS',
        #     message=f'Inspection {inspection_id} deleted',
        #     sso=sso)

    return msg


def delete_defect_dependencies(defect_id):
    with Session(db) as session: 
        evidence_file_list = session.scalars(
            select(RepairEvidenceFile).where(
                RepairEvidenceFile.defect_id == defect_id)
        ).all()

        if evidence_file_list is not None:
            for evidence_file in evidence_file_list:
                logging.info(
                    f'del evidence_file id# {evidence_file.id}')
                session.delete(evidence_file)
                session.commit()


# delete all measurement dependencies. Note that the measurement itself will not be deleted
def delete_measurement_dependencies(measurement_id):
    msg = ''
    with Session(db) as session:
        measurement_image_file_list = session.scalars(
            select(MeasurementImageFile).where(
                MeasurementImageFile.measurement_id == measurement_id)
        ).all()

        if measurement_image_file_list is not None:
            # logging.info(f'del measurement_image_file_list: {measurement_image_file_list}')
            for measurement_image_file in measurement_image_file_list:
                logging.info(
                    f'del measurement_image_file id {measurement_image_file.id}')
                session.delete(measurement_image_file)
                session.commit()

        # ---------------------------- del measurement_annotation_file -------------------------------
        measurement_annotation_file_list = session.scalars(
            select(MeasurementAnnotationFile).where(
                MeasurementAnnotationFile.measurement_id == measurement_id)
        ).all()

        if measurement_annotation_file_list is not None:
            # logging.info(f'del measurement_annotation_file_list: {measurement_annotation_file_list}')
            for measurement_annotation in measurement_annotation_file_list:
                logging.info(
                    f'del measurement_annotation id {measurement_annotation.id}')
                session.delete(measurement_annotation)
                session.commit()

        # ---------------------------- del defect_annotation_fragment -------------------------------
        defect_annotation_fragment_list = session.scalars(
            select(DefectAnnotationFragment).where(
                DefectAnnotationFragment.measurement_id == measurement_id)
        ).all()

        if defect_annotation_fragment_list is not None:
            # logging.info(f'del measurement_annotation_file_list: {measurement_annotation_file_list}')
            for defect_annotation in defect_annotation_fragment_list:
                logging.info(
                    f'del defect_fragment_annotation id# {defect_annotation.id}')
                session.delete(defect_annotation)
                session.commit()

        # ---------------------------- del defect -------------------------------
        defect_list = session.scalars(
            select(Defect).where(
                Defect.measurement_id == measurement_id)
        ).all()

        if defect_list is not None:
            # logging.info(f'del measurement_annotation_file_list: {measurement_annotation_file_list}')
            for defect in defect_list:
                delete_defect_dependencies(defect.id)
                logging.info(
                    f'del defect id# {defect.id}')
                session.delete(defect)
                session.commit()

        # ---------------------------- del validated_measurement_annotation_file -------------------------------
        validated_measurement_annotation_file_list = session.scalars(
            select(ValidatedMeasurementAnnotationFile).where(
                ValidatedMeasurementAnnotationFile.measurement_id == measurement_id)
        ).all()

        if validated_measurement_annotation_file_list is not None:
            for validated_measurement_annotation in validated_measurement_annotation_file_list:
                logging.info(
                    f'del validated_measurement_annotation id {validated_measurement_annotation.id}')
                session.delete(validated_measurement_annotation)
                session.commit()

        # ---------------------------- del original_measurement_annotation_file -------------------------------
        original_measurement_annotation_file_list = session.scalars(
            select(OriginalMeasurementAnnotationFile).where(
                OriginalMeasurementAnnotationFile.measurement_id == measurement_id)
        ).all()

        if original_measurement_annotation_file_list is not None:
            for original_measurement_annotation in original_measurement_annotation_file_list:
                logging.info(
                    f'del original_measurement_annotation id {original_measurement_annotation.id}')
                session.delete(original_measurement_annotation)
                session.commit()

    return msg


def delete_image_dependencies(image_id):
    msg = ''
    with Session(db) as session:
        image_file_list = session.scalars(
            select(ImageFile).where(
                ImageFile.image_id == image_id)
        ).all()

        # ---------------------------- measurements ---------------------------
        measurement_list = session.scalars(
            select(Measurement).where(
                Measurement.image_id == image_id)
        ).all()

        for measurement in measurement_list:
            measurement_id = measurement.id
            msg += delete_measurement_dependencies(measurement_id)

        # logging.info(f'del measurement_list: {measurement_list}')
        msg += f'del {len(measurement_list)} measurement entries for image_id: {image_id}. '
        for measurement in measurement_list:
            logging.info(f'del measurement id {measurement.id}')
            session.delete(measurement)
            session.commit()

        # ------------------------------- vtshots ----------------------------
        vtshot_list = session.scalars(
            select(VTShot).where(
                VTShot.image_id == image_id)
        ).all()

        for vtshot in vtshot_list:
            vtshot_id = vtshot.id

            vtshot_image_file_list = session.scalars(
                select(VTShotImageFile).where(
                    VTShotImageFile.vtshot_id == vtshot_id)
            ).all()

            if vtshot_image_file_list is not None:
                # logging.info(f'del measurement_image_file_list: {measurement_image_file_list}')
                for vtshot_image_file in vtshot_image_file_list:
                    logging.info(
                        f'del vtshot_image_file id {vtshot_image_file.id}')
                    session.delete(vtshot_image_file)
                    session.commit()

        # logging.info(f'del measurement_list: {measurement_list}')
        msg += f'del {len(vtshot_list)} vtshot entries for image_id: {image_id}. '
        for vtshot in vtshot_list:
            logging.info(f'del vtshot id {vtshot.id}')
            session.delete(vtshot)
            session.commit()

        # ------------------------------- del image file -----------------------

        # logging.info(f'del image_file_list: {image_file_list}')
        msg += f'del {len(image_file_list)} image_file entries from image id: {image_id}. '
        for image_file in image_file_list:
            session.delete(image_file)
            session.commit()

        return msg


def parse_manufacture_stage(input_manufacture_stage):
    parsed_stage = 'Other'
    if input_manufacture_stage is not None:
        stage = input_manufacture_stage.lower()

        if stage is not None:
            if 'post' in stage and 'molding' in stage:
                parsed_stage = 'Post Molding'
            elif 'post' in stage and 'rco' in stage:
                parsed_stage = 'Post Molding_RCO'
            elif 'final' in stage:
                parsed_stage = 'Final Release'
            elif 'post' in stage and 'shipping' in stage:
                parsed_stage = 'Post Shipping'
            elif 'uptower' in stage:
                parsed_stage = 'Uptower'

    logging.info(f'parsing manufacture_stage: {input_manufacture_stage} to: {parsed_stage}')
    return parsed_stage


class CreateInspectionAPI(Resource):

    def __init__(self):
        self.reqparse = reqparse.RequestParser()

        self.reqparse.add_argument('esn', type=str, location='json')
        self.reqparse.add_argument('customer_name', type=str, location='json')
        self.reqparse.add_argument('location', type=str, location='json')
        self.reqparse.add_argument('date', type=str, location='json')
        self.reqparse.add_argument('app_type', type=str, location='json')
        self.reqparse.add_argument('engine_type', type=str, location='json')
        self.reqparse.add_argument('sso', type=str, location='json')
        self.reqparse.add_argument('sect', type=str, location='json')
        self.reqparse.add_argument('disp', type=str, location='json')
        self.reqparse.add_argument('misc', type=str, location='json')
        self.reqparse.add_argument('status', type=str, location='json')
        self.reqparse.add_argument('annotation_status_comments', type=str, location='json')        

        self.reqparse.add_argument('blade_type', type=str, location='json')
        self.reqparse.add_argument(
            'manufacture_date', type=str, location='json')
        self.reqparse.add_argument('factory_name', type=str, location='json')
        self.reqparse.add_argument('inspector_name', type=str, location='json')
        self.reqparse.add_argument(
            'manufacture_stage', type=str, location='json')
        self.reqparse.add_argument(
            'certification_status', type=str, location='json')

        self.reqparse.add_argument('upload_date', type=str, location='json')
        self.reqparse.add_argument(
            'post_molding_date', type=str, location='json')
        self.reqparse.add_argument('d3_date', type=str, location='json')
        self.reqparse.add_argument('certificate_id', type=int, location='json')
        self.reqparse.add_argument('supplier', type=str, location='json')
        self.reqparse.add_argument('priority', type=str, location='json')

        super(CreateInspectionAPI, self).__init__()

    def post(self):
        """
        Create an inspection
        ---
        parameters:
          - name: InspectionBody
            in: body
            required: true
            schema:
              id: InspectionBody
              required:
                - customer_name
                - date
                - sso
                - esn
              properties: 
                esn:
                  type: string
                  example: "J80812"
                customer_name:
                  type: string
                  example: "Garden City Lay Down Yard"
                location:
                  type: string
                  example: "Lay Down Yard"
                date:
                  type: string
                  description: "ISO Date string with optional time"
                  example: "2023-12-12"
                app_type:
                  type: string
                  example: "crawler-thetav"
                engine_type:
                  type: string
                  example: "Blade Crawler-THETA-V"
                sso:
                  type: string
                  example: "212785809"
                sect:
                  type: string
                  example: "te_uw"
                disp:
                  type: string
                  example: "30"
                misc:
                  type: string
                  example: "J80812-07561-W860"
                status:
                  type: string
                  example: "Incomplete"
                blade_type:
                  type: string
                  example: "17m blade"
                manufacture_date:
                  type: string
                  description: ISO date when blade was manufactured
                  example: "2023-12-12"
                factory_name:
                  type: string
                  example: "TPI-MX4"
                inspector_name:
                  type: string
                  example: "Jose da Silva, Mario Soarez"
                manufacture_stage:
                  description: the stage in the factory where this inspection was performed
                  type: string
                  example: "Post-Molding Stage"
                certification_status:
                  type: string
                  example: "Pending QC Inspection / Pending Repair / Pending Repair Approval / Certified"
                upload_date:
                  type: string
                  description: ISO date when blade was manufactured
                  example: "2023-12-12"
                post_molding_date:
                  type: string
                  description: ISO date when blade was manufactured
                  example: "2023-12-12"
                d3_date:
                  type: string
                  description: ISO date when blade was manufactured
                  example: "2023-12-12"
                certificate_id:
                  type: number
                  description: Id of the certificate if it was issued
                supplier:
                  type: string
                  description: TPI or LM
                priority:
                  type: string
                  description: Default is None                  

        produces:
          - application/json
        responses:  
          200:
            description: Created inspection
            schema:
              id: InspectionObject
              properties:
                id:
                  type: number
                esn:
                  type: string
                  example: "J80812"
                customer_name:
                  type: string
                  example: "Garden City Lay Down Yard"
                location:
                  type: string
                  example: "Lay Down Yard"
                date:
                  type: string
                  description: "ISO Date string with optional time"
                  example: "2023-12-12"
                app_type:
                  type: string
                  example: "crawler-thetav"
                engine_type:
                  type: string
                  example: "Blade Crawler-THETA-V"
                sso:
                  type: string
                  example: "212785809"
                sect:
                  type: string
                  example: "te_uw"
                disp:
                  type: string
                  example: "30"
                misc:
                  type: string
                  example: "J80812-07561-W860"
                status:
                  type: string
                  example: "Incomplete"
                blade_type:
                  type: string
                  example: "17m blade"
                manufacture_date:
                  type: string
                  description: ISO date when blade was manufactured
                  example: "2023-12-12"
                factory_name:
                  type: string
                  example: "TPI-MX4"
                inspector_name:
                  type: string
                  example: "Jose da Silva, Mario Soarez"
                manufacture_stage:
                  description: the stage in the factory where this inspection was performed
                  type: string
                  example: "Post-Molding Stage"
                certification_status:
                  type: string
                  example: "Pending QC Inspection / Pending Repair / Pending Repair Approval / Certified"
                upload_date:
                  type: string
                  description: "ISO Date string with optional time"
                  example: "2023-12-12"
                post_molding_date:
                  type: string
                  description: "ISO Date string with optional time"
                  example: "2023-12-12"
                d3_date:
                  type: string
                  description: "ISO Date string with optional time"
                  example: "2023-12-12"
                certificate_id:
                  type: number
                  description: Id of the certificate if it was issued
                supplier:
                  type: string
                  description: name of the supplier TPI or LM
                  example: TPI
                priority:
                  type: string
                  description: priority for blades processing
                  example: High
                  

        """
        args = self.reqparse.parse_args()
        logger.info(args.items())

        # Dates are required for inspection, use current date if none is provided
        if args['date'] is None or args['date'] == '':
            # current_date = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            current_date = datetime.now().isoformat()
            args['date'] = current_date

        if args['manufacture_date'] and args['manufacture_date'] == '':
            args['manufacture_date'] = None

        if args['upload_date'] == None:
            args['upload_date'] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')

        if args['post_molding_date'] and args['post_molding_date'] == '':
            args['post_molding_date'] = None

        if args['d3_date'] and args['d3_date'] == '':
            args['d3_date'] = None

        if args['certificate_id'] and args['certificate_id'] <= 0:
            args['certificate_id'] = None

        inspection = Inspection(
            esn=args['esn'],
            customer_name=args['customer_name'],
            location=args['location'],
            date=args['date'],
            app_type=args['app_type'],
            engine_type=args['engine_type'],
            sso=args['sso'],
            sect=args['sect'],
            disp=args['disp'],
            misc=args['misc'],
            status=args['status'],
            annotation_status_comments=args['annotation_status_comments'],            

            blade_type=args['blade_type'],
            manufacture_date=args['manufacture_date'],
            factory_name=args['factory_name'],
            inspector_name=args['inspector_name'],
            manufacture_stage=parse_manufacture_stage(
                args['manufacture_stage']),
            certification_status=args['certification_status'],

            upload_date=args['upload_date'],
            post_molding_date=args['post_molding_date'],
            d3_date=args['d3_date'],
            certificate_id=args['certificate_id'],
            supplier=args['supplier'],
            priority='None'
        )

        with Session(db) as session:
            session.add(inspection)
            session.commit()
            session.refresh(inspection)  # to read the id

        # We do not update the logs here since creation of record is only the first
        # step in the import process. The next steps will be upload files.
        # only when all files are uploaded, we can update the logs.
        # the import UI function should do it

        # inspection_id = inspection.id
        # add_inspection_log(
        #     inspection_id=inspection_id,
        #     date=args['upload_date'],
        #     input_path=None,
        #     operation='CREATE',
        #     status='SUCCESS',
        #     message=f'Created Inspection id# {inspection.id}',
        #     sso=args['sso'])

        logger.info(f'Created inspection record: {str(inspection)}')

        # resp = jsonify(inspection.serialize())
        resp = inspection.toJson()
        logging.info(f'resp: {resp}')
        return resp


class UpdateInspectionAPI(Resource):

    def __init__(self):
        self.reqparse = reqparse.RequestParser()
        self.reqparse.add_argument('esn', type=str, location='json')
        self.reqparse.add_argument('customer_name', type=str, location='json')
        self.reqparse.add_argument('location', type=str, location='json')
        self.reqparse.add_argument('date', type=str, location='json')
        self.reqparse.add_argument('app_type', type=str, location='json')
        self.reqparse.add_argument('engine_type', type=str, location='json')
        self.reqparse.add_argument('sso', type=str, location='json')
        self.reqparse.add_argument('sect', type=str, location='json')
        self.reqparse.add_argument('disp', type=str, location='json')
        self.reqparse.add_argument('misc', type=str, location='json')
        self.reqparse.add_argument('status', type=str, location='json')
        self.reqparse.add_argument('annotation_status_comments', type=str, location='json')

        self.reqparse.add_argument('blade_type', type=str, location='json')
        self.reqparse.add_argument(
            'manufacture_date', type=str, location='json')
        self.reqparse.add_argument('factory_name', type=str, location='json')
        self.reqparse.add_argument('inspector_name', type=str, location='json')
        self.reqparse.add_argument(
            'manufacture_stage', type=str, location='json')
        self.reqparse.add_argument(
            'certification_status', type=str, location='json')

        self.reqparse.add_argument('upload_date', type=str, location='json')
        self.reqparse.add_argument(
            'post_molding_date', type=str, location='json')
        self.reqparse.add_argument('d3_date', type=str, location='json')
        self.reqparse.add_argument('certificate_id', type=int, location='json')

        self.reqparse.add_argument('priority', type=str, location='json')
        self.bm = UpdateBladeMonitoringAPI()        

        super(UpdateInspectionAPI, self).__init__()

    # Update an existing record
    def post(self, id):
        """
        Update an existing inspection
        ---
        parameters:
          - in: path
            name: id
            schema:
              type: number
            description: Existing inspection id
          - in: body
            name: UpdatedInspectionBody
            description: Object with updated inspection properties
            schema:
              id: InspectionBody

        produces:
          - application/json
        responses:  
          200:
            description: Updated inspection
            schema:
              id: InspectionObject

        """

        args = self.reqparse.parse_args()
        logger.info(f'args: {args.items()}')

        if args['date'] and not is_valid_date(args['date']):
            return {'message': f'Invalid date format: {args["date"]}. Please provide an ISO date string: yyyy-mm-dd'}

        if args['upload_date'] and not is_valid_date(args['upload_date']):
            return {'message': f'Invalid upload_date format: {args["upload_date"]}. Please provide an ISO date string: yyyy-mm-dd'}

        if args['post_molding_date'] and not is_valid_date(args['post_molding_date']):
            return {'message': f'Invalid post_molding_date format: {args["post_molding_date"]}. Please provide an ISO date string: yyyy-mm-dd'}

        if args['d3_date'] and not is_valid_date(args['d3_date']):
            return {'message': f'Invalid d3_date format: {args["d3_date"]}. Please provide an ISO date string: yyyy-mm-dd'}

        if args['certificate_id'] and args['certificate_id'] <= 0:
            args['certificate_id'] = None

        logging.info(f'Looking for existing inspection for id {id}')
        with Session(db) as session:
            inspection = session.query(Inspection).get(id)

            if inspection is None:
                return {'message': f'Inspection {id} not found'}

            if args['esn'] is not None:
                inspection.esn = args['esn']
            if args['customer_name'] is not None:
                inspection.customer_name = args['customer_name']
            if args['location'] is not None:
                inspection.location = args['location']
            if args['date'] is not None and args['date'] != '':
                inspection.date = args['date']
            if args['app_type'] is not None:
                inspection.app_type = args['app_type']
            if args['engine_type'] is not None:
                inspection.engine_type = args['engine_type']
            if args['sso'] is not None:
                inspection.sso = args['sso']
            if args['sect'] is not None:
                inspection.sect = args['sect']
            if args['disp'] is not None:
                inspection.disp = args['disp']
            if args['misc'] is not None:
                inspection.misc = args['misc']
            if args['status'] is not None:
                inspection.status = args['status']
            if args['annotation_status_comments'] is not None:
                inspection.annotation_status_comments = args['annotation_status_comments']


            if args['blade_type'] is not None:
                inspection.blade_type = args['blade_type']
            if args['manufacture_date'] is not None:
                inspection.manufacture_date = args['manufacture_date']
            if args['factory_name'] is not None:
                inspection.factory_name = args['factory_name']
            if args['inspector_name'] is not None:
                inspection.inspector_name = args['inspector_name']
            if args['manufacture_stage'] is not None:
                inspection.manufacture_stage = args['manufacture_stage']
            if args['certification_status'] is not None:
                inspection.certification_status = args['certification_status']

            if args['date'] is not None and args['date'] != '':
                inspection.date = args['date']

            if args['upload_date'] is not None and args['upload_date'] != '':
                inspection.upload_date = args['upload_date']

            if args['post_molding_date'] is not None and args['post_molding_date'] != '':
                inspection.post_molding_date = args['post_molding_date']

            if args['d3_date'] is not None and args['d3_date'] != '':
                inspection.d3_date = args['d3_date']

            if args['certificate_id'] is not None and args['certificate_id'] != '':
                inspection.certificate_id = args['certificate_id']

            if args['priority'] is not None and args['priority'] != '':
                inspection.priority = args['priority']
            # logging.info(f'Updating inspection to: {str(inspection)}')

            session.commit()
            # reload the inspection data from DB
            session.refresh(inspection)
            self.bm.update_annotation(args['esn'])
            # Will handle those in the UI level where we know the sso
            # add_inspection_log(
            #     inspection_id=inspection.id,
            #     date=None,
            #     input_path=None,
            #     operation='UPDATE',
            #     status='SUCCESS',
            #     message=f'Updated Inspection id# {inspection.id}',
            #     sso=args['sso'])

            logger.info(f'Updated inspection: {str(inspection)}')

            # resp = jsonify(inspection.serialize())
            resp = inspection.toJson()
            logging.info(f'resp: {resp}')
            return resp

# ------------------------------------- Certificate --------------------------------


class CertificateListAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):
        self.reqparse = reqparse.RequestParser()

        self.reqparse.add_argument('esn', type=str, location='args')
        self.reqparse.add_argument('offset', type=str, location='args')
        self.reqparse.add_argument('limit', type=str, location='args')

        super(CertificateListAPI, self).__init__()

    def get(self):
        """
        Read an existing certificate list
        ---
        parameters:
          - in: query
            required: false
            name: esn
            schema:
              type: string
            description: filter results by blade serial number
          - in: query
            required: false
            name: limit
            schema:
              type: number
            description: limit number of returned values from offset position
            default: 100000
          - in: query
            required: false
            name: offset
            schema:
              type: number
            description: initial index position for return values
            default: 0

        produces:
          - application/json
        responses:
          200:
            description: Existing certificate list
            schema:
              id: CertificateObjectList
              type: array
              items:
                schema:
                  id: CertificateObject

        """
        args = self.reqparse.parse_args()
        logger.info(args.items())

        offset = args['offset']
        limit = args['limit']
        esn = args['esn']

        if offset is None:
            offset = 0
        if limit is None:
            limit = 100000

        logger.info(f'offset: {offset}, limit: {limit}')
        with Session(db) as session:
            inspection_list = []
            if esn is None:
                inspection_list = session.scalars(
                    select(Certificate).offset(offset).limit(limit)
                ).all()
            else:
                inspection_list = session.scalars(
                    select(Certificate).where(
                        Certificate.blade_serial_number == esn).offset(offset).limit(limit)
                ).all()
            # logging.info(f'inspection_list: {inspection_list}')
            return Certificate.toJsonList(inspection_list)

        return {'message': 'No certificate records in DB.'}


class CertificateAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):
        super(CertificateAPI, self).__init__()

    def get(self, id):
        """
        Read an existing certificate
        ---
        parameters:
          - in: path
            required: true
            name: id
            schema:
              type: number
            description: Existing certificate id

        produces:
          - application/json
        responses:
          200:
            description: Existing certificate
            schema:
              id: CertificateObject


        """
        with Session(db) as session:
            certificate = session.query(Certificate).get(id)
            if certificate is None:
                return {'message': f'Certificate {id} not found.'}

            certificate_json = certificate.toJson()
            return certificate_json

    def delete(self, id):
        """
        Deletes an existing certificate
        ---
        parameters:
          - in: path
            required: true
            name: id
            schema:
              type: number
            description: Existing certificate id

        produces:
          - application/json
        responses:
          200:
            description: Success or failure message
            schema:
              properties:
                message:
                  type: string

        """
        logger.info(f'DELETE certificate id: {id}')
        msg = ''
        with Session(db) as session:
            certificate = session.query(Certificate).get(id)
            if certificate is None:
                return {'message': f'certificate id {id} not found'}

            # remove inspection table references first
            inspection_list = session.scalars(
                select(Inspection).filter(Inspection.certificate_id == id)
            ).all()
            id_list = []
            for inspection in inspection_list:
                inspection.certificate_id = None
                id_list.append(inspection.id)
            session.commit()

            session.delete(certificate)
            session.commit()
            msg += f'Certificate {id} db entry deleted. Updated inspection records: {id_list}'
            return {'message': msg}


class CreateCertificateAPI(Resource):

    def __init__(self):
        self.reqparse = reqparse.RequestParser()

        self.reqparse.add_argument('blade_type', type=str, location='json')
        self.reqparse.add_argument(
            'blade_serial_number', type=str, location='json')
        self.reqparse.add_argument('blade_model', type=str, location='json')
        self.reqparse.add_argument('supplier_name', type=str, location='json')
        self.reqparse.add_argument(
            'factory_location', type=str, location='json')
        self.reqparse.add_argument('factory_name', type=str, location='json')
        self.reqparse.add_argument(
            'manufacture_date', type=str, location='json')
        self.reqparse.add_argument(
            'inspection_modality', type=str, location='json')
        self.reqparse.add_argument(
            'inspection_equipment', type=str, location='json')
        self.reqparse.add_argument(
            'blade_areas_inspected', type=str, location='json')
        self.reqparse.add_argument(
            'inspection_date', type=str, location='json')

        self.reqparse.add_argument('inspector_name', type=str, location='json')
        self.reqparse.add_argument(
            'certification_date', type=str, location='json')
        self.reqparse.add_argument(
            'certificate_number', type=str, location='json')
        self.reqparse.add_argument('reason_for_deviation', type=str, location='json')                    
        self.bm = UpdateBladeMonitoringAPI()   
        super(CreateCertificateAPI, self).__init__()

    def post(self):
        """
        Create a certificate
        ---
        parameters:
          - name: CertificateBody
            in: body
            required: true
            schema:
              id: CertificateBody
              properties: 
                blade_type:
                  type: string
                  example: "68.7m"
                blade_serial_number:
                  type: string
                  example: "TPI-50693"
                blade_model:
                  type: string
                  example: "68.7m"
                supplier_name:
                  type: string
                  example: "TPI"
                factory_location:
                  type: string
                  example: "Mexico, El Paso"
                factory_name:
                  type: string
                  example: "TPI-MX4"
                manufacture_date:
                  type: string
                  example: "2024-02-03"
                inspection_modality:
                  type: string
                  example: "Visual"
                inspection_equipment:
                  type: string
                  example: "Aerones Crawler Gen 2"
                blade_areas_inspected:
                  type: string
                  example: "Trailing_Edge_Internal_Cavity\n Central Web_Internal_Cavity\n Leading_Edge_Internal_Cavity"
                inspection_date:
                  type: string
                  example: "2024-03-01"
                inspector_name:
                  type: string
                  example: "Alan Walker (SSO XXX)"
                manufacture_date:
                  type: string
                  description: ISO date when blade was manufactured
                  example: "2023-12-12"
                factory_name:
                  type: string
                  example: "TPI-MX4"
                inspector_name:
                  type: string
                  example: "Jose da Silva, Mario Soarez"
                certification_date:
                  type: string
                  example: "2024-03-12"
                certificate_number:
                  type: string
                  example: "CER_TPI-50693_1234"

        responses:  
          200:
            description: Created certificate recor
            schema:
              id: CertificateObject
              properties:
                id:
                  type: number
                blade_type:
                  type: string
                  example: "68.7m"
                blade_serial_number:
                  type: string
                  example: "TPI-50693"
                blade_model:
                  type: string
                  example: "68.7m"
                supplier_name:
                  type: string
                  example: "TPI"
                factory_location:
                  type: string
                  example: "Mexico, El Paso"
                factory_name:
                  type: string
                  example: "TPI-MX4"
                manufacture_date:
                  type: string
                  example: "2024-02-03"
                inspection_modality:
                  type: string
                  example: "Visual"
                inspection_equipment:
                  type: string
                  example: "Aerones Crawler Gen 2"
                blade_areas_inspected:
                  type: string
                  example: "Trailing_Edge_Internal_Cavity\n Central Web_Internal_Cavity\n Leading_Edge_Internal_Cavity"
                inspection_date:
                  type: string
                  example: "2024-03-01"
                inspector_name:
                  type: string
                  example: "Alan Walker (SSO XXX)"
                manufacture_date:
                  type: string
                  description: ISO date when blade was manufactured
                  example: "2023-12-12"
                factory_name:
                  type: string
                  example: "TPI-MX4"
                inspector_name:
                  type: string
                  example: "Jose da Silva, Mario Soarez"
                certification_date:
                  type: string
                  example: "2024-03-12"
                certificate_number:
                  type: string
                  example: "CER_TPI-50693_1234"  

        """
        args = self.reqparse.parse_args()
        logger.info(args.items())

        if args['blade_serial_number'] is None or args['blade_serial_number'] == '':
            return {'message': 'Error: a valid blade_serial_number is required'}

        # Dates are required for inspection, use current date if none is provided
        if args['certification_date'] is None or args['certification_date'] == '':
            # current_date = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            current_date = datetime.now().isoformat()
            args['certification_date'] = current_date

        if args['manufacture_date'] == '':
            args['manufacture_date'] = None

        if args['inspection_date'] == '':
            args['inspection_date'] = None

        certificate = Certificate(
            blade_type=args['blade_type'],
            blade_serial_number=args['blade_serial_number'],
            blade_model=args['blade_model'],
            supplier_name=args['supplier_name'],
            factory_location=args['factory_location'],
            factory_name=args['factory_name'],
            manufacture_date=args['manufacture_date'],
            inspection_modality=args['inspection_modality'],
            inspection_equipment=args['inspection_equipment'],
            blade_areas_inspected=args['blade_areas_inspected'],
            inspection_date=args['inspection_date'],

            inspector_name=args['inspector_name'],
            certification_date=args['certification_date'],
            certificate_number=args['certificate_number'],
            reason_for_deviation=args['reason_for_deviation']                       
        )

        with Session(db) as session:
            session.add(certificate)
            session.commit()
            session.refresh(certificate)  # to read the id

            # we need the record id in order to generate the certificate number (if not provided)
            if certificate.certificate_number == '':
                certificate.certificate_number = 'CER_'+ \
                    str(certificate.blade_serial_number)+'_'+str(certificate.id)
                session.add(certificate)
                session.commit()
                session.refresh(certificate)
            
            self.bm.update_certificate(certificate.blade_serial_number)    

                

        logger.info(f'Created certificate: {str(certificate)}')

        # resp = jsonify(inspection.serialize())
        resp = certificate.toJson()
        logging.info(f'resp: {resp}')
        return resp


class UpdateCertificateAPI(Resource):

    def __init__(self):
        self.reqparse = reqparse.RequestParser()

        self.reqparse.add_argument('blade_type', type=str, location='json')
        self.reqparse.add_argument(
            'blade_serial_number', type=str, location='json')
        self.reqparse.add_argument('blade_model', type=str, location='json')
        self.reqparse.add_argument('supplier_name', type=str, location='json')
        self.reqparse.add_argument(
            'factory_location', type=str, location='json')
        self.reqparse.add_argument('factory_name', type=str, location='json')
        self.reqparse.add_argument(
            'manufacture_date', type=str, location='json')
        self.reqparse.add_argument(
            'inspection_modality', type=str, location='json')
        self.reqparse.add_argument(
            'inspection_equipment', type=str, location='json')
        self.reqparse.add_argument(
            'blade_areas_inspected', type=str, location='json')
        self.reqparse.add_argument(
            'inspection_date', type=str, location='json')

        self.reqparse.add_argument('inspector_name', type=str, location='json')
        self.reqparse.add_argument(
            'certification_date', type=str, location='json')
        self.reqparse.add_argument(
            'certificate_number', type=str, location='json')

        self.bm = UpdateBladeMonitoringAPI()   
        super(UpdateCertificateAPI, self).__init__()

    # Update an existing record
    def post(self, id):
        """
        Update an existing certificate
        ---
        parameters:
          - in: path
            name: id
            schema:
              type: number
            description: Existing certificate id
          - in: body
            name: UpdateCertificateBody
            description: Object with updated certificate properties
            schema:
              id: CertificateBody

        responses:  
          200:
            description: Updated certificate
            schema:
              id: CertificateObject

        """

        args = self.reqparse.parse_args()
        logger.info(args.items())

        if args['manufacture_date'] and not is_valid_date(args['manufacture_date']):
            return {'message': 'Invalid manufacture_date format. Please provide an ISO date string: yyyy-mm-dd'}

        if args['inspection_date'] and not is_valid_date(args['inspection_date']):
            return {'message': 'Invalid inspection_date format. Please provide an ISO date string: yyyy-mm-dd'}

        if args['certification_date'] and not is_valid_date(args['certification_date']):
            return {'message': 'Invalid certification_date format. Please provide an ISO date string: yyyy-mm-dd'}

        logging.info(f'Looking for existing certificate for id {id}')
        with Session(db) as session:
            certificate = session.query(Certificate).get(id)

            if certificate is None:
                return {'message': f'Certificate {id} not found'}

            if args['blade_type'] is not None:
                certificate.blade_type = args['blade_type']
            if args['blade_serial_number'] is not None:
                certificate.blade_serial_number = args['blade_serial_number']
            if args['blade_model'] is not None:
                certificate.blade_model = args['blade_model']
            if args['supplier_name'] is not None and args['supplier_name'] != '':
                certificate.supplier_name = args['supplier_name']
            if args['factory_location'] is not None:
                certificate.factory_location = args['factory_location']
            if args['factory_name'] is not None:
                certificate.factory_name = args['factory_name']
            if args['manufacture_date'] is not None:
                certificate.manufacture_date = args['manufacture_date']
            if args['inspection_modality'] is not None:
                certificate.inspection_modality = args['inspection_modality']
            if args['inspection_equipment'] is not None:
                certificate.inspection_equipment = args['inspection_equipment']
            if args['blade_areas_inspected'] is not None is not None:
                certificate.blade_areas_inspected = args['blade_areas_inspected']
            if args['inspection_date'] is not None:
                certificate.inspection_date = args['inspection_date']
            if args['inspector_name'] is not None:
                certificate.inspector_name = args['inspector_name']
            if args['certification_date'] is not None:
                certificate.certification_date = args['certification_date']
            if args['certificate_number'] is not None:
                certificate.certificate_number = args['certificate_number']

            # logging.info(f'Updating certificate to: {str(certificate)}')

            session.commit()
            # reload the inspection data from DB
            session.refresh(certificate)

            self.bm.update_certificate(certificate.blade_serial_number)   

            logger.info(f'Updated certificate: {str(certificate)}')

            # resp = jsonify(inspection.serialize())
            resp = certificate.toJson()
            logging.info(f'resp: {resp}')
            return resp


# ----------------------------- Quality Certificate DOCX and PDF ---------------------------


class CertificateReportDocxAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):
        self.reqparse = reqparse.RequestParser()
        self.reqparse.add_argument('esn', type=str, location='args')
        super(CertificateReportDocxAPI, self).__init__()

    def get(self):
        """
        Generates a docx report with the certificate for an esn
        ---
        parameters:
          - in: query
            required: true 
            name: esn
            schema:
              type: string
            description: Blade serial number ESN
            example: J80812
        consumes:
          - application/json
        produces:
          - application/vnd.openxmlformats-officedocument
        responses:
          200:
            description: docx file with the certificate
            schema:
              type: file
        """
        # return {'message':'Not yet implemented'}
        args = self.reqparse.parse_args()
        logger.info(args.items())

        esn = args['esn']
        logging.info(f'esn: {esn}')

        with tempfile.TemporaryDirectory(dir=UPLOAD_FOLDER, prefix='temp-', suffix='-folder') as tempdirname:
            logging.info(f'created temporary directory {tempdirname}')
            docx_file_path, msg = generate_docx_certificate_report(
                esn, tempdirname)
            if docx_file_path is not None:
                return send_file(os.path.abspath(docx_file_path), as_attachment=True, download_name=f'Certificate-{esn}.docx')
            else:
                return msg


class CertificateReportPdfAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):
        self.reqparse = reqparse.RequestParser()
        self.reqparse.add_argument('esn', type=str, location='args')
        super(CertificateReportPdfAPI, self).__init__()

    def get(self):
        """
        Generates a pdf report with the certificate for an esn
        ---
        parameters:
          - in: query
            required: true 
            name: esn
            schema:
              type: string
            description: Blade serial number ESN
            example: J80812
        consumes:
          - application/json
        produces:
          - application/pdf
        responses:
          200:
            description: pdf file with the certificate
            schema:
              type: file
        """
        # return {'message':'Not yet implemented'}
        args = self.reqparse.parse_args()
        logger.info(args.items())

        esn = args['esn']
        logging.info(f'esn: {esn}')

        with tempfile.TemporaryDirectory(dir=UPLOAD_FOLDER, prefix='temp-', suffix='-folder') as tempdirname:
            logging.info(f'created temporary directory {tempdirname}')
            docx_file_path, msg = generate_docx_certificate_report(
                esn, tempdirname)
            if docx_file_path is not None:
                pdf_file_path = docx_file_path.replace('docx', 'pdf')
                doc2pdf.convert(docx_file_path, pdf_file_path)
                return send_file(os.path.abspath(pdf_file_path), as_attachment=True, download_name=f'Certificate-{esn}.pdf')
            else:
                return msg


def generate_docx_certificate_report(esn, tempdirname):
    # unique_name = str(uuid.uuid4());

    inspection_images_temp_path = tempdirname+'/images'
    logging.info(f'creating images folder: {inspection_images_temp_path}')
    os.makedirs(inspection_images_temp_path, exist_ok=True)

    unique_name = str(uuid.uuid4())
    report_filename = f'Certificate_{esn}_'+unique_name+'.docx'
    logger.info(f'Creating temp docx file: {report_filename}')

    docx_file_path = os.path.join(tempdirname, report_filename)

    with Session(db) as session:
        certificate_list = session.scalars(
            select(Certificate).filter(Certificate.blade_serial_number == esn)
        ).all()

        if certificate_list is None or len(certificate_list) == 0:
            return (None, {'message': f'No certificates for esn: {esn} were found.'})

        # certificate scope persists outside the loop. this will get the latest certificate if there is more than one.
        for certificate in certificate_list:
            logging.info(f'certificate.id: {certificate.id}')

        logging.info(f'selected certificate: {certificate}')

        # Add missing certificate number and save it...
        if certificate.certificate_number == '':
            certificate.certificate_number = 'CER_'+ \
                str(certificate.blade_serial_number)+'_'+str(certificate.id)
            session.add(certificate)
            session.commit()

        # Create a new document
        doc = Document()

        configure_styles(doc)

        # Add a title
        # title = doc.add_heading('Digital Blade Inspection Report', level=1)
        # title = add_heading(doc,  'Digital Blade Inspection Report', 1)
        # title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        add_certificate_report_header(doc)
        add_certificate_report_footer(doc, '')  # empty footer fo now

        # Add a paragraph with bold and italic text
        # paragraph = doc.add_paragraph('This is a sample document created using the python-docx library.')
        # run = paragraph.runs[0]
        # run.bold = True
        # run.italic = True

        # -------------------------------------------- Front Page ------------------------------------
        # Add a heading
        add_heading(doc, '', level=2)
        # first_inspection = inspection_list[0]

        # Note we excluded sect -> Section which we will use to distinguish each inspection
        certificate_table_labels = [
            "Blade Type",
            "Blade Serial Number",
            "Supplier Name",
            "Factory Location",
            "Factory Name",
            "Manufacture Date",
            "Inspection Modality",
            "Inspection Equipment",
            "Blade Areas Inspected",
            "Inspection Date",
            "Inspector Name (SSO)",
            "Certification Date",
            "Certificate Number"]
        certificate_props = [
            "blade_type",
            "blade_serial_number",
            "supplier_name",
            "factory_location",
            "factory_name",
            "manufacture_date",
            "inspection_modality",
            "inspection_equipment",
            "blade_areas_inspected",
            "inspection_date",
            "inspector_name",
            "certification_date",
            "certificate_number"]

        certificate_data_list = []
        for prop in certificate_props:
            # replace the sect property with the list of areas inspected
            att = str(getattr(certificate, prop))
            if 'date' in prop:
                att = att.split(' ')[0]
            certificate_data_list.append(att)

        if certificate.reason_for_deviation:
            certificate_table_labels.append('Reason For Deviation')
            certificate_data_list.append(certificate.reason_for_deviation) 

        now = datetime.now(timezone.utc)
        dt_string = now.strftime("%Y-%m-%d %H:%M:%S UTC")
        certificate_table_labels.append('Report Creation Date')
        certificate_data_list.append(dt_string)


        logging.info(f'certificate_table_labels: {certificate_table_labels}')
        logging.info(f'certificate_data_list: {certificate_data_list}')
        add_data_table(doc, certificate_table_labels, [certificate_data_list],
                       labels_on_top=False, label_width=200, data_width=400)

        add_heading(doc, '', level=2)
        certificate_seal_path = os.path.abspath(
            './src/report/certificate_seal.png')
        doc.add_picture(certificate_seal_path, width=Inches(3))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(docx_file_path)
    return (docx_file_path, {'message': 'document successfully generated'})


def add_certificate_report_header(doc):
    header = doc.sections[0].header
    logo_path = os.path.abspath('./src/report/GE-Vernova-Emblem.png')
    banner_path = os.path.abspath('./src/report/certificate_title.png')
    heading_table = header.add_table(1, 2, Inches(7))
    heading_cells = heading_table.rows[0].cells
    set_table_cell_bg_color(heading_cells[0], COLOR_EVER_GREEN)
    set_table_cell_bg_color(heading_cells[1], COLOR_EVER_GREEN)

    heading_cells[0].width = Inches(6)
    heading_cells[0].text = "\n\n"
    run0 = heading_cells[0].paragraphs[0].runs[0]
    run0.font.size = Pt(5)
    run0.add_picture(banner_path, width=Inches(3))
    heading_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # heading_cells[0].width = Inches(5)
    # heading_cells[0].text = f'Blade Quality Certificate'
    # run0 = heading_cells[0].paragraphs[0].runs[0]
    # run0.font.bold = True
    # run0.font.color.rgb = RGBColor(255, 255, 255)
    # run0.font.name = "Sons Condensed"
    # run0.font.size = Pt(24)
    # # run0.font.name = "Arial Nova"
    # # run0.font.size = Pt(14)
    # heading_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    heading_cells[1].width = Inches(1)
    heading_cells[1].text = "\n\n\n"
    # cell1_para = heading_cells[1].add_paragraph()
    # run1 = cell1_para.add_run()
    run1 = heading_cells[1].paragraphs[0].runs[0]
    run1.font.size = Pt(5)
    run1.add_picture(logo_path, width=Inches(1.25))
    heading_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    header.is_linked_to_previous = False  # only on front page


def add_certificate_report_footer(doc, footer_text=''):
    footer = doc.sections[0].footer
    paragraph = footer.paragraphs[0]
    run = paragraph.add_run(footer_text)
    run.font.color.rgb = RGBColor.from_string(COLOR_EVER_GREEN)
    run.font.size = Pt(10)
    # paragraph.style = document.styles["Footer"]

# --------------------------------------- Upload -----------------------------------


class UploadImageAndMetadataAPI(Resource):

    def __init__(self):
        super(UploadImageAndMetadataAPI, self).__init__()

    def allowed_file(self, filename):
        return '.' in filename and \
            filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

    def post(self, inspection_id, blade_id):
        """
        Uploads files: a metadata_file image_file to an existing inspection id
        ---
        consumes:
          - multipart/form-data
        parameters:
          - in: path
            required: true
            name: inspection_id
            schema:
              type: number
            description: Existing inspection id
          - in: path
            required: true
            name: blade_id
            schema:
              type: number
            description: Existing blade id
          - in: formData
            required: true
            name: image_file
            type: file
            description: image file
          - in: formData
            required: true
            name: metadata_file
            type: file
            description: image metadata json file
        produces:
          - application/json
        responses:
          200:
            description: returns the temp location of the file in the server
            schema:
              id: UploadResponse
              properties:
                message:
                  type: string
                  path: File path in destination

        """

        with Session(db) as session:

            logging.info(f'Inspection id: {inspection_id}')
            inspection = session.query(Inspection).get(inspection_id)
            if inspection is None:
                return {'message': f'Inspection {inspection_id} not found. Existing inspection id required.'}
            logger.info(f'using inspection: {inspection}')

            logging.info(f'Blade id: {blade_id}')
            blade = session.query(Blade).get(blade_id)
            if inspection is None:
                return {'message': f'Blade {blade_id} not found. Existing blade id required.'}
            logger.info(f'using blade: {blade}')

            form = request.form.to_dict()
            logger.info(f'Form data: {form}')

            if 'image_file' not in request.files:
                return {'message': 'File not found or no image_file provided'}

            msg = ''
            files_list = []

            image_id = -1
            image_file_id = -1

            # ---------------------------------- formData metadata_file -----------------------------
            # metadata must be read first, so Image can be created
            # some attributes as area and details are stored in the Image record
            metadata_file = request.files.get("metadata_file")
            image_json = {}
            if metadata_file and self.allowed_file(metadata_file.filename):
                metadata_filename = secure_filename(metadata_file.filename)
                unique_filename = str(uuid.uuid4())+'_'+metadata_filename
                file_path = os.path.join(
                    app.config['UPLOAD_FOLDER'], unique_filename)
                metadata_file.save(file_path)
                msg += f'Metadata file: {metadata_file.filename} uploaded successfully.'
                files_list.append(metadata_file.filename)

                image_json_content = None
                with open(file_path) as f:
                    image_json_content = json.load(f)
                    logging.info(
                        f'metadata_file_content: {image_json_content}')
                if os.path.isfile(file_path):
                    os.remove(file_path)

                if image_json_content['image_ts'] == '' or image_json_content['image_ts'] == 'None':
                    image_json_content['image_ts'] = None

                # frame is uded for measurement and is optional
                if not 'frame' in image_json_content:
                    image_json_content['frame'] = None

                new_image_rec = Image(
                    timestamp=image_json_content['image_ts'],
                    distance=image_json_content['image_distance'],
                    defect_severity=image_json_content['defect_severity'],
                    defect_location=image_json_content['defect_location'],
                    defect_size=image_json_content['defect_size'],
                    defect_desc=image_json_content['defect_desc'],
                    frame=None if (image_json_content['frame'] == None or image_json_content['frame'] == 'null') else json.dumps(
                        image_json_content['frame']),
                    inspection_id=inspection_id,
                    blade_id=blade_id)
                session.add(new_image_rec)
                session.commit()
                session.refresh(new_image_rec)  # to read the id
                image_id = new_image_rec.id
                logging.info(f'new_image_rec: {new_image_rec}')

                image_json = new_image_rec.toJson()
                logging.info(f'image_json: {image_json}')

            else:
                msg += 'Image Metadata file not found, or no image_metadata file content provided.'
                return {'message': msg}

            # -------------------------------- formData image_file ---------------------------------
            # Then we read the image_file associated with the image record
            image_file = request.files.get("image_file")
            image_file_json = {}
            if image_file and self.allowed_file(image_file.filename):
                image_filename = secure_filename(image_file.filename)
                unique_filename = str(uuid.uuid4())+'_'+image_filename
                image_file_path = os.path.join(
                    app.config['UPLOAD_FOLDER'], unique_filename)
                image_file.save(image_file_path)

                if image_id > 0:
                    logging.info(f'Processing image_file: {image_filename}')
                    bin_content = None
                    with open(image_file_path, mode="rb") as f:
                        bin_content = f.read()
                    thumb_bin_content = get_thumbnail_content(image_file_path)
                    if os.path.isfile(image_file_path):
                        os.remove(image_file_path)

                    if COMPRESS_INCOMING_IMAGES:
                        image_filename = image_filename.replace('.png', '.jpg')
                        bin_content = convert_png_content_to_jpg(bin_content)

                    new_image_file_rec = ImageFile(
                        image_id=image_id,
                        filename=image_filename,
                        content=bin_content,
                        thumbnail=thumb_bin_content)

                    if USE_S3:
                        s3key = get_inspection_s3key(
                            inspection)+f'/images/{image_filename}'
                        response = upload_content_to_s3(bin_content, s3key)
                        logging.info(f'upload to s3 resp: {response}')
                        if response.status_code != 200:
                            return {'message': f'Error uploading file: {image_filename} to S3'}
                        new_image_file_rec.s3key = s3key
                        new_image_file_rec.content = None

                    session.add(new_image_file_rec)
                    session.commit()
                    session.refresh(new_image_file_rec)
                    image_file_id = new_image_file_rec.id

                    # set values to null so we can print the log
                    new_image_file_rec.content = None
                    new_image_file_rec.thumbnail = None
                    logging.info(f'new_image_file_rec: {new_image_file_rec}')

                    image_file_json = new_image_file_rec.toJson()
                    logging.info(f'image_file_json: {image_file_json}')

                msg += f'Image file {image_file.filename} uploaded successfully. '
                files_list.append(image_file.filename)
            else:
                msg += 'Image file not found, or no image_file provided.'

            resp = {'message': msg,
                    'image_id': image_id,
                    'image_file_id': image_file_id,
                    'files_list': files_list}

            logging.info(f'resp: {resp}')

            return make_response(jsonify(resp), 200)

# ------------------------- Monitoring ---------------------------------------------------


class BladeMonitoringListAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):
        super(BladeMonitoringListAPI, self).__init__()

    def get(self):
        """
        List all the blades for monitoring
        ---
        produces:
          - application/json
        responses:
          200:
            description: List of blades for monitoring
            schema:
              properties:
                id:
                  type: integer
                blade_number:
                  type: string
                s3_bucket:
                  type: string
                s3_cvpl_input:
                  type: string

        """
        with Session(db) as session:
            if IDMGroupAPI.esn_prefix and all('tpi'.upper() in prfx for prfx in map(str.upper, IDMGroupAPI.esn_prefix)):
                bladeList = session.scalars(select(BladeMonitoring)
                 .filter(BladeMonitoring.blade_number.ilike('tpi-%'))
                 .order_by(BladeMonitoring.s3_cvpl_output.desc())).all()
            else:
                bladeList = session.scalars(
                    select(BladeMonitoring).order_by(BladeMonitoring.s3_cvpl_output.desc())).all()
            if bladeList is None:
                return {'message': f'No Blades in monitoring.'}
            
            return BladeMonitoring.toJsonList(bladeList)

# ----------------------- S3 Folder Upload APIs ----------------------- 
class S3UploadManager:
    def __init__(self):
        self.config = self.load_config()
        self.s3_client = self.create_s3_client()

    def load_config(self):
        return {
            'region': os.getenv('AWS_REGION'),
            'bucket_name': os.getenv('S3_BUCKET_NAME'),
            's3_base_folder': os.getenv('S3_BASE_FOLDER'),
            'kms_key_id': os.getenv('KMS_KEY_ID')
        }
        # Use below code for testing on local machine
        # return {
        #     'aws_access_key': os.getenv('AWS_ACCESS_KEY_ID'),
        #     'aws_secret_key': os.getenv('AWS_SECRET_ACCESS_KEY'),
        #     'region': os.getenv('AWS_REGION'),
        #     'bucket_name': os.getenv('S3_BUCKET_NAME'),
        #     's3_base_folder': os.getenv('S3_BASE_FOLDER'),
        #     'kms_key_id': os.getenv('KMS_KEY_ID')
        # }

    def create_s3_client(self):
        session = botocore.session.get_session()
        return session.create_client(
            's3',
            region_name=self.config['region'],
            config=Config(
                signature_version='s3v4',
                s3={'use_accelerate_endpoint': True}
            )
        )
        # #Use below section for testing on local machine
        # ca_bundle_path = os.getenv('CA_BUNDLE_PATH')
        # session.set_config_variable('ca_bundle', ca_bundle_path)
        # proxy_definitions = {
        #    'http': os.getenv('CA_HTTP_PROXY'),
        #    'https': os.getenv('CA_HTTP_PROXY')
        #   }
        # return session.create_client(
        #     's3',
        #     aws_access_key_id=self.config['aws_access_key'],
        #     aws_secret_access_key=self.config['aws_secret_key'],
        #     region_name=self.config['region'],
        #     config=Config(
        #         signature_version='s3v4',
        #         proxies=proxy_definitions,
        #         proxies_config={'proxy_ca_bundle': ca_bundle_path},
        #         s3={'use_accelerate_endpoint': True}
        #     )
        # )

class S3InitiateUploadAPI(Resource):
    def __init__(self):
        super(S3InitiateUploadAPI, self).__init__()
        self.s3_manager = S3UploadManager()
        

    def get_full_path(self, base_path, object_name, folderDate):
        # current_date = date.today().strftime("%Y-%m-%d")
        subfolder = "tpi" if 'TPI' in object_name else "lm"
        return f"{base_path}/{subfolder}/{folderDate}/{object_name}"

    def post(self):
        """
        Initiate a multipart upload for folder upload to S3
        ---
        parameters:
          - in: body
            name: body
            schema:
              type: object
              required:
                - object_name
                - file_size
              properties:
                object_name:
                  type: string
                file_size:
                  type: integer
        responses:
          200:
            description: Multipart upload initiated successfully
            schema:
              properties:
                upload_id:
                  type: string
                object_name:
                  type: string
          400:
            description: Missing required parameters
          500:
            description: Internal server error
        """
        data = request.json
        object_name = data.get('object_name')
        file_size = data.get('file_size')
        folderDate = data.get('folderDate')

        if not object_name or file_size is None:
            return {'error': 'object_name and file_size are required'}, 400
        logger.info(f"Initiating multipart upload: object_name: {object_name}, file_size: {file_size}")

        s3_base_path = self.s3_manager.config['s3_base_folder'].strip('/')
        object_name = self.get_full_path(s3_base_path, object_name, folderDate).replace("\\", "/")
        try:
            response = self.s3_manager.s3_client.create_multipart_upload(
                Bucket=self.s3_manager.config['bucket_name'],
                Key=object_name,
                ServerSideEncryption='aws:kms',
                SSEKMSKeyId=self.s3_manager.config['kms_key_id']
            )
            return {'upload_id': response['UploadId'], 'object_name': object_name}
        except Exception as e:
            logger.error(f"Error initiating multipart upload: {str(e)}", exc_info=True)
            return {'status': 'error', 'message': str(e)}, 500

class S3PresignedUrlsAPI(Resource):
    def __init__(self):
        super(S3PresignedUrlsAPI, self).__init__()
        self.s3_manager = S3UploadManager()
        self.bm = UpdateBladeMonitoringAPI()              

    def post(self):
        """
        Get presigned URLs for multipart upload
        ---
        parameters:
          - in: body
            name: body
            schema:
              type: object
              required:
                - object_name
                - upload_id
                - parts
              properties:
                object_name:
                  type: string
                upload_id:
                  type: string
                parts:
                  type: integer
        responses:
          200:
            description: Presigned URLs generated successfully
            schema:
              properties:
                presigned_urls:
                  type: array
                  items:
                    type: object
                    properties:
                      PartNumber:
                        type: integer
                      PresignedUrl:
                        type: string
          400:
            description: Missing required parameters
          500:
            description: Internal server error
        """
        data = request.json
        object_name = data.get('object_name')
        upload_id = data.get('upload_id')
        parts = data.get('parts')

        if not all([object_name, upload_id, parts]):
            return {'error': 'Missing required parameters'}, 400
        logger.info(f"Generating presigned URLs: object_name: {object_name}, upload_id: {upload_id}, parts: {parts}")
        self.bm.update_blade_upload(object_name,'In Progress')
        try:
            presigned_urls = [
                {
                    "PartNumber": part_number,
                    "PresignedUrl": self.s3_manager.s3_client.generate_presigned_url(
                        'upload_part',
                        Params={
                            'Bucket': self.s3_manager.config['bucket_name'],
                            'Key': object_name,
                            'UploadId': upload_id,
                            'PartNumber': part_number
                        },
                        ExpiresIn=3600
                    )
                }
                for part_number in range(1, parts + 1)
            ]
            return {'presigned_urls': presigned_urls}
        except Exception as e:
            logger.error(f"Error generating presigned URLs: {str(e)}", exc_info=True)
            return {'status': 'error', 'message': str(e)}, 500

class S3CompleteUploadAPI(Resource):
    def __init__(self):
        super(S3CompleteUploadAPI, self).__init__()
        self.s3_manager = S3UploadManager()
        self.bm = UpdateBladeMonitoringAPI()

    def post(self):
        """
        Complete a multipart upload
        ---
        parameters:
          - in: body
            name: body
            schema:
              type: object
              required:
                - object_name
                - upload_id
                - parts
              properties:
                object_name:
                  type: string
                upload_id:
                  type: string
                parts:
                  type: array
                  items:
                    type: object
                    properties:
                      PartNumber:
                        type: integer
                      ETag:
                        type: string
        responses:
          200:
            description: Multipart upload completed successfully
            schema:
              properties:
                status:
                  type: string
                response:
                  type: string
          400:
            description: Missing required parameters
          500:
            description: Internal server error
        """
        data = request.json
        object_name = data.get('object_name')
        upload_id = data.get('upload_id')
        parts = data.get('parts')

        if not all([object_name, upload_id, parts]):
            return {'error': 'Missing required parameters'}, 400
        logger.info(f"Completing multipart upload: object_name: {object_name}, upload_id: {upload_id}, parts: {len(parts)}")

        try:
            response = self.s3_manager.s3_client.complete_multipart_upload(
                Bucket=self.s3_manager.config['bucket_name'],
                Key=object_name,
                UploadId=upload_id,
                MultipartUpload={'Parts': parts}
            )
            self.bm.update_blade_upload(object_name,'Complete')
            return {'status': 'success', 'response': str(response)}
        except Exception as e:
            logger.error(f"Error completing multipart upload: {str(e)}", exc_info=True)
            return {'status': 'error', 'message': str(e)}, 500

class S3AbortUploadAPI(Resource):
    def __init__(self):
        super(S3AbortUploadAPI, self).__init__()
        self.s3_manager = S3UploadManager()
        self.bm = UpdateBladeMonitoringAPI()        

    def post(self):
        """
        Abort a multipart upload
        ---
        parameters:
          - in: body
            name: body
            schema:
              type: object
              required:
                - object_name
                - upload_id
              properties:
                object_name:
                  type: string
                upload_id:
                  type: string
        responses:
          200:
            description: Multipart upload aborted successfully
            schema:
              properties:
                status:
                  type: string
                message:
                  type: string
          400:
            description: Missing required parameters
          500:
            description: Internal server error
        """
        data = request.json
        object_name = data.get('object_name')
        upload_id = data.get('upload_id')

        if not all([object_name, upload_id]):
            return {'error': 'Missing required parameters'}, 400

        try:
            self.s3_manager.s3_client.abort_multipart_upload(
                Bucket=self.s3_manager.config['bucket_name'],
                Key=object_name,
                UploadId=upload_id
            )
            self.bm.update_blade_upload(object_name,'Aborted')            
            return {'status': 'success', 'message': 'Multipart upload aborted successfully'}
        except Exception as e:
            logger.error(f"Error aborting multipart upload: {str(e)}", exc_info=True)
            return {'status': 'error', 'message': str(e)}, 500

class UpdateBladeMonitoringAPI(Resource):

    def __init__(self):
        super(UpdateBladeMonitoringAPI, self).__init__()

    # Update an existing record
    def update_blade_upload(self, number, status ):
        """
        Update an existing inspection
        ---
        parameters:
          - in: path
            name: id
            schema:
              type: number
            description: Existing inspection id
          - in: body
            name: UpdatedInspectionBody
            description: Object with updated inspection properties
            schema:
              id: InspectionBody

        produces:
          - application/json
        responses:  
          200:
            description: Updated inspection
            schema:
              id: InspectionObject

        """

        blade_number = number
        folders = number.split("/")
        date_pattern = r"\d{4}-\d{2}-\d{2}"
        logging.info(f'Looking for existing blade for number {blade_number}')

        # Find date folder and the folder after the date
        for i, folder in enumerate(folders):
            if re.match(date_pattern, folder):
                if i + 1 < len(folders):
                    blade_number = folders[i + 1]      
        
        if '_' in blade_number:
            blade_number = blade_number.split('_')[1]
        
        print('blade_number ', blade_number)
        with Session(db) as session:
            blade = session.scalars(
              select(BladeMonitoring).filter(BladeMonitoring.blade_number == blade_number)).all()            
            current_date = datetime.now().strftime("%Y-%m-%d, %H:%M:%S")
            if not blade: ## make a new entry
                blade = BladeMonitoring(
                    blade_number = blade_number,
                    upload_date = current_date,
                    upload_status = status
                )
                session.add(blade)
                session.commit()                                                               
            else:
                session.query(BladeMonitoring).filter(BladeMonitoring.blade_number == blade_number).update({'upload_status': status ,'upload_end_date':current_date })

            session.commit()
            logger.info(f'Updated BladeMonitoring update_blade_upload: {str(blade_number)}')
            resp = blade_number
            logging.info(f'resp: {resp}')
            return resp

    def update_annotation(self, number ):
        """
        Update an existing inspection
        ---
        parameters:
          - in: path
            name: id
            schema:
              type: number
            description: Existing inspection id
          - in: body
            name: UpdatedInspectionBody
            description: Object with updated inspection properties
            schema:
              id: InspectionBody

        produces:
          - application/json
        responses:  
          200:
            description: Updated inspection
            schema:
              id: InspectionObject

        """

        blade_number = number
        
        print('blade_number ', blade_number)
        with Session(db) as session:
            blade = session.scalars(
              select(BladeMonitoring).filter(BladeMonitoring.blade_number == blade_number)).all()            
            current_datetime = datetime.now()

            Inspection_list = session.scalars(
              select(Inspection).filter(Inspection.esn == blade_number)).all()
            inspection_status = []
            
            if Inspection_list is not None:
                for inspection in Inspection_list:
                    inspection_status.append(inspection.status)

                print('inspection_status -- ',inspection_status)    

                if all('Complete' == status for status in inspection_status):
                    print('All are complete')
                    session.query(BladeMonitoring).filter(BladeMonitoring.blade_number == blade_number).update({'annot_end_time':current_datetime })

            session.commit()
            logger.info(f'Updated BladeMonitoring: {str(blade_number)}')
            resp = blade_number
            logging.info(f'resp: {resp}')
            return resp

    def update_certificate(self, number ):
        """
        Update an existing inspection
        ---
        parameters:
          - in: path
            name: id
            schema:
              type: number
            description: Existing inspection id
          - in: body
            name: UpdatedInspectionBody
            description: Object with updated inspection properties
            schema:
              id: InspectionBody

        produces:
          - application/json
        responses:  
          200:
            description: Updated inspection
            schema:
              id: InspectionObject

        """

        blade_number = number
        
        print('blade_number ', blade_number)
        with Session(db) as session:
            blade = session.scalars(
              select(BladeMonitoring).filter(BladeMonitoring.blade_number == blade_number)).all()            
            
            current_datetime = datetime.now()
            if blade:
                session.query(BladeMonitoring).filter(BladeMonitoring.blade_number == blade_number).update({'cert_issued':current_datetime })

            session.commit()
            logger.info(f'Updated BladeMonitoring cert_issued : {str(blade_number)}')
            resp = blade_number
            logging.info(f'resp: {resp}')
            return resp

    def post(self):
        data = request.json
        blade_number = data.get('blade_number')
        status = data.get('status')
        column_name = data.get('column_name')
        date = data.get('date')
        
        if not (blade_number or column_name):
            return {'error': 'blade_number and column_name are required'}, 400

        # Current datetime for application_ui
        current_datetime = date if date else datetime.now()

        try:
            with Session(db) as session:
                # Find the existing blade monitoring record by blade_number
                blade = session.scalars(select(BladeMonitoring).filter_by(blade_number=blade_number)).first()

                if not blade:
                    logging.info(f'No existing blade found for blade_number: {blade_number}')
                else:
                    # Update the existing blade monitoring record
                    logging.info(f'Updating existing blade record for blade_number: {blade_number}')
                    setattr(blade, column_name, current_datetime)
                    if column_name == 'application_ui' and blade.annot_start_time is None:
                        session.query(BladeMonitoring).filter(BladeMonitoring.blade_number == blade_number).update({'annot_start_time':current_datetime })

                session.commit()

                logging.info(f'Blade monitoring updated successfully for blade_number: {blade_number}')
            return {'blade_number': blade_number}
        except Exception as e:
            logging.error(f'Error updating blade monitoring for blade_number: {blade_number}')
            logging.error(f'Error : {e}')
            return {'error': 'Error updating blade monitoring'}, 500
        
class ValidateROSAPI(Resource):
    def __init__(self):
        super(ValidateROSAPI, self).__init__()

    def post(self):
        """
        Validate ROS bag files
        ---
        parameters:
          - in: formData
            name: files
            type: array
            items:
              type: file
            required: true
            description: ROS bag files to validate
        responses:
          200:
            description: ROS bag files validated successfully
            schema:
              properties:
                status:
                  type: string
                message:
                  type: string
          400:
            description: Missing required files or validation failed
          500:
            description: Internal server error
        """
        if 'files' not in request.files:
            return {'status': 'error', 'message': 'No files part in the request'}, 400

        files = request.files.getlist('files')
        cavityCount = request.form.get('cavityCount')

        logger.info(f'Validating {len(files)} files')
        
        if not files:
            return {'status': 'error', 'message': 'No files sent in the API request payload'}, 400

        processor = ros_bag_processor.ROSBagProcessor()

        try:
            # Use a temporary directory to process the bag files
            with tempfile.TemporaryDirectory() as temp_dir:
                rosbag_paths = []  # List to store paths of the saved rosbag files

                # Iterate through each file and save it to the temp_dir
                for file in files:
                    if file.filename == '':  # Skip empty filename uploads
                        continue
                    
                    # Secure the filename and save the file to the temp directory
                    logger.info(f"filename before secure_filename: {file.filename}")
                    filename = secure_filename(file.filename)
                    logger.info(f"filename after secure_filename: {filename}")
                    file_path = os.path.join(temp_dir, filename)
                    file.save(file_path)

                    # Add the saved file's path to the rosbag_paths list
                    if(file_path.endswith('.db3')):
                        filepath_without_extension = file_path[:-4]
                        rosbag_paths.append(filepath_without_extension)

                # Prepare and process ROS bag files in the temp directory
                processor.process_rosbag(temp_dir)

                logger.info(f"rosbag_paths: {rosbag_paths}")
                # If only one file is uploaded, no merging is needed
                if len(rosbag_paths) == 1:
                    logger.info(f"Only one file uploaded, returning: {rosbag_paths[0]}")
                    # Validate the ROS files inside temp_dir
                    if processor.rosbag_validity_check(temp_dir, cavityCount):
                        logger.info(f'No Validation Errors')
                        return {'errors': []}, 200
                    else:
                        logger.info(f'Validation Error found')
                        return {'errors': ["Rosbag files does not contain all three relevant sections, or has extra sections"]}, 200
                else :
                    return {'errors': []}, 200
                # To be used in future- If more than one file, proceed to merge them and create a new bag file
                    # logger.info(f'Multiple ros bag files uploaded, merging and saving to: {temp_dir}/merged_rosbag2.db3')
                    # output_rosbag_path = os.path.join(temp_dir, 'merged_rosbag2.db3')
                    # processor.merge_rosbags(rosbag_paths, output_rosbag_path)

        except Exception as e:
            logger.error(f"Error during rosbag files validation api call: {str(e)}", exc_info=True)
            return {'status': 'error', 'message': f'Error during validation: {str(e)}'}, 500

#------------------------------------Blade Defect Model AI List ------------------------------------
class BladeDefectModelListAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):
        super(BladeDefectModelListAPI, self).__init__()

    def get(self):
        """
        List the AI Models run by Power Sage Pipeline.
        ---
        produces:
          - application/json
        responses:
          200:
            description: List of blades for monitoring
            schema:
              properties:
                id:
                  type: integer
                blade_type_id:
                  type: integer
                defect_model_id:
                  type: integer
                ai_enabled:
                  type: boolean

        """
        with Session(db) as session:
            aiDefectModel = session.scalars(select(BladeDefectModel).order_by(BladeDefectModel.blade_type_id, BladeDefectModel.defect_model_id)).all()
            if aiDefectModel is None:
                return {'message': f'No AI Models listed.'}
            
            unique_bladeType = set()
            for blade in aiDefectModel:
                if blade.blade_type_id not in unique_bladeType:
                    unique_bladeType.add(blade.blade_type_id)

            aiDefectModel_json = serialize_BladeDefectModelList(aiDefectModel)            
            return aiDefectModel_json

class BladeDefectModelS3UploadAPI(Resource):
    # decorators = [auth.login_required]

    def __init__(self):
        super(BladeDefectModelS3UploadAPI, self).__init__()
        #self.s3_manager = S3UploadManager()

    def get(self):
        """
        List the AI Models run by Power Sage Pipeline.
        ---
        produces:
          - application/json
        responses:
          200:
            description: List of blades for monitoring
            schema:
              properties:
                id:
                  type: integer
                blade_type_id:
                  type: integer
                defect_model_id:
                  type: integer
                ai_enabled:
                  type: boolean

        """
        with Session(db) as session:
            aiDefectModel = session.scalars(select(BladeDefectModel)).all()
            if aiDefectModel is None:
                return {'message': f'No AI Models listed.'}
            
            unique_bladeType = set()
            for blade in aiDefectModel:
                if blade.blade_type_id not in unique_bladeType:
                    unique_bladeType.add(blade.blade_type_id)

            print("unique_bladeType  ", unique_bladeType)
            noAIBlades = session.scalars(select(BladeType).filter(BladeType.id.notin_(unique_bladeType))).all()
            print('frame_extractor_only ', noAIBlades)               
            frame_extractor_only = add_alias(noAIBlades)
            print('frame_extractor_only ', frame_extractor_only)            
            aiDefectModel_json = serialize_BladeDefectModelList(aiDefectModel)
            s3Json = generate_BladeDefectJSONForS3(aiDefectModel, frame_extractor_only)
            S3AIModelUploadAPI()
            return s3Json

def add_alias(input_array):

  output_array = []
  types = []

  for row in input_array:
      if row.alias:  
        types = [row.blade_type]+row.alias.split()
        output_array = output_array + types
      else:
        types = [row.blade_type]
        output_array = output_array + types

  return output_array
     
class DefectModelListAPI(Resource):

    def __init__(self):
        super(DefectModelListAPI, self).__init__()

    def get(self):
        """
        List the AI Models run by Power Sage Pipeline.
        ---
        produces:
          - application/json
        responses:
          200:
            description: List of blades for monitoring
            schema:
              properties:
                id:
                  type: integer
                blade_type_id:
                  type: integer
                defect_model_id:
                  type: integer
                ai_enabled:
                  type: boolean

        """
        with Session(db) as session:
            aiDefectModel = session.scalars(select(DefectModel)).all()
            if aiDefectModel is None:
                return {'message': f'No Defect Models listed.'}
            
            return serialize_DefectModelList(aiDefectModel)


class BladeTypeListAPI(Resource):

    def __init__(self):
        super(BladeTypeListAPI, self).__init__()

    def get(self):
        """
        List the AI Models run by Power Sage Pipeline.
        ---
        produces:
          - application/json
        responses:
          200:
            description: List of blades for monitoring
            schema:
              properties:
                id:
                  type: integer
                blade_type_id:
                  type: integer
                defect_model_id:
                  type: integer
                ai_enabled:
                  type: boolean

        """
        with Session(db) as session:
            aiBladeTypes = session.scalars(select(BladeType)).all()
            if aiBladeTypes is None:
                return {'message': f'No Blade Types listed.'}
            
            return serialize_BladeList(aiBladeTypes)

def serialize_BladeDefectModelList(bladeDefectModels):
    return [
        {
            "id": model.id,
            "blade_type_id": model.blade_type_id,
            "blade_type": model.bladeType.blade_type,
            "alias": model.bladeType.alias,
            "defect_model_id": model.defect_model_id, 
            "defect_model": model.defectModel.defect_name,             
            "ai_enabled": model.ai_enabled 
        }
        for model in bladeDefectModels
    ]

def serialize_BladeList(bladeTypeList):
    return [
        {
            "id": bladeType.id,
            "alias": bladeType.alias,
            "blade_type": bladeType.blade_type
        }
        for bladeType in bladeTypeList
    ]

def serialize_DefectModelList(defectModelList):
    return [
        {
            "id": defect.id,            
            "defect_name": defect.defect_name
        }
        for defect in defectModelList
    ]

def generate_BladeDefectJSONForS3(bladeDefectModels, noAIBlades):

    output = {"AiBlades": [], "FrameExtractorOnlyBlades":[], "VersionDate": datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
    blade_data = {}

    defect_model_mapping = {
        1: "Delamination",
        2: "AdhesiveVoids",
        3: "CoreGap",
        4: "DustAndDirt",
        5: "VoidOverlaminate",
        6: "ForeignObjectDetection",
        7: "RunWithMeasurement",
        8: "SemiDryGlass",
        9: "EntrainedAir"
    }

    for row in bladeDefectModels:
        blade_id = row.blade_type_id
        if blade_id not in blade_data:
            bladeNames = []
            if row.bladeType.alias:
                bladeNames = [row.bladeType.blade_type]+row.bladeType.alias.split()
            else:
                bladeNames = [row.bladeType.blade_type]
            blade_data[blade_id] = {
                "BladeName": bladeNames,
                "models": {name: False for name in defect_model_mapping.values()}
            }
        blade_data[blade_id]["models"][row.defectModel.defect_name] = row.ai_enabled

    # Convert processed data to list format for output
    output["AiBlades"] = list(blade_data.values())    
    output["FrameExtractorOnlyBlades"] = noAIBlades
    # Output JSON to file
    with open("./ai_pipeline_config.json", "w") as f:
        json.dump(output, f, indent=4)

    return output

def S3AIModelUploadAPI():
    
    # File details
    local_file_path = './ai_pipeline_config.json'  # Path to the local file
    region = os.getenv('AWS_REGION')
    bucket_name = os.getenv('S3_BUCKET_NAME')
    s3_base_folder = os.getenv('S3_BASE_FOLDER')
    s3_folder = f'{s3_base_folder}pipeline/ai_pipeline_config/'

    # Generate S3 object key
    file_name = local_file_path.split('/')[-1]
    s3_key = s3_folder + file_name

    try:
        print("local_file_path ", local_file_path)
        print("bucket_name ", bucket_name)
        print("s3_key ", s3_key)
        session = botocore.session.get_session()
        s3_client = session.create_client(
            's3',
            region_name=region,
            config=Config(
                signature_version='s3v4',
                s3={'use_accelerate_endpoint': True}
            )
        )
        url = s3_client.generate_presigned_url(
            "put_object",
            Params={"Bucket": bucket_name, "Key": s3_key},
            ExpiresIn=3600,
        )
        with open(local_file_path, 'rb') as file:
            response = httpx.put(url, data=file)        
            print(f"File uploaded successfully to s3://{bucket_name}/{s3_key} = {response}")

    except FileNotFoundError:
        print(f"File not found: {local_file_path}")
    except Exception as e:
        print(f"Error uploading file: {e}")    


class BladeDefectModelAddUpdateAPI(Resource):

    def __init__(self):
        super(BladeDefectModelAddUpdateAPI, self).__init__()

    def post(self):
        """
        Insert or Update AI Blade Defect Model
            schema:
              id: BladeType Id.        
        ---

        responses:
          200:
            description: AI Model Added.
            schema:
              id: BladeType Id.

        """

        data = request.json
        blade_type_id = data.get('bladeId')
        modelIds = data.get('modelIds')

        print('blade_type_id ', blade_type_id)
        print('modelIds ', modelIds)
        logging.info(f'BladeDefectModelAddUpdateAPI blade_type_id# ')

        with Session(db) as session:
            defectModelsList = session.execute(select(DefectModel)).scalars().all()  
            bladeDefectModelsList = session.execute(
                select(BladeDefectModel).filter(BladeDefectModel.blade_type_id == blade_type_id)
            ).scalars().all()                      
            print('bladeDefectModels ', bladeDefectModelsList)
            for bladeDefectModel in bladeDefectModelsList:
                session.delete(bladeDefectModel)
            print('else  ', bladeDefectModelsList)
            for model in defectModelsList:
                if model.id in modelIds:
                    newMapping = BladeDefectModel (
                        blade_type_id = blade_type_id,
                        defect_model_id = model.id,
                        ai_enabled = True
                    )
                else:
                    newMapping = BladeDefectModel (
                        blade_type_id = blade_type_id,
                        defect_model_id = model.id,
                        ai_enabled = False
                    )                                             
                session.add(newMapping)

            session.commit()            

        return {'message': f'AI Blade Defect Model is saved.'}
    
    def delete(self):
        """
        Delete AI Blade Defect Model
            schema:
              id: BladeType Id.        
        ---

        responses:
          200:
            description: AI Model Added.
            schema:
              id: BladeType Id.

        """

        data = request.json
        blade_type_id = data.get('bladeId')
        modelIds = data.get('modelIds')

        print('blade_type_id ', blade_type_id)
        logging.info(f'BladeDefectModelAddUpdateAPI delete# ')

        with Session(db) as session:
            bladeDefectModelsList = session.execute(
                select(BladeDefectModel).filter(BladeDefectModel.blade_type_id == blade_type_id)
            ).scalars().all()       

            for bladeDefectModel in bladeDefectModelsList:
                session.delete(bladeDefectModel)

            session.commit()

        return {'message': f'AI Blade Defect Model deleted.'}

class BladeDefectModelUpdateAPI(Resource):

    def __init__(self):
        super(BladeDefectModelUpdateAPI, self).__init__()

    def put(self):
        """
        Insert or Update AI Blade Defect Model
            schema:
              id: BladeType Id.        
        ---

        responses:
          200:
            description: AI Model Added.
            schema:
              id: BladeType Id.

        """

        bladeDefectModelsList = request.json
        print(bladeDefectModelsList)

        with Session(db) as session:
             for bladeDefectModel in bladeDefectModelsList:
                bladeDefect = session.query(BladeDefectModel).get(bladeDefectModel["id"])
                if bladeDefect is not None:
                    bladeDefect.ai_enabled = bladeDefectModel["ai_enabled"]

                    session.commit()
                    session.refresh(bladeDefect)


        return {'message': f'AI Blade Defect Model updated.'}

# --------------------------------- Filtered Reports generation -------------------------------------

def generate_filter(supplier, factory_name, location, blade_type, manufacture_stage, bladeCavity, start_date, end_date):
    filters = []
    if supplier:
        filters.append(Inspection.supplier.ilike(f"%{supplier}%"))
    if factory_name:
        filters.append(Inspection.factory_name.ilike(f"%{factory_name}%"))
    if location:
        filters.append(Inspection.location.ilike(f"%{location}%"))
    if blade_type:
        filters.append(Inspection.engine_type.ilike(f"%{blade_type}%"))
    if manufacture_stage:
        if manufacture_stage == "Final_Release_Inspection":
            filters.append(Inspection.manufacture_stage.ilike("Final%"))
        else:
            filters.append(Inspection.manufacture_stage.ilike(f"%{manufacture_stage}%"))
    if bladeCavity:
        filters.append(Inspection.sect.ilike(f"%{bladeCavity}%"))
    if start_date:
        if end_date:
            filters.append(Inspection.date.between(start_date, end_date))
        else:
            filters.append(Inspection.date >= start_date)
    elif end_date:
        filters.append(Inspection.date <= end_date)

    return filters

class InspectionDefectStatsReportAPI(Resource):
    def __init__(self):
        super(InspectionDefectStatsReportAPI, self).__init__()

    def inspection_defect_count(self, count_dict, defect_list):
        for defect in defect_list:
            finding_label = normalize_finding_label(defect.finding_type)
            label_suffix = " (M)" if defect.is_manual else " (AI)"
            finding_label += label_suffix
            disposition = get_disposition_key(defect.ge_disposition)
            disposition_finding_label = finding_label + ' - '+disposition
            count_dict.setdefault(disposition_finding_label, 0)
            count_dict[disposition_finding_label] += 1

    # count the defects of an inspection by disposition type, increment count_dict stats
    def inspection_disposition_count(self, count_dict, defect_list):
        for defect in defect_list:
            disposition_label = defect.ge_disposition or 'Total_Open_No_Disposition'
            if 'Out of' in disposition_label:
                disposition_label = 'Total_Repaired' if defect.status == 'Closed' else 'Total_Out_of_Tolerance'
            elif 'False' in disposition_label:
                disposition_label = 'Total_False_Positive'
            elif 'Within' in disposition_label:
                disposition_label = 'Total_Within_Tolerance'
            elif 'Duplicate' in disposition_label:
                disposition_label = 'Total_Duplicate'

            disposition_label += " (M)" if defect.is_manual else " (AI)"

            count_dict.setdefault(disposition_label, 0)
            count_dict[disposition_label] += 1

    def __get_defect_data(self, inspection, defect_disposition, defect_label_type):
        defectFilters = []
        if defect_disposition:
            defectFilters.append(Defect.ge_disposition.ilike(f"%{defect_disposition}%"))
        if defect_label_type:
            defectFilters.append(Defect.finding_type.ilike(f"%{defect_label_type}%"))

        with Session(db) as session:
            img_2d_shots_total = 0
            image_list = session.scalars(select(Image).where(Image.inspection_id == inspection.id)).all()
            defect_list = (session.scalars(session.query(Defect).join(Image, Image.id == Defect.image_id)
                                           .filter(Image.inspection_id == inspection.id, *defectFilters))).all()
            
            count_defects_dict = DEFECT_TOTAL_TEMPLATE_MANUAL_AI_DISPOSITION.copy()
            count_disposition_dict = DISPOSITION_TOTAL_TEMPLATE_MANUAL_AI.copy()

            if image_list:
                img_2d_shots_total = session.scalar(
                    select(func.count(MeasurementImageFile.id)).where(
                        MeasurementImageFile.image_id.in_(img.id for img in image_list))
                )

                self.inspection_defect_count(count_defects_dict, defect_list)
                self.inspection_disposition_count(count_disposition_dict, defect_list)

            row = [inspection.esn, inspection.blade_type, inspection.factory_name, get_section_name(
                inspection.sect), inspection.manufacture_stage, inspection.upload_date, inspection.date,
                inspection.sso, inspection.status, img_2d_shots_total]
            row.extend(count_defects_dict.values())
            row.extend(count_disposition_dict.values())
            return row

    def __run_generate_report(self, task_id, inspection_list, defect_disposition, defect_label_type):
        try:
            with tempfile.TemporaryDirectory(dir=UPLOAD_FOLDER, prefix='temp-', suffix='-folder') as tempdirname:
                output_filename = 'inspection_defect_stats_report.csv'

                # Handle async calls by saving to a unique directory
                tempdirname = TEMP_DOWNLOAD_FILES_DIR
                output_filename = task_id + '_' + output_filename

                csv_path = os.path.join(tempdirname, output_filename)
                logging.info(f'Creating report csv file: {csv_path}')
                
                with open(csv_path, 'w', newline='') as csvfile:
                    writer = csv.writer(csvfile)
                    # Create headers for the report
                    defect_headers = list(DEFECT_TOTAL_TEMPLATE_MANUAL_AI_DISPOSITION.keys())
                    disposition_headers = list(DISPOSITION_TOTAL_TEMPLATE_MANUAL_AI.keys())
                    headers = ['Blade ID', 'Blade Type', 'Factory Name', 'Blade Cavity', 'Manufacture Stage',
                               'Upload Date', 'Inspection Date', 'Inspector SSO', 'Annotation Status', '2D Images Count']
                    report_headers = headers + defect_headers + disposition_headers
                    writer.writerow(report_headers)
                    col_count = len(report_headers)

                    total_row = self.make_blank_row(col_count)
                    futures = []
                    
                    with ThreadPoolExecutor(THREAD_POOL_SIZE) as executor:
                        for inspection in inspection_list:
                            future = executor.submit(
                                self.__get_defect_data, inspection, defect_disposition, defect_label_type)
                            futures.append(future)
                        
                        for future in futures:
                            try:
                                row = future.result()
                                if row:
                                    self.add_total_row(total_row, row)
                                    writer.writerow(row)
                            except Exception as e:
                                logging.error(f'Error processing future: {e}')
                    
                    header = self.make_blank_row(col_count)
                    writer.writerow(header)

                    total_row[0] = "Total"
                    writer.writerow(total_row)

                logging.info(f'Done writing csv file: {csv_path}')

                TASK_STATUS_REGISTRY[task_id] = {
                    'id': task_id,
                    'timestamp': datetime.now().timestamp(),
                    'status': 'COMPLETE',
                    'filename': output_filename,
                    'path': os.path.abspath(csv_path)
                }
                logging.info(f'Set TASK_STATUS for task_id: {task_id} to: {TASK_STATUS_REGISTRY[task_id]}')
        except Exception as e:
            logging.error(f'Error generating report: {e}')
            TASK_STATUS_REGISTRY[task_id] = {
                'id': task_id,
                'timestamp': datetime.now().timestamp(),
                'status': 'FAILED',
                'filename': None,
                'path': None
            }
                
    def parse_date(self, date_str):
        try:
            return datetime.fromisoformat(date_str.replace("Z", "+00:00"))
        except ValueError:
            return None
    
    def make_blank_row(self, headers_length):
        row = headers_length * ['']
        return row
    
    def add_total_row(self, total_row, row):
        for idx, col in enumerate(row):
            col_str = str(col).strip()
            if col_str.isnumeric() and len(col_str) < 9:
                total_row[idx] = str(int(total_row[idx] or 0) + int(col_str))

    def post(self):
        """
        Generate a report of inspection defect stats
        ---
        parameters:
          - in: body
            required: true
            name: body
            schema:
              id: ReportRequest
              properties:
                supplier:
                  type: string
                  description: filter by supplier
                  example: "LM"
                factory_name:
                  type: string
                  description: filter by factory name
                  example: "VAD"
                location:
                  type: string
                  description: filter by location
                  example: "Vadodara"
                blade_type:
                  type: string
                  description: filter by blade type
                  example: "107"
                manufacture_stage:
                  type: string
                  description: filter by manufacture stage
                  example: "Final_Release_Inspection"
                bladeCavity:
                  type: string
                  description: filter by blade cavity type
                  example: "Central_Web"
                defect_disposition:
                  type: string
                  description: filter by defect disposition
                  example: "Duplicate"
                defect_label_type:
                  type: string
                  description: filter by defect label type
                  example: "Adhesive_Void"
                startDate:
                  type: string
                  description: filter by start date
                  example: "2024-01-01"
                endDate:
                  type: string
                  description: filter by end date
                  example: "2024-01-31"
                offset:
                  type: integer
                  description: offset of results
                  default: 0
                limit:
                  type: integer
                  description: number of results to return
                  default: 10000

        responses:
          202:
            description: Report generation task created
            schema:
              id: TaskStatus
              properties:
                id:
                  type: string
                  description: task id
                  example: "0f01f4a6-9f6a-4ad6-a90a-6eaf8f2e1b0a"
                status:
                  type: string
                  description: task status
                  example: "RUNNING"
        """
        data = request.json
        supplier = data.get('supplier')
        factory_name = data.get('factoryName')
        location = data.get('inspectionLocation')
        blade_type = data.get('bladeType')
        manufacture_stage = data.get('manufactureStage')
        bladeCavity = data.get('bladeCavityType')
        defect_disposition = data.get('defectDisposition')
        defect_label_type = data.get('defectLabelType')
        start_date = self.parse_date(data.get("startDate"))
        end_date = self.parse_date(data.get("endDate"))
        
        offset = data.get('offset', 0)
        limit = data.get('limit', 10000)

        # Construct filters dynamically based on non-null and non-empty values
        filters = generate_filter(supplier, factory_name, location, blade_type, manufacture_stage, bladeCavity, start_date, end_date)
        
        if filters:
            where_clause = reduce(lambda x, y: and_(x, y), filters) if len(filters) > 1 else filters[0]
        else:
            where_clause = True
        # Query the database with dynamic filters, offset, and limit
        with Session(db) as session:
            inspection_list = session.scalars(select(Inspection).where(where_clause).offset(offset).limit(limit)).all()

        logging.info(f'Number of inspections: {len(inspection_list)}')

        task_id = uuid.uuid4().hex
        TASK_STATUS_REGISTRY[task_id] = {
           'id': task_id,
           'timestamp': datetime.now().timestamp(),
           'status': 'RUNNING',
           'filename': None,
        }
        TASK_EXECUTOR.submit(self.__run_generate_report, task_id, inspection_list, defect_disposition, defect_label_type)
        return {'id': task_id,
               'status': 'RUNNING'}, 202

class InspectionListReportAPI(Resource):
    def __init__(self):
        super(InspectionListReportAPI, self).__init__()

    INSPECTION_ENTRY_TEMPLATE = {
    "customer_name": "",
    "app_type": "",
    "supplier": "",
    "factory_name": "",
    "location": "",
    "esn": "",
    "date": "",
    "blade_type": "",
    "engine_type": "",
    "upload_date": "",
    "sect": "",
    "manufacture_stage": "",
    "status": "",
    "inspector_name": "",
    "sso": ""
    }

    def parse_date(self, date_str):
        try:
            return datetime.fromisoformat(date_str.replace("Z", "+00:00"))
        except ValueError:
            return None

    def __run_generate_report_inspection_list(self, task_id, inspection_list, defect_disposition, defect_label_type):
        try:
            with tempfile.TemporaryDirectory(dir=UPLOAD_FOLDER, prefix="temp-", suffix="-folder") as tempdirname:
                output_filename = "inspection_list_report.csv"
                defectFilters = []
                if defect_disposition:
                    defectFilters.append(Defect.ge_disposition.ilike(f"%{defect_disposition}%"))
                if defect_label_type:
                    defectFilters.append(Defect.finding_type.ilike(f"%{defect_label_type}%"))

                tempdirname = TEMP_DOWNLOAD_FILES_DIR
                output_filename = task_id + "_" + output_filename

                csv_path = os.path.join(tempdirname, output_filename)

                with open(csv_path, "w", newline='') as csvfile:
                    writer = csv.writer(csvfile)
                    field_names = list(self.INSPECTION_ENTRY_TEMPLATE.keys())
                    headers = [
                        "blade_serial_number" if field == "esn"
                        else "factory_location" if field == "location"
                        else "scan_date" if field == "date"
                        else "inspection_status" if field == "status"
                        else "inspectior_SSO" if field == "sso"
                        else field
                        for field in field_names
                    ]

                    headers.extend(["number_of_defects", "certification_status", "certificate_number", "certification_date"])
                    writer.writerow(headers)

                    with Session(db) as session:
                        for inspection in inspection_list:
                            total_defects = (
                                session.query(func.count(Defect.id))
                                .join(Image)
                                .filter(
                                    Image.inspection_id == inspection.id,
                                    Image.id == Defect.image_id,
                                    *defectFilters
                                )
                                .scalar()
                            ) or 0

                            certificate_details = (
                                session.query(Certificate)
                                .filter(Certificate.id == inspection.certificate_id)
                                .first()
                            )

                            certification_status = "Complete" if certificate_details else "Incomplete"
                            certificate_number = certification_date = None
                            if certificate_details:
                                certificate_number = certificate_details.certificate_number
                                certification_date = certificate_details.certification_date

                            values_list = [getattr(inspection, field, None) for field in field_names]
                            values_list.extend([total_defects, certification_status, certificate_number, certification_date])
                            writer.writerow(values_list)

                TASK_STATUS_REGISTRY[task_id] = {
                    "id": task_id,
                    "timestamp": datetime.now().timestamp(),
                    "status": "COMPLETE",
                    "filename": output_filename,
                    "path": os.path.abspath(csv_path)
                }
        except Exception as e:
            logging.error(f"Error generating inspection list report: {e}")
            TASK_STATUS_REGISTRY[task_id] = {
                "id": task_id,
                "timestamp": datetime.now().timestamp(),
                "status": "FAILED",
                "filename": None,
                "path": None
            }

    def post(self):
        """
        Generate a report of inspection list
        ---
        parameters:
          - in: body
            required: true
            name: body
            schema:
              id: ReportRequest
              properties:
                supplier:
                  type: string
                  description: filter by supplier
                  example: "LM"
                factory_name:
                  type: string
                  description: filter by factory name
                  example: "VAD"
                location:
                  type: string
                  description: filter by location
                  example: "Vadodara"
                blade_type:
                  type: string
                  description: filter by blade type
                  example: "107"
                manufacture_stage:
                  type: string
                  description: filter by manufacture stage
                  example: "Final_Release_Inspection"
                bladeCavity:
                  type: string
                  description: filter by blade cavity type
                  example: "Central_Web"
                defect_disposition:
                  type: string
                  description: filter by defect disposition
                  example: "Duplicate"
                defect_label_type:
                  type: string
                  description: filter by defect label type
                  example: "Adhesive_Void"
                startDate:
                  type: string
                  description: filter by start date
                  example: "2024-01-01"
                endDate:
                  type: string
                  description: filter by end date
                  example: "2024-01-31"
                offset:
                  type: integer
                  description: offset of results
                  default: 0
                limit:
                  type: integer
                  description: number of results to return
                  default: 10000

        responses:
          202:
            description: Report generation task created
            schema:
              id: TaskStatus
              properties:
                id:
                  type: string
                  description: task id
                  example: "0f01f4a6-9f6a-4ad6-a90a-6eaf8f2e1b0a"
                status:
                  type: string
                  description: task status
                  example: "RUNNING"
        """
        data = request.json
        supplier = data.get("supplier")
        factory_name = data.get("factoryName")
        location = data.get("inspectionLocation")
        blade_type = data.get("bladeType")
        manufacture_stage = data.get("manufactureStage")
        bladeCavity = data.get("bladeCavityType")
        defect_disposition = data.get("defectDisposition")
        defect_label_type = data.get("defectLabelType")
        
        start_date_str = data.get("startDate")
        end_date_str = data.get("endDate")
        
        if start_date_str is None:
          start_date_str = "1970-01-01T00:00:00Z"
        if end_date_str is None:
          now = datetime.now()
          end_date_str = now.strftime("%Y-%m-%d %H:%M:%S")

        start_date = self.parse_date(start_date_str)    
        end_date = self.parse_date(end_date_str)

        offset = data.get("offset", 0)
        limit = data.get("limit", 10000)

        # Construct filters dynamically based on non-null and non-empty values
        filters = generate_filter(supplier, factory_name, location, blade_type, manufacture_stage, bladeCavity, start_date, end_date)

        if filters:
            where_clause = reduce(lambda x, y: and_(x, y), filters) if len(filters) > 1 else filters[0]
        else:
            where_clause = True
        # Query the database with dynamic filters, offset, and limit
        with Session(db) as session:
            inspection_list = session.scalars(select(Inspection).where(where_clause).offset(offset).limit(limit)).all()

        task_id = uuid.uuid4().hex
        TASK_STATUS_REGISTRY[task_id] = {
            "id": task_id,
            "timestamp": datetime.now().timestamp(),
            "status": "RUNNING",
            "filename": None,
        }
        TASK_EXECUTOR.submit(self.__run_generate_report_inspection_list, task_id, inspection_list, defect_disposition, defect_label_type)
        return {"id": task_id, "status": "RUNNING"}, 202

class BladeMonitoringReportAPI(Resource):
    def __init__(self):
        super(BladeMonitoringReportAPI, self).__init__()

    def parse_date(self, date_str):
        try:
            return datetime.fromisoformat(date_str.replace("Z", "+00:00"))
        except ValueError:
            return None
    
    def convert_to_gmt_iso(self, datetime_value):
        if isinstance(datetime_value, str):
            try:
                dt = datetime.fromisoformat(datetime_value)
            except ValueError:
                dt = datetime.strptime(datetime_value, '%Y-%m-%d %H:%M:%S.%f')
        elif isinstance(datetime_value, datetime):
            dt = datetime_value
        else:
            return None
        
        gmt_dt = dt.astimezone(timezone.utc)
        return gmt_dt.strftime('%Y-%m-%d %H:%M:%S GMT')

    def post(self):
        """
        Generate a report of blade monitoring data
        ---
        parameters:
          - in: body
            required: true
            name: body
            schema:
              id: ReportRequest
              properties:
                supplier:
                  type: string
                  description: filter by supplier
                  example: "LM"
                factory_name:
                  type: string
                  description: filter by factory name
                  example: "VAD"
                location:
                  type: string
                  description: filter by location
                  example: "Vadodara"
                blade_type:
                  type: string
                  description: filter by blade type
                  example: "107"
                manufacture_stage:
                  type: string
                  description: filter by manufacture stage
                  example: "Final_Release_Inspection"
                bladeCavity:
                  type: string
                  description: filter by blade cavity type
                  example: "Central_Web"
                startDate:
                  type: string
                  description: filter by start date
                  example: "2024-01-01"
                endDate:
                  type: string
                  description: filter by end date
                  example: "2024-01-31"
                offset:
                  type: integer
                  description: offset of results
                  default: 0
                limit:
                  type: integer
                  description: number of results to return
                  default: 10000

        responses:
          200:
            description: Report generated
        """
        data = request.json
        supplier = data.get("supplier")
        factory_name = data.get("factoryName")
        location = data.get("inspectionLocation")
        blade_type = data.get("bladeType")
        manufacture_stage = data.get("manufactureStage")
        bladeCavity = data.get("bladeCavityType")
        start_date = self.parse_date(data.get("startDate"))
        end_date = self.parse_date(data.get("endDate"))

        offset = data.get("offset", 0)
        limit = data.get("limit", 10000)

        filters = generate_filter(supplier, factory_name, location, blade_type, manufacture_stage, bladeCavity, start_date, end_date)

        if filters:
            where_clause = reduce(lambda x, y: and_(x, y), filters) if len(filters) > 1 else filters[0]
        else:
            where_clause = True
        # Query the database with dynamic filters, offset, and limit
        with Session(db) as session:
            inspection_list = session.scalars(select(Inspection).where(where_clause).offset(offset).limit(limit)).all()
            blade_monitoring_list = (
                session.query(BladeMonitoring)
                .filter(BladeMonitoring.blade_number.in_(inspection.esn for inspection in inspection_list))
                .all()
            )

        with tempfile.TemporaryDirectory(dir=UPLOAD_FOLDER, prefix="temp-", suffix="-folder") as tempdirname:
            tempdirname = TEMP_DOWNLOAD_FILES_DIR
            output_filename = 'blade_monitoring_list.csv'
            csv_path = os.path.join(tempdirname, output_filename)

            with open(csv_path, "w", newline='') as csvfile:
                fieldnames = ['blade_number', 'upload_date', 'upload_end_date', 'upload_status', 'file_size_in_GB', 'upload_time_taken_in_hrs', 'S3_input', 'S3_output', 'App_UI', 'annotation_start_time', 'annotation_end_time', 'annotation_total_time_in_hrs', 'cert_issue_date', 'total_time_in_hrs']
                writer = csv.DictWriter(csvfile, fieldnames=fieldnames)

                writer.writeheader()
                for blade_monitoring in sorted(blade_monitoring_list, key=lambda x: x.id, reverse=True):
                    writer.writerow({
                        'blade_number': blade_monitoring.blade_number,
                        'upload_date': self.convert_to_gmt_iso(blade_monitoring.upload_date),
                        'upload_end_date': self.convert_to_gmt_iso(blade_monitoring.upload_end_date),
                        'upload_status': blade_monitoring.upload_status,
                        'file_size_in_GB': blade_monitoring.file_size,
                        'upload_time_taken_in_hrs': blade_monitoring.upload_time_taken,
                        'S3_input': self.convert_to_gmt_iso(blade_monitoring.s3_cvpl_input),
                        'S3_output': self.convert_to_gmt_iso(blade_monitoring.s3_cvpl_output),
                        'App_UI': self.convert_to_gmt_iso(blade_monitoring.application_ui),
                        'annotation_start_time': self.convert_to_gmt_iso(blade_monitoring.annot_start_time),
                        'annotation_end_time': self.convert_to_gmt_iso(blade_monitoring.annot_end_time),
                        'annotation_total_time_in_hrs': blade_monitoring.annot_time_taken,
                        'cert_issue_date': self.convert_to_gmt_iso(blade_monitoring.cert_issued),
                        'total_time_in_hrs': blade_monitoring.total_time_taken
                    })
        return send_file(csv_path, as_attachment=True)

class InspectionDefectListReportAPI(Resource):
    def __init__(self):
        super(InspectionDefectListReportAPI, self).__init__()

    INSPECTION_DEFECT_HEADERS_TEMPLATE = {
    "id": "",
    "status": "",
    "esn": "",
    "location": "",
    "finding_type": "",
    "root_face_distance": "",
    "ge_disposition": "",
    "description": "",
    "disposition_provided_by": "",
    "repair_date": "",
    "repair_report_id": "",
    "repair_approved_by": "",
    "date": "",
    "sso": ""
    }

    def parse_date(self, date_str):
        try:
            return datetime.fromisoformat(date_str.replace("Z", "+00:00"))
        except ValueError:
            return None
    
    def __run_generate_report_insp_defect_list(self, task_id, inspection_list, defect_disposition, defect_label_type):
        try:
            with tempfile.TemporaryDirectory(dir=UPLOAD_FOLDER, prefix="temp-", suffix="-folder") as tempdirname:
                output_filename = "inspection_defect_list_report.csv"
                defectFilters = []
                if defect_disposition:
                    defectFilters.append(Defect.ge_disposition.ilike(f"%{defect_disposition}%"))
                if defect_label_type:
                    defectFilters.append(Defect.finding_type.ilike(f"%{defect_label_type}%"))

                tempdirname = TEMP_DOWNLOAD_FILES_DIR
                output_filename = task_id + "_" + output_filename

                csv_path = os.path.join(tempdirname, output_filename)

                with open(csv_path, "w", newline='') as csvfile:
                    writer = csv.writer(csvfile)
                    field_names = list(self.INSPECTION_DEFECT_HEADERS_TEMPLATE.keys())
                    headers = [
                        "blade_serial_number" if field == "esn"
                        else "defect_id" if field == "id"
                        else "z_distance" if field == "root_face_distance"
                        else "disposition" if field == "ge_disposition"
                        else "cavity" if field == "location"
                        else field
                        for field in field_names
                    ]
                    writer.writerow(headers)

                    with Session(db) as session:
                        for inspection in inspection_list:
                            defects = session.query(
                                Defect.id,
                                Defect.location,
                                Defect.finding_type,
                                Defect.root_face_distance,
                                Defect.ge_disposition,
                                Defect.disposition_provided_by,
                                Defect.status,
                                Defect.repair_date,
                                Defect.repair_report_id,
                                Defect.repair_approved_by,
                                Defect.description,
                                Defect.date,
                                Defect.sso
                            ).join(Image).filter(
                                Image.inspection_id == inspection.id,
                                Image.id == Defect.image_id,
                                *defectFilters
                            ).all()

                            for defect in defects:
                                values_list = [getattr(defect, field, None) for field in field_names]
                                values_list[field_names.index("esn")] = inspection.esn
                                writer.writerow(values_list)

                TASK_STATUS_REGISTRY[task_id] = {
                    "id": task_id,
                    "timestamp": datetime.now().timestamp(),
                    "status": "COMPLETE",
                    "filename": output_filename,
                    "path": os.path.abspath(csv_path)
                }
        except Exception as e:
            logging.error(f"Error generating inspection defect list report: {e}")
            TASK_STATUS_REGISTRY[task_id] = {
                "id": task_id,
                "timestamp": datetime.now().timestamp(),
                "status": "FAILED",
                "filename": None,
                "path": None
            }

    def post(self):
        """
        Generate a report of inspection defects
        ---
        parameters:
          - in: body
            required: true
            name: body
            schema:
              id: ReportRequest
              properties:
                supplier:
                  type: string
                  description: filter by supplier
                  example: "LM"
                factory_name:
                  type: string
                  description: filter by factory name
                  example: "VAD"
                location:
                  type: string
                  description: filter by location
                  example: "Vadodara"
                blade_type:
                  type: string
                  description: filter by blade type
                  example: "107"
                manufacture_stage:
                  type: string
                  description: filter by manufacture stage
                  example: "Final_Release_Inspection"
                bladeCavity:
                  type: string
                  description: filter by blade cavity type
                  example: "Central_Web"
                defect_disposition:
                  type: string
                  description: filter by defect disposition
                  example: "Duplicate"
                defect_label_type:
                  type: string
                  description: filter by defect label type
                  example: "Adhesive_Void"
                startDate:
                  type: string
                  description: filter by start date
                  example: "2024-01-01"
                endDate:
                  type: string
                  description: filter by end date
                  example: "2024-01-31"
                offset:
                  type: integer
                  description: offset of results
                  default: 0
                limit:
                  type: integer
                  description: number of results to return
                  default: 10000

        responses:
          200:
            description: Report generated
            schema:
              id: ReportResponse
              properties:
                inspection_id:
                  type: integer
                  description: id of the inspection
                defect_list:
                  type: array
                  description: list of defects
                  items:
                    type: object
                    properties:
                      id:
                        type: integer
                        description: id of the defect
                      location:
                        type: string
                        description: location of the defect
                      finding_type:
                        type: string
                        description: type of the defect
                      root_face_distance:
                        type: float
                        description: distance of the defect from root face
                      disposition:
                        type: string
                        description: disposition of the defect
                      repair_date:
                        type: string
                        description: date of the repair
                      repair_report_id:
                        type: string
                        description: id of the repair report
                      repair_approved_by:
                        type: string
                        description: who approved the repair
                      description:
                        type: string
                        description: description of the defect
                      date:
                        type: string
                        description: date of the defect
        """
        data = request.json
        supplier = data.get("supplier")
        factory_name = data.get("factoryName")
        location = data.get("inspectionLocation")
        blade_type = data.get("bladeType")
        manufacture_stage = data.get("manufactureStage")
        bladeCavity = data.get("bladeCavityType")
        defect_disposition = data.get("defectDisposition")
        defect_label_type = data.get("defectLabelType")
        
        start_date_str = data.get("startDate")
        end_date_str = data.get("endDate")
        
        if start_date_str is None:
          start_date_str = "1970-01-01T00:00:00Z"
        if end_date_str is None:
          now = datetime.now()
          end_date_str = now.strftime("%Y-%m-%d %H:%M:%S")

        start_date = self.parse_date(start_date_str)    
        end_date = self.parse_date(end_date_str)

        offset = data.get("offset", 0)
        limit = data.get("limit", 10000)

        # Construct filters dynamically based on non-null and non-empty values
        filters = generate_filter(supplier, factory_name, location, blade_type, manufacture_stage, bladeCavity, start_date, end_date)

        if filters:
            where_clause = reduce(lambda x, y: and_(x, y), filters) if len(filters) > 1 else filters[0]
        else:
            where_clause = True
        # Query the database with dynamic filters, offset, and limit
        with Session(db) as session:
            inspection_list = session.scalars(select(Inspection).where(where_clause).offset(offset).limit(limit)).all()

        task_id = uuid.uuid4().hex
        TASK_STATUS_REGISTRY[task_id] = {
            "id": task_id,
            "timestamp": datetime.now().timestamp(),
            "status": "RUNNING",
            "filename": None,
        }
        TASK_EXECUTOR.submit(self.__run_generate_report_insp_defect_list, task_id, inspection_list, defect_disposition, defect_label_type)
        return {"id": task_id, "status": "RUNNING"}, 202



# ========================== Service Now Integration ============================

class SNServiceRequestAPI(Resource):

    def __init__(self):
        super(SNServiceRequestAPI, self).__init__()

    def post(self):
        """
        Integrating with Service Now and create a service request. 
        ---
        parameters:
          - usersso: SSO of user raising the request.
          - shortdesc: Short description of the request. 
          - desc: Description of the request. 

        responses:
          200:
            "result": 
                "table": Table
                "number": Service Request Number
                "request_url": Service Now Request URL
                "cmdb_ci": Digital Blade Certification
                "assignment_group": Service Request Assigned to group 

        """

        data = request.json
        return create_service_request(data)


class SNIncidentCreateAPI(Resource):

    def __init__(self):
        super(SNIncidentCreateAPI, self).__init__()

    def post(self):
        """
        Integrating with Service Now and create a incident request. 
        ---
        parameters:
          - usersso: SSO of user raising the request.
          - shortdesc: Short description of the request. 
          - desc: Description of the request. 

        responses:
          200:
            "result": 
                "table": Table
                "number": Service Request Number
                "request_url": Service Now Request URL
                "cmdb_ci": Digital Blade Certification
                "assignment_group": Service Request Assigned to group 

        """

        data = request.json
        return create_service_incident(data)

class SNIncidentGetAPI(Resource):

    def __init__(self):
        super(SNIncidentGetAPI, self).__init__()

    def get(self, sso):
        """
        Request to get all incidents raised by the user with sso. 
        """

        logging.info(f'Get all the tickets for user {sso}')
        return get_all_service_incident(sso)

class SNTaskGetAPI(Resource):

    def __init__(self):
        super(SNTaskGetAPI, self).__init__()

    def get(self, sso):
        """
        Request to get all Tasks raised by the user with sso. 
        """

        logging.info(f'Get all the tickets for user {sso}')
        return get_all_service_tasks(sso)

# ========================== Input File Search & Download ============================

class InputFileURLsAPI(Resource):

    def __init__(self):
        super(InputFileURLsAPI, self).__init__()

    def get(self, blade_number):
        """
        Searching the AWS S3 for the blade_number. 
        """

        logging.info(f'Searching the AWS S3 for the blade_number.  {blade_number}')
        return download_s3_input_files(blade_number)

class InputFileSearchAPI(Resource):

    def __init__(self):
        super(InputFileSearchAPI, self).__init__()

    def get(self):
        """
        Request to get all Blade List. 
        """

        logging.info(f'Get all Blade List.')
        bladeList = []
        with Session(db) as session:
            if IDMGroupAPI.esn_prefix and all('tpi'.upper() in prfx for prfx in map(str.upper, IDMGroupAPI.esn_prefix)):
                bladeList = session.scalars(select(BladeMonitoring.blade_number)
                 .filter(BladeMonitoring.blade_number.ilike('tpi-%'))
                 .order_by(BladeMonitoring.s3_cvpl_output.desc())).all()
            else:
                bladeList = session.scalars(
                    select(BladeMonitoring.blade_number).order_by(BladeMonitoring.s3_cvpl_output.desc())).all()
            if bladeList is None:
                return {'message': f'No Blades in monitoring.'}

        obj_dict = {"blades": bladeList}
        return obj_dict

